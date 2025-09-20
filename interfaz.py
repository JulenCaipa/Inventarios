import customtkinter as ctk
import customtkinter
import tkinter
import traceback
from tkinter import messagebox
import os
import docx
import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import Inches, Cm
#___________________________
from PIL import ImageTk, Image
import tkinter.ttk as ttk
customtkinter.set_appearance_mode("Dark")
customtkinter.set_default_color_theme("blue")

ruta_assets=('Assets')
def set_opacity(window, alpha):
    #Check if the system supports window transparency
     if window.winfo_ismapped():
        window.attributes("-alpha", alpha)
def mostras_Ariba():
    global Matricula_Garaje_,Chip_Garaje_,Numero_Garaje_,Deposito,Numero_Deposito_
    global nombre_deudor_1_,tipo_di_deudor_1_no__,Numero_id_deudor1_,di_deudor_1_ciudad_expedicion_di_deudor_1_,direccion_deudor_1_,telefono_deudor_1_,correo_deudor_1_,Numero_Deudores
    global nombre_deudor_2_,tipo_di_deudor_2_no__,Numero_id_deudor2_,di_deudor_2_ciudad_expedicion_di_deudor_2_,direccion_deudor_2_,telefono_deudor_2_,correo_deudor_2_
    global nombre_deudor_3_,tipo_di_deudor_3_no__,Numero_id_deudor3_,di_deudor_3_ciudad_expedicion_di_deudor_2_,direccion_deudor_3_,telefono_deudor_3_,correo_deudor_3_,Nombre_Archivo
    print("Deposito",Deposito)
    Surcursal_letra=Sucursal_.get()
    Surcursal2_letra=Representante.get()
    Respuesta1=Garaje.get()
    Respuesta3=Admin.get()
    depo=Deposito.get()
    Direccion_Inmueble=Direccion.get()
    Apartamento_y_torre=Apartamento_Torre.get()
    Nombre_edificio=NOMBRE_EDIFICIO.get()
    Ciudad=Ciudad_.get()
    valor_canon=Valor_Canon_.get()
    Canon_Letras=Valor_Canon_letra.get()
    valor_administracion=Valor_Admin_des.get()
    Administracion_Letras=Valor_Admin_des_letra.get()
    valor_total=valor_total_.get()
    valor_total_letra=valor_total_letras.get()
    numero_matricula=numero_matricula_.get()
    Chip=Chip_.get()
    valor_pleno_administracion=valor_pleno_administracion_.get()
    administracion_plena=administracion_plena_.get()
    nombre_arrendatario=nombre_arrendatario_.get()
    tipodiarrendatariono_=tipodiarrendatariono__.get()
    Numero_id_arrendatario=Numero_id_arrendatario_.get()
    ciudadexpediciondiarrendatario=ciudadexpediciondiarrendatario_.get()
    direccion_arrendatario =direccion_arrendatario_.get()
    telefono_arrendatario =telefono_arrendatario_.get()
    correo_arrendatario =correo_arrendatario_.get() 
    Cuenta_Acueducto=Cuenta_Acueducto_.get()
    Cuenta_Alcatarillado=Cuenta_Alcatarillado_.get()
    Cuenta_Basuras=Cuenta_Basuras_.get()
    Cuenta_Energia=Cuenta_Energia_.get()
    Cuenta_Gas=Cuenta_Gas_.get()
    Rta=Rta_.get()
    if 'Numero_Deposito_' in globals() and hasattr(Numero_Deposito_, 'get'):
        Numero_deposito = Numero_Deposito_.get()
    else:
        Numero_deposito = 0
    print("aribas_adentro",Rta)
#__________________________________________________________________________________________________________________
    try:
        if Surcursal_letra=='Sucursal salitre':
            Surcursal=1
            Matricula_Mercantil_Promotora="01941761"
            direccion_promotora="Calle 116 #23-06 oficina 307 edificio Bussines 116"
            telefono_promotora="47385114 ,3143589973 "
            correo_promotora="Salitre@rgimobiliaria.com.co"
        elif Surcursal_letra=="Sucursal Lago":
            Surcursal=2
            Matricula_Mercantil_Promotora="01305015"
            direccion_promotora="Calle 116 #23-06 oficina 314 edificio Bussines 116"
            telefono_promotora="6367050 ,3006777830"
            correo_promotora="Lago@rgimobiliaria.com.co" 
        elif Surcursal_letra=="Sucursal Unicentro ":
            Surcursal=3
            Matricula_Mercantil_Promotora="00611238"
            direccion_promotora="Calle 116 #23-06 oficina 313 edificio Bussines 116"
            telefono_promotora="2155303 ,3212333298 "
            correo_promotora="Unicentro@rgimobiliaria.com.co"
        elif Surcursal_letra=="Sucursal multicentro":
            Surcursal=4
            Matricula_Mercantil_Promotora="02896739"
            direccion_promotora="Calle 116 #23-06 oficina 314 edificio Bussines 116"
            telefono_promotora="3057737008 "
            correo_promotora="Multicentro@rgimobiliaria.com.co"
        print("Surcursal:",Surcursal)
        if Surcursal2_letra=='FREDDY HERNANDO GUERRERO RIVERA':
            Surcursal2=1
            Nombre_Promotora="FREDY HERNANDO GUERRERO RIVERA"
            cc_rep_promotora="79.509.071 "
        elif Surcursal2_letra=="NELSON ALEJANDRO GUERRERO RIVERA":
            Surcursal2=2
            Nombre_Promotora="NELSON ALEJANDRO GUERRERO RIVERA"
            cc_rep_promotora="79.946.997 "
        elif Surcursal2_letra=="CINDY JOHANNA GUERRERO RIVERA":
            Surcursal2=3
            Nombre_Promotora="CINDY JOHANNA GUERRERO RIVERA "
            cc_rep_promotora="1.030.552.945 "
        print("Representasnte:",Surcursal2)
        if Respuesta1== "No":
            Matricula_Garaje=("")
            Chip_Garaje=("")
            Numero_Garaje=("") 
        else:
            Matricula_Garaje=Matricula_Garaje_.get()
            Chip_Garaje=Chip_Garaje_.get()
            Numero_Garaje=Numero_Garaje_.get()
        print("Matricula:",Matricula_Garaje) 
        if tipodiarrendatariono_=="NIT"or tipodiarrendatariono_=="Nit" or tipodiarrendatariono_=="nit":
            Arrendatario_Empresa="si"
        else:
            Arrendatario_Empresa="no" 
#____________________________________

        if Rta== "No":
            print("entro No")
            Numero_Deudores=()
            print("Numero de deudores",Numero_Deudores)
            x1="se obliga"
            x2=""
            
        elif Rta== "Si" :
            x1="y los deudores se obligan"
            x2="DEUDORES SOLIDARIOS"
            Numero_Deudores_letra=Numero_Deudores_.get()
            print("entro",Numero_Deudores_letra)
            if Numero_Deudores_letra=="1":
                Numero_Deudores=1
                nombre_deudor_1 =nombre_deudor_1_.get()
                print("1 deudor",nombre_deudor_1)
                tipo_di_deudor_1_no_=tipo_di_deudor_1_no__.get()
                Numero_id_deudor1=Numero_id_deudor1_.get()
                di_deudor_1_ciudad_expedicion_di_deudor_1=di_deudor_1_ciudad_expedicion_di_deudor_1_.get()
                direccion_deudor_1 =direccion_deudor_1_.get()
                telefono_deudor_1 =telefono_deudor_1_.get()
                correo_deudor_1 =correo_deudor_1_.get()
                ##Array
                Deudor_nombre=[nombre_deudor_1]
                Deudor_di=[tipo_di_deudor_1_no_]
                Deudor_idnumm=[ Numero_id_deudor1]
                Deudor_ciudad=[di_deudor_1_ciudad_expedicion_di_deudor_1]
                Deudor_direccion=[direccion_deudor_1]
                Deudor_telefono=[ telefono_deudor_1]
                Deudor_correo=[correo_deudor_1 ]
                print("Salio Bien aca")


            elif Numero_Deudores_letra=="2":
                Numero_Deudores=2
                nombre_deudor_1 =nombre_deudor_1_.get()
                tipo_di_deudor_1_no_=tipo_di_deudor_1_no__.get()
                Numero_id_deudor1=Numero_id_deudor1_.get()
                di_deudor_1_ciudad_expedicion_di_deudor_1=di_deudor_1_ciudad_expedicion_di_deudor_1_.get()
                direccion_deudor_1 =direccion_deudor_1_.get()
                telefono_deudor_1 =telefono_deudor_1_.get()
                correo_deudor_1 =correo_deudor_1_.get()

                nombre_deudor_2 =nombre_deudor_2_.get()
                tipo_di_deudor_2_no_=tipo_di_deudor_2_no__.get()
                Numero_id_deudor2=Numero_id_deudor2_.get()
                di_deudor_2_ciudad_expedicion_di_deudor_2=di_deudor_2_ciudad_expedicion_di_deudor_2_.get()
                direccion_deudor_2 =direccion_deudor_2_.get()
                telefono_deudor_2 =telefono_deudor_2_.get()
                correo_deudor_2 =correo_deudor_2_.get()
                ##Array
                Deudor_nombre=[nombre_deudor_1," y ",nombre_deudor_2]
                Deudor_di=[tipo_di_deudor_1_no_," y ",tipo_di_deudor_2_no_]
                Deudor_idnumm=[ Numero_id_deudor1," y ",Numero_id_deudor2]
                Deudor_ciudad=[di_deudor_1_ciudad_expedicion_di_deudor_1," y ",di_deudor_2_ciudad_expedicion_di_deudor_2]
                Deudor_direccion=[direccion_deudor_1," y ",direccion_deudor_2]
                Deudor_telefono=[ telefono_deudor_1," y ",telefono_deudor_2]
                Deudor_correo=[correo_deudor_1 ," y ",correo_deudor_2]
            elif Numero_Deudores_letra=="3":
                Numero_Deudores=3
                nombre_deudor_1 =nombre_deudor_1_.get()
                tipo_di_deudor_1_no_=tipo_di_deudor_1_no__.get()
                Numero_id_deudor1=Numero_id_deudor1_.get()
                di_deudor_1_ciudad_expedicion_di_deudor_1=di_deudor_1_ciudad_expedicion_di_deudor_1_.get()
                direccion_deudor_1 =direccion_deudor_1_.get()
                telefono_deudor_1 =telefono_deudor_1_.get()
                correo_deudor_1 =correo_deudor_1_.get()

                nombre_deudor_2 =nombre_deudor_2_.get()
                tipo_di_deudor_2_no_=tipo_di_deudor_2_no__.get()
                Numero_id_deudor2=Numero_id_deudor2_.get()
                di_deudor_2_ciudad_expedicion_di_deudor_2=di_deudor_2_ciudad_expedicion_di_deudor_2_.get()
                direccion_deudor_2 =direccion_deudor_2_.get()
                telefono_deudor_2 =telefono_deudor_2_.get()
                correo_deudor_2 =correo_deudor_2_.get()

                nombre_deudor_3 =nombre_deudor_3_.get()
                tipo_di_deudor_3_no_=tipo_di_deudor_3_no__.get()
                Numero_id_deudor3=Numero_id_deudor3_.get()
                di_deudor_3_ciudad_expedicion_di_deudor_2=di_deudor_3_ciudad_expedicion_di_deudor_2_.get()
                direccion_deudor_3 =direccion_deudor_3_.get()
                telefono_deudor_3 =telefono_deudor_3_.get()
                correo_deudor_3 =correo_deudor_3_.get()
                    ##Array
                Deudor_nombre=[nombre_deudor_1," , ",nombre_deudor_2,"y ",nombre_deudor_3]
                Deudor_di=[tipo_di_deudor_1_no_," , ",tipo_di_deudor_2_no_,"y ",tipo_di_deudor_3_no_]
                Deudor_idnumm=[ Numero_id_deudor1," , ",Numero_id_deudor2,"y ",Numero_id_deudor3]
                Deudor_ciudad=[di_deudor_1_ciudad_expedicion_di_deudor_1," , ",di_deudor_2_ciudad_expedicion_di_deudor_2,"y ",di_deudor_3_ciudad_expedicion_di_deudor_2]
                Deudor_direccion=[direccion_deudor_1," , ",direccion_deudor_2,"y ",direccion_deudor_3]
                Deudor_telefono=[ telefono_deudor_1," , ",telefono_deudor_2,"y ",telefono_deudor_3]
                Deudor_correo=[correo_deudor_1 ," , ",correo_deudor_2,"y ",correo_deudor_3]

        Numero_de_Contrato=Numero_de_Contrato_.get()
        Fecha_inicio_Contrato=Fecha_inicio_Contrato_.get()
        Fecha_Firma_Contrato=Fecha_Firma_Contrato_.get()
        Vigencia_Contrato=Vigencia_Contrato_.get()
        Clausula=Clausula_.get()
        print("salio bien otra vez")
#Codigo word_____________________________________________

        documentoObjeto=docx.Document()


        #Fuente docmento
        style = documentoObjeto.styles['Normal']
        style.font.name = 'Arial Narrow'
        style.font.size = Pt(11)

        #Margenes
        sections = documentoObjeto.sections
        for section in sections:
            section.top_margin = Cm(3)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(3)

        #Sangria
        style.paragraph_format.right_indent = Cm(-0.66)
            
        #Titulo1
        if(Surcursal==1):
            parrafo = documentoObjeto.add_paragraph('')
            parrafo .add_run('CONTRATO DE ARRENDAMIENTO DE VIVIENDA URBANA SS\n').bold = True
            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif(Surcursal==2):
            parrafo = documentoObjeto.add_paragraph('')
            parrafo .add_run('CONTRATO DE ARRENDAMIENTO DE VIVIENDA URBANA SL\n').bold = True
            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif(Surcursal==3):
            parrafo = documentoObjeto.add_paragraph('')
            parrafo .add_run('CONTRATO DE ARRENDAMIENTO DE VIVIENDA URBANA SU\n').bold = True
            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif(Surcursal==4):
            parrafo = documentoObjeto.add_paragraph('')
            parrafo .add_run('CONTRATO DE ARRENDAMIENTO DE VIVIENDA URBANA SM\n').bold = True
            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in parrafo.runs:
            run.font.size=Pt(11)
        p0 = documentoObjeto.add_paragraph('')
        p0.add_run('DIRECCIÓN DEL INMUEBLE           : ').bold = True 
        p0.add_run( Direccion_Inmueble+"  "+ Apartamento_y_torre + "  "+Nombre_edificio+", "+Ciudad+"\n\n")
        # documentoObjeto.add_paragraph('ARRENDADOR/ADMINISTRADOR: PROMOTORA INMOBILIARIA R&G S.A.S. \n''IDENTIFICACIÓN		: N.I.T.: 800.239.928-9'+'REPRESENTANTE LEGAL	: ' +Nombre_Promotora+'\n'+ 'IDENTIFICACIÓN		: '+cc_rep_promotora+'\n'+'TELÉFONO DE CONTACTO	: '+ telefono_promotora+'\n'+'CORREO ELECTRÓNICO	 	: ' +correo_promotora)
        p0.add_run('ARRENDADOR/ADMINISTRADOR :').bold = True 
        p0.add_run('PROMOTORA INMOBILIARIA R&G S.A.S.\n')
        p0.add_run('IDENTIFICACIÓN                              : ').bold = True
        p0.add_run('N.I.T.: 800.239.928-9 \n')
        p0.add_run('REPRESENTANTE LEGAL              : ').bold = True
        p0.add_run(Nombre_Promotora+'\n')
        p0.add_run('IDENTIFICACIÓN                             : ').bold = True
        p0.add_run(' C.C. '+ cc_rep_promotora+' BOGOTA D.C\n')
        p0.add_run('DIRECCIÓN DE CONTACTO           : ').bold = True
        p0.add_run(direccion_promotora+'\n')
        p0.add_run('TELÉFONO DE CONTACTO           : ').bold = True
        p0.add_run(telefono_promotora+'\n')
        p0.add_run('CORREO ELECTRÓNICO               : ').bold = True
        p0.add_run(correo_promotora+"\n\n")
        p0.add_run('ARRENDATARIO                             : ').bold = True
        p0.add_run(nombre_arrendatario+'\n')
        p0.add_run('IDENTIFICACIÓN                             : ').bold = True
        p0.add_run(tipodiarrendatariono_+' '+Numero_id_arrendatario)
        p0.add_run(' '+ciudadexpediciondiarrendatario+'\n')
        p0.add_run('DIRECCIÓN DE CONTACTO           : ').bold = True
        p0.add_run(direccion_arrendatario+"\n")
        p0.add_run('TELÉFONO DE CONTACTO           : ').bold = True
        p0.add_run(telefono_arrendatario+'\n')
        p0.add_run('CORREO ELECTRÓNICO               : ').bold = True
        p0.add_run( correo_arrendatario+'\t\n\n')

        if Rta== "No" or Rta=="no"or Rta=="NO"  :
            print("Salio Noooo")
        
        elif Rta== "si" or Rta=="Si"or Rta=="SI"  :
            if Numero_Deudores==1:
                p0.add_run('DEUDOR SOLIDARIO                     : ').bold = True
                p0.add_run(nombre_deudor_1+"\n")
                p0.add_run('IDENTIFICACIÓN                             : ').bold = True
                p0.add_run(tipo_di_deudor_1_no_+' '+Numero_id_deudor1)
                p0.add_run(" "+di_deudor_1_ciudad_expedicion_di_deudor_1+'\n')
                p0.add_run('DIRECCIÓN DE CONTACTO           : ').bold = True
                p0.add_run(direccion_deudor_1+'\n')
                p0.add_run('TELÉFONO DE CONTACTO           : ').bold = True
                p0.add_run(telefono_deudor_1+"\n")
                p0.add_run('CORREO ELECTRÓNICO               : ').bold = True
                p0.add_run(correo_deudor_1+'\t\n\n')
            elif Numero_Deudores==2:

                p0.add_run('DEUDOR SOLIDARIO                     : ').bold = True
                p0.add_run(nombre_deudor_1+"\n")
                p0.add_run('IDENTIFICACIÓN                             : ').bold = True
                p0.add_run(tipo_di_deudor_1_no_+' '+Numero_id_deudor1)
                p0.add_run(" "+di_deudor_1_ciudad_expedicion_di_deudor_1+'\n')
                p0.add_run('DIRECCIÓN DE CONTACTO           : ').bold = True
                p0.add_run(direccion_deudor_1+'\n')
                p0.add_run('TELÉFONO DE CONTACTO           : ').bold = True
                p0.add_run(telefono_deudor_1+"\n")
                p0.add_run('CORREO ELECTRÓNICO               : ').bold = True
                p0.add_run(correo_deudor_1+'\t\n\n')

                p0.add_run('DEUDOR SOLIDARIO                     : ').bold = True
                p0.add_run(nombre_deudor_2+"\n")
                p0.add_run('IDENTIFICACIÓN                             : ').bold = True
                p0.add_run(tipo_di_deudor_2_no_+' '+Numero_id_deudor2)
                p0.add_run(" "+di_deudor_2_ciudad_expedicion_di_deudor_2+'\n')
                p0.add_run('DIRECCIÓN DE CONTACTO           : ').bold = True
                p0.add_run(direccion_deudor_2+'\n')
                p0.add_run('TELÉFONO DE CONTACTO           : ').bold = True
                p0.add_run(telefono_deudor_2+"\n")
                p0.add_run('CORREO ELECTRÓNICO               : ').bold = True
                p0.add_run(correo_deudor_2+'\t\n\n')
            elif Numero_Deudores==3:

                p0.add_run('DEUDOR SOLIDARIO                     : ').bold = True
                p0.add_run(nombre_deudor_1+"\n")
                p0.add_run('IDENTIFICACIÓN                             : ').bold = True
                p0.add_run(tipo_di_deudor_1_no_+' '+Numero_id_deudor1)
                p0.add_run(" "+di_deudor_1_ciudad_expedicion_di_deudor_1+'\n')
                p0.add_run('DIRECCIÓN DE CONTACTO           : ').bold = True
                p0.add_run(direccion_deudor_1+'\n')
                p0.add_run('TELÉFONO DE CONTACTO           : ').bold = True
                p0.add_run(telefono_deudor_1+"\n")
                p0.add_run('CORREO ELECTRÓNICO               : ').bold = True
                p0.add_run(correo_deudor_1+'\t\n\n')

                p0.add_run('DEUDOR SOLIDARIO                     : ').bold = True
                p0.add_run(nombre_deudor_2+"\n")
                p0.add_run('IDENTIFICACIÓN                             : ').bold = True
                p0.add_run(tipo_di_deudor_2_no_+' '+Numero_id_deudor2)
                p0.add_run(" "+di_deudor_2_ciudad_expedicion_di_deudor_2+'\n')
                p0.add_run('DIRECCIÓN DE CONTACTO           : ').bold = True
                p0.add_run(direccion_deudor_2+'\n')
                p0.add_run('TELÉFONO DE CONTACTO           : ').bold = True
                p0.add_run(telefono_deudor_2+"\n")
                p0.add_run('CORREO ELECTRÓNICO               : ').bold = True
                p0.add_run(correo_deudor_2+'\t\n\n')

                p0.add_run('DEUDOR SOLIDARIO                     : ').bold = True
                p0.add_run(nombre_deudor_3+"\n")
                p0.add_run('IDENTIFICACIÓN                             : ').bold = True
                p0.add_run(tipo_di_deudor_3_no_+' '+Numero_id_deudor3)
                p0.add_run(" "+di_deudor_3_ciudad_expedicion_di_deudor_2+'\n')
                p0.add_run('DIRECCIÓN DE CONTACTO           : ').bold = True
                p0.add_run(direccion_deudor_3+'\n')
                p0.add_run('TELÉFONO DE CONTACTO           : ').bold = True
                p0.add_run(telefono_deudor_3+"\n")
                p0.add_run('CORREO ELECTRÓNICO               : ').bold = True
                p0.add_run(correo_deudor_3+'\t\n\n')

        p0.add_run('VALOR CANON                               : ').bold = True
        p0.add_run(valor_canon+"\n") 
        p0.add_run('VALOR ADMINISTRACIÒN             : ').bold = True
        p0.add_run(valor_administracion+'\n')
        p0.add_run('TOTAL                                              : ').bold = True
        p0.add_run(valor_total+'\t\n')

        #Titulo2
        parrafo2 = documentoObjeto.add_paragraph('')
        for run in parrafo2.runs:
            run.font.size=Pt(11)
        if Arrendatario_Empresa=="no":
            parrafo2 .add_run('CONDICIONES GENERALES').bold = True
            parrafo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2 = documentoObjeto.add_paragraph('')
            p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p2.add_run('PROMOTORA INMOBILIARIA R&G S.A.S., con NIT: 800.239.928-9, con matrícula mercantil '+Matricula_Mercantil_Promotora+' y matricula arrendadora: 1417, la cual se encuentra sometida a control de la Subsecretaria de Inspección, Vigilancia y Control de Vivienda de la Secretaría Distrital de Hábitat, representada legalmente por '+Nombre_Promotora+ ' , identificado(a) con cédula de ciudadanía No. '+cc_rep_promotora+ ' expedida en Bogotá, de nacionalidad Colombiana, con domicilio en la ciudad de Bogotá en '+direccion_promotora+ ', actuando en calidad de MANDATARIO según contrato de administración y/o mandato previamente suscrito, quien para efectos de este contrato se denominará EL ARRENDADOR y/o EL ADMINISTRADOR, quien ha sido contratado por el propietario para administrar el inmueble de su propiedad, por una parte, y por la otra, '+nombre_arrendatario+', mayor de edad, identificado con '+tipodiarrendatariono_ +' No. '+Numero_id_arrendatario+' de '+ciudadexpediciondiarrendatario+', quien para efectos de este contrato obra en nombre propio y se denominará EL ARRENDATARIO; manifestaron que han decidido celebrar un contrato de arrendamiento de bien inmueble destinado a vivienda, en adelante ¨EL CONTRATO¨, el cual se rige por las siguientes cláusulas: ')  
            p2.add_run("")

        elif (Arrendatario_Empresa=="si") and(Numero_Deudores==1):
            parrafo2 .add_run('CONDICIONES GENERALES').bold = True
            parrafo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2 = documentoObjeto.add_paragraph('')
            p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p2.add_run('PROMOTORA INMOBILIARIA R&G S.A.S., con NIT: 800.239.928-9, con matrícula mercantil '+Matricula_Mercantil_Promotora+' y matricula arrendadora: 1417, la cual se encuentra sometida a control de la Subsecretaria de Inspección, Vigilancia y Control de Vivienda de la Secretaría Distrital de Hábitat, representada legalmente por '+Nombre_Promotora+ ' , identificado(a) con cédula de ciudadanía No. '+cc_rep_promotora+ ' expedida en Bogotá, de nacionalidad Colombiana, con domicilio en la ciudad de Bogotá en '+direccion_promotora+ ', actuando en calidad de MANDATARIO según contrato de administración y/o mandato previamente suscrito, quien para efectos de este contrato se denominará EL ARRENDADOR y/o EL ADMINISTRADOR, quien ha sido contratado por el propietario para administrar el inmueble de su propiedad, por una parte, y por la otra, '+nombre_arrendatario+', Sociedad identificada con '+tipodiarrendatariono_ +' No. '+Numero_id_arrendatario+',representada legalmente por '+Deudor_nombre[0]+', identificada con '+Deudor_di[0]+' No. '+Deudor_idnumm[0]+' de '+Deudor_ciudad[0]+'; manifestaron que han decidido celebrar un contrato de arrendamiento de bien inmueble destinado a vivienda, en adelante ¨EL CONTRATO¨, el cual se rige por las siguientes cláusulas: ')  
            p2.add_run("") 
        elif (Arrendatario_Empresa=="si") and(Numero_Deudores==2):
            parrafo2 .add_run('CONDICIONES GENERALES').bold = True
            parrafo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2 = documentoObjeto.add_paragraph('')
            p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p2.add_run('PROMOTORA INMOBILIARIA R&G S.A.S., con NIT: 800.239.928-9, con matrícula mercantil '+Matricula_Mercantil_Promotora+' y matricula arrendadora: 1417, la cual se encuentra sometida a control de la Subsecretaria de Inspección, Vigilancia y Control de Vivienda de la Secretaría Distrital de Hábitat, representada legalmente por '+Nombre_Promotora+ ' , identificado(a) con cédula de ciudadanía No. '+cc_rep_promotora+ ' expedida en Bogotá, de nacionalidad Colombiana, con domicilio en la ciudad de Bogotá en '+direccion_promotora+ ', actuando en calidad de MANDATARIO según contrato de administración y/o mandato previamente suscrito, quien para efectos de este contrato se denominará EL ARRENDADOR y/o EL ADMINISTRADOR, quien ha sido contratado por el propietario para administrar el inmueble de su propiedad, por una parte, y por la otra, '+nombre_arrendatario+', Sociedad identificada con '+tipodiarrendatariono_ +' No. '+Numero_id_arrendatario+',representada legalmente por '+Deudor_nombre[2]+', identificada con '+Deudor_di[2]+' No. '+Deudor_idnumm[2]+' de '+Deudor_ciudad[2]+'; manifestaron que han decidido celebrar un contrato de arrendamiento de bien inmueble destinado a vivienda, en adelante ¨EL CONTRATO¨, el cual se rige por las siguientes cláusulas: ')  
            p2.add_run("") 
        elif (Arrendatario_Empresa=="si") and(Numero_Deudores==3):
            parrafo2 .add_run('CONDICIONES GENERALES').bold = True
            parrafo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2 = documentoObjeto.add_paragraph('')
            p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p2.add_run('PROMOTORA INMOBILIARIA R&G S.A.S., con NIT: 800.239.928-9, con matrícula mercantil '+Matricula_Mercantil_Promotora+' y matricula arrendadora: 1417, la cual se encuentra sometida a control de la Subsecretaria de Inspección, Vigilancia y Control de Vivienda de la Secretaría Distrital de Hábitat, representada legalmente por '+Nombre_Promotora+ ' , identificado(a) con cédula de ciudadanía No. '+cc_rep_promotora+ ' expedida en Bogotá, de nacionalidad Colombiana, con domicilio en la ciudad de Bogotá en '+direccion_promotora+ ', actuando en calidad de MANDATARIO según contrato de administración y/o mandato previamente suscrito, quien para efectos de este contrato se denominará EL ARRENDADOR y/o EL ADMINISTRADOR, quien ha sido contratado por el propietario para administrar el inmueble de su propiedad, por una parte, y por la otra, '+nombre_arrendatario+', Sociedad identificada con '+tipodiarrendatariono_ +' No. '+Numero_id_arrendatario+',representada legalmente por '+Deudor_nombre[4]+', identificada con '+Deudor_di[4]+' No. '+Deudor_idnumm[4]+' de '+Deudor_ciudad[4]+'; manifestaron que han decidido celebrar un contrato de arrendamiento de bien inmueble destinado a vivienda, en adelante ¨EL CONTRATO¨, el cual se rige por las siguientes cláusulas: ')  
            p2.add_run("") 
        elif(Arrendatario_Empresa=="si") and(Rta== "No" or Rta=="no"or Rta=="NO" ):
            parrafo2 .add_run('CONDICIONES GENERALES').bold = True
            parrafo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2 = documentoObjeto.add_paragraph('')
            p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p2.add_run('PROMOTORA INMOBILIARIA R&G S.A.S., con NIT: 800.239.928-9, con matrícula mercantil '+Matricula_Mercantil_Promotora+' y matricula arrendadora: 1417, la cual se encuentra sometida a control de la Subsecretaria de Inspección, Vigilancia y Control de Vivienda de la Secretaría Distrital de Hábitat, representada legalmente por '+Nombre_Promotora+ ' , identificado(a) con cédula de ciudadanía No. '+cc_rep_promotora+ ' expedida en Bogotá, de nacionalidad Colombiana, con domicilio en la ciudad de Bogotá en '+direccion_promotora+ ', actuando en calidad de MANDATARIO según contrato de administración y/o mandato previamente suscrito, quien para efectos de este contrato se denominará EL ARRENDADOR y/o EL ADMINISTRADOR, quien ha sido contratado por el propietario para administrar el inmueble de su propiedad, por una parte, y por la otra, '+nombre_arrendatario+', Sociedad identificada con '+tipodiarrendatariono_ +' No. '+Numero_id_arrendatario+'; manifestaron que han decidido celebrar un contrato de arrendamiento de bien inmueble destinado a vivienda, en adelante ¨EL CONTRATO¨, el cual se rige por las siguientes cláusulas: ')  
            p2.add_run("")  

        p = documentoObjeto.add_paragraph('')
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        titulo1='CLÁUSULA PRIMERA. OBJETO: '
        p.add_run(titulo1).bold = True
        p.add_run('Mediante el presente contrato, EL ARRENDADOR concede a EL ARRENDATARIO el uso y goce del inmueble identificado como '+Apartamento_y_torre+", ubicado en la ciudad de "+Ciudad+"  en la "+Direccion_Inmueble+", identificado con Matricula Inmobiliaria No. "+numero_matricula+" de la oficina de registro de instrumentos públicos de "+Ciudad+", CHIP: "+Chip+", cuya descripción cabida y linderos del inmueble se encuentran contenidos en la escritura de adquisición del mismo, y conforme con el inventario que las partes firman por separado, el cual hace parte integral de este documento como ANEXO 1.") 


        if Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO" :
            p1 = documentoObjeto.add_paragraph('')
            p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif(Deposito=="No"):
            p1 = documentoObjeto.add_paragraph('')
            p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            titulo2='CLÀUSULA SEGUNDA. COSAS O USOS CONEXOS: '
            p1.add_run(titulo2).bold = True
            p1.add_run('Además del inmueble identificado y descrito anteriormente, tendrá EL ARRENDATARIO derecho de goce sobre las siguientes cosas y usos: Un garaje No '+Numero_Garaje+' con matrícula inmobiliaria No. '+Matricula_Garaje+", CHIP: "+Chip_Garaje+"." )
        else:
            p1 = documentoObjeto.add_paragraph('')
            p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            titulo2='CLÀUSULA SEGUNDA. COSAS O USOS CONEXOS: '
            p1.add_run(titulo2).bold = True
            p1.add_run('Además del inmueble identificado y descrito anteriormente, tendrá EL ARRENDATARIO derecho de goce sobre las siguientes cosas y usos: Un garaje No '+Numero_Garaje+' con matrícula inmobiliaria No. '+Matricula_Garaje+", CHIP: "+Chip_Garaje+" y deposito identificado con el número "+Numero_deposito+".")
       
        p3 = documentoObjeto.add_paragraph('')
        p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo3='CLÀUSULA SEGUNDA. RÉGIMEN DE PROPIEDAD HORIZONTAL: '
            p3.add_run(titulo3).bold = True
            p3.add_run('El Inmueble objeto del contrato de arrendamiento forma parte del '+Nombre_edificio+", el cual se encuentra ubicado en el área urbana de la ciudad de "+Ciudad+" en la "+Direccion_Inmueble+", sometido al Régimen de Propiedad Horizontal, según consta en la Escritura Pública del Reglamento de Propiedad Horizontal, debidamente registrada en la oficina de Instrumentos Públicos, el cual el ARRENDATARIO acepta, conoce y se obliga a cumplir.")
            
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo3='CLÀUSULA TERCERA. RÉGIMEN DE PROPIEDAD HORIZONTAL: '
            p3.add_run(titulo3).bold = True
            p3.add_run('El Inmueble objeto del contrato de arrendamiento forma parte del '+Nombre_edificio+", el cual se encuentra ubicado en el área urbana de la ciudad de "+Ciudad+" en la "+Direccion_Inmueble+", sometido al Régimen de Propiedad Horizontal, según consta en la Escritura Pública del Reglamento de Propiedad Horizontal, debidamente registrada en la oficina de Instrumentos Públicos, el cual el ARRENDATARIO acepta, conoce y se obliga a cumplir.")
            
        elif (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            p3.add_run('')

        p3 = documentoObjeto.add_paragraph('')
        p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo4='CLÁUSULA TERCERA. DESTINACIÓN:'
            p3.add_run(titulo4).bold = True
            p3.add_run(' EL ARRENDATARIO, durante la vigencia del Contrato, destinará el Inmueble única y exclusivamente para su vivienda y la de su familia. En ningún caso EL ARRENDATARIO podrá subarrendar o ceder en todo o en parte este arrendamiento, so pena de que EL ARRENDADOR pueda dar por terminado válidamente el Contrato en forma inmediata, sin lugar a indemnización alguna en favor de EL ARRENDATARIO y podrá exigir la devolución del Inmueble sin necesidad de ningún tipo de requerimiento previo por parte de EL ARRENDADOR. Igualmente, EL ARRENDATARIO se abstendrá de guardar o permitir que dentro del Inmueble se guarden semovientes y/o elementos inflamables, tóxicos, insalubres, explosivos o dañosos para la conservación, higiene, seguridad y estética del inmueble y en general de sus ocupantes permanentes o transitorios. ')
            p3.add_run('PARÁGRAFO.').bold = True 
            p3.add_run(' EL ARRENDADOR declara expresa y terminantemente prohibida la destinación del inmueble a los fines contemplados en el literal b) del Parágrafo del Artículo 34 de la Ley 30 de 1986, modificado por el artículo 18 de la ley 365 de 1997 y en consecuencia EL ARRENDATARIO se obliga a no usar, el Inmueble para el ocultamiento de personas, depósito de armas o explosivos y dinero de los grupos terroristas. No destinará el inmueble para la elaboración, almacenamiento o venta de sustancias alucinógenas tales como marihuana, hachís, cocaína, metacualona y similares. EL ARRENDATARIO faculta a EL ARRENDADOR para que, directamente o a través de sus funcionarios debidamente autorizados por escrito y previo aviso a EL ARRENDATARIO, visiten el Inmueble para verificar el cumplimiento de las obligaciones de EL ARRENDATARIO. PARÁGRAFO. En caso que el ARRENDATARIO posea animales domésticos y estos generen daños en el inmueble, el ARRENDATARIO responderá por la totalidad de los daños y perjuicios que se causen.')
            
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo4='CLÁUSULA CUARTA. DESTINACIÓN:'
            p3.add_run(titulo4).bold = True
            p3.add_run(' EL ARRENDATARIO, durante la vigencia del Contrato, destinará el Inmueble única y exclusivamente para su vivienda y la de su familia. En ningún caso EL ARRENDATARIO podrá subarrendar o ceder en todo o en parte este arrendamiento, so pena de que EL ARRENDADOR pueda dar por terminado válidamente el Contrato en forma inmediata, sin lugar a indemnización alguna en favor de EL ARRENDATARIO y podrá exigir la devolución del Inmueble sin necesidad de ningún tipo de requerimiento previo por parte de EL ARRENDADOR. Igualmente, EL ARRENDATARIO se abstendrá de guardar o permitir que dentro del Inmueble se guarden semovientes y/o elementos inflamables, tóxicos, insalubres, explosivos o dañosos para la conservación, higiene, seguridad y estética del inmueble y en general de sus ocupantes permanentes o transitorios. ')
            p3.add_run('PARÁGRAFO.').bold = True 
            p3.add_run(' EL ARRENDADOR declara expresa y terminantemente prohibida la destinación del inmueble a los fines contemplados en el literal b) del Parágrafo del Artículo 34 de la Ley 30 de 1986, modificado por el artículo 18 de la ley 365 de 1997 y en consecuencia EL ARRENDATARIO se obliga a no usar, el Inmueble para el ocultamiento de personas, depósito de armas o explosivos y dinero de los grupos terroristas. No destinará el inmueble para la elaboración, almacenamiento o venta de sustancias alucinógenas tales como marihuana, hachís, cocaína, metacualona y similares. EL ARRENDATARIO faculta a EL ARRENDADOR para que, directamente o a través de sus funcionarios debidamente autorizados por escrito y previo aviso a EL ARRENDATARIO, visiten el Inmueble para verificar el cumplimiento de las obligaciones de EL ARRENDATARIO. PARÁGRAFO. En caso que el ARRENDATARIO posea animales domésticos y estos generen daños en el inmueble, el ARRENDATARIO responderá por la totalidad de los daños y perjuicios que se causen.')
            
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo4='CLÁUSULA TERCERA. DESTINACIÓN:'
            p3.add_run(titulo4).bold = True
            p3.add_run(' EL ARRENDATARIO, durante la vigencia del Contrato, destinará el Inmueble única y exclusivamente para su vivienda y la de su familia. En ningún caso EL ARRENDATARIO podrá subarrendar o ceder en todo o en parte este arrendamiento, so pena de que EL ARRENDADOR pueda dar por terminado válidamente el Contrato en forma inmediata, sin lugar a indemnización alguna en favor de EL ARRENDATARIO y podrá exigir la devolución del Inmueble sin necesidad de ningún tipo de requerimiento previo por parte de EL ARRENDADOR. Igualmente, EL ARRENDATARIO se abstendrá de guardar o permitir que dentro del Inmueble se guarden semovientes y/o elementos inflamables, tóxicos, insalubres, explosivos o dañosos para la conservación, higiene, seguridad y estética del inmueble y en general de sus ocupantes permanentes o transitorios. ')
            p3.add_run('PARÁGRAFO.').bold = True 
            p3.add_run(' EL ARRENDADOR declara expresa y terminantemente prohibida la destinación del inmueble a los fines contemplados en el literal b) del Parágrafo del Artículo 34 de la Ley 30 de 1986, modificado por el artículo 18 de la ley 365 de 1997 y en consecuencia EL ARRENDATARIO se obliga a no usar, el Inmueble para el ocultamiento de personas, depósito de armas o explosivos y dinero de los grupos terroristas. No destinará el inmueble para la elaboración, almacenamiento o venta de sustancias alucinógenas tales como marihuana, hachís, cocaína, metacualona y similares. EL ARRENDATARIO faculta a EL ARRENDADOR para que, directamente o a través de sus funcionarios debidamente autorizados por escrito y previo aviso a EL ARRENDATARIO, visiten el Inmueble para verificar el cumplimiento de las obligaciones de EL ARRENDATARIO. PARÁGRAFO. En caso que el ARRENDATARIO posea animales domésticos y estos generen daños en el inmueble, el ARRENDATARIO responderá por la totalidad de los daños y perjuicios que se causen.')
            
        elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo4='CLÁUSULA SEGUNDA. DESTINACIÓN:'
            p3.add_run(titulo4).bold = True
            p3.add_run(' EL ARRENDATARIO, durante la vigencia del Contrato, destinará el Inmueble única y exclusivamente para su vivienda y la de su familia. En ningún caso EL ARRENDATARIO podrá subarrendar o ceder en todo o en parte este arrendamiento, so pena de que EL ARRENDADOR pueda dar por terminado válidamente el Contrato en forma inmediata, sin lugar a indemnización alguna en favor de EL ARRENDATARIO y podrá exigir la devolución del Inmueble sin necesidad de ningún tipo de requerimiento previo por parte de EL ARRENDADOR. Igualmente, EL ARRENDATARIO se abstendrá de guardar o permitir que dentro del Inmueble se guarden semovientes y/o elementos inflamables, tóxicos, insalubres, explosivos o dañosos para la conservación, higiene, seguridad y estética del inmueble y en general de sus ocupantes permanentes o transitorios. ')
            p3.add_run('PARÁGRAFO.').bold = True 
            p3.add_run(' EL ARRENDADOR declara expresa y terminantemente prohibida la destinación del inmueble a los fines contemplados en el literal b) del Parágrafo del Artículo 34 de la Ley 30 de 1986, modificado por el artículo 18 de la ley 365 de 1997 y en consecuencia EL ARRENDATARIO se obliga a no usar, el Inmueble para el ocultamiento de personas, depósito de armas o explosivos y dinero de los grupos terroristas. No destinará el inmueble para la elaboración, almacenamiento o venta de sustancias alucinógenas tales como marihuana, hachís, cocaína, metacualona y similares. EL ARRENDATARIO faculta a EL ARRENDADOR para que, directamente o a través de sus funcionarios debidamente autorizados por escrito y previo aviso a EL ARRENDATARIO, visiten el Inmueble para verificar el cumplimiento de las obligaciones de EL ARRENDATARIO. PARÁGRAFO. En caso que el ARRENDATARIO posea animales domésticos y estos generen daños en el inmueble, el ARRENDATARIO responderá por la totalidad de los daños y perjuicios que se causen.')

        p4 = documentoObjeto.add_paragraph('')
        p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo5='CLÁUSULA CUARTA.OBLIGACIONES DEL ARRENDADOR: '
            p4.add_run(titulo5).bold = True
            p4.add_run('El ARRENDADOR se encuentra obligado a lo siguiente:\t\n'+"1. Entregar a EL ARRENDATARIO en la fecha convenida el inmueble dado en arrendamiento en buen estado de servicio, seguridad y sanidad y poner a su disposición los servicios, cosas o usos conexos y los adicionales aquí convenido.\t\n"+"2. Mantener en el inmueble los servicios, las cosas y los usos conexos y adicionales en buen estado de servir para el fin convenido en el contrato.\t\n"+"3. Las demás obligaciones consagradas para los arrendadores en el capítulo II, título XXVI, libro 4º del Código Civil y Ley 820 del año 2003.")
            
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo5='CLÁUSULA QUINTA.OBLIGACIONES DEL ARRENDADOR: '
            p4.add_run(titulo5).bold = True
            p4.add_run('El ARRENDADOR se encuentra obligado a lo siguiente:\t\n'+"1. Entregar a EL ARRENDATARIO en la fecha convenida el inmueble dado en arrendamiento en buen estado de servicio, seguridad y sanidad y poner a su disposición los servicios, cosas o usos conexos y los adicionales aquí convenido.\t\n"+"2. Mantener en el inmueble los servicios, las cosas y los usos conexos y adicionales en buen estado de servir para el fin convenido en el contrato.\t\n"+"3. Las demás obligaciones consagradas para los arrendadores en el capítulo II, título XXVI, libro 4º del Código Civil y Ley 820 del año 2003.")
        
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo5='CLÁUSULA CUARTA.OBLIGACIONES DEL ARRENDADOR: '
            p4.add_run(titulo5).bold = True
            p4.add_run('El ARRENDADOR se encuentra obligado a lo siguiente:\t\n'+"1. Entregar a EL ARRENDATARIO en la fecha convenida el inmueble dado en arrendamiento en buen estado de servicio, seguridad y sanidad y poner a su disposición los servicios, cosas o usos conexos y los adicionales aquí convenido.\t\n"+"2. Mantener en el inmueble los servicios, las cosas y los usos conexos y adicionales en buen estado de servir para el fin convenido en el contrato.\t\n"+"3. Las demás obligaciones consagradas para los arrendadores en el capítulo II, título XXVI, libro 4º del Código Civil y Ley 820 del año 2003.")

        elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo5='CLÁUSULA TERCERA.OBLIGACIONES DEL ARRENDADOR: '
            p4.add_run(titulo5).bold = True
            p4.add_run('El ARRENDADOR se encuentra obligado a lo siguiente:\t\n'+"1. Entregar a EL ARRENDATARIO en la fecha convenida el inmueble dado en arrendamiento en buen estado de servicio, seguridad y sanidad y poner a su disposición los servicios, cosas o usos conexos y los adicionales aquí convenido.\t\n"+"2. Mantener en el inmueble los servicios, las cosas y los usos conexos y adicionales en buen estado de servir para el fin convenido en el contrato.\t\n"+"3. Las demás obligaciones consagradas para los arrendadores en el capítulo II, título XXVI, libro 4º del Código Civil y Ley 820 del año 2003.")

        p5 = documentoObjeto.add_paragraph('')
        p5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo6='CLÁUSULA QUINTA.OBLIGACIONES DEL ARRENDATARIO: '
            p5.add_run(titulo6).bold = True
            p5.add_run('Son obligaciones del ARRENDATARIO.\t\n'+"1. Pagar al ARRENDADOR el precio del arrendamiento convenido en el presente contrato.\t\n"+"2. Cuidar el inmueble y las cosas recibidas en arrendamiento. En caso de daños o deterioros distintos derivados del uso normal o de la acción del tiempo y que fueren imputables al mal uso del inmueble o a su propia culpa, efectuar oportunamente y por su cuenta las reparaciones o sustituciones necesarias.\t\n"+"3. Pagar a tiempo los servicios, cosas o usos conexos y adicionales, así como las expensas comunes en los casos en que haya lugar, de conformidad con lo establecido en el contrato.\t\n"+"4. Cumplir las normas consagradas en los reglamentos de propiedad horizontal y las que expida el Gobierno Nacional dirigidas a la protección de los derechos de todos los vecinos. En caso de vivienda compartida y de pensión, el arrendatario está obligado además, de cuidar las zonas y servicios de uso común y efectuar por su cuenta las reparaciones o sustituciones necesarias, cuando sean atribuibles a su propia culpa o a la de sus dependientes.\t\n"+"5. El ARRENDATARIO "+x1+" a contactar a la administración para conocer el reglamento de propiedad horizontal del bien dado en arrendamiento y se comprometen a cumplir y respetar cabalmente todas y cada una de las normas establecidas en él, al cual está sometido el inmueble objeto del presente contrato de arrendamiento.\n")
            p5.add_run('6. El ARRENDATARIO y sus deudores solidarios autorizan expresamente al arrendador para el envío de correos con información relevante para su buena presentación de servicio (entiéndase como publicidad, notas importantes, etc)\n')
            p5.add_run('7. Las demás obligaciones consagradas para los arrendatarios en el capítulo III, título XXVI, libro 4º del Código Civil y Ley 820 del año 2003.')
            
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo6='CLÁUSULA SEXTA.OBLIGACIONES DEL ARRENDATARIO: '
            p5.add_run(titulo6).bold = True
            p5.add_run('Son obligaciones del ARRENDATARIO.\t\n'+"1. Pagar al ARRENDADOR el precio del arrendamiento convenido en el presente contrato.\t\n"+"2. Cuidar el inmueble y las cosas recibidas en arrendamiento. En caso de daños o deterioros distintos derivados del uso normal o de la acción del tiempo y que fueren imputables al mal uso del inmueble o a su propia culpa, efectuar oportunamente y por su cuenta las reparaciones o sustituciones necesarias.\t\n"+"3. Pagar a tiempo los servicios, cosas o usos conexos y adicionales, así como las expensas comunes en los casos en que haya lugar, de conformidad con lo establecido en el contrato.\t\n"+"4. Cumplir las normas consagradas en los reglamentos de propiedad horizontal y las que expida el Gobierno Nacional dirigidas a la protección de los derechos de todos los vecinos. En caso de vivienda compartida y de pensión, el arrendatario está obligado además, de cuidar las zonas y servicios de uso común y efectuar por su cuenta las reparaciones o sustituciones necesarias, cuando sean atribuibles a su propia culpa o a la de sus dependientes.\t\n"+"5. El ARRENDATARIO "+x1+" a contactar a la administración para conocer el reglamento de propiedad horizontal del bien dado en arrendamiento y se comprometen a cumplir y respetar cabalmente todas y cada una de las normas establecidas en él, al cual está sometido el inmueble objeto del presente contrato de arrendamiento.\n") 
            p5.add_run('6. El ARRENDATARIO y sus deudores solidarios autorizan expresamente al arrendador para el envío de correos con información relevante para su buena presentación de servicio (entiéndase como publicidad, notas importantes, etc)\n')
            p5.add_run('7. Las demás obligaciones consagradas para los arrendatarios en el capítulo III, título XXVI, libro 4º del Código Civil y Ley 820 del año 2003.')

        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo6='CLÁUSULA QUINTA.OBLIGACIONES DEL ARRENDATARIO: '
            p5.add_run(titulo6).bold = True
            p5.add_run('Son obligaciones del ARRENDATARIO.\t\n'+"1. Pagar al ARRENDADOR el precio del arrendamiento convenido en el presente contrato.\t\n"+"2. Cuidar el inmueble y las cosas recibidas en arrendamiento. En caso de daños o deterioros distintos derivados del uso normal o de la acción del tiempo y que fueren imputables al mal uso del inmueble o a su propia culpa, efectuar oportunamente y por su cuenta las reparaciones o sustituciones necesarias.\t\n"+"3. Pagar a tiempo los servicios, cosas o usos conexos y adicionales, así como las expensas comunes en los casos en que haya lugar, de conformidad con lo establecido en el contrato.\t\n"+"4. Cumplir las normas consagradas en los reglamentos de propiedad horizontal y las que expida el Gobierno Nacional dirigidas a la protección de los derechos de todos los vecinos. En caso de vivienda compartida y de pensión, el arrendatario está obligado además, de cuidar las zonas y servicios de uso común y efectuar por su cuenta las reparaciones o sustituciones necesarias, cuando sean atribuibles a su propia culpa o a la de sus dependientes.\t\n"+"5. El ARRENDATARIO "+x1+" a contactar a la administración para conocer el reglamento de propiedad horizontal del bien dado en arrendamiento y se comprometen a cumplir y respetar cabalmente todas y cada una de las normas establecidas en él, al cual está sometido el inmueble objeto del presente contrato de arrendamiento.\n")
            p5.add_run('6. El ARRENDATARIO y sus deudores solidarios autorizan expresamente al arrendador para el envío de correos con información relevante para su buena presentación de servicio (entiéndase como publicidad, notas importantes, etc)\n')
            p5.add_run('7. Las demás obligaciones consagradas para los arrendatarios en el capítulo III, título XXVI, libro 4º del Código Civil y Ley 820 del año 2003.')

        elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo6='CLÁUSULA CUARTA.OBLIGACIONES DEL ARRENDATARIO: '
            p5.add_run(titulo6).bold = True
            p5.add_run('Son obligaciones del ARRENDATARIO.\t\n'+"1. Pagar al ARRENDADOR el precio del arrendamiento convenido en el presente contrato.\t\n"+"2. Cuidar el inmueble y las cosas recibidas en arrendamiento. En caso de daños o deterioros distintos derivados del uso normal o de la acción del tiempo y que fueren imputables al mal uso del inmueble o a su propia culpa, efectuar oportunamente y por su cuenta las reparaciones o sustituciones necesarias.\t\n"+"3. Pagar a tiempo los servicios, cosas o usos conexos y adicionales, así como las expensas comunes en los casos en que haya lugar, de conformidad con lo establecido en el contrato.\t\n"+"4. Cumplir las normas consagradas en los reglamentos de propiedad horizontal y las que expida el Gobierno Nacional dirigidas a la protección de los derechos de todos los vecinos. En caso de vivienda compartida y de pensión, el arrendatario está obligado además, de cuidar las zonas y servicios de uso común y efectuar por su cuenta las reparaciones o sustituciones necesarias, cuando sean atribuibles a su propia culpa o a la de sus dependientes.\t\n"+"5. El ARRENDATARIO "+x1+" a contactar a la administración para conocer el reglamento de propiedad horizontal del bien dado en arrendamiento y se comprometen a cumplir y respetar cabalmente todas y cada una de las normas establecidas en él, al cual está sometido el inmueble objeto del presente contrato de arrendamiento.\n")
            p5.add_run('6. El ARRENDATARIO y sus deudores solidarios autorizan expresamente al arrendador para el envío de correos con información relevante para su buena presentación de servicio (entiéndase como publicidad, notas importantes, etc)\n')
            p5.add_run('7. Las demás obligaciones consagradas para los arrendatarios en el capítulo III, título XXVI, libro 4º del Código Civil y Ley 820 del año 2003.')

        ##
        p6 = documentoObjeto.add_paragraph('')
        p6.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if  (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo7='CLÁUSULA SEXTA. VALOR DEL CANON DE ARRENDAMIENTO Y FECHA DE PAGO: '
            p6.add_run(titulo7).bold = True
            p6.add_run('El valor del canon mensual de arrendamiento será de '+Canon_Letras+" PESOS  M/C. ($ "+valor_canon+") moneda legal vigente colombiana, más "+Administracion_Letras+" PESOS M/C. ($ "+valor_administracion+") que corresponde a cuota de administración, para un total a pagar de "+valor_total_letra+" PESOS M/C. ($ " +valor_total+" ), pagaderos dentro de los primeros cinco (05) días calendario de cada periodo mensual por anticipado al ARRENDADOR ")
            p6.add_run('PARÁGRAFO 1. ').bold = True
            p6.add_run('En caso de incumplimiento del pago del canon de arrendamiento dentro de los primeros cinco (5) dìas del mes, el ARRENDADOR Y/O ADMINISTRADOR podrá dar por terminado unilateralmente el contrato con justa causa y exigir la entrega inmediata del inmueble, para lo cual el ARRENDADOR podrá iniciar el procedimiento correspondiente ante la aseguradora, afianzadora u oficina de cobro jurídico.')
            
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo7='CLÁUSULA SÉPTIMA. VALOR DEL CANON DE ARRENDAMIENTO Y FECHA DE PAGO: '
            p6.add_run(titulo7).bold = True
            p6.add_run('El valor del canon mensual de arrendamiento será de '+Canon_Letras+" PESOS  M/C. ($ "+valor_canon+") moneda legal vigente colombiana, más "+Administracion_Letras+" PESOS M/C. ($ "+valor_administracion+") que corresponde a cuota de administración, para un total a pagar de "+valor_total_letra+" PESOS M/C. ($ " +valor_total+" ), pagaderos dentro de los primeros cinco (05) días calendario de cada periodo mensual por anticipado al ARRENDADOR ")
            p6.add_run('PARÁGRAFO 1. ').bold = True
            p6.add_run('En caso de incumplimiento del pago del canon de arrendamiento dentro de los primeros cinco (5) dìas del mes, el ARRENDADOR Y/O ADMINISTRADOR podrá dar por terminado unilateralmente el contrato con justa causa y exigir la entrega inmediata del inmueble, para lo cual el ARRENDADOR podrá iniciar el procedimiento correspondiente ante la aseguradora, afianzadora u oficina de cobro jurídico.')

        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo7='CLÁUSULA SEXTA. VALOR DEL CANON DE ARRENDAMIENTO Y FECHA DE PAGO: '
            p6.add_run(titulo7).bold = True
            p6.add_run('El valor del canon mensual de arrendamiento será de '+Canon_Letras+" PESOS  M/C. ($ "+valor_canon+") moneda legal vigente colombiana"+valor_total_letra+" PESOS M/C. ($ " +valor_total+" ), pagaderos dentro de los primeros cinco (05) días calendario de cada periodo mensual por anticipado al ARRENDADOR ")
            p6.add_run('PARÁGRAFO 1. ').bold = True
            p6.add_run('En caso de incumplimiento del pago del canon de arrendamiento dentro de los primeros cinco (5) dìas del mes, el ARRENDADOR Y/O ADMINISTRADOR podrá dar por terminado unilateralmente el contrato con justa causa y exigir la entrega inmediata del inmueble, para lo cual el ARRENDADOR podrá iniciar el procedimiento correspondiente ante la aseguradora, afianzadora u oficina de cobro jurídico.')

        elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo7='CLÁUSULA QUINTA. VALOR DEL CANON DE ARRENDAMIENTO Y FECHA DE PAGO: '
            p6.add_run(titulo7).bold = True
            p6.add_run('El valor del canon mensual de arrendamiento será de '+Canon_Letras+" PESOS  M/C. ($ "+valor_canon+") moneda legal vigente colombiana"+valor_total_letra+" PESOS M/C. ($ " +valor_total+" ), pagaderos dentro de los primeros cinco (05) días calendario de cada periodo mensual por anticipado al ARRENDADOR ")
            p6.add_run('PARÁGRAFO 1. ').bold = True
            p6.add_run('En caso de incumplimiento del pago del canon de arrendamiento dentro de los primeros cinco (5) dìas del mes, el ARRENDADOR Y/O ADMINISTRADOR podrá dar por terminado unilateralmente el contrato con justa causa y exigir la entrega inmediata del inmueble, para lo cual el ARRENDADOR podrá iniciar el procedimiento correspondiente ante la aseguradora, afianzadora u oficina de cobro jurídico.')
            ####
        p7 = documentoObjeto.add_paragraph('')
        p7.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo8='CLÁUSULA SÉPTIMA. LUGAR PARA EL PAGO – Y GASTOS: '
            p7.add_run(titulo8).bold = True
            p7.add_run('El pago del canon de arrendamiento será cancelado a través de cuenta de recaudo por el sistema nacional de código de barras y/o PSE de acuerdo con la factura que EL ARRENDATARIO podrá descargar de la página web ')
            p7.add_run('www.rginmobiliaria.com.co').underline = True
            p7.add_run('para su pago oportuno.Paragrafo: los gastos como uso de medios de pago electrónico, recaudo bancario o de cualquier naturaleza imputable al contrato, corresponderán al ARRENDATARIO.')
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo8='CLÁUSULA OCTAVA. LUGAR PARA EL PAGO – Y GASTOS: '
            p7.add_run(titulo8).bold = True
            p7.add_run('El pago del canon de arrendamiento será cancelado a través de cuenta de recaudo por el sistema nacional de código de barras y/o PSE de acuerdo con la factura que EL ARRENDATARIO podrá descargar de la página web')
            p7.add_run('www.rginmobiliaria.com.co').underline = True
            p7.add_run('para su pago oportuno.Paragrafo: los gastos como uso de medios de pago electrónico, recaudo bancario o de cualquier naturaleza imputable al contrato, corresponderán al ARRENDATARIO.')

        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo8='CLÁUSULA SÉPTIMA. LUGAR PARA EL PAGO – Y GASTOS: '
            p7.add_run(titulo8).bold = True
            p7.add_run('El pago del canon de arrendamiento será cancelado a través de cuenta de recaudo por el sistema nacional de código de barras y/o PSE de acuerdo con la factura que EL ARRENDATARIO podrá descargar de la página web')
            p7.add_run('www.rginmobiliaria.com.co').underline = True
            p7.add_run('para su pago oportuno.Paragrafo: los gastos como uso de medios de pago electrónico, recaudo bancario o de cualquier naturaleza imputable al contrato, corresponderán al ARRENDATARIO.')
        elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo8='CLÁUSULA SEXTA. LUGAR PARA EL PAGO – Y GASTOS: '
            p7.add_run(titulo8).bold = True
            p7.add_run('El pago del canon de arrendamiento será cancelado a través de cuenta de recaudo por el sistema nacional de código de barras y/o PSE de acuerdo con la factura que EL ARRENDATARIO podrá descargar de la página web')
            p7.add_run('www.rginmobiliaria.com.co').underline = True
            p7.add_run('para su pago oportuno.Paragrafo: los gastos como uso de medios de pago electrónico, recaudo bancario o de cualquier naturaleza imputable al contrato, corresponderán al ARRENDATARIO.')

        ####

        p8 = documentoObjeto.add_paragraph('')
        p8.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo9='CLÁUSULA OCTAVA.  TÉRMINO DE VIGENCIA DEL CONTRATO: '
            p8.add_run(titulo9).bold = True
            p8.add_run('El Contrato de arrendamiento tendrá una vigencia de ('+Vigencia_Contrato+') contados desde el '+Fecha_inicio_Contrato)
            
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo9='CLÁUSULA NOVENA.  TÉRMINO DE VIGENCIA DEL CONTRATO: '
            p8.add_run(titulo9).bold = True
            p8.add_run('El Contrato de arrendamiento tendrá una vigencia de ('+Vigencia_Contrato+') contados desde el '+Fecha_inicio_Contrato)

        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo9='CLÁUSULA OCTAVA.  TÉRMINO DE VIGENCIA DEL CONTRATO: '
            p8.add_run(titulo9).bold = True
            p8.add_run('El Contrato de arrendamiento tendrá una vigencia de ('+Vigencia_Contrato+') contados desde el '+Fecha_inicio_Contrato)

        elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo9='CLÁUSULA SÉPTIMA.  TÉRMINO DE VIGENCIA DEL CONTRATO: '
            p8.add_run(titulo9).bold = True
            p8.add_run('El Contrato de arrendamiento tendrá una vigencia de ('+Vigencia_Contrato+') contados desde el '+Fecha_inicio_Contrato)

        ####
        p9 = documentoObjeto.add_paragraph('')
        p9.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI")and (Rta=="Si" or Rta=="si"or Rta=="SI"):
            titulo10='CLÁUSULA NOVENA. PRÓRROGAS: '
            p9.add_run(titulo10).bold = True
            p9.add_run('Si a la fecha del vencimiento del término inicial o cualquiera de sus prórrogas, ninguna de las partes ha dado aviso con una antelación no menor a tres (3) meses a la fecha de vencimiento, su intención de dar por terminado el presente contrato de arrendamiento, se entenderá prorrogado de forma automática en las mismas condiciones y por el mismo término inicial pactado, siempre y cuando cada una de las partes haya cumplido con las obligaciones a su cargo, y el ARRENDATARIO se avenga a los reajustes autorizados por ley, conforme al artículo 6o de la Ley 820 de 2003 y el numeral 4. art. 8 del Decreto 051 de 2004.')
            p9.add_run(' Parágrafo').bold=True
            p9.add_run(': Conforme a lo dispuesto en el artículo 20 de la Ley 820 de 2003 se remitirá al correo electrónico '+correo_arrendatario+' el aviso de incremento y fecha de aplicación del mismo al arrendatario, en donde los Deudores Solidarios realizarán aceptación tácita del mismo y forma de notificación.')

        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI")and (Rta=="Si" or Rta=="si"or Rta=="SI"):
            titulo10='CLÁUSULA DÉCIMA. PRÓRROGAS: '
            p9.add_run(titulo10).bold = True
            p9.add_run('Si a la fecha del vencimiento del término inicial o cualquiera de sus prórrogas, ninguna de las partes ha dado aviso con una antelación no menor a tres (3) meses a la fecha de vencimiento, su intención de dar por terminado el presente contrato de arrendamiento, se entenderá prorrogado de forma automática en las mismas condiciones y por el mismo término inicial pactado, siempre y cuando cada una de las partes haya cumplido con las obligaciones a su cargo, y el ARRENDATARIO se avenga a los reajustes autorizados por ley, conforme al artículo 6o de la Ley 820 de 2003 y el numeral 4. art. 8 del Decreto 051 de 2004.')
            p9.add_run(' Parágrafo').bold=True
            p9.add_run(': Conforme a lo dispuesto en el artículo 20 de la Ley 820 de 2003 se remitirá al correo electrónico '+correo_arrendatario+' el aviso de incremento y fecha de aplicación del mismo al arrendatario, en donde los Deudores Solidarios realizarán aceptación tácita del mismo y forma de notificación.')

        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO")and (Rta=="Si" or Rta=="si"or Rta=="SI"):
            titulo10='CLÁUSULA NOVENA. PRÓRROGAS: '
            p9.add_run(titulo10).bold = True
            p9.add_run('Si a la fecha del vencimiento del término inicial o cualquiera de sus prórrogas, ninguna de las partes ha dado aviso con una antelación no menor a tres (3) meses a la fecha de vencimiento, su intención de dar por terminado el presente contrato de arrendamiento, se entenderá prorrogado de forma automática en las mismas condiciones y por el mismo término inicial pactado, siempre y cuando cada una de las partes haya cumplido con las obligaciones a su cargo, y el ARRENDATARIO se avenga a los reajustes autorizados por ley, conforme al artículo 6o de la Ley 820 de 2003 y el numeral 4. art. 8 del Decreto 051 de 2004.')
            p9.add_run(' Parágrafo').bold=True
            p9.add_run(': Conforme a lo dispuesto en el artículo 20 de la Ley 820 de 2003 se remitirá al correo electrónico '+correo_arrendatario+' el aviso de incremento y fecha de aplicación del mismo al arrendatario, en donde los Deudores Solidarios realizarán aceptación tácita del mismo y forma de notificación.')

        elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO")and (Rta=="Si" or Rta=="si"or Rta=="SI"):
            titulo10='CLÁUSULA OCTAVA. PRÓRROGAS: '
            p9.add_run(titulo10).bold = True
            p9.add_run('Si a la fecha del vencimiento del término inicial o cualquiera de sus prórrogas, ninguna de las partes ha dado aviso con una antelación no menor a tres (3) meses a la fecha de vencimiento, su intención de dar por terminado el presente contrato de arrendamiento, se entenderá prorrogado de forma automática en las mismas condiciones y por el mismo término inicial pactado, siempre y cuando cada una de las partes haya cumplido con las obligaciones a su cargo, y el ARRENDATARIO se avenga a los reajustes autorizados por ley, conforme al artículo 6o de la Ley 820 de 2003 y el numeral 4. art. 8 del Decreto 051 de 2004.')
            p9.add_run(' Parágrafo').bold=True
            p9.add_run(': Conforme a lo dispuesto en el artículo 20 de la Ley 820 de 2003 se remitirá al correo electrónico '+correo_arrendatario+' el aviso de incremento y fecha de aplicación del mismo al arrendatario, en donde los Deudores Solidarios realizarán aceptación tácita del mismo y forma de notificación.')
        
        else:
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI")and (Rta=="No" or Rta=="no"or Rta=="NO"):
                titulo10='CLÁUSULA NOVENA. PRÓRROGAS: '
                p9.add_run(titulo10).bold = True
                p9.add_run('Si a la fecha del vencimiento del término inicial o cualquiera de sus prórrogas, ninguna de las partes ha dado aviso con una antelación no menor a tres (3) meses a la fecha de vencimiento, su intención de dar por terminado el presente contrato de arrendamiento, se entenderá prorrogado de forma automática en las mismas condiciones y por el mismo término inicial pactado, siempre y cuando cada una de las partes haya cumplido con las obligaciones a su cargo, y el ARRENDATARIO se avenga a los reajustes autorizados por ley, conforme al artículo 6o de la Ley 820 de 2003 y el numeral 4. art. 8 del Decreto 051 de 2004.')
                p9.add_run(' Parágrafo').bold=True
                p9.add_run(': Conforme a lo dispuesto en el artículo 20 de la Ley 820 de 2003 se remitirá al correo electrónico '+correo_arrendatario+' el aviso de incremento y fecha de aplicación del mismo al arrendatario, en donde los Deudores Solidarios realizarán aceptación tácita del mismo y forma de notificación.')

            elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI")and (Rta=="No" or Rta=="no"or Rta=="NO"):
                titulo10='CLÁUSULA DÉCIMA. PRÓRROGAS: '
                p9.add_run(titulo10).bold = True
                p9.add_run('Si a la fecha del vencimiento del término inicial o cualquiera de sus prórrogas, ninguna de las partes ha dado aviso con una antelación no menor a tres (3) meses a la fecha de vencimiento, su intención de dar por terminado el presente contrato de arrendamiento, se entenderá prorrogado de forma automática en las mismas condiciones y por el mismo término inicial pactado, siempre y cuando cada una de las partes haya cumplido con las obligaciones a su cargo, y el ARRENDATARIO se avenga a los reajustes autorizados por ley, conforme al artículo 6o de la Ley 820 de 2003 y el numeral 4. art. 8 del Decreto 051 de 2004.')
                p9.add_run(' Parágrafo').bold=True
                p9.add_run(': Conforme a lo dispuesto en el artículo 20 de la Ley 820 de 2003 se remitirá al correo electrónico '+correo_arrendatario+' el aviso de incremento y fecha de aplicación del mismo al arrendatario, en donde los Deudores Solidarios realizarán aceptación tácita del mismo y forma de notificación.')

            elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO"):
                titulo10='CLÁUSULA NOVENA. PRÓRROGAS: '
                p9.add_run(titulo10).bold = True
                p9.add_run('Si a la fecha del vencimiento del término inicial o cualquiera de sus prórrogas, ninguna de las partes ha dado aviso con una antelación no menor a tres (3) meses a la fecha de vencimiento, su intención de dar por terminado el presente contrato de arrendamiento, se entenderá prorrogado de forma automática en las mismas condiciones y por el mismo término inicial pactado, siempre y cuando cada una de las partes haya cumplido con las obligaciones a su cargo, y el ARRENDATARIO se avenga a los reajustes autorizados por ley, conforme al artículo 6o de la Ley 820 de 2003 y el numeral 4. art. 8 del Decreto 051 de 2004.')
                p9.add_run(' Parágrafo').bold=True
                p9.add_run(': Conforme a lo dispuesto en el artículo 20 de la Ley 820 de 2003 se remitirá al correo electrónico '+correo_arrendatario+' el aviso de incremento y fecha de aplicación del mismo al arrendatario, en donde los Deudores Solidarios realizarán aceptación tácita del mismo y forma de notificación.')

            elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO"):
                titulo10='CLÁUSULA OCTAVA. PRÓRROGAS: '
                p9.add_run(titulo10).bold = True
                p9.add_run('Si a la fecha del vencimiento del término inicial o cualquiera de sus prórrogas, ninguna de las partes ha dado aviso con una antelación no menor a tres (3) meses a la fecha de vencimiento, su intención de dar por terminado el presente contrato de arrendamiento, se entenderá prorrogado de forma automática en las mismas condiciones y por el mismo término inicial pactado, siempre y cuando cada una de las partes haya cumplido con las obligaciones a su cargo, y el ARRENDATARIO se avenga a los reajustes autorizados por ley, conforme al artículo 6o de la Ley 820 de 2003 y el numeral 4. art. 8 del Decreto 051 de 2004.')
                p9.add_run(' Parágrafo').bold=True
                p9.add_run(': Conforme a lo dispuesto en el artículo 20 de la Ley 820 de 2003 se remitirá al correo electrónico '+correo_arrendatario+' el aviso de incremento y fecha de aplicación del mismo al arrendatario, en donde los Deudores Solidarios realizarán aceptación tácita del mismo y forma de notificación.')
            
        ####....

        p11 = documentoObjeto.add_paragraph('')
        p11.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo11='CLÁUSULA DÉCIMA. INCREMENTOS DEL VALOR DEL CANON DE ARRENDAMIENTO: '
            p11.add_run(titulo11).bold = True
            p11.add_run('Vencida la primer vigencia de este contrato, y así sucesivamente cada doce (12) mensualidades en caso de prórroga tácita o expresa en forma automática y sin necesidad de requerimiento alguno entre las partes, el precio mensual del arrendamiento se reajustará en una proporción igual al cien por ciento (100%) del incremento que haya tenido el índice de precios al consumidor (I.P.C.), certificado por el DANE para el año calendario inmediatamente anterior a aquel que se efectúe el incremento. ')
            p11.add_run('PARÁGRAFO.').bold = True
            p11.add_run(' Al suscribir el presente contrato EL ARRENDATARIO, y sus DEUDORES SOLIDARIOS quedan plenamente notificados de todos los reajustes automáticos pactados en este contrato y que han de operar durante la vigencia del mismo. ')
        
        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo11='CLÁUSULA DÉCIMA. INCREMENTOS DEL VALOR DEL CANON DE ARRENDAMIENTO: '
            p11.add_run(titulo11).bold = True
            p11.add_run('Vencida la primer vigencia de este contrato, y así sucesivamente cada doce (12) mensualidades en caso de prórroga tácita o expresa en forma automática y sin necesidad de requerimiento alguno entre las partes, el precio mensual del arrendamiento se reajustará en una proporción igual al cien por ciento (100%) del incremento que haya tenido el índice de precios al consumidor (I.P.C.), certificado por el DANE para el año calendario inmediatamente anterior a aquel que se efectúe el incremento. ')
            p11.add_run('PARÁGRAFO.').bold = True
            p11.add_run(' Al suscribir el presente contrato EL ARRENDATARIO, y sus DEUDORES SOLIDARIOS quedan plenamente notificados de todos los reajustes automáticos pactados en este contrato y que han de operar durante la vigencia del mismo. ')
            ##
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo11='CLÁUSULA UNDÉCIMA. INCREMENTOS DEL VALOR DEL CANON DE ARRENDAMIENTO: '
            p11.add_run(titulo11).bold = True
            p11.add_run('Vencida la primer vigencia de este contrato, y así sucesivamente cada doce (12) mensualidades en caso de prórroga tácita o expresa en forma automática y sin necesidad de requerimiento alguno entre las partes, el precio mensual del arrendamiento se reajustará en una proporción igual al cien por ciento (100%) del incremento que haya tenido el índice de precios al consumidor (I.P.C.), certificado por el DANE para el año calendario inmediatamente anterior a aquel que se efectúe el incremento. ')
            p11.add_run('PARÁGRAFO.').bold = True
            p11.add_run(' Al suscribir el presente contrato EL ARRENDATARIO, y sus DEUDORES SOLIDARIOS quedan plenamente notificados de todos los reajustes automáticos pactados en este contrato y que han de operar durante la vigencia del mismo. ')
        
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo11='CLÁUSULA UNDÉCIMA. INCREMENTOS DEL VALOR DEL CANON DE ARRENDAMIENTO: '
            p11.add_run(titulo11).bold = True
            p11.add_run('Vencida la primer vigencia de este contrato, y así sucesivamente cada doce (12) mensualidades en caso de prórroga tácita o expresa en forma automática y sin necesidad de requerimiento alguno entre las partes, el precio mensual del arrendamiento se reajustará en una proporción igual al cien por ciento (100%) del incremento que haya tenido el índice de precios al consumidor (I.P.C.), certificado por el DANE para el año calendario inmediatamente anterior a aquel que se efectúe el incremento. ')
            p11.add_run('PARÁGRAFO.').bold = True
            p11.add_run(' Al suscribir el presente contrato EL ARRENDATARIO, y sus DEUDORES SOLIDARIOS quedan plenamente notificados de todos los reajustes automáticos pactados en este contrato y que han de operar durante la vigencia del mismo. ')
            
        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo11='CLÁUSULA NOVENA. INCREMENTOS DEL VALOR DEL CANON DE ARRENDAMIENTO: '
            p11.add_run(titulo11).bold = True
            p11.add_run('Vencida la primer vigencia de este contrato, y así sucesivamente cada doce (12) mensualidades en caso de prórroga tácita o expresa en forma automática y sin necesidad de requerimiento alguno entre las partes, el precio mensual del arrendamiento se reajustará en una proporción igual al cien por ciento (100%) del incremento que haya tenido el índice de precios al consumidor (I.P.C.), certificado por el DANE para el año calendario inmediatamente anterior a aquel que se efectúe el incremento. ')
            p11.add_run('PARÁGRAFO.').bold = True
            p11.add_run(' Al suscribir el presente contrato EL ARRENDATARIO, y sus DEUDORES SOLIDARIOS quedan plenamente notificados de todos los reajustes automáticos pactados en este contrato y que han de operar durante la vigencia del mismo. ')
        
        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo11='CLÁUSULA NOVENA. INCREMENTOS DEL VALOR DEL CANON DE ARRENDAMIENTO: '
            p11.add_run(titulo11).bold = True
            p11.add_run('Vencida la primer vigencia de este contrato, y así sucesivamente cada doce (12) mensualidades en caso de prórroga tácita o expresa en forma automática y sin necesidad de requerimiento alguno entre las partes, el precio mensual del arrendamiento se reajustará en una proporción igual al cien por ciento (100%) del incremento que haya tenido el índice de precios al consumidor (I.P.C.), certificado por el DANE para el año calendario inmediatamente anterior a aquel que se efectúe el incremento. ')
            p11.add_run('PARÁGRAFO.').bold = True
            p11.add_run(' Al suscribir el presente contrato EL ARRENDATARIO, y sus DEUDORES SOLIDARIOS quedan plenamente notificados de todos los reajustes automáticos pactados en este contrato y que han de operar durante la vigencia del mismo. ')
            ##
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo11='CLÁUSULA DÉCIMA. INCREMENTOS DEL VALOR DEL CANON DE ARRENDAMIENTO: '
            p11.add_run(titulo11).bold = True
            p11.add_run('Vencida la primer vigencia de este contrato, y así sucesivamente cada doce (12) mensualidades en caso de prórroga tácita o expresa en forma automática y sin necesidad de requerimiento alguno entre las partes, el precio mensual del arrendamiento se reajustará en una proporción igual al cien por ciento (100%) del incremento que haya tenido el índice de precios al consumidor (I.P.C.), certificado por el DANE para el año calendario inmediatamente anterior a aquel que se efectúe el incremento. ')
            p11.add_run('PARÁGRAFO.').bold = True
            p11.add_run(' Al suscribir el presente contrato EL ARRENDATARIO, y sus DEUDORES SOLIDARIOS quedan plenamente notificados de todos los reajustes automáticos pactados en este contrato y que han de operar durante la vigencia del mismo. ')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo11='CLÁUSULA DÉCIMA. INCREMENTOS DEL VALOR DEL CANON DE ARRENDAMIENTO: '
            p11.add_run(titulo11).bold = True
            p11.add_run('Vencida la primer vigencia de este contrato, y así sucesivamente cada doce (12) mensualidades en caso de prórroga tácita o expresa en forma automática y sin necesidad de requerimiento alguno entre las partes, el precio mensual del arrendamiento se reajustará en una proporción igual al cien por ciento (100%) del incremento que haya tenido el índice de precios al consumidor (I.P.C.), certificado por el DANE para el año calendario inmediatamente anterior a aquel que se efectúe el incremento. ')
            p11.add_run('PARÁGRAFO.').bold = True
            p11.add_run(' Al suscribir el presente contrato EL ARRENDATARIO, y sus DEUDORES SOLIDARIOS quedan plenamente notificados de todos los reajustes automáticos pactados en este contrato y que han de operar durante la vigencia del mismo. ')
            ##______________________
        p12 = documentoObjeto.add_paragraph('')
        p12.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            titulo12='CLÁUSULA UNDÉCIMA. SERVICIOS PÚBLICOS DOMICILIARIOS: '
            p12.add_run(titulo12).bold = True
            p12.add_run('A partir del momento en que el inmueble arrendado sea entregado al ARRENDATARIO y hasta la fecha de su entrega al ARRENDADOR, estará a cargo del ARRENDATARIO, el pago oportuno de los servicios públicos de acueducto ( '+Cuenta_Acueducto+' ), alcantarillado ( '+Cuenta_Alcatarillado+' ), recolección de basuras ( '+Cuenta_Basuras+' ), energía eléctrica ( '+Cuenta_Energia+' ), gas ( '+Cuenta_Gas+' ) y demás instalados en el inmueble, de acuerdo a la respectiva facturación. ')
            p12.add_run('PARÁGRAFO 1. ').bold = True
            p12.add_run('Las reclamaciones respecto de la óptima prestación o facturación de los servicios públicos serán tramitadas directamente por el ARRENDATARIO ante las respectivas empresas prestadoras del servicio. Cualquier otro servicio adicional o suntuario al que pretenda acceder el ARRENDATARIO, deberá ser autorizado por el ARRENDADOR. ')
            p12.add_run('PARÁGRAFO 2. ').bold = True
            p12.add_run('El incumplimiento de EL ARRENDATARIO en el pago oportuno de los servicios públicos del Inmueble se tendrá como incumplimiento del Contrato, pudiendo el ARRENDADOR darlo por terminado unilateralmente sin necesidad de requerimientos privado y judiciales previstos en la ley. ')
            p12.add_run('PARÁGRAFO 3. ').bold = True
            p12.add_run('Igualmente, si como consecuencia del no pago oportuno de las empresas respectivas se suspende el servicio, se retira el contador o la línea telefónica, será a cargo del ARRENDATARIO el pago de los intereses de mora y los gastos que demanden su reconexión. ')
            p12.add_run('PARÁGRAFO 4. ').bold = True 
            p12.add_run('El presente documento, así como los recibos cancelados por el ARRENDADOR, constituyen mérito ejecutivo para cobrar judicialmente al ARRENDATARIO y sus deudores solidarios, los servicios públicos o cuotas de administración que dejare de pagar. ')
            p12.add_run('PARÁGRAFO 5. ').bold = True 
            p12.add_run('En el evento que el inmueble cuente con el servicio de caldera, el ARRENDATARIO se obliga al pago del mismo de forma directa a la administración.  ')
            p12.add_run('PARÁGRAFO 6. ').bold = True 
            p12.add_run('EL ARRENDATARIO declara que ha recibido en buen estado de funcionamiento y conservación las instalaciones para uso de los servicios públicos domiciliarios del inmueble, y que se abstendrá de modificarlas sin permiso previo y escrito del ARRENDADOR. Asimismo, responderá por daños y/o violaciones de los reglamentos de las correspondientes empresas de servicios públicos. ')
            p12.add_run('PARÁGRAFO 7. ').bold = True 
            p12.add_run('Se prohíbe al ARRENDATARIO solicitar cualquier tipo de crédito de consumo, comercial, seguro, publicaciones, o publicidad sobre el inmueble arrendado o sus facturaciones. En caso incumplimiento, el ARRENDATARIO se sujetará a las sanciones legales correspondientes, y EL ARRENDADOR podrá abstenerse de recibir el inmueble hasta tanto se retire o transfiera la obligación y continuará el cobro de cánones de arrendamiento hasta la restitución efectiva del inmueble.  ')
            p12.add_run('PARÁGRAFO 8. ').bold = True 
            p12.add_run('El Inmueble se entrega en arrendamiento sin línea telefónica, internet o planes similares. EL ARRENDATARIO no podrá instalar en el Inmueble ninguna línea sin la aprobación previa y escrita de EL ARRENDADOR, y en caso de aprobación, EL ARRENDATARIO se obliga a cancelarla o trasladarla al término de la vigencia del presente contrato o de sus prórrogas. Sin la constancia de cancelación o del traslado EL ARRENDADOR se abstendrá de entregar a EL ARRENDATARIO paz y salvo correspondiente, ahora bien, en el evento que las empresas prestadoras del servicio de telefonica e internet no cuenten con el servicio en el inmueble o se demore su instalación, el ARRENDATARIO reconoce que estas eventualidades son ajenas al ARRENDADOR y bajo ninguna circunstancia será causal de terminación anticipada de contrato.. ')
            p12.add_run('PARÁGRAFO 9. ').bold = True 
            p12.add_run('EL ARRENDADOR en cualquier tiempo durante la vigencia de este Contrato, podrá exigir del ARRENDATARIO la presentación de las facturas de los servicios públicos del inmueble a fin de verificar la cancelación de los mismos. En el evento que EL ARRENDADOR llegare a comprobar que alguna de las facturas no ha sido pagada por EL ARRENDATARIO encontrándose vencido el plazo para el pago previsto en la respectiva factura, EL ARRENDADOR podrá terminar de manera inmediata este Contrato y exigir del ARRENDATARIO el pago de las sumas a que hubiere lugar.')
            
        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo12='CLÁUSULA UNDÉCIMA. SERVICIOS PÚBLICOS DOMICILIARIOS: '
            p12.add_run(titulo12).bold = True
            p12.add_run('A partir del momento en que el inmueble arrendado sea entregado al ARRENDATARIO y hasta la fecha de su entrega al ARRENDADOR, estará a cargo del ARRENDATARIO, el pago oportuno de los servicios públicos de acueducto ( '+Cuenta_Acueducto+' ), alcantarillado ( '+Cuenta_Alcatarillado+' ), recolección de basuras ( '+Cuenta_Basuras+' ), energía eléctrica ( '+Cuenta_Energia+' ), gas ( '+Cuenta_Gas+' ) y demás instalados en el inmueble, de acuerdo a la respectiva facturación. ')
            p12.add_run('PARÁGRAFO 1. ').bold = True
            p12.add_run('Las reclamaciones respecto de la óptima prestación o facturación de los servicios públicos serán tramitadas directamente por el ARRENDATARIO ante las respectivas empresas prestadoras del servicio. Cualquier otro servicio adicional o suntuario al que pretenda acceder el ARRENDATARIO, deberá ser autorizado por el ARRENDADOR. ')
            p12.add_run('PARÁGRAFO 2. ').bold = True
            p12.add_run('El incumplimiento de EL ARRENDATARIO en el pago oportuno de los servicios públicos del Inmueble se tendrá como incumplimiento del Contrato, pudiendo el ARRENDADOR darlo por terminado unilateralmente sin necesidad de requerimientos privado y judiciales previstos en la ley. ')
            p12.add_run('PARÁGRAFO 3. ').bold = True
            p12.add_run('Igualmente, si como consecuencia del no pago oportuno de las empresas respectivas se suspende el servicio, se retira el contador o la línea telefónica, será a cargo del ARRENDATARIO el pago de los intereses de mora y los gastos que demanden su reconexión. ')
            p12.add_run('PARÁGRAFO 4. ').bold = True 
            p12.add_run('El presente documento, así como los recibos cancelados por el ARRENDADOR, constituyen mérito ejecutivo para cobrar judicialmente al ARRENDATARIO, los servicios públicos o cuotas de administración que dejare de pagar. ')
            p12.add_run('PARÁGRAFO 5. ').bold = True 
            p12.add_run('En el evento que el inmueble cuente con el servicio de caldera, el ARRENDATARIO se obliga al pago del mismo de forma directa a la administración.  ')
            p12.add_run('PARÁGRAFO 6. ').bold = True 
            p12.add_run('EL ARRENDATARIO declara que ha recibido en buen estado de funcionamiento y conservación las instalaciones para uso de los servicios públicos domiciliarios del inmueble, y que se abstendrá de modificarlas sin permiso previo y escrito del ARRENDADOR. Asimismo, responderá por daños y/o violaciones de los reglamentos de las correspondientes empresas de servicios públicos. ')
            p12.add_run('PARÁGRAFO 7. ').bold = True 
            p12.add_run('Se prohíbe al ARRENDATARIO solicitar cualquier tipo de crédito de consumo, comercial, seguro, publicaciones, o publicidad sobre el inmueble arrendado o sus facturaciones. En caso incumplimiento, el ARRENDATARIO se sujetará a las sanciones legales correspondientes, y EL ARRENDADOR podrá abstenerse de recibir el inmueble hasta tanto se retire o transfiera la obligación y continuará el cobro de cánones de arrendamiento hasta la restitución efectiva del inmueble.  ')
            p12.add_run('PARÁGRAFO 8. ').bold = True 
            p12.add_run('El Inmueble se entrega en arrendamiento sin línea telefónica, internet o planes similares. EL ARRENDATARIO no podrá instalar en el Inmueble ninguna línea sin la aprobación previa y escrita de EL ARRENDADOR, y en caso de aprobación, EL ARRENDATARIO se obliga a cancelarla o trasladarla al término de la vigencia del presente contrato o de sus prórrogas. Sin la constancia de cancelación o del traslado EL ARRENDADOR se abstendrá de entregar a EL ARRENDATARIO paz y salvo correspondiente, ahora bien, en el evento que las empresas prestadoras del servicio de telefonica e internet no cuenten con el servicio en el inmueble o se demore su instalación, el ARRENDATARIO reconoce que estas eventualidades son ajenas al ARRENDADOR y bajo ninguna circunstancia será causal de terminación anticipada de contrato.. ')
            p12.add_run('PARÁGRAFO 9. ').bold = True 
            p12.add_run('EL ARRENDADOR en cualquier tiempo durante la vigencia de este Contrato, podrá exigir del ARRENDATARIO la presentación de las facturas de los servicios públicos del inmueble a fin de verificar la cancelación de los mismos. En el evento que EL ARRENDADOR llegare a comprobar que alguna de las facturas no ha sido pagada por EL ARRENDATARIO encontrándose vencido el plazo para el pago previsto en la respectiva factura, EL ARRENDADOR podrá terminar de manera inmediata este Contrato y exigir del ARRENDATARIO el pago de las sumas a que hubiere lugar.')

        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo12='CLÁUSULA DUODÉCIMA. SERVICIOS PÚBLICOS DOMICILIARIOS: '
            p12.add_run(titulo12).bold = True
            p12.add_run('A partir del momento en que el inmueble arrendado sea entregado al ARRENDATARIO y hasta la fecha de su entrega al ARRENDADOR, estará a cargo del ARRENDATARIO, el pago oportuno de los servicios públicos de acueducto ( '+Cuenta_Acueducto+' ), alcantarillado ( '+Cuenta_Alcatarillado+' ), recolección de basuras ( '+Cuenta_Basuras+' ), energía eléctrica ( '+Cuenta_Energia+' ), gas ( '+Cuenta_Gas+' ) y demás instalados en el inmueble, de acuerdo a la respectiva facturación. ')
            p12.add_run('PARÁGRAFO 1. ').bold = True
            p12.add_run('Las reclamaciones respecto de la óptima prestación o facturación de los servicios públicos serán tramitadas directamente por el ARRENDATARIO ante las respectivas empresas prestadoras del servicio. Cualquier otro servicio adicional o suntuario al que pretenda acceder el ARRENDATARIO, deberá ser autorizado por el ARRENDADOR. ')
            p12.add_run('PARÁGRAFO 2. ').bold = True
            p12.add_run('El incumplimiento de EL ARRENDATARIO en el pago oportuno de los servicios públicos del Inmueble se tendrá como incumplimiento del Contrato, pudiendo el ARRENDADOR darlo por terminado unilateralmente sin necesidad de requerimientos privado y judiciales previstos en la ley. ')
            p12.add_run('PARÁGRAFO 3. ').bold = True
            p12.add_run('Igualmente, si como consecuencia del no pago oportuno de las empresas respectivas se suspende el servicio, se retira el contador o la línea telefónica, será a cargo del ARRENDATARIO el pago de los intereses de mora y los gastos que demanden su reconexión. ')
            p12.add_run('PARÁGRAFO 4. ').bold = True 
            p12.add_run('El presente documento, así como los recibos cancelados por el ARRENDADOR, constituyen mérito ejecutivo para cobrar judicialmente al ARRENDATARIO, los servicios públicos o cuotas de administración que dejare de pagar. ')
            p12.add_run('PARÁGRAFO 5. ').bold = True 
            p12.add_run('En el evento que el inmueble cuente con el servicio de caldera, el ARRENDATARIO se obliga al pago del mismo de forma directa a la administración.  ')
            p12.add_run('PARÁGRAFO 6. ').bold = True 
            p12.add_run('EL ARRENDATARIO declara que ha recibido en buen estado de funcionamiento y conservación las instalaciones para uso de los servicios públicos domiciliarios del inmueble, y que se abstendrá de modificarlas sin permiso previo y escrito del ARRENDADOR. Asimismo, responderá por daños y/o violaciones de los reglamentos de las correspondientes empresas de servicios públicos. ')
            p12.add_run('PARÁGRAFO 7. ').bold = True 
            p12.add_run('Se prohíbe al ARRENDATARIO solicitar cualquier tipo de crédito de consumo, comercial, seguro, publicaciones, o publicidad sobre el inmueble arrendado o sus facturaciones. En caso incumplimiento, el ARRENDATARIO se sujetará a las sanciones legales correspondientes, y EL ARRENDADOR podrá abstenerse de recibir el inmueble hasta tanto se retire o transfiera la obligación y continuará el cobro de cánones de arrendamiento hasta la restitución efectiva del inmueble.  ')
            p12.add_run('PARÁGRAFO 8. ').bold = True 
            p12.add_run('El Inmueble se entrega en arrendamiento sin línea telefónica, internet o planes similares. EL ARRENDATARIO no podrá instalar en el Inmueble ninguna línea sin la aprobación previa y escrita de EL ARRENDADOR, y en caso de aprobación, EL ARRENDATARIO se obliga a cancelarla o trasladarla al término de la vigencia del presente contrato o de sus prórrogas. Sin la constancia de cancelación o del traslado EL ARRENDADOR se abstendrá de entregar a EL ARRENDATARIO paz y salvo correspondiente, ahora bien, en el evento que las empresas prestadoras del servicio de telefonica e internet no cuenten con el servicio en el inmueble o se demore su instalación, el ARRENDATARIO reconoce que estas eventualidades son ajenas al ARRENDADOR y bajo ninguna circunstancia será causal de terminación anticipada de contrato.. ')
            p12.add_run('PARÁGRAFO 9. ').bold = True 
            p12.add_run('EL ARRENDADOR en cualquier tiempo durante la vigencia de este Contrato, podrá exigir del ARRENDATARIO la presentación de las facturas de los servicios públicos del inmueble a fin de verificar la cancelación de los mismos. En el evento que EL ARRENDADOR llegare a comprobar que alguna de las facturas no ha sido pagada por EL ARRENDATARIO encontrándose vencido el plazo para el pago previsto en la respectiva factura, EL ARRENDADOR podrá terminar de manera inmediata este Contrato y exigir del ARRENDATARIO el pago de las sumas a que hubiere lugar.')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo12='CLÁUSULA DUODÉCIMA. SERVICIOS PÚBLICOS DOMICILIARIOS: '
            p12.add_run(titulo12).bold = True
            p12.add_run('A partir del momento en que el inmueble arrendado sea entregado al ARRENDATARIO y hasta la fecha de su entrega al ARRENDADOR, estará a cargo del ARRENDATARIO, el pago oportuno de los servicios públicos de acueducto ( '+Cuenta_Acueducto+' ), alcantarillado ( '+Cuenta_Alcatarillado+' ), recolección de basuras ( '+Cuenta_Basuras+' ), energía eléctrica ( '+Cuenta_Energia+' ), gas ( '+Cuenta_Gas+' ) y demás instalados en el inmueble, de acuerdo a la respectiva facturación. ')
            p12.add_run('PARÁGRAFO 1. ').bold = True
            p12.add_run('Las reclamaciones respecto de la óptima prestación o facturación de los servicios públicos serán tramitadas directamente por el ARRENDATARIO ante las respectivas empresas prestadoras del servicio. Cualquier otro servicio adicional o suntuario al que pretenda acceder el ARRENDATARIO, deberá ser autorizado por el ARRENDADOR. ')
            p12.add_run('PARÁGRAFO 2. ').bold = True
            p12.add_run('El incumplimiento de EL ARRENDATARIO en el pago oportuno de los servicios públicos del Inmueble se tendrá como incumplimiento del Contrato, pudiendo el ARRENDADOR darlo por terminado unilateralmente sin necesidad de requerimientos privado y judiciales previstos en la ley. ')
            p12.add_run('PARÁGRAFO 3. ').bold = True
            p12.add_run('Igualmente, si como consecuencia del no pago oportuno de las empresas respectivas se suspende el servicio, se retira el contador o la línea telefónica, será a cargo del ARRENDATARIO el pago de los intereses de mora y los gastos que demanden su reconexión. ')
            p12.add_run('PARÁGRAFO 4. ').bold = True 
            p12.add_run('El presente documento, así como los recibos cancelados por el ARRENDADOR, constituyen mérito ejecutivo para cobrar judicialmente al ARRENDATARIO y sus deudores solidarios, los servicios públicos o cuotas de administración que dejare de pagar. ')
            p12.add_run('PARÁGRAFO 5. ').bold = True 
            p12.add_run('En el evento que el inmueble cuente con el servicio de caldera, el ARRENDATARIO se obliga al pago del mismo de forma directa a la administración.  ')
            p12.add_run('PARÁGRAFO 6. ').bold = True 
            p12.add_run('EL ARRENDATARIO declara que ha recibido en buen estado de funcionamiento y conservación las instalaciones para uso de los servicios públicos domiciliarios del inmueble, y que se abstendrá de modificarlas sin permiso previo y escrito del ARRENDADOR. Asimismo, responderá por daños y/o violaciones de los reglamentos de las correspondientes empresas de servicios públicos. ')
            p12.add_run('PARÁGRAFO 7. ').bold = True 
            p12.add_run('Se prohíbe al ARRENDATARIO solicitar cualquier tipo de crédito de consumo, comercial, seguro, publicaciones, o publicidad sobre el inmueble arrendado o sus facturaciones. En caso incumplimiento, el ARRENDATARIO se sujetará a las sanciones legales correspondientes, y EL ARRENDADOR podrá abstenerse de recibir el inmueble hasta tanto se retire o transfiera la obligación y continuará el cobro de cánones de arrendamiento hasta la restitución efectiva del inmueble.  ')
            p12.add_run('PARÁGRAFO 8. ').bold = True 
            p12.add_run('El Inmueble se entrega en arrendamiento sin línea telefónica, internet o planes similares. EL ARRENDATARIO no podrá instalar en el Inmueble ninguna línea sin la aprobación previa y escrita de EL ARRENDADOR, y en caso de aprobación, EL ARRENDATARIO se obliga a cancelarla o trasladarla al término de la vigencia del presente contrato o de sus prórrogas. Sin la constancia de cancelación o del traslado EL ARRENDADOR se abstendrá de entregar a EL ARRENDATARIO paz y salvo correspondiente, ahora bien, en el evento que las empresas prestadoras del servicio de telefonica e internet no cuenten con el servicio en el inmueble o se demore su instalación, el ARRENDATARIO reconoce que estas eventualidades son ajenas al ARRENDADOR y bajo ninguna circunstancia será causal de terminación anticipada de contrato.. ')
            p12.add_run('PARÁGRAFO 9. ').bold = True 
            p12.add_run('EL ARRENDADOR en cualquier tiempo durante la vigencia de este Contrato, podrá exigir del ARRENDATARIO la presentación de las facturas de los servicios públicos del inmueble a fin de verificar la cancelación de los mismos. En el evento que EL ARRENDADOR llegare a comprobar que alguna de las facturas no ha sido pagada por EL ARRENDATARIO encontrándose vencido el plazo para el pago previsto en la respectiva factura, EL ARRENDADOR podrá terminar de manera inmediata este Contrato y exigir del ARRENDATARIO el pago de las sumas a que hubiere lugar.')
            ##Sin administrador
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
            titulo12='CLÁUSULA DÉCIMA. SERVICIOS PÚBLICOS DOMICILIARIOS: '
            p12.add_run(titulo12).bold = True
            p12.add_run('A partir del momento en que el inmueble arrendado sea entregado al ARRENDATARIO y hasta la fecha de su entrega al ARRENDADOR, estará a cargo del ARRENDATARIO, el pago oportuno de los servicios públicos de acueducto ( '+Cuenta_Acueducto+' ), alcantarillado ( '+Cuenta_Alcatarillado+' ), recolección de basuras ( '+Cuenta_Basuras+' ), energía eléctrica ( '+Cuenta_Energia+' ), gas ( '+Cuenta_Gas+' ) y demás instalados en el inmueble, de acuerdo a la respectiva facturación. ')
            p12.add_run('PARÁGRAFO 1. ').bold = True
            p12.add_run('Las reclamaciones respecto de la óptima prestación o facturación de los servicios públicos serán tramitadas directamente por el ARRENDATARIO ante las respectivas empresas prestadoras del servicio. Cualquier otro servicio adicional o suntuario al que pretenda acceder el ARRENDATARIO, deberá ser autorizado por el ARRENDADOR. ')
            p12.add_run('PARÁGRAFO 2. ').bold = True
            p12.add_run('El incumplimiento de EL ARRENDATARIO en el pago oportuno de los servicios públicos del Inmueble se tendrá como incumplimiento del Contrato, pudiendo el ARRENDADOR darlo por terminado unilateralmente sin necesidad de requerimientos privado y judiciales previstos en la ley. ')
            p12.add_run('PARÁGRAFO 3. ').bold = True
            p12.add_run('Igualmente, si como consecuencia del no pago oportuno de las empresas respectivas se suspende el servicio, se retira el contador o la línea telefónica, será a cargo del ARRENDATARIO el pago de los intereses de mora y los gastos que demanden su reconexión. ')
            p12.add_run('PARÁGRAFO 4. ').bold = True 
            p12.add_run('El presente documento, así como los recibos cancelados por el ARRENDADOR, constituyen mérito ejecutivo para cobrar judicialmente al ARRENDATARIO y sus deudores solidarios, los servicios públicos o cuotas de administración que dejare de pagar. ')
            p12.add_run('PARÁGRAFO 5. ').bold = True 
            p12.add_run('En el evento que el inmueble cuente con el servicio de caldera, el ARRENDATARIO se obliga al pago del mismo de forma directa a la administración.  ')
            p12.add_run('PARÁGRAFO 6. ').bold = True 
            p12.add_run('EL ARRENDATARIO declara que ha recibido en buen estado de funcionamiento y conservación las instalaciones para uso de los servicios públicos domiciliarios del inmueble, y que se abstendrá de modificarlas sin permiso previo y escrito del ARRENDADOR. Asimismo, responderá por daños y/o violaciones de los reglamentos de las correspondientes empresas de servicios públicos. ')
            p12.add_run('PARÁGRAFO 7. ').bold = True 
            p12.add_run('Se prohíbe al ARRENDATARIO solicitar cualquier tipo de crédito de consumo, comercial, seguro, publicaciones, o publicidad sobre el inmueble arrendado o sus facturaciones. En caso incumplimiento, el ARRENDATARIO se sujetará a las sanciones legales correspondientes, y EL ARRENDADOR podrá abstenerse de recibir el inmueble hasta tanto se retire o transfiera la obligación y continuará el cobro de cánones de arrendamiento hasta la restitución efectiva del inmueble.  ')
            p12.add_run('PARÁGRAFO 8. ').bold = True 
            p12.add_run('El Inmueble se entrega en arrendamiento sin línea telefónica, internet o planes similares. EL ARRENDATARIO no podrá instalar en el Inmueble ninguna línea sin la aprobación previa y escrita de EL ARRENDADOR, y en caso de aprobación, EL ARRENDATARIO se obliga a cancelarla o trasladarla al término de la vigencia del presente contrato o de sus prórrogas. Sin la constancia de cancelación o del traslado EL ARRENDADOR se abstendrá de entregar a EL ARRENDATARIO paz y salvo correspondiente, ahora bien, en el evento que las empresas prestadoras del servicio de telefonica e internet no cuenten con el servicio en el inmueble o se demore su instalación, el ARRENDATARIO reconoce que estas eventualidades son ajenas al ARRENDADOR y bajo ninguna circunstancia será causal de terminación anticipada de contrato.. ')
            p12.add_run('PARÁGRAFO 9. ').bold = True 
            p12.add_run('EL ARRENDADOR en cualquier tiempo durante la vigencia de este Contrato, podrá exigir del ARRENDATARIO la presentación de las facturas de los servicios públicos del inmueble a fin de verificar la cancelación de los mismos. En el evento que EL ARRENDADOR llegare a comprobar que alguna de las facturas no ha sido pagada por EL ARRENDATARIO encontrándose vencido el plazo para el pago previsto en la respectiva factura, EL ARRENDADOR podrá terminar de manera inmediata este Contrato y exigir del ARRENDATARIO el pago de las sumas a que hubiere lugar.')
            
        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo12='CLÁUSULA DÉCIMA. SERVICIOS PÚBLICOS DOMICILIARIOS: '
            p12.add_run(titulo12).bold = True
            p12.add_run('A partir del momento en que el inmueble arrendado sea entregado al ARRENDATARIO y hasta la fecha de su entrega al ARRENDADOR, estará a cargo del ARRENDATARIO, el pago oportuno de los servicios públicos de acueducto ( '+Cuenta_Acueducto+' ), alcantarillado ( '+Cuenta_Alcatarillado+' ), recolección de basuras ( '+Cuenta_Basuras+' ), energía eléctrica ( '+Cuenta_Energia+' ), gas ( '+Cuenta_Gas+' ) y demás instalados en el inmueble, de acuerdo a la respectiva facturación. ')
            p12.add_run('PARÁGRAFO 1. ').bold = True
            p12.add_run('Las reclamaciones respecto de la óptima prestación o facturación de los servicios públicos serán tramitadas directamente por el ARRENDATARIO ante las respectivas empresas prestadoras del servicio. Cualquier otro servicio adicional o suntuario al que pretenda acceder el ARRENDATARIO, deberá ser autorizado por el ARRENDADOR. ')
            p12.add_run('PARÁGRAFO 2. ').bold = True
            p12.add_run('El incumplimiento de EL ARRENDATARIO en el pago oportuno de los servicios públicos del Inmueble se tendrá como incumplimiento del Contrato, pudiendo el ARRENDADOR darlo por terminado unilateralmente sin necesidad de requerimientos privado y judiciales previstos en la ley. ')
            p12.add_run('PARÁGRAFO 3. ').bold = True
            p12.add_run('Igualmente, si como consecuencia del no pago oportuno de las empresas respectivas se suspende el servicio, se retira el contador o la línea telefónica, será a cargo del ARRENDATARIO el pago de los intereses de mora y los gastos que demanden su reconexión. ')
            p12.add_run('PARÁGRAFO 4. ').bold = True 
            p12.add_run('El presente documento, así como los recibos cancelados por el ARRENDADOR, constituyen mérito ejecutivo para cobrar judicialmente al ARRENDATARIO, los servicios públicos o cuotas de administración que dejare de pagar. ')
            p12.add_run('PARÁGRAFO 5. ').bold = True 
            p12.add_run('En el evento que el inmueble cuente con el servicio de caldera, el ARRENDATARIO se obliga al pago del mismo de forma directa a la administración.  ')
            p12.add_run('PARÁGRAFO 6. ').bold = True 
            p12.add_run('EL ARRENDATARIO declara que ha recibido en buen estado de funcionamiento y conservación las instalaciones para uso de los servicios públicos domiciliarios del inmueble, y que se abstendrá de modificarlas sin permiso previo y escrito del ARRENDADOR. Asimismo, responderá por daños y/o violaciones de los reglamentos de las correspondientes empresas de servicios públicos. ')
            p12.add_run('PARÁGRAFO 7. ').bold = True 
            p12.add_run('Se prohíbe al ARRENDATARIO solicitar cualquier tipo de crédito de consumo, comercial, seguro, publicaciones, o publicidad sobre el inmueble arrendado o sus facturaciones. En caso incumplimiento, el ARRENDATARIO se sujetará a las sanciones legales correspondientes, y EL ARRENDADOR podrá abstenerse de recibir el inmueble hasta tanto se retire o transfiera la obligación y continuará el cobro de cánones de arrendamiento hasta la restitución efectiva del inmueble.  ')
            p12.add_run('PARÁGRAFO 8. ').bold = True 
            p12.add_run('El Inmueble se entrega en arrendamiento sin línea telefónica, internet o planes similares. EL ARRENDATARIO no podrá instalar en el Inmueble ninguna línea sin la aprobación previa y escrita de EL ARRENDADOR, y en caso de aprobación, EL ARRENDATARIO se obliga a cancelarla o trasladarla al término de la vigencia del presente contrato o de sus prórrogas. Sin la constancia de cancelación o del traslado EL ARRENDADOR se abstendrá de entregar a EL ARRENDATARIO paz y salvo correspondiente, ahora bien, en el evento que las empresas prestadoras del servicio de telefonica e internet no cuenten con el servicio en el inmueble o se demore su instalación, el ARRENDATARIO reconoce que estas eventualidades son ajenas al ARRENDADOR y bajo ninguna circunstancia será causal de terminación anticipada de contrato.. ')
            p12.add_run('PARÁGRAFO 9. ').bold = True 
            p12.add_run('EL ARRENDADOR en cualquier tiempo durante la vigencia de este Contrato, podrá exigir del ARRENDATARIO la presentación de las facturas de los servicios públicos del inmueble a fin de verificar la cancelación de los mismos. En el evento que EL ARRENDADOR llegare a comprobar que alguna de las facturas no ha sido pagada por EL ARRENDATARIO encontrándose vencido el plazo para el pago previsto en la respectiva factura, EL ARRENDADOR podrá terminar de manera inmediata este Contrato y exigir del ARRENDATARIO el pago de las sumas a que hubiere lugar.')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo12='CLÁUSULA UNDÉCIMA. SERVICIOS PÚBLICOS DOMICILIARIOS: '
            p12.add_run(titulo12).bold = True
            p12.add_run('A partir del momento en que el inmueble arrendado sea entregado al ARRENDATARIO y hasta la fecha de su entrega al ARRENDADOR, estará a cargo del ARRENDATARIO, el pago oportuno de los servicios públicos de acueducto ( '+Cuenta_Acueducto+' ), alcantarillado ( '+Cuenta_Alcatarillado+' ), recolección de basuras ( '+Cuenta_Basuras+' ), energía eléctrica ( '+Cuenta_Energia+' ), gas ( '+Cuenta_Gas+' ) y demás instalados en el inmueble, de acuerdo a la respectiva facturación. ')
            p12.add_run('PARÁGRAFO 1. ').bold = True
            p12.add_run('Las reclamaciones respecto de la óptima prestación o facturación de los servicios públicos serán tramitadas directamente por el ARRENDATARIO ante las respectivas empresas prestadoras del servicio. Cualquier otro servicio adicional o suntuario al que pretenda acceder el ARRENDATARIO, deberá ser autorizado por el ARRENDADOR. ')
            p12.add_run('PARÁGRAFO 2. ').bold = True
            p12.add_run('El incumplimiento de EL ARRENDATARIO en el pago oportuno de los servicios públicos del Inmueble se tendrá como incumplimiento del Contrato, pudiendo el ARRENDADOR darlo por terminado unilateralmente sin necesidad de requerimientos privado y judiciales previstos en la ley. ')
            p12.add_run('PARÁGRAFO 3. ').bold = True
            p12.add_run('Igualmente, si como consecuencia del no pago oportuno de las empresas respectivas se suspende el servicio, se retira el contador o la línea telefónica, será a cargo del ARRENDATARIO el pago de los intereses de mora y los gastos que demanden su reconexión. ')
            p12.add_run('PARÁGRAFO 4. ').bold = True 
            p12.add_run('El presente documento, así como los recibos cancelados por el ARRENDADOR, constituyen mérito ejecutivo para cobrar judicialmente al ARRENDATARIO, los servicios públicos o cuotas de administración que dejare de pagar. ')
            p12.add_run('PARÁGRAFO 5. ').bold = True 
            p12.add_run('En el evento que el inmueble cuente con el servicio de caldera, el ARRENDATARIO se obliga al pago del mismo de forma directa a la administración.  ')
            p12.add_run('PARÁGRAFO 6. ').bold = True 
            p12.add_run('EL ARRENDATARIO declara que ha recibido en buen estado de funcionamiento y conservación las instalaciones para uso de los servicios públicos domiciliarios del inmueble, y que se abstendrá de modificarlas sin permiso previo y escrito del ARRENDADOR. Asimismo, responderá por daños y/o violaciones de los reglamentos de las correspondientes empresas de servicios públicos. ')
            p12.add_run('PARÁGRAFO 7. ').bold = True 
            p12.add_run('Se prohíbe al ARRENDATARIO solicitar cualquier tipo de crédito de consumo, comercial, seguro, publicaciones, o publicidad sobre el inmueble arrendado o sus facturaciones. En caso incumplimiento, el ARRENDATARIO se sujetará a las sanciones legales correspondientes, y EL ARRENDADOR podrá abstenerse de recibir el inmueble hasta tanto se retire o transfiera la obligación y continuará el cobro de cánones de arrendamiento hasta la restitución efectiva del inmueble.  ')
            p12.add_run('PARÁGRAFO 8. ').bold = True 
            p12.add_run('El Inmueble se entrega en arrendamiento sin línea telefónica, internet o planes similares. EL ARRENDATARIO no podrá instalar en el Inmueble ninguna línea sin la aprobación previa y escrita de EL ARRENDADOR, y en caso de aprobación, EL ARRENDATARIO se obliga a cancelarla o trasladarla al término de la vigencia del presente contrato o de sus prórrogas. Sin la constancia de cancelación o del traslado EL ARRENDADOR se abstendrá de entregar a EL ARRENDATARIO paz y salvo correspondiente, ahora bien, en el evento que las empresas prestadoras del servicio de telefonica e internet no cuenten con el servicio en el inmueble o se demore su instalación, el ARRENDATARIO reconoce que estas eventualidades son ajenas al ARRENDADOR y bajo ninguna circunstancia será causal de terminación anticipada de contrato.. ')
            p12.add_run('PARÁGRAFO 9. ').bold = True 
            p12.add_run('EL ARRENDADOR en cualquier tiempo durante la vigencia de este Contrato, podrá exigir del ARRENDATARIO la presentación de las facturas de los servicios públicos del inmueble a fin de verificar la cancelación de los mismos. En el evento que EL ARRENDADOR llegare a comprobar que alguna de las facturas no ha sido pagada por EL ARRENDATARIO encontrándose vencido el plazo para el pago previsto en la respectiva factura, EL ARRENDADOR podrá terminar de manera inmediata este Contrato y exigir del ARRENDATARIO el pago de las sumas a que hubiere lugar.')
        
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo12='CLÁUSULA UNDÉCIMA. SERVICIOS PÚBLICOS DOMICILIARIOS: '
            p12.add_run(titulo12).bold = True
            p12.add_run('A partir del momento en que el inmueble arrendado sea entregado al ARRENDATARIO y hasta la fecha de su entrega al ARRENDADOR, estará a cargo del ARRENDATARIO, el pago oportuno de los servicios públicos de acueducto ( '+Cuenta_Acueducto+' ), alcantarillado ( '+Cuenta_Alcatarillado+' ), recolección de basuras ( '+Cuenta_Basuras+' ), energía eléctrica ( '+Cuenta_Energia+' ), gas ( '+Cuenta_Gas+' ) y demás instalados en el inmueble, de acuerdo a la respectiva facturación. ')
            p12.add_run('PARÁGRAFO 1. ').bold = True
            p12.add_run('Las reclamaciones respecto de la óptima prestación o facturación de los servicios públicos serán tramitadas directamente por el ARRENDATARIO ante las respectivas empresas prestadoras del servicio. Cualquier otro servicio adicional o suntuario al que pretenda acceder el ARRENDATARIO, deberá ser autorizado por el ARRENDADOR. ')
            p12.add_run('PARÁGRAFO 2. ').bold = True
            p12.add_run('El incumplimiento de EL ARRENDATARIO en el pago oportuno de los servicios públicos del Inmueble se tendrá como incumplimiento del Contrato, pudiendo el ARRENDADOR darlo por terminado unilateralmente sin necesidad de requerimientos privado y judiciales previstos en la ley. ')
            p12.add_run('PARÁGRAFO 3. ').bold = True
            p12.add_run('Igualmente, si como consecuencia del no pago oportuno de las empresas respectivas se suspende el servicio, se retira el contador o la línea telefónica, será a cargo del ARRENDATARIO el pago de los intereses de mora y los gastos que demanden su reconexión. ')
            p12.add_run('PARÁGRAFO 4. ').bold = True 
            p12.add_run('El presente documento, así como los recibos cancelados por el ARRENDADOR, constituyen mérito ejecutivo para cobrar judicialmente al ARRENDATARIO y sus deudores solidarios, los servicios públicos o cuotas de administración que dejare de pagar. ')
            p12.add_run('PARÁGRAFO 5. ').bold = True 
            p12.add_run('En el evento que el inmueble cuente con el servicio de caldera, el ARRENDATARIO se obliga al pago del mismo de forma directa a la administración.  ')
            p12.add_run('PARÁGRAFO 6. ').bold = True 
            p12.add_run('EL ARRENDATARIO declara que ha recibido en buen estado de funcionamiento y conservación las instalaciones para uso de los servicios públicos domiciliarios del inmueble, y que se abstendrá de modificarlas sin permiso previo y escrito del ARRENDADOR. Asimismo, responderá por daños y/o violaciones de los reglamentos de las correspondientes empresas de servicios públicos. ')
            p12.add_run('PARÁGRAFO 7. ').bold = True 
            p12.add_run('Se prohíbe al ARRENDATARIO solicitar cualquier tipo de crédito de consumo, comercial, seguro, publicaciones, o publicidad sobre el inmueble arrendado o sus facturaciones. En caso incumplimiento, el ARRENDATARIO se sujetará a las sanciones legales correspondientes, y EL ARRENDADOR podrá abstenerse de recibir el inmueble hasta tanto se retire o transfiera la obligación y continuará el cobro de cánones de arrendamiento hasta la restitución efectiva del inmueble.  ')
            p12.add_run('PARÁGRAFO 8. ').bold = True 
            p12.add_run('El Inmueble se entrega en arrendamiento sin línea telefónica, internet o planes similares. EL ARRENDATARIO no podrá instalar en el Inmueble ninguna línea sin la aprobación previa y escrita de EL ARRENDADOR, y en caso de aprobación, EL ARRENDATARIO se obliga a cancelarla o trasladarla al término de la vigencia del presente contrato o de sus prórrogas. Sin la constancia de cancelación o del traslado EL ARRENDADOR se abstendrá de entregar a EL ARRENDATARIO paz y salvo correspondiente, ahora bien, en el evento que las empresas prestadoras del servicio de telefonica e internet no cuenten con el servicio en el inmueble o se demore su instalación, el ARRENDATARIO reconoce que estas eventualidades son ajenas al ARRENDADOR y bajo ninguna circunstancia será causal de terminación anticipada de contrato.. ')
            p12.add_run('PARÁGRAFO 9. ').bold = True 
            p12.add_run('EL ARRENDADOR en cualquier tiempo durante la vigencia de este Contrato, podrá exigir del ARRENDATARIO la presentación de las facturas de los servicios públicos del inmueble a fin de verificar la cancelación de los mismos. En el evento que EL ARRENDADOR llegare a comprobar que alguna de las facturas no ha sido pagada por EL ARRENDATARIO encontrándose vencido el plazo para el pago previsto en la respectiva factura, EL ARRENDADOR podrá terminar de manera inmediata este Contrato y exigir del ARRENDATARIO el pago de las sumas a que hubiere lugar.')

        ##___________________

        if(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            p13 = documentoObjeto.add_paragraph('')
            p13.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
            titulo13='CLÁUSULA DUODÉCIMA. CUOTAS DE ADMINISTRACIÓN: '
            p13.add_run(titulo13).bold = True
            p13.add_run('EL ARRENDATARIO se obliga a cancelar al ARRENDADOR junto con el canon de arrendamiento descrito en el presente contrato, la suma de '+Administracion_Letras+"  PESOS M/C. ($ "+valor_administracion+") por concepto de cuota de administración y los reajustes que señale la Asamblea de Propietarios, de acuerdo con el reglamento de propiedad horizontal que rige la copropiedad. ")
            p13.add_run('PARÁGRAFO 1. ').bold = True
            p13.add_run('Si el pago se realiza por fuera de los plazos de descuento que otorga la administración, el ARRENDATARIO perderá dicho beneficio y cancelará la tarifa plena estipulada por valor de '+administracion_plena+" PESOS M/C. ($ "+valor_pleno_administracion+"). ")
            p13.add_run('PARÁGRAFO 2. ').bold = True
            p13.add_run('EL ARRENDATARIO y los deudores solidarios renuncian expresamente a los requerimientos para su constitución en mora respecto a esta obligación pecuniaria. ')
            p13.add_run('PARÁGRAFO 3. ').bold = True
            p13.add_run('EL ARRENDATARIO asumirá las sanciones, multas o recargos que cualquier incumplimiento  genere y tenga contemplada la administración de la copropiedad. ')
            p13.add_run('PARÁGRAFO 4. ').bold = True
            p13.add_run('EL ARRENDATARIO, deberá informar por un medio comprobable toda comunicación de la administración y/o cualquier otra entidad que le competa al ARRENDADOR, respecto a los reportes de asambleas ordinarias y/o extraordinarias de propietarios, así como cualquier otro requerimiento tales como impuestos y/o comunicados que sean de su competencia, so pena que EL ARRENDATARIO responda por las multas impuestas derivadas de la falta de información y/o participación del ARRENDADOR en las mismas. Los comunicados se entenderán recibidos por el ARRENDATARIO si fuesen entregados personalmente, por aviso o en casillero del bien inmueble.')
            
        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            p13 = documentoObjeto.add_paragraph('')
            p13.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            titulo13='CLÁUSULA DUODÉCIMA. CUOTAS DE ADMINISTRACIÓN: '
            p13.add_run(titulo13).bold = True
            p13.add_run('EL ARRENDATARIO se obliga a cancelar al ARRENDADOR junto con el canon de arrendamiento descrito en el presente contrato, la suma de '+Administracion_Letras+"  PESOS M/C. ($ "+valor_administracion+") por concepto de cuota de administración y los reajustes que señale la Asamblea de Propietarios, de acuerdo con el reglamento de propiedad horizontal que rige la copropiedad. ")
            p13.add_run('PARÁGRAFO 1. ').bold = True
            p13.add_run('Si el pago se realiza por fuera de los plazos de descuento que otorga la administración, el ARRENDATARIO perderá dicho beneficio y cancelará la tarifa plena estipulada por valor de '+administracion_plena+" PESOS M/C. ($ "+valor_pleno_administracion+"). ")
            p13.add_run('PARÁGRAFO 2. ').bold = True
            p13.add_run('EL ARRENDATARIO  renuncia expresamente a los requerimientos para su constitución en mora respecto a esta obligación pecuniaria. ')
            p13.add_run('PARÁGRAFO 3. ').bold = True
            p13.add_run('EL ARRENDATARIO asumirá las sanciones, multas o recargos que cualquier incumplimiento  genere y tenga contemplada la administración de la copropiedad. ')
            p13.add_run('PARÁGRAFO 4. ').bold = True
            p13.add_run('EL ARRENDATARIO, deberá informar por un medio comprobable toda comunicación de la administración y/o cualquier otra entidad que le competa al ARRENDADOR, respecto a los reportes de asambleas ordinarias y/o extraordinarias de propietarios, así como cualquier otro requerimiento tales como impuestos y/o comunicados que sean de su competencia, so pena que EL ARRENDATARIO responda por las multas impuestas derivadas de la falta de información y/o participación del ARRENDADOR en las mismas. Los comunicados se entenderán recibidos por el ARRENDATARIO si fuesen entregados personalmente, por aviso o en casillero del bien inmueble.')

        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            p13 = documentoObjeto.add_paragraph('')
            p13.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            titulo13='CLÁUSULA DÉCIMO TERCERA. CUOTAS DE ADMINISTRACIÓN: '
            p13.add_run(titulo13).bold = True
            p13.add_run('EL ARRENDATARIO se obliga a cancelar al ARRENDADOR junto con el canon de arrendamiento descrito en el presente contrato, la suma de '+Administracion_Letras+"  PESOS M/C. ($ "+valor_administracion+") por concepto de cuota de administración y los reajustes que señale la Asamblea de Propietarios, de acuerdo con el reglamento de propiedad horizontal que rige la copropiedad. ")
            p13.add_run('PARÁGRAFO 1. ').bold = True
            p13.add_run('Si el pago se realiza por fuera de los plazos de descuento que otorga la administración, el ARRENDATARIO perderá dicho beneficio y cancelará la tarifa plena estipulada por valor de '+administracion_plena+" PESOS M/C. ($ "+valor_pleno_administracion+"). ")
            p13.add_run('PARÁGRAFO 2. ').bold = True
            p13.add_run('EL ARRENDATARIO  renuncia expresamente a los requerimientos para su constitución en mora respecto a esta obligación pecuniaria. ')
            p13.add_run('PARÁGRAFO 3. ').bold = True
            p13.add_run('EL ARRENDATARIO asumirá las sanciones, multas o recargos que cualquier incumplimiento  genere y tenga contemplada la administración de la copropiedad. ')
            p13.add_run('PARÁGRAFO 4. ').bold = True
            p13.add_run('EL ARRENDATARIO, deberá informar por un medio comprobable toda comunicación de la administración y/o cualquier otra entidad que le competa al ARRENDADOR, respecto a los reportes de asambleas ordinarias y/o extraordinarias de propietarios, así como cualquier otro requerimiento tales como impuestos y/o comunicados que sean de su competencia, so pena que EL ARRENDATARIO responda por las multas impuestas derivadas de la falta de información y/o participación del ARRENDADOR en las mismas. Los comunicados se entenderán recibidos por el ARRENDATARIO si fuesen entregados personalmente, por aviso o en casillero del bien inmueble.')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            p13 = documentoObjeto.add_paragraph('')
            p13.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            titulo13='CLÁUSULA DÉCIMO TERCERA. CUOTAS DE ADMINISTRACIÓN: '
            p13.add_run(titulo13).bold = True
            p13.add_run('EL ARRENDATARIO se obliga a cancelar al ARRENDADOR junto con el canon de arrendamiento descrito en el presente contrato, la suma de '+Administracion_Letras+"  PESOS M/C. ($ "+valor_administracion+") por concepto de cuota de administración y los reajustes que señale la Asamblea de Propietarios, de acuerdo con el reglamento de propiedad horizontal que rige la copropiedad. ")
            p13.add_run('PARÁGRAFO 1. ').bold = True
            p13.add_run('Si el pago se realiza por fuera de los plazos de descuento que otorga la administración, el ARRENDATARIO perderá dicho beneficio y cancelará la tarifa plena estipulada por valor de '+administracion_plena+" PESOS M/C. ($ "+valor_pleno_administracion+"). ")
            p13.add_run('PARÁGRAFO 2. ').bold = True
            p13.add_run('EL ARRENDATARIO y los deudores solidarios renuncian expresamente a los requerimientos para su constitución en mora respecto a esta obligación pecuniaria. ')
            p13.add_run('PARÁGRAFO 3. ').bold = True
            p13.add_run('EL ARRENDATARIO asumirá las sanciones, multas o recargos que cualquier incumplimiento  genere y tenga contemplada la administración de la copropiedad. ')
            p13.add_run('PARÁGRAFO 4. ').bold = True
            p13.add_run('EL ARRENDATARIO, deberá informar por un medio comprobable toda comunicación de la administración y/o cualquier otra entidad que le competa al ARRENDADOR, respecto a los reportes de asambleas ordinarias y/o extraordinarias de propietarios, así como cualquier otro requerimiento tales como impuestos y/o comunicados que sean de su competencia, so pena que EL ARRENDATARIO responda por las multas impuestas derivadas de la falta de información y/o participación del ARRENDADOR en las mismas. Los comunicados se entenderán recibidos por el ARRENDATARIO si fuesen entregados personalmente, por aviso o en casillero del bien inmueble.')

        ##Sin administrador
        if(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
        
            pass
        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            pass
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            pass
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            pass
        #____________

        p14 = documentoObjeto.add_paragraph('')
        p14.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            titulo14='CLÁUSULA DÉCIMO TERCERA. TERMINACIÓN POR PARTE DEL ARRENDADOR: '
            p14.add_run(titulo14).bold = True
            p14.add_run('Son causales para que el arrendador pueda pedir unilateralmente la terminación del contrato, las siguientes: \n'+"1. La no cancelación por parte del arrendatario de las rentas y reajustes dentro del término estipulado en el contrato.\n"+"2. La no cancelación de los servicios públicos, que cause la desconexión o pérdida del servicio, o el incumplimiento del pago de las expensas comunes cuando estuviere a cargo del arrendatario.\n"+"3. El subarriendo total o parcial del inmueble, la cesión del contrato o del goce del inmueble o el cambio de destinación del mismo por parte del arrendatario, sin expresa autorización del arrendador\n"+"4. La incursión reiterada del arrendatario en procederes que afecten la tranquilidad ciudadana de los vecinos, o la destinación del inmueble para actos delictivos o que impliquen contravención, debidamente comprobados ante la autoridad policiva.\n"+"5. La realización de mejoras, cambios o ampliaciones del inmueble, sin expresa autorización del arrendador o la destrucción total o parcial del inmueble o área arrendada por parte del arrendatario.\n"+"6. La violación por el arrendatario a las normas del respectivo reglamento de propiedad horizontal cuando se trate de viviendas sometidas a ese régimen.\n"+"7. El arrendador podrá dar por terminado unilateralmente el contrato de arrendamiento durante las prórrogas, previo aviso escrito dirigido al arrendatario a través del servicio postal autorizado, con una antelación no menor de tres (3) meses y el pago de una indemnización equivalente al precio de tres (3) meses de arrendamiento. Cumplidas estas condiciones el arrendatario estará obligado a restituir el inmueble.\t\n")
            p14.add_run("8. El arrendador podrá dar por terminado unilateralmente el contrato de arrendamiento a la fecha de vencimiento del término inicial o de sus prórrogas invocando cualquiera de las siguientes causales especiales de restitución, previo aviso escrito al arrendatario a través del servicio postal autorizado con una antelación no menor a tres (3) meses a la referida fecha de vencimiento:\n"+"a) Cuando el propietario o poseedor del inmueble necesitare ocuparlo para su propia habitación, por un término no menor de un (1) año;\n"+" b) Cuando el inmueble haya de demolerse para efectuar una nueva construcción, o cuando se requiere desocuparlo con el fin de ejecutar obras independientes para su reparación;\n c) Cuando haya de entregarse en cumplimiento de las obligaciones originadas en un contrato de compraventa;\n d) La plena voluntad de dar por terminado el contrato, siempre y cuando, el contrato de arrendamiento cumpliere como mínimo cuatro (4) años de ejecución. El arrendador deberá indemnizar al arrendatario con una suma equivalente al precio de uno punto cinco (1.5) meses de arrendamiento.\t\n")
            p14.add_run("Cuando se trate de las causales previstas en los literales a), b) y c), el arrendador acompañará al aviso escrito la constancia de haber constituido una caución en dinero, bancaria u otorgada por compañía de seguros legalmente reconocida, constituida a favor del arrendatario por un valor equivalente a seis (6) meses del precio del arrendamiento vigente, para garantizar el cumplimiento de la causal invocada dentro de los seis (6) meses siguientes a la fecha de la restitución.\t\n"+"Cuando se trate de la causal prevista en el literal d), el pago de la indemnización se realizará mediante el mismo procedimiento establecido en el artículo 23 de la Ley 820 de 2003.\t\n  De no mediar constancia por escrito del preaviso, el contrato de arrendamiento se entenderá renovado automáticamente por un término igual al inicialmente pactado.")
            
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo14='CLÁUSULA DÉCIMO CUARTA. TERMINACIÓN POR PARTE DEL ARRENDADOR: '
            p14.add_run(titulo14).bold = True
            p14.add_run('Son causales para que el arrendador pueda pedir unilateralmente la terminación del contrato, las siguientes: \n'+"1. La no cancelación por parte del arrendatario de las rentas y reajustes dentro del término estipulado en el contrato.\n"+"2. La no cancelación de los servicios públicos, que cause la desconexión o pérdida del servicio, o el incumplimiento del pago de las expensas comunes cuando estuviere a cargo del arrendatario.\n"+"3. El subarriendo total o parcial del inmueble, la cesión del contrato o del goce del inmueble o el cambio de destinación del mismo por parte del arrendatario, sin expresa autorización del arrendador\n"+"4. La incursión reiterada del arrendatario en procederes que afecten la tranquilidad ciudadana de los vecinos, o la destinación del inmueble para actos delictivos o que impliquen contravención, debidamente comprobados ante la autoridad policiva.\n"+"5. La realización de mejoras, cambios o ampliaciones del inmueble, sin expresa autorización del arrendador o la destrucción total o parcial del inmueble o área arrendada por parte del arrendatario.\n"+"6. La violación por el arrendatario a las normas del respectivo reglamento de propiedad horizontal cuando se trate de viviendas sometidas a ese régimen.\n"+"7. El arrendador podrá dar por terminado unilateralmente el contrato de arrendamiento durante las prórrogas, previo aviso escrito dirigido al arrendatario a través del servicio postal autorizado, con una antelación no menor de tres (3) meses y el pago de una indemnización equivalente al precio de tres (3) meses de arrendamiento. Cumplidas estas condiciones el arrendatario estará obligado a restituir el inmueble.\t\n")
            p14.add_run("8. El arrendador podrá dar por terminado unilateralmente el contrato de arrendamiento a la fecha de vencimiento del término inicial o de sus prórrogas invocando cualquiera de las siguientes causales especiales de restitución, previo aviso escrito al arrendatario a través del servicio postal autorizado con una antelación no menor a tres (3) meses a la referida fecha de vencimiento:\n"+"a) Cuando el propietario o poseedor del inmueble necesitare ocuparlo para su propia habitación, por un término no menor de un (1) año;\n"+" b) Cuando el inmueble haya de demolerse para efectuar una nueva construcción, o cuando se requiere desocuparlo con el fin de ejecutar obras independientes para su reparación;\n c) Cuando haya de entregarse en cumplimiento de las obligaciones originadas en un contrato de compraventa;\n d) La plena voluntad de dar por terminado el contrato, siempre y cuando, el contrato de arrendamiento cumpliere como mínimo cuatro (4) años de ejecución. El arrendador deberá indemnizar al arrendatario con una suma equivalente al precio de uno punto cinco (1.5) meses de arrendamiento.\t\n")
            p14.add_run("Cuando se trate de las causales previstas en los literales a), b) y c), el arrendador acompañará al aviso escrito la constancia de haber constituido una caución en dinero, bancaria u otorgada por compañía de seguros legalmente reconocida, constituida a favor del arrendatario por un valor equivalente a seis (6) meses del precio del arrendamiento vigente, para garantizar el cumplimiento de la causal invocada dentro de los seis (6) meses siguientes a la fecha de la restitución.\t\n"+"Cuando se trate de la causal prevista en el literal d), el pago de la indemnización se realizará mediante el mismo procedimiento establecido en el artículo 23 de la Ley 820 de 2003.\t\n  De no mediar constancia por escrito del preaviso, el contrato de arrendamiento se entenderá renovado automáticamente por un término igual al inicialmente pactado.")

        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "No" or Respuesta3=="no" or Respuesta3=="NO"):
            titulo14='CLÁUSULA DUODÉCIMA. TERMINACIÓN POR PARTE DEL ARRENDADOR: '
            p14.add_run(titulo14).bold = True
            p14.add_run('Son causales para que el arrendador pueda pedir unilateralmente la terminación del contrato, las siguientes: \n'+"1. La no cancelación por parte del arrendatario de las rentas y reajustes dentro del término estipulado en el contrato.\n"+"2. La no cancelación de los servicios públicos, que cause la desconexión o pérdida del servicio, o el incumplimiento del pago de las expensas comunes cuando estuviere a cargo del arrendatario.\n"+"3. El subarriendo total o parcial del inmueble, la cesión del contrato o del goce del inmueble o el cambio de destinación del mismo por parte del arrendatario, sin expresa autorización del arrendador\n"+"4. La incursión reiterada del arrendatario en procederes que afecten la tranquilidad ciudadana de los vecinos, o la destinación del inmueble para actos delictivos o que impliquen contravención, debidamente comprobados ante la autoridad policiva.\n"+"5. La realización de mejoras, cambios o ampliaciones del inmueble, sin expresa autorización del arrendador o la destrucción total o parcial del inmueble o área arrendada por parte del arrendatario.\n"+"6. La violación por el arrendatario a las normas del respectivo reglamento de propiedad horizontal cuando se trate de viviendas sometidas a ese régimen.\n"+"7. El arrendador podrá dar por terminado unilateralmente el contrato de arrendamiento durante las prórrogas, previo aviso escrito dirigido al arrendatario a través del servicio postal autorizado, con una antelación no menor de tres (3) meses y el pago de una indemnización equivalente al precio de tres (3) meses de arrendamiento. Cumplidas estas condiciones el arrendatario estará obligado a restituir el inmueble.\t\n")
            p14.add_run("8. El arrendador podrá dar por terminado unilateralmente el contrato de arrendamiento a la fecha de vencimiento del término inicial o de sus prórrogas invocando cualquiera de las siguientes causales especiales de restitución, previo aviso escrito al arrendatario a través del servicio postal autorizado con una antelación no menor a tres (3) meses a la referida fecha de vencimiento:\n"+"a) Cuando el propietario o poseedor del inmueble necesitare ocuparlo para su propia habitación, por un término no menor de un (1) año;\n"+" b) Cuando el inmueble haya de demolerse para efectuar una nueva construcción, o cuando se requiere desocuparlo con el fin de ejecutar obras independientes para su reparación;\n c) Cuando haya de entregarse en cumplimiento de las obligaciones originadas en un contrato de compraventa;\n d) La plena voluntad de dar por terminado el contrato, siempre y cuando, el contrato de arrendamiento cumpliere como mínimo cuatro (4) años de ejecución. El arrendador deberá indemnizar al arrendatario con una suma equivalente al precio de uno punto cinco (1.5) meses de arrendamiento.\t\n")
            p14.add_run("Cuando se trate de las causales previstas en los literales a), b) y c), el arrendador acompañará al aviso escrito la constancia de haber constituido una caución en dinero, bancaria u otorgada por compañía de seguros legalmente reconocida, constituida a favor del arrendatario por un valor equivalente a seis (6) meses del precio del arrendamiento vigente, para garantizar el cumplimiento de la causal invocada dentro de los seis (6) meses siguientes a la fecha de la restitución.\t\n"+"Cuando se trate de la causal prevista en el literal d), el pago de la indemnización se realizará mediante el mismo procedimiento establecido en el artículo 23 de la Ley 820 de 2003.\t\n  De no mediar constancia por escrito del preaviso, el contrato de arrendamiento se entenderá renovado automáticamente por un término igual al inicialmente pactado.")

        elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo14='CLÁUSULA UNDÉCIMA.  TERMINACIÓN POR PARTE DEL ARRENDADOR: '
            p14.add_run(titulo14).bold = True
            p14.add_run('Son causales para que el arrendador pueda pedir unilateralmente la terminación del contrato, las siguientes: \n'+"1. La no cancelación por parte del arrendatario de las rentas y reajustes dentro del término estipulado en el contrato.\n"+"2. La no cancelación de los servicios públicos, que cause la desconexión o pérdida del servicio, o el incumplimiento del pago de las expensas comunes cuando estuviere a cargo del arrendatario.\n"+"3. El subarriendo total o parcial del inmueble, la cesión del contrato o del goce del inmueble o el cambio de destinación del mismo por parte del arrendatario, sin expresa autorización del arrendador\n"+"4. La incursión reiterada del arrendatario en procederes que afecten la tranquilidad ciudadana de los vecinos, o la destinación del inmueble para actos delictivos o que impliquen contravención, debidamente comprobados ante la autoridad policiva.\n"+"5. La realización de mejoras, cambios o ampliaciones del inmueble, sin expresa autorización del arrendador o la destrucción total o parcial del inmueble o área arrendada por parte del arrendatario.\n"+"6. La violación por el arrendatario a las normas del respectivo reglamento de propiedad horizontal cuando se trate de viviendas sometidas a ese régimen.\n"+"7. El arrendador podrá dar por terminado unilateralmente el contrato de arrendamiento durante las prórrogas, previo aviso escrito dirigido al arrendatario a través del servicio postal autorizado, con una antelación no menor de tres (3) meses y el pago de una indemnización equivalente al precio de tres (3) meses de arrendamiento. Cumplidas estas condiciones el arrendatario estará obligado a restituir el inmueble.\t\n")
            p14.add_run("8. El arrendador podrá dar por terminado unilateralmente el contrato de arrendamiento a la fecha de vencimiento del término inicial o de sus prórrogas invocando cualquiera de las siguientes causales especiales de restitución, previo aviso escrito al arrendatario a través del servicio postal autorizado con una antelación no menor a tres (3) meses a la referida fecha de vencimiento:\n"+"a) Cuando el propietario o poseedor del inmueble necesitare ocuparlo para su propia habitación, por un término no menor de un (1) año;\n"+" b) Cuando el inmueble haya de demolerse para efectuar una nueva construcción, o cuando se requiere desocuparlo con el fin de ejecutar obras independientes para su reparación;\n c) Cuando haya de entregarse en cumplimiento de las obligaciones originadas en un contrato de compraventa;\n d) La plena voluntad de dar por terminado el contrato, siempre y cuando, el contrato de arrendamiento cumpliere como mínimo cuatro (4) años de ejecución. El arrendador deberá indemnizar al arrendatario con una suma equivalente al precio de uno punto cinco (1.5) meses de arrendamiento.\t\n")
            p14.add_run("Cuando se trate de las causales previstas en los literales a), b) y c), el arrendador acompañará al aviso escrito la constancia de haber constituido una caución en dinero, bancaria u otorgada por compañía de seguros legalmente reconocida, constituida a favor del arrendatario por un valor equivalente a seis (6) meses del precio del arrendamiento vigente, para garantizar el cumplimiento de la causal invocada dentro de los seis (6) meses siguientes a la fecha de la restitución.\t\n"+"Cuando se trate de la causal prevista en el literal d), el pago de la indemnización se realizará mediante el mismo procedimiento establecido en el artículo 23 de la Ley 820 de 2003.\t\n  De no mediar constancia por escrito del preaviso, el contrato de arrendamiento se entenderá renovado automáticamente por un término igual al inicialmente pactado.")

        ####
        p15 = documentoObjeto.add_paragraph('')
        p15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI")  :
            titulo15='CLÁUSULA DÉCIMO CUARTA. TERMINACIÓN UNILATERAL POR PARTE DEL ARRENDADOR MEDIANTE PREAVISO CON INDEMNIZACIÓN: '
            p15.add_run(titulo15).bold = True
            p15.add_run(' Para que el arrendador pueda dar por terminado unilateralmente el contrato de arrendamiento en el evento previsto en el numeral 7 del artículo anterior, deberá cumplir con los siguientes requisitos:\t\t\t\na) Comunicar a través del servicio postal autorizado al arrendatario o a su representante legal, con la antelación allí prevista, indicando la fecha para la terminación del contrato y, manifestando que se pagará la indemnización de ley.\t\nb) Consignar a favor del arrendatario y a órdenes de la autoridad competente, la indemnización de que trata el artículo anterior de la presente ley, dentro de los tres (3) meses anteriores a la fecha señalada para la terminación unilateral del contrato. La consignación se efectuará en las entidades autorizadas por el Gobierno Nacional para tal efecto y la autoridad competente allegará copia del título respectivo al arrendatario o le enviará comunicación en que se haga constar tal circunstancia, inmediatamente tenga conocimiento de la misma.El valor de la indemnización se hará con base en la renta vigente a la fecha del preaviso;\t\n c) Al momento de efectuar la consignación se dejará constancia en los respectivos títulos de las causas de la misma como también el nombre y dirección precisa del arrendatario o su representante;\t\n d) Si el arrendatario cumple con la obligación de entregar el inmueble en la fecha señalada, recibirá el pago de la indemnización, de conformidad con la autorización que expida la autoridad competente.')
            p15.add_run(' PARÁGRAFO 1. ').bold = True
            p15.add_run('En caso de que el arrendatario no entregue el inmueble, el arrendador tendrá derecho a que se le devuelva la indemnización consignada, sin perjuicio de que pueda iniciar el correspondiente proceso de restitución del inmueble.')
            p15.add_run(' PARÁGRAFO 2. ').bold = True
            p15.add_run('Si el arrendador con la aceptación del arrendatario desiste de dar por terminado el contrato de arrendamiento, podrá solicitar a la autoridad competente, la autorización para la devolución de la suma consignada.')

        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo15='CLÁUSULA DÉCIMO QUINTA. TERMINACIÓN UNILATERAL POR PARTE DEL ARRENDADOR MEDIANTE PREAVISO CON INDEMNIZACIÓN: '
            p15.add_run(titulo15).bold = True
            p15.add_run(' Para que el arrendador pueda dar por terminado unilateralmente el contrato de arrendamiento en el evento previsto en el numeral 7 del artículo anterior, deberá cumplir con los siguientes requisitos:\t\t\t\na) Comunicar a través del servicio postal autorizado al arrendatario o a su representante legal, con la antelación allí prevista, indicando la fecha para la terminación del contrato y, manifestando que se pagará la indemnización de ley.\t\nb) Consignar a favor del arrendatario y a órdenes de la autoridad competente, la indemnización de que trata el artículo anterior de la presente ley, dentro de los tres (3) meses anteriores a la fecha señalada para la terminación unilateral del contrato. La consignación se efectuará en las entidades autorizadas por el Gobierno Nacional para tal efecto y la autoridad competente allegará copia del título respectivo al arrendatario o le enviará comunicación en que se haga constar tal circunstancia, inmediatamente tenga conocimiento de la misma.El valor de la indemnización se hará con base en la renta vigente a la fecha del preaviso;\t\n c) Al momento de efectuar la consignación se dejará constancia en los respectivos títulos de las causas de la misma como también el nombre y dirección precisa del arrendatario o su representante;\t\n d) Si el arrendatario cumple con la obligación de entregar el inmueble en la fecha señalada, recibirá el pago de la indemnización, de conformidad con la autorización que expida la autoridad competente.')
            p15.add_run(' PARÁGRAFO 1. ').bold = True
            p15.add_run('En caso de que el arrendatario no entregue el inmueble, el arrendador tendrá derecho a que se le devuelva la indemnización consignada, sin perjuicio de que pueda iniciar el correspondiente proceso de restitución del inmueble.')
            p15.add_run(' PARÁGRAFO 2. ').bold = True
            p15.add_run('Si el arrendador con la aceptación del arrendatario desiste de dar por terminado el contrato de arrendamiento, podrá solicitar a la autoridad competente, la autorización para la devolución de la suma consignada.')

        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo15='CLÁUSULA DÉCIMO TERCERA. TERMINACIÓN UNILATERAL POR PARTE DEL ARRENDADOR MEDIANTE PREAVISO CON INDEMNIZACIÓN: '
            p15.add_run(titulo15).bold = True
            p15.add_run(' Para que el arrendador pueda dar por terminado unilateralmente el contrato de arrendamiento en el evento previsto en el numeral 7 del artículo anterior, deberá cumplir con los siguientes requisitos:\t\t\t\na) Comunicar a través del servicio postal autorizado al arrendatario o a su representante legal, con la antelación allí prevista, indicando la fecha para la terminación del contrato y, manifestando que se pagará la indemnización de ley.\t\nb) Consignar a favor del arrendatario y a órdenes de la autoridad competente, la indemnización de que trata el artículo anterior de la presente ley, dentro de los tres (3) meses anteriores a la fecha señalada para la terminación unilateral del contrato. La consignación se efectuará en las entidades autorizadas por el Gobierno Nacional para tal efecto y la autoridad competente allegará copia del título respectivo al arrendatario o le enviará comunicación en que se haga constar tal circunstancia, inmediatamente tenga conocimiento de la misma.El valor de la indemnización se hará con base en la renta vigente a la fecha del preaviso;\t\n c) Al momento de efectuar la consignación se dejará constancia en los respectivos títulos de las causas de la misma como también el nombre y dirección precisa del arrendatario o su representante;\t\n d) Si el arrendatario cumple con la obligación de entregar el inmueble en la fecha señalada, recibirá el pago de la indemnización, de conformidad con la autorización que expida la autoridad competente.')
            p15.add_run(' PARÁGRAFO 1. ').bold = True
            p15.add_run('En caso de que el arrendatario no entregue el inmueble, el arrendador tendrá derecho a que se le devuelva la indemnización consignada, sin perjuicio de que pueda iniciar el correspondiente proceso de restitución del inmueble.')
            p15.add_run(' PARÁGRAFO 2. ').bold = True
            p15.add_run('Si el arrendador con la aceptación del arrendatario desiste de dar por terminado el contrato de arrendamiento, podrá solicitar a la autoridad competente, la autorización para la devolución de la suma consignada.')

        elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo15='CLÁUSULA DUODÉCIMA. TERMINACIÓN UNILATERAL POR PARTE DEL ARRENDADOR MEDIANTE PREAVISO CON INDEMNIZACIÓN: '
            p15.add_run(titulo15).bold = True
            p15.add_run(' Para que el arrendador pueda dar por terminado unilateralmente el contrato de arrendamiento en el evento previsto en el numeral 7 del artículo anterior, deberá cumplir con los siguientes requisitos:\t\t\t\na) Comunicar a través del servicio postal autorizado al arrendatario o a su representante legal, con la antelación allí prevista, indicando la fecha para la terminación del contrato y, manifestando que se pagará la indemnización de ley.\t\nb) Consignar a favor del arrendatario y a órdenes de la autoridad competente, la indemnización de que trata el artículo anterior de la presente ley, dentro de los tres (3) meses anteriores a la fecha señalada para la terminación unilateral del contrato. La consignación se efectuará en las entidades autorizadas por el Gobierno Nacional para tal efecto y la autoridad competente allegará copia del título respectivo al arrendatario o le enviará comunicación en que se haga constar tal circunstancia, inmediatamente tenga conocimiento de la misma.El valor de la indemnización se hará con base en la renta vigente a la fecha del preaviso;\t\n c) Al momento de efectuar la consignación se dejará constancia en los respectivos títulos de las causas de la misma como también el nombre y dirección precisa del arrendatario o su representante;\t\n d) Si el arrendatario cumple con la obligación de entregar el inmueble en la fecha señalada, recibirá el pago de la indemnización, de conformidad con la autorización que expida la autoridad competente.')
            p15.add_run(' PARÁGRAFO 1. ').bold = True
            p15.add_run('En caso de que el arrendatario no entregue el inmueble, el arrendador tendrá derecho a que se le devuelva la indemnización consignada, sin perjuicio de que pueda iniciar el correspondiente proceso de restitución del inmueble.')
            p15.add_run(' PARÁGRAFO 2. ').bold = True
            p15.add_run('Si el arrendador con la aceptación del arrendatario desiste de dar por terminado el contrato de arrendamiento, podrá solicitar a la autoridad competente, la autorización para la devolución de la suma consignada.')

        ####

        p16 = documentoObjeto.add_paragraph('')
        p16.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI")  :
            titulo16='CLÁUSULA DÉCIMO QUINTA. TERMINACIÓN POR PARTE DEL ARRENDATARIO: '
            p16.add_run(titulo16).bold = True
            p16.add_run(' Son causales para que el arrendatario pueda pedir unilateralmente la terminación del contrato, las siguientes:\t\n 1. La suspensión de la prestación de los servicios públicos al inmueble, por acción premeditada del arrendador o porque incurra en mora en pagos que estuvieren a su cargo. En estos casos el arrendatario podrá optar por asumir el costo del restablecimiento del servicio y descontarlo de los pagos que le corresponda hacer como arrendatario.\n 2. La incursión reiterada del arrendador en procederes que afecten gravemente el disfrute cabal por el arrendatario del inmueble arrendado, debidamente comprobada ante la autoridad policiva.\t\n 3. El desconocimiento por parte del arrendador de derechos reconocidos al arrendatario por la Ley o contractualmente.\t\n4. El arrendatario podrá dar por terminado unilateralmente el contrato de arrendamiento dentro del término inicial o durante sus prórrogas, previo aviso escrito dirigido al arrendador a través del servicio postal autorizado, con una antelación no menor de tres (3) meses y el pago de una indemnización equivalente al precio de tres (3) meses de arrendamiento.Cumplidas estas condiciones el arrendador estará obligado a recibir el inmueble; si no lo hiciere, el arrendatario podrá hacer entrega provisional mediante la intervención de la autoridad competente, sin prejuicio de acudir a la acción judicial correspondiente.\t\n5. El arrendatario podrá dar por terminado unilateralmente el contrato de arrendamiento a la fecha de vencimiento del término inicial o de sus prórrogas, siempre y cuando dé previo aviso escrito al arrendador a través del servicio postal autorizado, con una antelación no menor de tres (3) meses a la referida fecha de vencimiento. En este caso el arrendatario no estará obligado a invocar causal alguna diferente a la de su plena voluntad, ni deberá indemnizar al arrendador.\t\nDe no mediar constancia por escrito del preaviso, el contrato de arrendamiento se entenderá renovado automáticamente por un término igual al inicialmente pactado.\t\n PARÁGRAFO. Para efectos de la entrega provisional de que trata este artículo, la autoridad competente, a solicitud escrita del arrendatario y una vez acreditado por parte del mismo el cumplimiento de las condiciones allí previstas, procederá a señalar fecha y hora para llevar, a cabo la entrega del inmueble. ')
            p16.add_run('\t\nCumplido lo anterior se citará al arrendador y al arrendatario mediante comunicación enviada por el servicio postal autorizado, a fin de que comparezcan el día y hora señalada al lugar de ubicación del inmueble para efectuar la entrega al arrendador.')
            p16.add_run('\t\nSi el arrendador no acudiere a recibir el inmueble el día de la diligencia, el funcionario competente para tal efecto hará entrega del inmueble a un secuestre que para su custodia designare de la lista de auxiliares de la justicia hasta la entrega al arrendador a cuyo cargo corren los gastos del secuestre. ')
            p16.add_run('\t\nDe todo lo anterior se levantará un acta que será suscrita por las personas que intervinieron en la diligencia.')
            
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo16='CLÁUSULA DÉCIMO SEXTA. TERMINACIÓN POR PARTE DEL ARRENDATARIO: '
            p16.add_run(titulo16).bold = True
            p16.add_run(' Son causales para que el arrendatario pueda pedir unilateralmente la terminación del contrato, las siguientes:\t\n 1. La suspensión de la prestación de los servicios públicos al inmueble, por acción premeditada del arrendador o porque incurra en mora en pagos que estuvieren a su cargo. En estos casos el arrendatario podrá optar por asumir el costo del restablecimiento del servicio y descontarlo de los pagos que le corresponda hacer como arrendatario.\n 2. La incursión reiterada del arrendador en procederes que afecten gravemente el disfrute cabal por el arrendatario del inmueble arrendado, debidamente comprobada ante la autoridad policiva.\t\n 3. El desconocimiento por parte del arrendador de derechos reconocidos al arrendatario por la Ley o contractualmente.\t\n4. El arrendatario podrá dar por terminado unilateralmente el contrato de arrendamiento dentro del término inicial o durante sus prórrogas, previo aviso escrito dirigido al arrendador a través del servicio postal autorizado, con una antelación no menor de tres (3) meses y el pago de una indemnización equivalente al precio de tres (3) meses de arrendamiento.Cumplidas estas condiciones el arrendador estará obligado a recibir el inmueble; si no lo hiciere, el arrendatario podrá hacer entrega provisional mediante la intervención de la autoridad competente, sin prejuicio de acudir a la acción judicial correspondiente.\t\n5. El arrendatario podrá dar por terminado unilateralmente el contrato de arrendamiento a la fecha de vencimiento del término inicial o de sus prórrogas, siempre y cuando dé previo aviso escrito al arrendador a través del servicio postal autorizado, con una antelación no menor de tres (3) meses a la referida fecha de vencimiento. En este caso el arrendatario no estará obligado a invocar causal alguna diferente a la de su plena voluntad, ni deberá indemnizar al arrendador.\t\n De no mediar constancia por escrito del preaviso, el contrato de arrendamiento se entenderá renovado automáticamente por un término igual al inicialmente pactado.\t\n PARÁGRAFO. Para efectos de la entrega provisional de que trata este artículo, la autoridad competente, a solicitud escrita del arrendatario y una vez acreditado por parte del mismo el cumplimiento de las condiciones allí previstas, procederá a señalar fecha y hora para llevar, a cabo la entrega del inmueble. ')
            p16.add_run('\t\nCumplido lo anterior se citará al arrendador y al arrendatario mediante comunicación enviada por el servicio postal autorizado, a fin de que comparezcan el día y hora señalada al lugar de ubicación del inmueble para efectuar la entrega al arrendador.')
            p16.add_run('\t\nSi el arrendador no acudiere a recibir el inmueble el día de la diligencia, el funcionario competente para tal efecto hará entrega del inmueble a un secuestre que para su custodia designare de la lista de auxiliares de la justicia hasta la entrega al arrendador a cuyo cargo corren los gastos del secuestre. ')
            p16.add_run('\t\nDe todo lo anterior se levantará un acta que será suscrita por las personas que intervinieron en la diligencia.')

        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo16='CLÁUSULA DÉCIMO CUARTA. TERMINACIÓN POR PARTE DEL ARRENDATARIO: '
            p16.add_run(titulo16).bold = True
            p16.add_run(' Son causales para que el arrendatario pueda pedir unilateralmente la terminación del contrato, las siguientes:\t\n 1. La suspensión de la prestación de los servicios públicos al inmueble, por acción premeditada del arrendador o porque incurra en mora en pagos que estuvieren a su cargo. En estos casos el arrendatario podrá optar por asumir el costo del restablecimiento del servicio y descontarlo de los pagos que le corresponda hacer como arrendatario.\n 2. La incursión reiterada del arrendador en procederes que afecten gravemente el disfrute cabal por el arrendatario del inmueble arrendado, debidamente comprobada ante la autoridad policiva.\t\n 3. El desconocimiento por parte del arrendador de derechos reconocidos al arrendatario por la Ley o contractualmente.\t\n4. El arrendatario podrá dar por terminado unilateralmente el contrato de arrendamiento dentro del término inicial o durante sus prórrogas, previo aviso escrito dirigido al arrendador a través del servicio postal autorizado, con una antelación no menor de tres (3) meses y el pago de una indemnización equivalente al precio de tres (3) meses de arrendamiento.Cumplidas estas condiciones el arrendador estará obligado a recibir el inmueble; si no lo hiciere, el arrendatario podrá hacer entrega provisional mediante la intervención de la autoridad competente, sin prejuicio de acudir a la acción judicial correspondiente.\t\n5. El arrendatario podrá dar por terminado unilateralmente el contrato de arrendamiento a la fecha de vencimiento del término inicial o de sus prórrogas, siempre y cuando dé previo aviso escrito al arrendador a través del servicio postal autorizado, con una antelación no menor de tres (3) meses a la referida fecha de vencimiento. En este caso el arrendatario no estará obligado a invocar causal alguna diferente a la de su plena voluntad, ni deberá indemnizar al arrendador.\t\n De no mediar constancia por escrito del preaviso, el contrato de arrendamiento se entenderá renovado automáticamente por un término igual al inicialmente pactado.\t\n PARÁGRAFO. Para efectos de la entrega provisional de que trata este artículo, la autoridad competente, a solicitud escrita del arrendatario y una vez acreditado por parte del mismo el cumplimiento de las condiciones allí previstas, procederá a señalar fecha y hora para llevar, a cabo la entrega del inmueble. ')
            p16.add_run('\t\nCumplido lo anterior se citará al arrendador y al arrendatario mediante comunicación enviada por el servicio postal autorizado, a fin de que comparezcan el día y hora señalada al lugar de ubicación del inmueble para efectuar la entrega al arrendador.')
            p16.add_run('\t\nSi el arrendador no acudiere a recibir el inmueble el día de la diligencia, el funcionario competente para tal efecto hará entrega del inmueble a un secuestre que para su custodia designare de la lista de auxiliares de la justicia hasta la entrega al arrendador a cuyo cargo corren los gastos del secuestre. ')
            p16.add_run('\t\nDe todo lo anterior se levantará un acta que será suscrita por las personas que intervinieron en la diligencia.')

        elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo16='CLÁUSULA DÉCIMO TERCERA. TERMINACIÓN POR PARTE DEL ARRENDATARIO: '
            p16.add_run(titulo16).bold = True
            p16.add_run(' Son causales para que el arrendatario pueda pedir unilateralmente la terminación del contrato, las siguientes:\t\n 1. La suspensión de la prestación de los servicios públicos al inmueble, por acción premeditada del arrendador o porque incurra en mora en pagos que estuvieren a su cargo. En estos casos el arrendatario podrá optar por asumir el costo del restablecimiento del servicio y descontarlo de los pagos que le corresponda hacer como arrendatario.\n 2. La incursión reiterada del arrendador en procederes que afecten gravemente el disfrute cabal por el arrendatario del inmueble arrendado, debidamente comprobada ante la autoridad policiva.\t\n 3. El desconocimiento por parte del arrendador de derechos reconocidos al arrendatario por la Ley o contractualmente.\t\n4. El arrendatario podrá dar por terminado unilateralmente el contrato de arrendamiento dentro del término inicial o durante sus prórrogas, previo aviso escrito dirigido al arrendador a través del servicio postal autorizado, con una antelación no menor de tres (3) meses y el pago de una indemnización equivalente al precio de tres (3) meses de arrendamiento.Cumplidas estas condiciones el arrendador estará obligado a recibir el inmueble; si no lo hiciere, el arrendatario podrá hacer entrega provisional mediante la intervención de la autoridad competente, sin prejuicio de acudir a la acción judicial correspondiente.\t\n5. El arrendatario podrá dar por terminado unilateralmente el contrato de arrendamiento a la fecha de vencimiento del término inicial o de sus prórrogas, siempre y cuando dé previo aviso escrito al arrendador a través del servicio postal autorizado, con una antelación no menor de tres (3) meses a la referida fecha de vencimiento. En este caso el arrendatario no estará obligado a invocar causal alguna diferente a la de su plena voluntad, ni deberá indemnizar al arrendador.\t\n De no mediar constancia por escrito del preaviso, el contrato de arrendamiento se entenderá renovado automáticamente por un término igual al inicialmente pactado.\t\n PARÁGRAFO. Para efectos de la entrega provisional de que trata este artículo, la autoridad competente, a solicitud escrita del arrendatario y una vez acreditado por parte del mismo el cumplimiento de las condiciones allí previstas, procederá a señalar fecha y hora para llevar, a cabo la entrega del inmueble. ')
            p16.add_run('\t\nCumplido lo anterior se citará al arrendador y al arrendatario mediante comunicación enviada por el servicio postal autorizado, a fin de que comparezcan el día y hora señalada al lugar de ubicación del inmueble para efectuar la entrega al arrendador.')
            p16.add_run('\t\nSi el arrendador no acudiere a recibir el inmueble el día de la diligencia, el funcionario competente para tal efecto hará entrega del inmueble a un secuestre que para su custodia designare de la lista de auxiliares de la justicia hasta la entrega al arrendador a cuyo cargo corren los gastos del secuestre. ')
            p16.add_run('\t\nDe todo lo anterior se levantará un acta que será suscrita por las personas que intervinieron en la diligencia.')

        ####
        p23 = documentoObjeto.add_paragraph('')
        p23.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI")  :
            titulo17='CLÁUSULA DÉCIMO SEXTA. LA TERMINACIÓN UNILATERAL POR PARTE DEL ARRENDATARIO MEDIANTE PREAVISO CON INDEMNIZACIÓN: '
            p23.add_run(titulo17).bold = True
            p23.add_run('Para que el arrendatario pueda dar por terminado unilateralmente el contrato de arrendamiento en el evento previsto en el numeral 4 del artículo anterior, deberá cumplir con los siguientes requisitos:\t\n ')

            p23.add_run('a) Comunicar a través del servicio postal autorizado al arrendador o a su representante legal, con la antelación allí prevista, indicando la fecha para la terminación del contrato y, manifestando que se pagará la indemnización de ley.\t\nb) Consignar a favor del arrendador y a órdenes de la autoridad competente, la indemnización de que trata el artículo anterior de la presente ley, dentro de los tres (3) meses anteriores a la fecha señalada para la terminación unilateral del contrato. La consignación se efectuará en las entidades autorizadas por el Gobierno Nacional para tal efecto y la autoridad competente allegará copia del título respectivo al arrendador o le enviará comunicación en que se haga constar tal circunstancia, inmediatamente tenga conocimiento de la misma. El valor de la indemnización se hará con base en la renta vigente a la fecha del preaviso.\t\nc) Al momento de efectuar la consignación se dejará constancia en los respectivos títulos de las causas de la misma como también el nombre y dirección precisa del arrendatario o su representante.\t\n d) Si el arrendador cumple con la obligación de entregar el inmueble en la fecha señalada, recibirá el pago de la indemnización, de conformidad con la autorización que expida la autoridad competente.\t\n')               

            p23.add_run('PARÁGRAFO 1o. En caso de que el arrendador no reciba el inmueble, el arrendatario tendrá derecho a que se le devuelva la indemnización consignada, sin perjuicio de que pueda realizar la entrega provisional del inmueble de conformidad con lo previsto en el artículo anterior. PARÁGRAFO 2o. Si el arrendatario con la aceptación del arrendador desiste de dar por terminado el contrato de arrendamiento, podrá solicitar a la autoridad competente, la autorización para la devolución de la suma consignada.')
            
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo17='CLÁUSULA DÉCIMO SÉPTIMA. LA TERMINACIÓN UNILATERAL POR PARTE DEL ARRENDATARIO MEDIANTE PREAVISO CON INDEMNIZACIÓN: '
            p23.add_run(titulo17).bold = True
            p23.add_run('Para que el arrendatario pueda dar por terminado unilateralmente el contrato de arrendamiento en el evento previsto en el numeral 4 del artículo anterior, deberá cumplir con los siguientes requisitos:\t\n ')

            p23.add_run('a) Comunicar a través del servicio postal autorizado al arrendador o a su representante legal, con la antelación allí prevista, indicando la fecha para la terminación del contrato y, manifestando que se pagará la indemnización de ley.\t\nb) Consignar a favor del arrendador y a órdenes de la autoridad competente, la indemnización de que trata el artículo anterior de la presente ley, dentro de los tres (3) meses anteriores a la fecha señalada para la terminación unilateral del contrato. La consignación se efectuará en las entidades autorizadas por el Gobierno Nacional para tal efecto y la autoridad competente allegará copia del título respectivo al arrendador o le enviará comunicación en que se haga constar tal circunstancia, inmediatamente tenga conocimiento de la misma. El valor de la indemnización se hará con base en la renta vigente a la fecha del preaviso.\t\nc) Al momento de efectuar la consignación se dejará constancia en los respectivos títulos de las causas de la misma como también el nombre y dirección precisa del arrendatario o su representante.\t\n d) Si el arrendador cumple con la obligación de entregar el inmueble en la fecha señalada, recibirá el pago de la indemnización, de conformidad con la autorización que expida la autoridad competente.\t\n')               

            p23.add_run('PARÁGRAFO 1o. En caso de que el arrendador no reciba el inmueble, el arrendatario tendrá derecho a que se le devuelva la indemnización consignada, sin perjuicio de que pueda realizar la entrega provisional del inmueble de conformidad con lo previsto en el artículo anterior. PARÁGRAFO 2o. Si el arrendatario con la aceptación del arrendador desiste de dar por terminado el contrato de arrendamiento, podrá solicitar a la autoridad competente, la autorización para la devolución de la suma consignada.')

        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo17='CLÁUSULA DÉCIMO QUINTA. LA TERMINACIÓN UNILATERAL POR PARTE DEL ARRENDATARIO MEDIANTE PREAVISO CON INDEMNIZACIÓN: '
            p23.add_run(titulo17).bold = True
            p23.add_run('Para que el arrendatario pueda dar por terminado unilateralmente el contrato de arrendamiento en el evento previsto en el numeral 4 del artículo anterior, deberá cumplir con los siguientes requisitos:\t\n ')

            p23.add_run('a) Comunicar a través del servicio postal autorizado al arrendador o a su representante legal, con la antelación allí prevista, indicando la fecha para la terminación del contrato y, manifestando que se pagará la indemnización de ley.\t\nb) Consignar a favor del arrendador y a órdenes de la autoridad competente, la indemnización de que trata el artículo anterior de la presente ley, dentro de los tres (3) meses anteriores a la fecha señalada para la terminación unilateral del contrato. La consignación se efectuará en las entidades autorizadas por el Gobierno Nacional para tal efecto y la autoridad competente allegará copia del título respectivo al arrendador o le enviará comunicación en que se haga constar tal circunstancia, inmediatamente tenga conocimiento de la misma. El valor de la indemnización se hará con base en la renta vigente a la fecha del preaviso.\t\nc) Al momento de efectuar la consignación se dejará constancia en los respectivos títulos de las causas de la misma como también el nombre y dirección precisa del arrendatario o su representante.\t\n d) Si el arrendador cumple con la obligación de entregar el inmueble en la fecha señalada, recibirá el pago de la indemnización, de conformidad con la autorización que expida la autoridad competente.\t\n')               

            p23.add_run('PARÁGRAFO 1o. En caso de que el arrendador no reciba el inmueble, el arrendatario tendrá derecho a que se le devuelva la indemnización consignada, sin perjuicio de que pueda realizar la entrega provisional del inmueble de conformidad con lo previsto en el artículo anterior. PARÁGRAFO 2o. Si el arrendatario con la aceptación del arrendador desiste de dar por terminado el contrato de arrendamiento, podrá solicitar a la autoridad competente, la autorización para la devolución de la suma consignada.')

        elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo17='CLÁUSULA DÉCIMO CUARTA. LA TERMINACIÓN UNILATERAL POR PARTE DEL ARRENDATARIO MEDIANTE PREAVISO CON INDEMNIZACIÓN: '
            p23.add_run(titulo17).bold = True
            p23.add_run('Para que el arrendatario pueda dar por terminado unilateralmente el contrato de arrendamiento en el evento previsto en el numeral 4 del artículo anterior, deberá cumplir con los siguientes requisitos:\t\n ')

            p23.add_run('a) Comunicar a través del servicio postal autorizado al arrendador o a su representante legal, con la antelación allí prevista, indicando la fecha para la terminación del contrato y, manifestando que se pagará la indemnización de ley.\t\nb) Consignar a favor del arrendador y a órdenes de la autoridad competente, la indemnización de que trata el artículo anterior de la presente ley, dentro de los tres (3) meses anteriores a la fecha señalada para la terminación unilateral del contrato. La consignación se efectuará en las entidades autorizadas por el Gobierno Nacional para tal efecto y la autoridad competente allegará copia del título respectivo al arrendador o le enviará comunicación en que se haga constar tal circunstancia, inmediatamente tenga conocimiento de la misma. El valor de la indemnización se hará con base en la renta vigente a la fecha del preaviso.\t\nc) Al momento de efectuar la consignación se dejará constancia en los respectivos títulos de las causas de la misma como también el nombre y dirección precisa del arrendatario o su representante.\t\n d) Si el arrendador cumple con la obligación de entregar el inmueble en la fecha señalada, recibirá el pago de la indemnización, de conformidad con la autorización que expida la autoridad competente.\t\n')               

            p23.add_run('PARÁGRAFO 1o. En caso de que el arrendador no reciba el inmueble, el arrendatario tendrá derecho a que se le devuelva la indemnización consignada, sin perjuicio de que pueda realizar la entrega provisional del inmueble de conformidad con lo previsto en el artículo anterior. PARÁGRAFO 2o. Si el arrendatario con la aceptación del arrendador desiste de dar por terminado el contrato de arrendamiento, podrá solicitar a la autoridad competente, la autorización para la devolución de la suma consignada.')

        ####
        p24 = documentoObjeto.add_paragraph('')
        p24.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI")  :
            titulo18='CLÁUSULA DÉCIMO SÉPTIMA. PREAVISOS PARA LA ENTREGA: '
            p24.add_run(titulo18).bold = True
            p24.add_run('Las partes se obligan, en caso de terminación del contrato, a efectuar el correspondiente preaviso a través del servicio postal autorizado y/o radicación en la oficina, con tres (3) meses de anticipación a la finalización del plazo original o de sus prórrogas; subsistiendo durante dichas prórrogas, todas las garantías, compromisos y estipulaciones de este contrato. ')
            
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo18='CLÁUSULA DÉCIMO OCTAVA. PREAVISOS PARA LA ENTREGA: '
            p24.add_run(titulo18).bold = True
            p24.add_run('Las partes se obligan, en caso de terminación del contrato, a efectuar el correspondiente preaviso a través del servicio postal autorizado y/o radicación en la oficina, con tres (3) meses de anticipación a la finalización del plazo original o de sus prórrogas; subsistiendo durante dichas prórrogas, todas las garantías, compromisos y estipulaciones de este contrato. ')

        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo18='CLÁUSULA DÉCIMO SEXTA. PREAVISOS PARA LA ENTREGA: '
            p24.add_run(titulo18).bold = True
            p24.add_run('Las partes se obligan, en caso de terminación del contrato, a efectuar el correspondiente preaviso a través del servicio postal autorizado y/o radicación en la oficina, con tres (3) meses de anticipación a la finalización del plazo original o de sus prórrogas; subsistiendo durante dichas prórrogas, todas las garantías, compromisos y estipulaciones de este contrato. ')

        elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo18='CLÁUSULA DÉCIMO QUINTA. PREAVISOS PARA LA ENTREGA: '
            p24.add_run(titulo18).bold = True
            p24.add_run('Las partes se obligan, en caso de terminación del contrato, a efectuar el correspondiente preaviso a través del servicio postal autorizado y/o radicación en la oficina, con tres (3) meses de anticipación a la finalización del plazo original o de sus prórrogas; subsistiendo durante dichas prórrogas, todas las garantías, compromisos y estipulaciones de este contrato. ')

        ####

        p25 = documentoObjeto.add_paragraph('')
        p25.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI")  :
            titulo19='CLÁUSULA DÉCIMO OCTAVA. FECHA DE ENTREGA: '
            p25.add_run(titulo19).bold = True
            p25.add_run('La fecha de entrega real y material del inmueble, constará en la respectiva acta de entrega que suscriban las partes y que se anexa al presente contrato como ANEXO 2. ')
            
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo19='CLÁUSULA DÉCIMO NOVENA. FECHA DE ENTREGA: '
            p25.add_run(titulo19).bold = True
            p25.add_run('La fecha de entrega real y material del inmueble, constará en la respectiva acta de entrega que suscriban las partes y que se anexa al presente contrato como ANEXO 2. ')
            
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo19='CLÁUSULA DÉCIMO SÉPTIMA. FECHA DE ENTREGA: '
            p25.add_run(titulo19).bold = True
            p25.add_run('La fecha de entrega real y material del inmueble, constará en la respectiva acta de entrega que suscriban las partes y que se anexa al presente contrato como ANEXO 2. ')
            
        elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo19='CLÁUSULA DÉCIMO SEXTA. FECHA DE ENTREGA: '
            p25.add_run(titulo19).bold = True
            p25.add_run('La fecha de entrega real y material del inmueble, constará en la respectiva acta de entrega que suscriban las partes y que se anexa al presente contrato como ANEXO 2. ')
            
        #### 

        p27 = documentoObjeto.add_paragraph("")
        p27.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI")  :
            titulo20='CLÁUSULA DÉCIMO NOVENA. ESTADO E INVENTARIO: '
            p27.add_run(titulo20).bold = True
            p27.add_run('El ARRENDATARIO declara recibir el Inmueble de manos del ARRENDADOR en buen estado, de conformidad con el inventario elaborado por las partes y que forma parte integrante de este contrato en calidad de ANEXO 1. Asimismo se obliga a cuidarlo, conservarlo y mantenerlo en el mismo estado que lo restituirá al ARRENDADOR, salvo por el deterioro o paso del tiempo y su uso legítimo.')
            
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo20='CLÁUSULA VIGÉSIMA. ESTADO E INVENTARIO: '
            p27.add_run(titulo20).bold = True
            p27.add_run('El ARRENDATARIO declara recibir el Inmueble de manos del ARRENDADOR en buen estado, de conformidad con el inventario elaborado por las partes y que forma parte integrante de este contrato en calidad de ANEXO 1. Asimismo se obliga a cuidarlo, conservarlo y mantenerlo en el mismo estado que lo restituirá al ARRENDADOR, salvo por el deterioro o paso del tiempo y su uso legítimo.')
            
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo20='CLÁUSULA DÉCIMO OCTAVA. ESTADO E INVENTARIO: '
            p27.add_run(titulo20).bold = True
            p27.add_run('El ARRENDATARIO declara recibir el Inmueble de manos del ARRENDADOR en buen estado, de conformidad con el inventario elaborado por las partes y que forma parte integrante de este contrato en calidad de ANEXO 1. Asimismo se obliga a cuidarlo, conservarlo y mantenerlo en el mismo estado que lo restituirá al ARRENDADOR, salvo por el deterioro o paso del tiempo y su uso legítimo.')
            
        elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo20='CLÁUSULA DÉCIMO SEPTIMA. ESTADO E INVENTARIO: '
            p27.add_run(titulo20).bold = True
            p27.add_run('El ARRENDATARIO declara recibir el Inmueble de manos del ARRENDADOR en buen estado, de conformidad con el inventario elaborado por las partes y que forma parte integrante de este contrato en calidad de ANEXO 1. Asimismo se obliga a cuidarlo, conservarlo y mantenerlo en el mismo estado que lo restituirá al ARRENDADOR, salvo por el deterioro o paso del tiempo y su uso legítimo.')

        #### 
        p28 = documentoObjeto.add_paragraph("")
        p28.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            titulo21='CLÁUSULA VIGÉSIMA. REPARACIONES LOCATIVAS A CARGO DEL ARRENDATARIO: '
            p28.add_run(titulo21).bold = True
            p28.add_run('El  ARRENDATARIO se encuentra obligado a efectuar las reparaciones locativas que sean necesarias para mantener el buen estado de conservación del inmueble, y aquellas que sean imputables a su uso y goce o el de sus dependientes, sin que ello implique indemnización alguna, derecho a reclamar o a ejercer derecho de retención. ')
            p28.add_run('PARÁGRAFO1. ').bold = True
            p28.add_run('Los daños del inmueble derivados del maltrato o descuido durante su tenencia estarán a cargo del ARRENDATARIO, quien está obligado a efectuar las reparaciones locativas, es decir a mantener el inmueble en el estado que lo recibió, estando especialmente obligado al cumplimiento de lo estipulado en los artículos 2028, 2029 y 2030 del Código Civil. ')
            p28.add_run('PARÁGRAFO2. ').bold = True
            p28.add_run('Asimismo, el ARRENDATARIO deberá: 1) Conservar la integridad interior de las paredes, techos, pavimentos y cañerías, reponiendo las que durante el arrendamiento se quiebren o desencajen. 2) Reponer los cristales quebrados en las ventanas, puertas y tabiques. 3) Mantener en estado de servicio las ventanas, puertas y cerraduras, pisos y demás partes interiores y exteriores del inmueble debidamente aseado. 4) Conservar las llaves de agua, arreglo de grifos o salidas de acueducto, baños e instalaciones sanitarias, etc. 5) Dar el manejo adecuado a los servicios, cosas y usos conexos y/o adicionales que le sean entregados. ')
            p28.add_run('PARÁGRAFO3. ').bold = True
            p28.add_run('El ARRENDADOR no responderá por daños al calentador si estos fueren ocasionados por indebida manipulación por parte del ARRENDATARIO. ')
            p28.add_run('PARÁGRAFO4. ').bold = True
            p28.add_run('EL ARRENDATARIO renuncia expresamente a descontar del canon de arrendamiento el valor de las reparaciones indispensables, a que se refiere el artículo 27 de la ley 820 de 2.003. ')

        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo21='CLÁUSULA VIGÉSIMA PRIMERA. REPARACIONES LOCATIVAS A CARGO DEL ARRENDATARIO: '
            p28.add_run(titulo21).bold = True
            p28.add_run('El  ARRENDATARIO se encuentra obligado a efectuar las reparaciones locativas que sean necesarias para mantener el buen estado de conservación del inmueble, y aquellas que sean imputables a su uso y goce o el de sus dependientes, sin que ello implique indemnización alguna, derecho a reclamar o a ejercer derecho de retención. ')
            p28.add_run('PARÁGRAFO1. ').bold = True
            p28.add_run('Los daños del inmueble derivados del maltrato o descuido durante su tenencia estarán a cargo del ARRENDATARIO, quien está obligado a efectuar las reparaciones locativas, es decir a mantener el inmueble en el estado que lo recibió, estando especialmente obligado al cumplimiento de lo estipulado en los artículos 2028, 2029 y 2030 del Código Civil. ')
            p28.add_run('PARÁGRAFO2. ').bold = True
            p28.add_run('Asimismo, el ARRENDATARIO deberá: 1) Conservar la integridad interior de las paredes, techos, pavimentos y cañerías, reponiendo las que durante el arrendamiento se quiebren o desencajen. 2) Reponer los cristales quebrados en las ventanas, puertas y tabiques. 3) Mantener en estado de servicio las ventanas, puertas y cerraduras, pisos y demás partes interiores y exteriores del inmueble debidamente aseado. 4) Conservar las llaves de agua, arreglo de grifos o salidas de acueducto, baños e instalaciones sanitarias, etc. 5) Dar el manejo adecuado a los servicios, cosas y usos conexos y/o adicionales que le sean entregados. ')
            p28.add_run('PARÁGRAFO3. ').bold = True
            p28.add_run('El ARRENDADOR no responderá por daños al calentador si estos fueren ocasionados por indebida manipulación por parte del ARRENDATARIO. ')
            p28.add_run('PARÁGRAFO4. ').bold = True
            p28.add_run('EL ARRENDATARIO renuncia expresamente a descontar del canon de arrendamiento el valor de las reparaciones indispensables, a que se refiere el artículo 27 de la ley 820 de 2.003. ')
            
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo21='CLÁUSULA DÉCIMO NOVENA. REPARACIONES LOCATIVAS A CARGO DEL ARRENDATARIO: '
            p28.add_run(titulo21).bold = True
            p28.add_run('El  ARRENDATARIO se encuentra obligado a efectuar las reparaciones locativas que sean necesarias para mantener el buen estado de conservación del inmueble, y aquellas que sean imputables a su uso y goce o el de sus dependientes, sin que ello implique indemnización alguna, derecho a reclamar o a ejercer derecho de retención. ')
            p28.add_run('PARÁGRAFO1. ').bold = True
            p28.add_run('Los daños del inmueble derivados del maltrato o descuido durante su tenencia estarán a cargo del ARRENDATARIO, quien está obligado a efectuar las reparaciones locativas, es decir a mantener el inmueble en el estado que lo recibió, estando especialmente obligado al cumplimiento de lo estipulado en los artículos 2028, 2029 y 2030 del Código Civil. ')
            p28.add_run('PARÁGRAFO2. ').bold = True
            p28.add_run('Asimismo, el ARRENDATARIO deberá: 1) Conservar la integridad interior de las paredes, techos, pavimentos y cañerías, reponiendo las que durante el arrendamiento se quiebren o desencajen. 2) Reponer los cristales quebrados en las ventanas, puertas y tabiques. 3) Mantener en estado de servicio las ventanas, puertas y cerraduras, pisos y demás partes interiores y exteriores del inmueble debidamente aseado. 4) Conservar las llaves de agua, arreglo de grifos o salidas de acueducto, baños e instalaciones sanitarias, etc. 5) Dar el manejo adecuado a los servicios, cosas y usos conexos y/o adicionales que le sean entregados. ')
            p28.add_run('PARÁGRAFO3. ').bold = True
            p28.add_run('El ARRENDADOR no responderá por daños al calentador si estos fueren ocasionados por indebida manipulación por parte del ARRENDATARIO. ')
            p28.add_run('PARÁGRAFO4. ').bold = True
            p28.add_run('EL ARRENDATARIO renuncia expresamente a descontar del canon de arrendamiento el valor de las reparaciones indispensables, a que se refiere el artículo 27 de la ley 820 de 2.003. ')
            
        elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo21='CLÁUSULA DÉCIMO OCTAVA. REPARACIONES LOCATIVAS A CARGO DEL ARRENDATARIO: '
            p28.add_run(titulo21).bold = True
            p28.add_run('El  ARRENDATARIO se encuentra obligado a efectuar las reparaciones locativas que sean necesarias para mantener el buen estado de conservación del inmueble, y aquellas que sean imputables a su uso y goce o el de sus dependientes, sin que ello implique indemnización alguna, derecho a reclamar o a ejercer derecho de retención. ')
            p28.add_run('PARÁGRAFO1. ').bold = True
            p28.add_run('Los daños del inmueble derivados del maltrato o descuido durante su tenencia estarán a cargo del ARRENDATARIO, quien está obligado a efectuar las reparaciones locativas, es decir a mantener el inmueble en el estado que lo recibió, estando especialmente obligado al cumplimiento de lo estipulado en los artículos 2028, 2029 y 2030 del Código Civil. ')
            p28.add_run('PARÁGRAFO2. ').bold = True
            p28.add_run('Asimismo, el ARRENDATARIO deberá: 1) Conservar la integridad interior de las paredes, techos, pavimentos y cañerías, reponiendo las que durante el arrendamiento se quiebren o desencajen. 2) Reponer los cristales quebrados en las ventanas, puertas y tabiques. 3) Mantener en estado de servicio las ventanas, puertas y cerraduras, pisos y demás partes interiores y exteriores del inmueble debidamente aseado. 4) Conservar las llaves de agua, arreglo de grifos o salidas de acueducto, baños e instalaciones sanitarias, etc. 5) Dar el manejo adecuado a los servicios, cosas y usos conexos y/o adicionales que le sean entregados. ')
            p28.add_run('PARÁGRAFO3. ').bold = True
            p28.add_run('El ARRENDADOR no responderá por daños al calentador si estos fueren ocasionados por indebida manipulación por parte del ARRENDATARIO. ')
            p28.add_run('PARÁGRAFO4. ').bold = True
            p28.add_run('EL ARRENDATARIO renuncia expresamente a descontar del canon de arrendamiento el valor de las reparaciones indispensables, a que se refiere el artículo 27 de la ley 820 de 2.003. ')

        #### 
        p30 = documentoObjeto.add_paragraph("")
        p30.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            titulo22='CLÁUSULA VIGÉSIMA PRIMERA. REPARACIONES LOCATIVAS A CARGO DEL ARRENDADOR: '
            p30.add_run(titulo22).bold = True
            p30.add_run('El ARRENDADOR estará a cargo de las reparaciones que provengan de deterioros que provengan de fuerza mayor o caso fortuito, mala calidad del edificio, por su vetustez, por la naturaleza del suelo, por defectos de construcción, asentamientos del inmueble, humedades que provengan de la edificación, filtraciones imperceptibles de redes hidráulicas, deterioros de redes internas eléctricas y de gas.')
            
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo22='CLÁUSULA VIGÉSIMA SEGUNDA. REPARACIONES LOCATIVAS A CARGO DEL ARRENDADOR: '
            p30.add_run(titulo22).bold = True
            p30.add_run('El ARRENDADOR estará a cargo de las reparaciones que provengan de deterioros que provengan de fuerza mayor o caso fortuito, mala calidad del edificio, por su vetustez, por la naturaleza del suelo, por defectos de construcción, asentamientos del inmueble, humedades que provengan de la edificación, filtraciones imperceptibles de redes hidráulicas, deterioros de redes internas eléctricas y de gas.')
            
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo22='CLÁUSULA VIGÉSIMA . REPARACIONES LOCATIVAS A CARGO DEL ARRENDADOR: '
            p30.add_run(titulo22).bold = True
            p30.add_run('El ARRENDADOR estará a cargo de las reparaciones que provengan de deterioros que provengan de fuerza mayor o caso fortuito, mala calidad del edificio, por su vetustez, por la naturaleza del suelo, por defectos de construcción, asentamientos del inmueble, humedades que provengan de la edificación, filtraciones imperceptibles de redes hidráulicas, deterioros de redes internas eléctricas y de gas.')
            
        elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo22='CLÁUSULA DÉCIMO NOVENA. REPARACIONES LOCATIVAS A CARGO DEL ARRENDADOR: '
            p30.add_run(titulo22).bold = True
            p30.add_run('El ARRENDADOR estará a cargo de las reparaciones que provengan de deterioros que provengan de fuerza mayor o caso fortuito, mala calidad del edificio, por su vetustez, por la naturaleza del suelo, por defectos de construcción, asentamientos del inmueble, humedades que provengan de la edificación, filtraciones imperceptibles de redes hidráulicas, deterioros de redes internas eléctricas y de gas.')

        #### 
        p31 = documentoObjeto.add_paragraph("")
        p31.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI")  :
            titulo23='CLÁUSULA VIGÉSIMA SEGUNDA. MEJORAS: '
            p31.add_run(titulo23).bold = True
            p31.add_run('EL ARRENDATARIO no podrá ejecutar en el inmueble mejoras de ninguna especie sin permiso escrito del ARRENDADOR excepto reparaciones locativas. Si se ejecutaren, accederán al propietario del inmueble sin indemnización para quien las efectúe. Son responsabilidad del ARRENDATARIO y sus deudores solidarios, las multas que impongan las autoridades competentes para la realización de cualquier intervención en el inmueble que requiera licencia de construcción.  ')
            p31.add_run('PARÁGRAFO 1. ').bold = True
            p31.add_run('Si EL ARRENDATARIO requiere efectuar modificaciones que se adecuan a su uso y goce, las hará por su propia cuenta, siempre y cuando estas a la terminación del contrato, sean retiradas sin menoscabar la estructura inicial del inmueble, ceñidas a las disposiciones legales vigentes, sin indemnización por haberlas efectuado ni derecho a reclamo por ellas. ')
            p31.add_run('PARÁGRAFO 2. ').bold = True
            p31.add_run('En caso de que se efectúen mejoras, estas no podrán retirarse salvo que EL ARRENDADOR lo exija por escrito, para lo cual EL ARRENDATARIO accederá inmediatamente a su costa, dejando el Inmueble en el mismo estado en que lo recibió, salvo el deterioro natural por el uso legítimo. ')
            
        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            titulo23='CLÁUSULA VIGÉSIMA SEGUNDA. MEJORAS: '
            p31.add_run(titulo23).bold = True
            p31.add_run('EL ARRENDATARIO no podrá ejecutar en el inmueble mejoras de ninguna especie sin permiso escrito del ARRENDADOR excepto reparaciones locativas. Si se ejecutaren, accederán al propietario del inmueble sin indemnización para quien las efectúe. Son responsabilidad del ARRENDATARIO, las multas que impongan las autoridades competentes para la realización de cualquier intervención en el inmueble que requiera licencia de construcción.  ')
            p31.add_run('PARÁGRAFO 1. ').bold = True
            p31.add_run('Si EL ARRENDATARIO requiere efectuar modificaciones que se adecuan a su uso y goce, las hará por su propia cuenta, siempre y cuando estas a la terminación del contrato, sean retiradas sin menoscabar la estructura inicial del inmueble, ceñidas a las disposiciones legales vigentes, sin indemnización por haberlas efectuado ni derecho a reclamo por ellas. ')
            p31.add_run('PARÁGRAFO 2. ').bold = True
            p31.add_run('En caso de que se efectúen mejoras, estas no podrán retirarse salvo que EL ARRENDADOR lo exija por escrito, para lo cual EL ARRENDATARIO accederá inmediatamente a su costa, dejando el Inmueble en el mismo estado en que lo recibió, salvo el deterioro natural por el uso legítimo.')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            titulo23='CLÁUSULA VIGÉSIMA TERCERA. MEJORAS: '
            p31.add_run(titulo23).bold = True
            p31.add_run('EL ARRENDATARIO no podrá ejecutar en el inmueble mejoras de ninguna especie sin permiso escrito del ARRENDADOR excepto reparaciones locativas. Si se ejecutaren, accederán al propietario del inmueble sin indemnización para quien las efectúe. Son responsabilidad del ARRENDATARIO, las multas que impongan las autoridades competentes para la realización de cualquier intervención en el inmueble que requiera licencia de construcción.  ')
            p31.add_run('PARÁGRAFO 1. ').bold = True
            p31.add_run('Si EL ARRENDATARIO requiere efectuar modificaciones que se adecuan a su uso y goce, las hará por su propia cuenta, siempre y cuando estas a la terminación del contrato, sean retiradas sin menoscabar la estructura inicial del inmueble, ceñidas a las disposiciones legales vigentes, sin indemnización por haberlas efectuado ni derecho a reclamo por ellas. ')
            p31.add_run('PARÁGRAFO 2. ').bold = True
            p31.add_run('En caso de que se efectúen mejoras, estas no podrán retirarse salvo que EL ARRENDADOR lo exija por escrito, para lo cual EL ARRENDATARIO accederá inmediatamente a su costa, dejando el Inmueble en el mismo estado en que lo recibió, salvo el deterioro natural por el uso legítimo.')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            titulo23='CLÁUSULA VIGÉSIMA TERCERA. MEJORAS: '
            p31.add_run(titulo23).bold = True
            p31.add_run('EL ARRENDATARIO no podrá ejecutar en el inmueble mejoras de ninguna especie sin permiso escrito del ARRENDADOR excepto reparaciones locativas. Si se ejecutaren, accederán al propietario del inmueble sin indemnización para quien las efectúe. Son responsabilidad del ARRENDATARIO y sus deudores solidarios, las multas que impongan las autoridades competentes para la realización de cualquier intervención en el inmueble que requiera licencia de construcción.  ')
            p31.add_run('PARÁGRAFO 1. ').bold = True
            p31.add_run('Si EL ARRENDATARIO requiere efectuar modificaciones que se adecuan a su uso y goce, las hará por su propia cuenta, siempre y cuando estas a la terminación del contrato, sean retiradas sin menoscabar la estructura inicial del inmueble, ceñidas a las disposiciones legales vigentes, sin indemnización por haberlas efectuado ni derecho a reclamo por ellas. ')
            p31.add_run('PARÁGRAFO 2. ').bold = True
            p31.add_run('En caso de que se efectúen mejoras, estas no podrán retirarse salvo que EL ARRENDADOR lo exija por escrito, para lo cual EL ARRENDATARIO accederá inmediatamente a su costa, dejando el Inmueble en el mismo estado en que lo recibió, salvo el deterioro natural por el uso legítimo. ')
            ##SIN ADMINISTRADOR
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO")  :
            titulo23='CLÁUSULA VIGÉSIMA . MEJORAS: '
            p31.add_run(titulo23).bold = True
            p31.add_run('EL ARRENDATARIO no podrá ejecutar en el inmueble mejoras de ninguna especie sin permiso escrito del ARRENDADOR excepto reparaciones locativas. Si se ejecutaren, accederán al propietario del inmueble sin indemnización para quien las efectúe. Son responsabilidad del ARRENDATARIO y sus deudores solidarios, las multas que impongan las autoridades competentes para la realización de cualquier intervención en el inmueble que requiera licencia de construcción.  ')
            p31.add_run('PARÁGRAFO 1. ').bold = True
            p31.add_run('Si EL ARRENDATARIO requiere efectuar modificaciones que se adecuan a su uso y goce, las hará por su propia cuenta, siempre y cuando estas a la terminación del contrato, sean retiradas sin menoscabar la estructura inicial del inmueble, ceñidas a las disposiciones legales vigentes, sin indemnización por haberlas efectuado ni derecho a reclamo por ellas. ')
            p31.add_run('PARÁGRAFO 2. ').bold = True
            p31.add_run('En caso de que se efectúen mejoras, estas no podrán retirarse salvo que EL ARRENDADOR lo exija por escrito, para lo cual EL ARRENDATARIO accederá inmediatamente a su costa, dejando el Inmueble en el mismo estado en que lo recibió, salvo el deterioro natural por el uso legítimo. ')
            
        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
            titulo23='CLÁUSULA VIGÉSIMA . MEJORAS: '
            p31.add_run(titulo23).bold = True
            p31.add_run('EL ARRENDATARIO no podrá ejecutar en el inmueble mejoras de ninguna especie sin permiso escrito del ARRENDADOR excepto reparaciones locativas. Si se ejecutaren, accederán al propietario del inmueble sin indemnización para quien las efectúe. Son responsabilidad del ARRENDATARIO, las multas que impongan las autoridades competentes para la realización de cualquier intervención en el inmueble que requiera licencia de construcción.  ')
            p31.add_run('PARÁGRAFO 1. ').bold = True
            p31.add_run('Si EL ARRENDATARIO requiere efectuar modificaciones que se adecuan a su uso y goce, las hará por su propia cuenta, siempre y cuando estas a la terminación del contrato, sean retiradas sin menoscabar la estructura inicial del inmueble, ceñidas a las disposiciones legales vigentes, sin indemnización por haberlas efectuado ni derecho a reclamo por ellas. ')
            p31.add_run('PARÁGRAFO 2. ').bold = True
            p31.add_run('En caso de que se efectúen mejoras, estas no podrán retirarse salvo que EL ARRENDADOR lo exija por escrito, para lo cual EL ARRENDATARIO accederá inmediatamente a su costa, dejando el Inmueble en el mismo estado en que lo recibió, salvo el deterioro natural por el uso legítimo.')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
            titulo23='CLÁUSULA VIGÉSIMA PRIMERA. MEJORAS: '
            p31.add_run(titulo23).bold = True
            p31.add_run('EL ARRENDATARIO no podrá ejecutar en el inmueble mejoras de ninguna especie sin permiso escrito del ARRENDADOR excepto reparaciones locativas. Si se ejecutaren, accederán al propietario del inmueble sin indemnización para quien las efectúe. Son responsabilidad del ARRENDATARIO, las multas que impongan las autoridades competentes para la realización de cualquier intervención en el inmueble que requiera licencia de construcción.  ')
            p31.add_run('PARÁGRAFO 1. ').bold = True
            p31.add_run('Si EL ARRENDATARIO requiere efectuar modificaciones que se adecuan a su uso y goce, las hará por su propia cuenta, siempre y cuando estas a la terminación del contrato, sean retiradas sin menoscabar la estructura inicial del inmueble, ceñidas a las disposiciones legales vigentes, sin indemnización por haberlas efectuado ni derecho a reclamo por ellas. ')
            p31.add_run('PARÁGRAFO 2. ').bold = True
            p31.add_run('En caso de que se efectúen mejoras, estas no podrán retirarse salvo que EL ARRENDADOR lo exija por escrito, para lo cual EL ARRENDATARIO accederá inmediatamente a su costa, dejando el Inmueble en el mismo estado en que lo recibió, salvo el deterioro natural por el uso legítimo.')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
            titulo23='CLÁUSULA VIGÉSIMA PRIMERA. MEJORAS: '
            p31.add_run(titulo23).bold = True
            p31.add_run('EL ARRENDATARIO no podrá ejecutar en el inmueble mejoras de ninguna especie sin permiso escrito del ARRENDADOR excepto reparaciones locativas. Si se ejecutaren, accederán al propietario del inmueble sin indemnización para quien las efectúe. Son responsabilidad del ARRENDATARIO y sus deudores solidarios, las multas que impongan las autoridades competentes para la realización de cualquier intervención en el inmueble que requiera licencia de construcción.  ')
            p31.add_run('PARÁGRAFO 1. ').bold = True
            p31.add_run('Si EL ARRENDATARIO requiere efectuar modificaciones que se adecuan a su uso y goce, las hará por su propia cuenta, siempre y cuando estas a la terminación del contrato, sean retiradas sin menoscabar la estructura inicial del inmueble, ceñidas a las disposiciones legales vigentes, sin indemnización por haberlas efectuado ni derecho a reclamo por ellas. ')
            p31.add_run('PARÁGRAFO 2. ').bold = True
            p31.add_run('En caso de que se efectúen mejoras, estas no podrán retirarse salvo que EL ARRENDADOR lo exija por escrito, para lo cual EL ARRENDATARIO accederá inmediatamente a su costa, dejando el Inmueble en el mismo estado en que lo recibió, salvo el deterioro natural por el uso legítimo. ')


        ##___________________

        p33 = documentoObjeto.add_paragraph("")
        p33.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            titulo24='CLÁUSULA VIGÉSIMA TERCERA. REQUERIMIENTOS: '
            p33.add_run(titulo24).bold = True
            p33.add_run('Las partes que suscriben este contrato renunciarán expresamente a los requerimientos de que tratan la ley 820 de 2003 y los artículos 2007 del código civil y el 423 del Código General del Proceso; y en general a los que consagre cualquier norma sustancial o procesal para efectos de la constitución en mora. ')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo24='CLÁUSULA VIGÉSIMA CUARTA. REQUERIMIENTOS: '
            p33.add_run(titulo24).bold = True
            p33.add_run('Las partes que suscriben este contrato renunciarán expresamente a los requerimientos de que tratan la ley 820 de 2003 y los artículos 2007 del código civil y el 423 del Código General del Proceso; y en general a los que consagre cualquier norma sustancial o procesal para efectos de la constitución en mora. ')
            
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo24='CLÁUSULA VIGÉSIMA SEGUNDA. REQUERIMIENTOS: '
            p33.add_run(titulo24).bold = True
            p33.add_run('Las partes que suscriben este contrato renunciarán expresamente a los requerimientos de que tratan la ley 820 de 2003 y los artículos 2007 del código civil y el 423 del Código General del Proceso; y en general a los que consagre cualquier norma sustancial o procesal para efectos de la constitución en mora. ')
            
        elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo24='CLÁUSULA VIGÉSIMA PRIMERA. REQUERIMIENTOS: '
            p33.add_run(titulo24).bold = True
            p33.add_run('Las partes que suscriben este contrato renunciarán expresamente a los requerimientos de que tratan la ley 820 de 2003 y los artículos 2007 del código civil y el 423 del Código General del Proceso; y en general a los que consagre cualquier norma sustancial o procesal para efectos de la constitución en mora. ')

        #### 
        p35 = documentoObjeto.add_paragraph("")
        p35.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            titulo25='CLÁUSULA VIGÉSIMA CUARTA. SUBARRIENDO Y CESIÓN: '
            p35.add_run(titulo25).bold = True
            p35.add_run('El ARRENDATARIO se obliga expresamente a no ceder o subarrendar el inmueble, ni transferir su tenencia. En caso de incumplimiento, el ARRENDADOR podrá dar por terminado el contrato de arrendamiento y exigir la entrega del inmueble. PARÁGRAFO 1. El ARRENDATARIO acepta desde ahora cualquier cesión que realice el ARRENDADOR del presente contrato cuando este se haya notificado al ARRENDATARIO  y a sus deudores solidarios, mediante comunicación enviada por correo certificado. La notificación se entenderá surtida desde la fecha de envío de la citada comunicación. PARÁGRAFO 2. En todo caso, el cesionario del contrato deberá dar cumplimiento a los presupuestos contenidos en el artículo 28 de la Ley 820 de 2003, salvo que la cesión recaiga sobre el propietario del inmueble. ')
            
        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            titulo25='CLÁUSULA VIGÉSIMA CUARTA. SUBARRIENDO Y CESIÓN: '
            p35.add_run(titulo25).bold = True
            p35.add_run('El ARRENDATARIO se obliga expresamente a no ceder o subarrendar el inmueble, ni transferir su tenencia. En caso de incumplimiento, el ARRENDADOR podrá dar por terminado el contrato de arrendamiento y exigir la entrega del inmueble. PARÁGRAFO 1. El ARRENDATARIO acepta desde ahora cualquier cesión que realice el ARRENDADOR del presente contrato cuando este se haya notificado al ARRENDATARIO, mediante comunicación enviada por correo certificado. La notificación se entenderá surtida desde la fecha de envío de la citada comunicación. PARÁGRAFO 2. En todo caso, el cesionario del contrato deberá dar cumplimiento a los presupuestos contenidos en el artículo 28 de la Ley 820 de 2003, salvo que la cesión recaiga sobre el propietario del inmueble.  ')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            titulo25='CLÁUSULA VIGÉSIMA QUINTA. SUBARRIENDO Y CESIÓN: '
            p35.add_run(titulo25).bold = True
            p35.add_run('El ARRENDATARIO se obliga expresamente a no ceder o subarrendar el inmueble, ni transferir su tenencia. En caso de incumplimiento, el ARRENDADOR podrá dar por terminado el contrato de arrendamiento y exigir la entrega del inmueble. PARÁGRAFO 1. El ARRENDATARIO acepta desde ahora cualquier cesión que realice el ARRENDADOR del presente contrato cuando este se haya notificado al ARRENDATARIO, mediante comunicación enviada por correo certificado. La notificación se entenderá surtida desde la fecha de envío de la citada comunicación. PARÁGRAFO 2. En todo caso, el cesionario del contrato deberá dar cumplimiento a los presupuestos contenidos en el artículo 28 de la Ley 820 de 2003, salvo que la cesión recaiga sobre el propietario del inmueble.  ')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            titulo25='CLÁUSULA VIGÉSIMA QUINTA. SUBARRIENDO Y CESIÓN: '
            p35.add_run(titulo25).bold = True
            p35.add_run('El ARRENDATARIO se obliga expresamente a no ceder o subarrendar el inmueble, ni transferir su tenencia. En caso de incumplimiento, el ARRENDADOR podrá dar por terminado el contrato de arrendamiento y exigir la entrega del inmueble. PARÁGRAFO 1. El ARRENDATARIO acepta desde ahora cualquier cesión que realice el ARRENDADOR del presente contrato cuando este se haya notificado al ARRENDATARIO  y a sus deudores solidarios, mediante comunicación enviada por correo certificado. La notificación se entenderá surtida desde la fecha de envío de la citada comunicación. PARÁGRAFO 2. En todo caso, el cesionario del contrato deberá dar cumplimiento a los presupuestos contenidos en el artículo 28 de la Ley 820 de 2003, salvo que la cesión recaiga sobre el propietario del inmueble. ')
            ##SI ADMINISTRADOR
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
            titulo25='CLÁUSULA VIGÉSIMA SEGUNDA. SUBARRIENDO Y CESIÓN: '
            p35.add_run(titulo25).bold = True
            p35.add_run('El ARRENDATARIO se obliga expresamente a no ceder o subarrendar el inmueble, ni transferir su tenencia. En caso de incumplimiento, el ARRENDADOR podrá dar por terminado el contrato de arrendamiento y exigir la entrega del inmueble. PARÁGRAFO 1. El ARRENDATARIO acepta desde ahora cualquier cesión que realice el ARRENDADOR del presente contrato cuando este se haya notificado al ARRENDATARIO  y a sus deudores solidarios, mediante comunicación enviada por correo certificado. La notificación se entenderá surtida desde la fecha de envío de la citada comunicación. PARÁGRAFO 2. En todo caso, el cesionario del contrato deberá dar cumplimiento a los presupuestos contenidos en el artículo 28 de la Ley 820 de 2003, salvo que la cesión recaiga sobre el propietario del inmueble. ')
            
        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
            titulo25='CLÁUSULA VIGÉSIMA SEGUNDA. SUBARRIENDO Y CESIÓN: '
            p35.add_run(titulo25).bold = True
            p35.add_run('El ARRENDATARIO se obliga expresamente a no ceder o subarrendar el inmueble, ni transferir su tenencia. En caso de incumplimiento, el ARRENDADOR podrá dar por terminado el contrato de arrendamiento y exigir la entrega del inmueble. PARÁGRAFO 1. El ARRENDATARIO acepta desde ahora cualquier cesión que realice el ARRENDADOR del presente contrato cuando este se haya notificado al ARRENDATARIO, mediante comunicación enviada por correo certificado. La notificación se entenderá surtida desde la fecha de envío de la citada comunicación. PARÁGRAFO 2. En todo caso, el cesionario del contrato deberá dar cumplimiento a los presupuestos contenidos en el artículo 28 de la Ley 820 de 2003, salvo que la cesión recaiga sobre el propietario del inmueble.  ')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
            titulo25='CLÁUSULA VIGÉSIMA TERCERA. SUBARRIENDO Y CESIÓN: '
            p35.add_run(titulo25).bold = True
            p35.add_run('El ARRENDATARIO se obliga expresamente a no ceder o subarrendar el inmueble, ni transferir su tenencia. En caso de incumplimiento, el ARRENDADOR podrá dar por terminado el contrato de arrendamiento y exigir la entrega del inmueble. PARÁGRAFO 1. El ARRENDATARIO acepta desde ahora cualquier cesión que realice el ARRENDADOR del presente contrato cuando este se haya notificado al ARRENDATARIO, mediante comunicación enviada por correo certificado. La notificación se entenderá surtida desde la fecha de envío de la citada comunicación. PARÁGRAFO 2. En todo caso, el cesionario del contrato deberá dar cumplimiento a los presupuestos contenidos en el artículo 28 de la Ley 820 de 2003, salvo que la cesión recaiga sobre el propietario del inmueble.  ')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
            titulo25='CLÁUSULA VIGÉSIMA TERCERA. SUBARRIENDO Y CESIÓN: '
            p35.add_run(titulo25).bold = True
            p35.add_run('El ARRENDATARIO se obliga expresamente a no ceder o subarrendar el inmueble, ni transferir su tenencia. En caso de incumplimiento, el ARRENDADOR podrá dar por terminado el contrato de arrendamiento y exigir la entrega del inmueble. PARÁGRAFO 1. El ARRENDATARIO acepta desde ahora cualquier cesión que realice el ARRENDADOR del presente contrato cuando este se haya notificado al ARRENDATARIO  y a sus deudores solidarios, mediante comunicación enviada por correo certificado. La notificación se entenderá surtida desde la fecha de envío de la citada comunicación. PARÁGRAFO 2. En todo caso, el cesionario del contrato deberá dar cumplimiento a los presupuestos contenidos en el artículo 28 de la Ley 820 de 2003, salvo que la cesión recaiga sobre el propietario del inmueble. ')

        ##______________________
        p37 = documentoObjeto.add_paragraph("")
        p37.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            titulo26='CLÁUSULA VIGÉSIMA QUINTA. ABANDONO DEL INMUEBLE: '
            p37.add_run(titulo26).bold = True
            p37.add_run('EL ARRENDATARIO autoriza de manera expresa e irrevocable al ARRENDADOR para ingresar al Inmueble y recuperar su tenencia, con el solo requisito de la presencia de dos (2) testigos, en procura de evitar el deterioro o desmantelamiento del Inmueble, en el evento que por cualquier causa o circunstancia el Inmueble permanezca abandonado o deshabitado por el término de dos (2) meses o más y que la exposición al riesgo sea tal que amenace la integridad física del bien o la seguridad del vecindario.')

        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo26='CLÁUSULA VIGÉSIMA SEXTA. ABANDONO DEL INMUEBLE: '
            p37.add_run(titulo26).bold = True
            p37.add_run('EL ARRENDATARIO autoriza de manera expresa e irrevocable al ARRENDADOR para ingresar al Inmueble y recuperar su tenencia, con el solo requisito de la presencia de dos (2) testigos, en procura de evitar el deterioro o desmantelamiento del Inmueble, en el evento que por cualquier causa o circunstancia el Inmueble permanezca abandonado o deshabitado por el término de dos (2) meses o más y que la exposición al riesgo sea tal que amenace la integridad física del bien o la seguridad del vecindario.')
    
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo26='CLÁUSULA VIGÉSIMA CUARTA. ABANDONO DEL INMUEBLE: '
            p37.add_run(titulo26).bold = True
            p37.add_run('EL ARRENDATARIO autoriza de manera expresa e irrevocable al ARRENDADOR para ingresar al Inmueble y recuperar su tenencia, con el solo requisito de la presencia de dos (2) testigos, en procura de evitar el deterioro o desmantelamiento del Inmueble, en el evento que por cualquier causa o circunstancia el Inmueble permanezca abandonado o deshabitado por el término de dos (2) meses o más y que la exposición al riesgo sea tal que amenace la integridad física del bien o la seguridad del vecindario.')

        elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo26='CLÁUSULA VIGÉSIMA TERCERA. ABANDONO DEL INMUEBLE: '
            p37.add_run(titulo26).bold = True
            p37.add_run('EL ARRENDATARIO autoriza de manera expresa e irrevocable al ARRENDADOR para ingresar al Inmueble y recuperar su tenencia, con el solo requisito de la presencia de dos (2) testigos, en procura de evitar el deterioro o desmantelamiento del Inmueble, en el evento que por cualquier causa o circunstancia el Inmueble permanezca abandonado o deshabitado por el término de dos (2) meses o más y que la exposición al riesgo sea tal que amenace la integridad física del bien o la seguridad del vecindario.')

        ####

        p39 = documentoObjeto.add_paragraph("")
        p39.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            titulo27='CLÁUSULA VIGÉSIMA SEXTA. EXENCIÓN DE RESPONSABILIDAD: '
            p39.add_run(titulo27).bold = True
            p39.add_run('El ARRENDADOR no asume responsabilidad alguna por los daños o perjuicios que el ARRENDATARIO  pueda sufrir por causas atribuibles a terceros, ni por robos, hurtos, o siniestros causados por incendios, inundación o terrorismo. Serán de cargo del ARRENDATARIO las medidas, dirección y manejo tomadas para la seguridad del bien. ')
        
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo27='CLÁUSULA VIGÉSIMA SÉPTIMA. EXENCIÓN DE RESPONSABILIDAD: '
            p39.add_run(titulo27).bold = True
            p39.add_run('El ARRENDADOR no asume responsabilidad alguna por los daños o perjuicios que el ARRENDATARIO  pueda sufrir por causas atribuibles a terceros, ni por robos, hurtos, o siniestros causados por incendios, inundación o terrorismo. Serán de cargo del ARRENDATARIO las medidas, dirección y manejo tomadas para la seguridad del bien. ')
            
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo27='CLÁUSULA VIGÉSIMA QUINTA. EXENCIÓN DE RESPONSABILIDAD: '
            p39.add_run(titulo27).bold = True
            p39.add_run('El ARRENDADOR no asume responsabilidad alguna por los daños o perjuicios que el ARRENDATARIO  pueda sufrir por causas atribuibles a terceros, ni por robos, hurtos, o siniestros causados por incendios, inundación o terrorismo. Serán de cargo del ARRENDATARIO las medidas, dirección y manejo tomadas para la seguridad del bien. ')
            
        elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo27='CLÁUSULA VIGÉSIMA CUARTA. EXENCIÓN DE RESPONSABILIDAD: '
            p39.add_run(titulo27).bold = True
            p39.add_run('El ARRENDADOR no asume responsabilidad alguna por los daños o perjuicios que el ARRENDATARIO  pueda sufrir por causas atribuibles a terceros, ni por robos, hurtos, o siniestros causados por incendios, inundación o terrorismo. Serán de cargo del ARRENDATARIO las medidas, dirección y manejo tomadas para la seguridad del bien. ')

        ####
        p40 = documentoObjeto.add_paragraph("")
        p40.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            titulo28='CLÁUSULA VIGÉSIMA SÉPTIMA. MÉRITO EJECUTIVO: '
            p40.add_run(titulo28).bold = True
            p40.add_run('EL ARRENDATARIO declara de manera expresa que reconoce y acepta que este contrato presta mérito ejecutivo para exigir por parte del ARRENDADOR las sumas causadas y no pagadas por cánones de arrendamiento, multas y sanciones, servicios públicos y cualquier otra suma de dinero que por cualquier concepto deba ser pagada por EL ARRENDATARIO en cualquiera de las obligaciones a su cargo en virtud de la ley o de este contrato, para lo cual bastará la sola afirmación de incumplimiento por parte del ARRENDADOR, afirmación que sólo podrá ser desvirtuada por EL ARRENDATARIO con la presentación de los respectivos recibos de pago.')
            p40.add_run(' PARÁGRAFO. ').bold = True
            p40.add_run('Las Partes acuerdan que cualquier copia autenticada ante Notario de este Contrato tendrá el mismo valor que el original para efectos judiciales y extrajudiciales. ')

        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo28='CLÁUSULA VIGÉSIMA OCTAVA. MÉRITO EJECUTIVO: '
            p40.add_run(titulo28).bold = True
            p40.add_run('EL ARRENDATARIO declara de manera expresa que reconoce y acepta que este contrato presta mérito ejecutivo para exigir por parte del ARRENDADOR las sumas causadas y no pagadas por cánones de arrendamiento, multas y sanciones, servicios públicos y cualquier otra suma de dinero que por cualquier concepto deba ser pagada por EL ARRENDATARIO en cualquiera de las obligaciones a su cargo en virtud de la ley o de este contrato, para lo cual bastará la sola afirmación de incumplimiento por parte del ARRENDADOR, afirmación que sólo podrá ser desvirtuada por EL ARRENDATARIO con la presentación de los respectivos recibos de pago.')
            p40.add_run(' PARÁGRAFO. ').bold = True
            p40.add_run('Las Partes acuerdan que cualquier copia autenticada ante Notario de este Contrato tendrá el mismo valor que el original para efectos judiciales y extrajudiciales. ')
            
        elif(Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo28='CLÁUSULA VIGÉSIMA SEXTA. MÉRITO EJECUTIVO: '
            p40.add_run(titulo28).bold = True
            p40.add_run('EL ARRENDATARIO declara de manera expresa que reconoce y acepta que este contrato presta mérito ejecutivo para exigir por parte del ARRENDADOR las sumas causadas y no pagadas por cánones de arrendamiento, multas y sanciones, servicios públicos y cualquier otra suma de dinero que por cualquier concepto deba ser pagada por EL ARRENDATARIO en cualquiera de las obligaciones a su cargo en virtud de la ley o de este contrato, para lo cual bastará la sola afirmación de incumplimiento por parte del ARRENDADOR, afirmación que sólo podrá ser desvirtuada por EL ARRENDATARIO con la presentación de los respectivos recibos de pago.')
            p40.add_run(' PARÁGRAFO. ').bold = True
            p40.add_run('Las Partes acuerdan que cualquier copia autenticada ante Notario de este Contrato tendrá el mismo valor que el original para efectos judiciales y extrajudiciales. ')
            
        elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo28='CLÁUSULA VIGÉSIMA QUINTA. MÉRITO EJECUTIVO: '
            p40.add_run(titulo28).bold = True
            p40.add_run('EL ARRENDATARIO declara de manera expresa que reconoce y acepta que este contrato presta mérito ejecutivo para exigir por parte del ARRENDADOR las sumas causadas y no pagadas por cánones de arrendamiento, multas y sanciones, servicios públicos y cualquier otra suma de dinero que por cualquier concepto deba ser pagada por EL ARRENDATARIO en cualquiera de las obligaciones a su cargo en virtud de la ley o de este contrato, para lo cual bastará la sola afirmación de incumplimiento por parte del ARRENDADOR, afirmación que sólo podrá ser desvirtuada por EL ARRENDATARIO con la presentación de los respectivos recibos de pago.')
            p40.add_run(' PARÁGRAFO. ').bold = True
            p40.add_run('Las Partes acuerdan que cualquier copia autenticada ante Notario de este Contrato tendrá el mismo valor que el original para efectos judiciales y extrajudiciales. ')

        ####
        p41 = documentoObjeto.add_paragraph("")
        p41.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            titulo29='CLÁUSULA VIGÉSIMA OCTAVA. AUTORIZACIÓN DE CONSULTA EN CENTRALES DE RIESGOS: '
            p41.add_run(titulo29).bold = True
            p41.add_run('EL ARRENDATARIO autoriza expresamente e irrevocablemente a EL ARRENDADOR Y/O AL CESIONARIO de este Contrato a consultar información de EL ARRENDATARIO que obre en las bases de datos de información del comportamiento financiero y crediticio o centrales de riesgo que existan en el país, así como a reportar a dichas bases de datos cualquier incumplimiento de EL ARRENDATARIO a este Contrato.')
            p41.add_run(' PARÁGRAFO 1. AUTORIZACIONES: ').bold = True
            p41.add_run('EL ARRENDATARIO Y DEUDORES SOLIDARIOS, autorizan de manera irrevocable a EL ARRENDADOR O A QUIEN REPRESENTE SUS DERECHOS U OSTENTE EN EL FUTURO LA CALIDAD DE ACREEDOR, a consultar, solicitar, suministrar, reportar, procesar y divulgar toda la información que se refiera al comportamiento crediticio, financiero, comercial, de servicios, a CIFIN, DATACREDITO, o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos, y para que en el evento en que se constituyan en mora en el pago de cualquier servicio público, arrendamiento, cualquier otro concepto que sean a su cargo, durante el término inicial o el de sus prórrogas o a la terminación del contrato, se incorporen sus nombres, apellidos y documento de identificación a los archivos de deudores morosos o con referencias negativas que lleve CIFIN, DATACREDITO, o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos. EL ARRENDATARIO y DEUDORES SOLIDARIOS exoneran de toda responsabilidad por la inclusión de tales datos tanto a EL ARRENDADOR como a la entidad que produzca el correspondiente archivo. EL ARRENDATARIO y DEUDORES SOLIDARIOS conocen que el alcance de esta autorización implica, que el comportamiento frente a sus obligaciones será registrado con el objeto de suministrar información suficiente y adecuada al mercado sobre el estado de sus obligaciones financieras, comerciales, crediticias, de servicios etc. En consecuencia, quienes se encuentren afiliados y/o tengan acceso a las centrales de información y entidades aquí relacionadas o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos, podrán conocer esta información, de conformidad con la legislación y jurisprudencia aplicable. La información podrá ser igualmente utilizada para efectos estadísticos. Los derechos y obligaciones de EL ARRENDATARIO así como la permanencia de su información en las bases de datos corresponden a lo determinado por el ordenamiento jurídico aplicable del cual, por ser de carácter público, manifiesta que está enterado. Así mismo, manifiestan EL ARRENDATARIO y DEUDORES SOLIDARIOS que conocen el contenido del reglamento de las citadas entidades. En caso de que en el futuro, EL ARRENDADOR, efectúe, a favor de un tercero, una venta de cartera, o una cesión a cualquier título de las obligaciones a cargo de EL ARRENDATARIO, los efectos de la anterior autorización se extenderán a este en los mismos términos y condiciones. Así mismo, autoriza a las entidades encargadas del manejo de la información, a que, en su calidad de operadores, pongan la información a disposición de otros operadores nacionales o extranjeros, en los términos que establece la ley, siempre y cuando su objeto sea similar al aquí establecido.')
        
        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo29='CLÁUSULA VIGÉSIMA OCTAVA. AUTORIZACIÓN DE CONSULTA EN CENTRALES DE RIESGOS: '
            p41.add_run(titulo29).bold = True
            p41.add_run('EL ARRENDATARIO autoriza expresamente e irrevocablemente a EL ARRENDADOR Y/O AL CESIONARIO de este Contrato a consultar información de EL ARRENDATARIO que obre en las bases de datos de información del comportamiento financiero y crediticio o centrales de riesgo que existan en el país, así como a reportar a dichas bases de datos cualquier incumplimiento de EL ARRENDATARIO a este Contrato.')
            p41.add_run(' PARÁGRAFO 1. AUTORIZACIONES: ').bold = True
            p41.add_run('EL ARRENDATARIO , autoriza de manera irrevocable a EL ARRENDADOR O A QUIEN REPRESENTE SUS DERECHOS U OSTENTE EN EL FUTURO LA CALIDAD DE ACREEDOR, a consultar, solicitar, suministrar, reportar, procesar y divulgar toda la información que se refiera al comportamiento crediticio, financiero, comercial, de servicios, a CIFIN, DATACREDITO, o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos, y para que en el evento en que se constituyan en mora en el pago de cualquier servicio público, arrendamiento, cualquier otro concepto que sean a su cargo, durante el término inicial o el de sus prórrogas o a la terminación del contrato, se incorporen sus nombres, apellidos y documento de identificación a los archivos  o con referencias negativas que lleve CIFIN, DATACREDITO, o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos. EL ARRENDATARIO exonera de toda responsabilidad por la inclusión de tales datos tanto a EL ARRENDADOR como a la entidad que produzca el correspondiente archivo. EL ARRENDATARIO  conoce que el alcance de esta autorización implica, que el comportamiento frente a sus obligaciones será registrado con el objeto de suministrar información suficiente y adecuada al mercado sobre el estado de sus obligaciones financieras, comerciales, crediticias, de servicios etc. En consecuencia, quienes se encuentren afiliados y/o tengan acceso a las centrales de información y entidades aquí relacionadas o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos, podrán conocer esta información, de conformidad con la legislación y jurisprudencia aplicable. La información podrá ser igualmente utilizada para efectos estadísticos. Los derechos y obligaciones de EL ARRENDATARIO así como la permanencia de su información en las bases de datos corresponden a lo determinado por el ordenamiento jurídico aplicable del cual, por ser de carácter público, manifiesta que está enterado. Así mismo, manifiestan EL ARRENDATARIO  que conoce el contenido del reglamento de las citadas entidades. En caso de que en el futuro, EL ARRENDADOR, efectúe, a favor de un tercero, una venta de cartera, o una cesión a cualquier título de las obligaciones a cargo de EL ARRENDATARIO, los efectos de la anterior autorización se extenderán a este en los mismos términos y condiciones. Así mismo, autoriza a las entidades encargadas del manejo de la información, a que, en su calidad de operadores, pongan la información a disposición de otros operadores nacionales o extranjeros, en los términos que establece la ley, siempre y cuando su objeto sea similar al aquí establecido.')
        
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo29='CLÁUSULA VIGÉSIMA NOVENA. AUTORIZACIÓN DE CONSULTA EN CENTRALES DE RIESGOS: '
            p41.add_run(titulo29).bold = True
            p41.add_run('EL ARRENDATARIO autoriza expresamente e irrevocablemente a EL ARRENDADOR Y/O AL CESIONARIO de este Contrato a consultar información de EL ARRENDATARIO que obre en las bases de datos de información del comportamiento financiero y crediticio o centrales de riesgo que existan en el país, así como a reportar a dichas bases de datos cualquier incumplimiento de EL ARRENDATARIO a este Contrato.')
            p41.add_run(' PARÁGRAFO 1. AUTORIZACIONES: ').bold = True
            p41.add_run('EL ARRENDATARIO , autoriza de manera irrevocable a EL ARRENDADOR O A QUIEN REPRESENTE SUS DERECHOS U OSTENTE EN EL FUTURO LA CALIDAD DE ACREEDOR, a consultar, solicitar, suministrar, reportar, procesar y divulgar toda la información que se refiera al comportamiento crediticio, financiero, comercial, de servicios, a CIFIN, DATACREDITO, o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos, y para que en el evento en que se constituyan en mora en el pago de cualquier servicio público, arrendamiento, cualquier otro concepto que sean a su cargo, durante el término inicial o el de sus prórrogas o a la terminación del contrato, se incorporen sus nombres, apellidos y documento de identificación a los archivos  o con referencias negativas que lleve CIFIN, DATACREDITO, o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos. EL ARRENDATARIO exonera de toda responsabilidad por la inclusión de tales datos tanto a EL ARRENDADOR como a la entidad que produzca el correspondiente archivo. EL ARRENDATARIO  conoce que el alcance de esta autorización implica, que el comportamiento frente a sus obligaciones será registrado con el objeto de suministrar información suficiente y adecuada al mercado sobre el estado de sus obligaciones financieras, comerciales, crediticias, de servicios etc. En consecuencia, quienes se encuentren afiliados y/o tengan acceso a las centrales de información y entidades aquí relacionadas o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos, podrán conocer esta información, de conformidad con la legislación y jurisprudencia aplicable. La información podrá ser igualmente utilizada para efectos estadísticos. Los derechos y obligaciones de EL ARRENDATARIO así como la permanencia de su información en las bases de datos corresponden a lo determinado por el ordenamiento jurídico aplicable del cual, por ser de carácter público, manifiesta que está enterado. Así mismo, manifiestan EL ARRENDATARIO  que conoce el contenido del reglamento de las citadas entidades. En caso de que en el futuro, EL ARRENDADOR, efectúe, a favor de un tercero, una venta de cartera, o una cesión a cualquier título de las obligaciones a cargo de EL ARRENDATARIO, los efectos de la anterior autorización se extenderán a este en los mismos términos y condiciones. Así mismo, autoriza a las entidades encargadas del manejo de la información, a que, en su calidad de operadores, pongan la información a disposición de otros operadores nacionales o extranjeros, en los términos que establece la ley, siempre y cuando su objeto sea similar al aquí establecido.')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo29='CLÁUSULA VIGÉSIMA NOVENA. AUTORIZACIÓN DE CONSULTA EN CENTRALES DE RIESGOS: '
            p41.add_run(titulo29).bold = True
            p41.add_run('EL ARRENDATARIO autoriza expresamente e irrevocablemente a EL ARRENDADOR Y/O AL CESIONARIO de este Contrato a consultar información de EL ARRENDATARIO que obre en las bases de datos de información del comportamiento financiero y crediticio o centrales de riesgo que existan en el país, así como a reportar a dichas bases de datos cualquier incumplimiento de EL ARRENDATARIO a este Contrato.')
            p41.add_run(' PARÁGRAFO 1. AUTORIZACIONES: ').bold = True
            p41.add_run('EL ARRENDATARIO Y DEUDORES SOLIDARIOS, autorizan de manera irrevocable a EL ARRENDADOR O A QUIEN REPRESENTE SUS DERECHOS U OSTENTE EN EL FUTURO LA CALIDAD DE ACREEDOR, a consultar, solicitar, suministrar, reportar, procesar y divulgar toda la información que se refiera al comportamiento crediticio, financiero, comercial, de servicios, a CIFIN, DATACREDITO, o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos, y para que en el evento en que se constituyan en mora en el pago de cualquier servicio público, arrendamiento, cualquier otro concepto que sean a su cargo, durante el término inicial o el de sus prórrogas o a la terminación del contrato, se incorporen sus nombres, apellidos y documento de identificación a los archivos de deudores morosos o con referencias negativas que lleve CIFIN, DATACREDITO, o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos. EL ARRENDATARIO y DEUDORES SOLIDARIOS exoneran de toda responsabilidad por la inclusión de tales datos tanto a EL ARRENDADOR como a la entidad que produzca el correspondiente archivo. EL ARRENDATARIO y DEUDORES SOLIDARIOS conocen que el alcance de esta autorización implica, que el comportamiento frente a sus obligaciones será registrado con el objeto de suministrar información suficiente y adecuada al mercado sobre el estado de sus obligaciones financieras, comerciales, crediticias, de servicios etc. En consecuencia, quienes se encuentren afiliados y/o tengan acceso a las centrales de información y entidades aquí relacionadas o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos, podrán conocer esta información, de conformidad con la legislación y jurisprudencia aplicable. La información podrá ser igualmente utilizada para efectos estadísticos. Los derechos y obligaciones de EL ARRENDATARIO así como la permanencia de su información en las bases de datos corresponden a lo determinado por el ordenamiento jurídico aplicable del cual, por ser de carácter público, manifiesta que está enterado. Así mismo, manifiestan EL ARRENDATARIO y DEUDORES SOLIDARIOS que conocen el contenido del reglamento de las citadas entidades. En caso de que en el futuro, EL ARRENDADOR, efectúe, a favor de un tercero, una venta de cartera, o una cesión a cualquier título de las obligaciones a cargo de EL ARRENDATARIO, los efectos de la anterior autorización se extenderán a este en los mismos términos y condiciones. Así mismo, autoriza a las entidades encargadas del manejo de la información, a que, en su calidad de operadores, pongan la información a disposición de otros operadores nacionales o extranjeros, en los términos que establece la ley, siempre y cuando su objeto sea similar al aquí establecido.')
            ##SIN ADMINISTRADOR
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
            titulo29='CLÁUSULA VIGÉSIMA SEXTA. AUTORIZACIÓN DE CONSULTA EN CENTRALES DE RIESGOS: '
            p41.add_run(titulo29).bold = True
            p41.add_run('EL ARRENDATARIO autoriza expresamente e irrevocablemente a EL ARRENDADOR Y/O AL CESIONARIO de este Contrato a consultar información de EL ARRENDATARIO que obre en las bases de datos de información del comportamiento financiero y crediticio o centrales de riesgo que existan en el país, así como a reportar a dichas bases de datos cualquier incumplimiento de EL ARRENDATARIO a este Contrato.')
            p41.add_run(' PARÁGRAFO 1. AUTORIZACIONES: ').bold = True
            p41.add_run('EL ARRENDATARIO Y DEUDORES SOLIDARIOS, autorizan de manera irrevocable a EL ARRENDADOR O A QUIEN REPRESENTE SUS DERECHOS U OSTENTE EN EL FUTURO LA CALIDAD DE ACREEDOR, a consultar, solicitar, suministrar, reportar, procesar y divulgar toda la información que se refiera al comportamiento crediticio, financiero, comercial, de servicios, a CIFIN, DATACREDITO, o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos, y para que en el evento en que se constituyan en mora en el pago de cualquier servicio público, arrendamiento, cualquier otro concepto que sean a su cargo, durante el término inicial o el de sus prórrogas o a la terminación del contrato, se incorporen sus nombres, apellidos y documento de identificación a los archivos de deudores morosos o con referencias negativas que lleve CIFIN, DATACREDITO, o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos. EL ARRENDATARIO y DEUDORES SOLIDARIOS exoneran de toda responsabilidad por la inclusión de tales datos tanto a EL ARRENDADOR como a la entidad que produzca el correspondiente archivo. EL ARRENDATARIO y DEUDORES SOLIDARIOS conocen que el alcance de esta autorización implica, que el comportamiento frente a sus obligaciones será registrado con el objeto de suministrar información suficiente y adecuada al mercado sobre el estado de sus obligaciones financieras, comerciales, crediticias, de servicios etc. En consecuencia, quienes se encuentren afiliados y/o tengan acceso a las centrales de información y entidades aquí relacionadas o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos, podrán conocer esta información, de conformidad con la legislación y jurisprudencia aplicable. La información podrá ser igualmente utilizada para efectos estadísticos. Los derechos y obligaciones de EL ARRENDATARIO así como la permanencia de su información en las bases de datos corresponden a lo determinado por el ordenamiento jurídico aplicable del cual, por ser de carácter público, manifiesta que está enterado. Así mismo, manifiestan EL ARRENDATARIO y DEUDORES SOLIDARIOS que conocen el contenido del reglamento de las citadas entidades. En caso de que en el futuro, EL ARRENDADOR, efectúe, a favor de un tercero, una venta de cartera, o una cesión a cualquier título de las obligaciones a cargo de EL ARRENDATARIO, los efectos de la anterior autorización se extenderán a este en los mismos términos y condiciones. Así mismo, autoriza a las entidades encargadas del manejo de la información, a que, en su calidad de operadores, pongan la información a disposición de otros operadores nacionales o extranjeros, en los términos que establece la ley, siempre y cuando su objeto sea similar al aquí establecido.')
            
        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo29='CLÁUSULA VIGÉSIMA SEXTA. AUTORIZACIÓN DE CONSULTA EN CENTRALES DE RIESGOS: '
            p41.add_run(titulo29).bold = True
            p41.add_run('EL ARRENDATARIO autoriza expresamente e irrevocablemente a EL ARRENDADOR Y/O AL CESIONARIO de este Contrato a consultar información de EL ARRENDATARIO que obre en las bases de datos de información del comportamiento financiero y crediticio o centrales de riesgo que existan en el país, así como a reportar a dichas bases de datos cualquier incumplimiento de EL ARRENDATARIO a este Contrato.')
            p41.add_run(' PARÁGRAFO 1. AUTORIZACIONES: ').bold = True
            p41.add_run('EL ARRENDATARIO , autoriza de manera irrevocable a EL ARRENDADOR O A QUIEN REPRESENTE SUS DERECHOS U OSTENTE EN EL FUTURO LA CALIDAD DE ACREEDOR, a consultar, solicitar, suministrar, reportar, procesar y divulgar toda la información que se refiera al comportamiento crediticio, financiero, comercial, de servicios, a CIFIN, DATACREDITO, o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos, y para que en el evento en que se constituyan en mora en el pago de cualquier servicio público, arrendamiento, cualquier otro concepto que sean a su cargo, durante el término inicial o el de sus prórrogas o a la terminación del contrato, se incorporen sus nombres, apellidos y documento de identificación a los archivos  o con referencias negativas que lleve CIFIN, DATACREDITO, o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos. EL ARRENDATARIO exonera de toda responsabilidad por la inclusión de tales datos tanto a EL ARRENDADOR como a la entidad que produzca el correspondiente archivo. EL ARRENDATARIO  conoce que el alcance de esta autorización implica, que el comportamiento frente a sus obligaciones será registrado con el objeto de suministrar información suficiente y adecuada al mercado sobre el estado de sus obligaciones financieras, comerciales, crediticias, de servicios etc. En consecuencia, quienes se encuentren afiliados y/o tengan acceso a las centrales de información y entidades aquí relacionadas o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos, podrán conocer esta información, de conformidad con la legislación y jurisprudencia aplicable. La información podrá ser igualmente utilizada para efectos estadísticos. Los derechos y obligaciones de EL ARRENDATARIO así como la permanencia de su información en las bases de datos corresponden a lo determinado por el ordenamiento jurídico aplicable del cual, por ser de carácter público, manifiesta que está enterado. Así mismo, manifiestan EL ARRENDATARIO  que conoce el contenido del reglamento de las citadas entidades. En caso de que en el futuro, EL ARRENDADOR, efectúe, a favor de un tercero, una venta de cartera, o una cesión a cualquier título de las obligaciones a cargo de EL ARRENDATARIO, los efectos de la anterior autorización se extenderán a este en los mismos términos y condiciones. Así mismo, autoriza a las entidades encargadas del manejo de la información, a que, en su calidad de operadores, pongan la información a disposición de otros operadores nacionales o extranjeros, en los términos que establece la ley, siempre y cuando su objeto sea similar al aquí establecido.')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo29='CLÁUSULA VIGÉSIMA SÉPTIMA. AUTORIZACIÓN DE CONSULTA EN CENTRALES DE RIESGOS: '
            p41.add_run(titulo29).bold = True
            p41.add_run('EL ARRENDATARIO autoriza expresamente e irrevocablemente a EL ARRENDADOR Y/O AL CESIONARIO de este Contrato a consultar información de EL ARRENDATARIO que obre en las bases de datos de información del comportamiento financiero y crediticio o centrales de riesgo que existan en el país, así como a reportar a dichas bases de datos cualquier incumplimiento de EL ARRENDATARIO a este Contrato.')
            p41.add_run(' PARÁGRAFO 1. AUTORIZACIONES: ').bold = True
            p41.add_run('EL ARRENDATARIO , autoriza de manera irrevocable a EL ARRENDADOR O A QUIEN REPRESENTE SUS DERECHOS U OSTENTE EN EL FUTURO LA CALIDAD DE ACREEDOR, a consultar, solicitar, suministrar, reportar, procesar y divulgar toda la información que se refiera al comportamiento crediticio, financiero, comercial, de servicios, a CIFIN, DATACREDITO, o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos, y para que en el evento en que se constituyan en mora en el pago de cualquier servicio público, arrendamiento, cualquier otro concepto que sean a su cargo, durante el término inicial o el de sus prórrogas o a la terminación del contrato, se incorporen sus nombres, apellidos y documento de identificación a los archivos  o con referencias negativas que lleve CIFIN, DATACREDITO, o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos. EL ARRENDATARIO exonera de toda responsabilidad por la inclusión de tales datos tanto a EL ARRENDADOR como a la entidad que produzca el correspondiente archivo. EL ARRENDATARIO  conoce que el alcance de esta autorización implica, que el comportamiento frente a sus obligaciones será registrado con el objeto de suministrar información suficiente y adecuada al mercado sobre el estado de sus obligaciones financieras, comerciales, crediticias, de servicios etc. En consecuencia, quienes se encuentren afiliados y/o tengan acceso a las centrales de información y entidades aquí relacionadas o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos, podrán conocer esta información, de conformidad con la legislación y jurisprudencia aplicable. La información podrá ser igualmente utilizada para efectos estadísticos. Los derechos y obligaciones de EL ARRENDATARIO así como la permanencia de su información en las bases de datos corresponden a lo determinado por el ordenamiento jurídico aplicable del cual, por ser de carácter público, manifiesta que está enterado. Así mismo, manifiestan EL ARRENDATARIO  que conoce el contenido del reglamento de las citadas entidades. En caso de que en el futuro, EL ARRENDADOR, efectúe, a favor de un tercero, una venta de cartera, o una cesión a cualquier título de las obligaciones a cargo de EL ARRENDATARIO, los efectos de la anterior autorización se extenderán a este en los mismos términos y condiciones. Así mismo, autoriza a las entidades encargadas del manejo de la información, a que, en su calidad de operadores, pongan la información a disposición de otros operadores nacionales o extranjeros, en los términos que establece la ley, siempre y cuando su objeto sea similar al aquí establecido.')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo29='CLÁUSULA VIGÉSIMA SÉPTIMA. AUTORIZACIÓN DE CONSULTA EN CENTRALES DE RIESGOS: '
            p41.add_run(titulo29).bold = True
            p41.add_run('EL ARRENDATARIO autoriza expresamente e irrevocablemente a EL ARRENDADOR Y/O AL CESIONARIO de este Contrato a consultar información de EL ARRENDATARIO que obre en las bases de datos de información del comportamiento financiero y crediticio o centrales de riesgo que existan en el país, así como a reportar a dichas bases de datos cualquier incumplimiento de EL ARRENDATARIO a este Contrato.')
            p41.add_run(' PARÁGRAFO 1. AUTORIZACIONES: ').bold = True
            p41.add_run('EL ARRENDATARIO Y DEUDORES SOLIDARIOS, autorizan de manera irrevocable a EL ARRENDADOR O A QUIEN REPRESENTE SUS DERECHOS U OSTENTE EN EL FUTURO LA CALIDAD DE ACREEDOR, a consultar, solicitar, suministrar, reportar, procesar y divulgar toda la información que se refiera al comportamiento crediticio, financiero, comercial, de servicios, a CIFIN, DATACREDITO, o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos, y para que en el evento en que se constituyan en mora en el pago de cualquier servicio público, arrendamiento, cualquier otro concepto que sean a su cargo, durante el término inicial o el de sus prórrogas o a la terminación del contrato, se incorporen sus nombres, apellidos y documento de identificación a los archivos de deudores morosos o con referencias negativas que lleve CIFIN, DATACREDITO, o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos. EL ARRENDATARIO y DEUDORES SOLIDARIOS exoneran de toda responsabilidad por la inclusión de tales datos tanto a EL ARRENDADOR como a la entidad que produzca el correspondiente archivo. EL ARRENDATARIO y DEUDORES SOLIDARIOS conocen que el alcance de esta autorización implica, que el comportamiento frente a sus obligaciones será registrado con el objeto de suministrar información suficiente y adecuada al mercado sobre el estado de sus obligaciones financieras, comerciales, crediticias, de servicios etc. En consecuencia, quienes se encuentren afiliados y/o tengan acceso a las centrales de información y entidades aquí relacionadas o cualquier otra entidad encargada del manejo de datos comerciales, personales o económicos, podrán conocer esta información, de conformidad con la legislación y jurisprudencia aplicable. La información podrá ser igualmente utilizada para efectos estadísticos. Los derechos y obligaciones de EL ARRENDATARIO así como la permanencia de su información en las bases de datos corresponden a lo determinado por el ordenamiento jurídico aplicable del cual, por ser de carácter público, manifiesta que está enterado. Así mismo, manifiestan EL ARRENDATARIO y DEUDORES SOLIDARIOS que conocen el contenido del reglamento de las citadas entidades. En caso de que en el futuro, EL ARRENDADOR, efectúe, a favor de un tercero, una venta de cartera, o una cesión a cualquier título de las obligaciones a cargo de EL ARRENDATARIO, los efectos de la anterior autorización se extenderán a este en los mismos términos y condiciones. Así mismo, autoriza a las entidades encargadas del manejo de la información, a que, en su calidad de operadores, pongan la información a disposición de otros operadores nacionales o extranjeros, en los términos que establece la ley, siempre y cuando su objeto sea similar al aquí establecido.')


        #_______________________________________________
        p42 = documentoObjeto.add_paragraph("")
        p42.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo30='CLÁUSULA VIGÉSIMA NOVENA. INSPECCIÓN DEL INMUEBLE: '
            p42.add_run(titulo30).bold = True
            p42.add_run('EL ARRENDATARIO y los DEUDORES SOLIDARIOS, autorizan desde ahora al ARRENDADOR o a su representante, para visitar el inmueble arrendado en cualquier momento, previa comunicación y coordinación, cuándo las circunstancias así lo ameriten, en caso de venta, de arreglos o reparaciones locativas o simples inspecciones.') 
            
        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo30='CLÁUSULA VIGÉSIMA NOVENA. INSPECCIÓN DEL INMUEBLE: '
            p42.add_run(titulo30).bold = True
            p42.add_run('EL ARRENDATARIO, autoriza desde ahora al ARRENDADOR o a su representante, para visitar el inmueble arrendado en cualquier momento, previa comunicación y coordinación, cuándo las circunstancias así lo ameriten, en caso de venta, de arreglos o reparaciones locativas o simples inspecciones.') 
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo30='CLÁUSULA TRIGÉSIMA. INSPECCIÓN DEL INMUEBLE: '
            p42.add_run(titulo30).bold = True
            p42.add_run('EL ARRENDATARIO, autoriza desde ahora al ARRENDADOR o a su representante, para visitar el inmueble arrendado en cualquier momento, previa comunicación y coordinación, cuándo las circunstancias así lo ameriten, en caso de venta, de arreglos o reparaciones locativas o simples inspecciones.') 
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo30='CLÁUSULA TRIGÉSIMA. INSPECCIÓN DEL INMUEBLE: '
            p42.add_run(titulo30).bold = True
            p42.add_run('EL ARRENDATARIO y los DEUDORES SOLIDARIOS, autorizan desde ahora al ARRENDADOR o a su representante, para visitar el inmueble arrendado en cualquier momento, previa comunicación y coordinación, cuándo las circunstancias así lo ameriten, en caso de venta, de arreglos o reparaciones locativas o simples inspecciones.') 
            ##SIN ADMINISTRADOR
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and   (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo30='CLÁUSULA VIGÉSIMA SEPTIMA. INSPECCIÓN DEL INMUEBLE: '
            p42.add_run(titulo30).bold = True
            p42.add_run('EL ARRENDATARIO y los DEUDORES SOLIDARIOS, autorizan desde ahora al ARRENDADOR o a su representante, para visitar el inmueble arrendado en cualquier momento, previa comunicación y coordinación, cuándo las circunstancias así lo ameriten, en caso de venta, de arreglos o reparaciones locativas o simples inspecciones.') 
            
        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO")and   (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo30='CLÁUSULA VIGÉSIMA SEPTIMA. INSPECCIÓN DEL INMUEBLE: '
            p42.add_run(titulo30).bold = True
            p42.add_run('EL ARRENDATARIO, autoriza desde ahora al ARRENDADOR o a su representante, para visitar el inmueble arrendado en cualquier momento, previa comunicación y coordinación, cuándo las circunstancias así lo ameriten, en caso de venta, de arreglos o reparaciones locativas o simples inspecciones.') 
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="No" or Rta=="no"or Rta=="NO")and   (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo30='CLÁUSULA VIGÉSIMA OCTAVA. INSPECCIÓN DEL INMUEBLE: '
            p42.add_run(titulo30).bold = True
            p42.add_run('EL ARRENDATARIO, autoriza desde ahora al ARRENDADOR o a su representante, para visitar el inmueble arrendado en cualquier momento, previa comunicación y coordinación, cuándo las circunstancias así lo ameriten, en caso de venta, de arreglos o reparaciones locativas o simples inspecciones.') 
        
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="Si" or Rta=="si"or Rta=="SI")and   (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo30='CLÁUSULA VIGÉSIMA OCTAVA. INSPECCIÓN DEL INMUEBLE: '
            p42.add_run(titulo30).bold = True
            p42.add_run('EL ARRENDATARIO y los DEUDORES SOLIDARIOS, autorizan desde ahora al ARRENDADOR o a su representante, para visitar el inmueble arrendado en cualquier momento, previa comunicación y coordinación, cuándo las circunstancias así lo ameriten, en caso de venta, de arreglos o reparaciones locativas o simples inspecciones.') 


        #______________________________________

        p44 = documentoObjeto.add_paragraph("")
        p44.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if  (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo31='CLÁUSULA TTRIGÉSIMA. FALLECIMIENTO DEL ARRENDATARIO: '
            p44.add_run(titulo31).bold = True
            p44.add_run('En caso de muerte de EL ARRENDATARIO y/o DEUDORES SOLIDARIOS, los restantes sobrevivientes están en la obligación de dar aviso aL ARRENDADOR por carta certificada, anexando la correspondiente partida de defunción. El ARRENDADOR podrá cumplir con la norma del artículo 1434 del C.C. respecto de uno de los herederos sin necesidad de notificar o demandar a los demás.') 
            
        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo31='CLÁUSULA TTRIGÉSIMA. FALLECIMIENTO DEL ARRENDATARIO: '
            p44.add_run(titulo31).bold = True
            p44.add_run('En caso de muerte de EL ARRENDATARIO , los restantes sobrevivientes están en la obligación de dar aviso aL ARRENDADOR por carta certificada, anexando la correspondiente partida de defunción. El ARRENDADOR podrá cumplir con la norma del artículo 1434 del C.C. respecto de uno de los herederos sin necesidad de notificar o demandar a los demás.') 
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo31='CLÁUSULA TRIGÉSIMA PRIMERA. FALLECIMIENTO DEL ARRENDATARIO: '
            p44.add_run(titulo31).bold = True
            p44.add_run('En caso de muerte de EL ARRENDATARIO , los restantes sobrevivientes están en la obligación de dar aviso aL ARRENDADOR por carta certificada, anexando la correspondiente partida de defunción. El ARRENDADOR podrá cumplir con la norma del artículo 1434 del C.C. respecto de uno de los herederos sin necesidad de notificar o demandar a los demás.') 
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo31='CLÁUSULA TRIGÉSIMA PRIMERA. FALLECIMIENTO DEL ARRENDATARIO: '
            p44.add_run(titulo31).bold = True
            p44.add_run('En caso de muerte de EL ARRENDATARIO y/o DEUDORES SOLIDARIOS, los restantes sobrevivientes están en la obligación de dar aviso aL ARRENDADOR por carta certificada, anexando la correspondiente partida de defunción. El ARRENDADOR podrá cumplir con la norma del artículo 1434 del C.C. respecto de uno de los herederos sin necesidad de notificar o demandar a los demás.') 
            ##SIN ADMINISTRADOR
        if  (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo31='CLÁUSULA VIGÉSIMA OCTAVA. FALLECIMIENTO DEL ARRENDATARIO: '
            p44.add_run(titulo31).bold = True
            p44.add_run('En caso de muerte de EL ARRENDATARIO y/o DEUDORES SOLIDARIOS, los restantes sobrevivientes están en la obligación de dar aviso aL ARRENDADOR por carta certificada, anexando la correspondiente partida de defunción. El ARRENDADOR podrá cumplir con la norma del artículo 1434 del C.C. respecto de uno de los herederos sin necesidad de notificar o demandar a los demás.') 
            
        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo31='CLÁUSULA VIGÉSIMA OCTAVA. FALLECIMIENTO DEL ARRENDATARIO: '
            p44.add_run(titulo31).bold = True
            p44.add_run('En caso de muerte de EL ARRENDATARIO , los restantes sobrevivientes están en la obligación de dar aviso aL ARRENDADOR por carta certificada, anexando la correspondiente partida de defunción. El ARRENDADOR podrá cumplir con la norma del artículo 1434 del C.C. respecto de uno de los herederos sin necesidad de notificar o demandar a los demás.') 
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo31='CLÁUSULA VIGÉSIMA NOVENA. FALLECIMIENTO DEL ARRENDATARIO: '
            p44.add_run(titulo31).bold = True
            p44.add_run('En caso de muerte de EL ARRENDATARIO , los restantes sobrevivientes están en la obligación de dar aviso aL ARRENDADOR por carta certificada, anexando la correspondiente partida de defunción. El ARRENDADOR podrá cumplir con la norma del artículo 1434 del C.C. respecto de uno de los herederos sin necesidad de notificar o demandar a los demás.') 
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo31='CLÁUSULA VIGÉSIMA NOVENA . FALLECIMIENTO DEL ARRENDATARIO: '
            p44.add_run(titulo31).bold = True
            p44.add_run('En caso de muerte de EL ARRENDATARIO y/o DEUDORES SOLIDARIOS, los restantes sobrevivientes están en la obligación de dar aviso aL ARRENDADOR por carta certificada, anexando la correspondiente partida de defunción. El ARRENDADOR podrá cumplir con la norma del artículo 1434 del C.C. respecto de uno de los herederos sin necesidad de notificar o demandar a los demás.') 


        #____________________________________________##COTINUAR......


        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            p46 = documentoObjeto.add_paragraph("")
            p46.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            titulo32='CLÁUSULA TRIGÉSIMA PRIMERA. DEUDORES SOLIDARIOS: '
            p46.add_run(titulo32).bold = True
            if Numero_Deudores==1:
                p46.add_run('Para garantizar a ARRENDADOR el cumplimiento de las obligaciones a cargo de EL ARRENDATARIO, tiene como deudor solidario a '+Deudor_nombre[0]+", mayor de edad, identificado con "+Deudor_di[0]+" No "+Deudor_idnumm[0]+" de "+Deudor_ciudad[0]+", quien declará que se obliga de manera solidaria con EL ARRENDATARIO y frente a EL ARRENDADOR durante el término de duración de este Contrato y hasta que el Inmueble sea devuelto al ARRENDADOR a su entera satisfacción. Manifiesta que se declara deudor solidario de EL ARRENDATARIO en forma solidaria e indivisible indicado al inicio de este documento, de todas las cargas y obligaciones contenidas en el presente contrato, tanto durante el término inicialmente pactado como durante sus prórrogas o renovaciones expresas o tácitas y hasta la restitución material del inmueble a EL ARRENDADOR. Responderá por el cumplimiento y pago por concepto de arrendamientos, cuotas de administración, servicios públicos, indemnizaciones, daños en el inmueble, cláusulas penales, gastos de cobranza, costas procesales y cualquier otra derivada del contrato, las cuales podrán ser exigidas por EL ARRENDADOR al obligado por la vía ejecutiva, sin necesidad de requerimientos privados o judiciales y sin que por esta calidad asuman el caracter de arrendatarios. En caso de abandono del inmueble EL ARRENDATARIO y/o deudor solidario podrán hacer entrega válidamente del inmueble a EL ARRENDADOR o a quien éste señale en la forma y términos señalados en la cláusula vigésima de este documento. Para este exclusivo efecto EL ARRENDATARIO otorga poder amplio y suficiente al deudor solidario en este mismo acto al suscribir el presente contrato.")
            elif Numero_Deudores==2:
                p46.add_run('Para garantizar a ARRENDADOR el cumplimiento de las obligaciones a cargo de EL ARRENDATARIO, tiene como deudores solidarios a '+Deudor_nombre[0]+", mayor de edad, identificado con "+Deudor_di[0]+" No "+Deudor_idnumm[0]+" de "+Deudor_ciudad[0]+" y "+Deudor_nombre[2]+" mayor de edad, identificado(a) con "+Deudor_di[2]+" No. "+Deudor_idnumm[2]+" de "+Deudor_ciudad[2]+" ,quienes declaran que se obligan de manera solidaria con EL ARRENDATARIO y frente a EL ARRENDADOR durante el término de duración de este Contrato y hasta que el Inmueble sea devuelto al ARRENDADOR a su entera satisfacción. Manifiestan que se declaran deudores solidarios de EL ARRENDATARIO en forma solidaria e indivisible indicado al inicio de este documento, de todas las cargas y obligaciones contenidas en el presente contrato, tanto durante el término inicialmente pactado como durante sus prórrogas o renovaciones expresas o tácitas y hasta la restitución material del inmueble a EL ARRENDADOR. Responderán por el cumplimiento y pago por concepto de arrendamientos, cuotas de administración, servicios públicos, indemnizaciones, daños en el inmueble, cláusulas penales, gastos de cobranza, costas procesales y cualquier otra derivada del contrato, las cuales podrán ser exigidas por EL ARRENDADOR a cualquiera de los obligados por la vía ejecutiva, sin necesidad de requerimientos privados o judiciales y sin que por esta calidad asuman el caracter de arrendatarios. En caso de abandono del inmueble EL ARRENDATARIO y/o deudores solidarios podrán hacer entrega válidamente del inmueble a EL ARRENDADOR o a quien éste señale en la forma y términos señalados en la cláusula vigésima de este documento. Para este exclusivo efecto EL ARRENDATARIO otorga poder amplio y suficiente a los deudores solidarios en este mismo acto al suscribir el presente contrato.")
            elif Numero_Deudores==3:
                p46.add_run('Para garantizar a ARRENDADOR el cumplimiento de las obligaciones a cargo de EL ARRENDATARIO, tiene como deudores solidarios a '+Deudor_nombre[0]+", mayor de edad, identificado con "+Deudor_di[0]+" No "+Deudor_idnumm[0]+" de "+Deudor_ciudad[0]+" , "+Deudor_nombre[2]+" mayor de edad, identificado(a) con "+Deudor_di[2]+" No. "+Deudor_idnumm[2]+" de "+Deudor_ciudad[2]+" y "+Deudor_nombre[4]+" mayor de edad, identificado(a) con "+Deudor_di[4]+" No. "+Deudor_idnumm[4]+" de "+Deudor_ciudad[4]+" ,quienes declaran que se obligan de manera solidaria con EL ARRENDATARIO y frente a EL ARRENDADOR durante el término de duración de este Contrato y hasta que el Inmueble sea devuelto al ARRENDADOR a su entera satisfacción. Manifiestan que se declaran deudores solidarios de EL ARRENDATARIO en forma solidaria e indivisible indicado al inicio de este documento, de todas las cargas y obligaciones contenidas en el presente contrato, tanto durante el término inicialmente pactado como durante sus prórrogas o renovaciones expresas o tácitas y hasta la restitución material del inmueble a EL ARRENDADOR. Responderán por el cumplimiento y pago por concepto de arrendamientos, cuotas de administración, servicios públicos, indemnizaciones, daños en el inmueble, cláusulas penales, gastos de cobranza, costas procesales y cualquier otra derivada del contrato, las cuales podrán ser exigidas por EL ARRENDADOR a cualquiera de los obligados por la vía ejecutiva, sin necesidad de requerimientos privados o judiciales y sin que por esta calidad asuman el caracter de arrendatarios. En caso de abandono del inmueble EL ARRENDATARIO y/o deudores solidarios podrán hacer entrega válidamente del inmueble a EL ARRENDADOR o a quien éste señale en la forma y términos señalados en la cláusula vigésima de este documento. Para este exclusivo efecto EL ARRENDATARIO otorga poder amplio y suficiente a los deudores solidarios en este mismo acto al suscribir el presente contrato.")

        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            p46 = documentoObjeto.add_paragraph("")
            p46.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
            titulo32='CLÁUSULA TRIGÉSIMA SEGUNDA. DEUDORES SOLIDARIOS: '
            p46.add_run(titulo32).bold = True
        
            if Numero_Deudores==1:
                p46.add_run('Para garantizar a ARRENDADOR el cumplimiento de las obligaciones a cargo de EL ARRENDATARIO, tiene como deudor solidario a '+Deudor_nombre[0]+", mayor de edad, identificado con "+Deudor_di[0]+" No "+Deudor_idnumm[0]+" de "+Deudor_ciudad[0]+", quien declará que se obliga de manera solidaria con EL ARRENDATARIO y frente a EL ARRENDADOR durante el término de duración de este Contrato y hasta que el Inmueble sea devuelto al ARRENDADOR a su entera satisfacción. Manifiesta que se declara deudor solidario de EL ARRENDATARIO en forma solidaria e indivisible indicado al inicio de este documento, de todas las cargas y obligaciones contenidas en el presente contrato, tanto durante el término inicialmente pactado como durante sus prórrogas o renovaciones expresas o tácitas y hasta la restitución material del inmueble a EL ARRENDADOR. Responderá por el cumplimiento y pago por concepto de arrendamientos, cuotas de administración, servicios públicos, indemnizaciones, daños en el inmueble, cláusulas penales, gastos de cobranza, costas procesales y cualquier otra derivada del contrato, las cuales podrán ser exigidas por EL ARRENDADOR al obligado por la vía ejecutiva, sin necesidad de requerimientos privados o judiciales y sin que por esta calidad asuman el caracter de arrendatarios. En caso de abandono del inmueble EL ARRENDATARIO y/o deudor solidario podrán hacer entrega válidamente del inmueble a EL ARRENDADOR o a quien éste señale en la forma y términos señalados en la cláusula vigésima de este documento. Para este exclusivo efecto EL ARRENDATARIO otorga poder amplio y suficiente al deudor solidario en este mismo acto al suscribir el presente contrato.")
            elif Numero_Deudores==2:
                p46.add_run('Para garantizar a ARRENDADOR el cumplimiento de las obligaciones a cargo de EL ARRENDATARIO, tiene como deudores solidarios a '+Deudor_nombre[0]+", mayor de edad, identificado con "+Deudor_di[0]+" No "+Deudor_idnumm[0]+" de "+Deudor_ciudad[0]+" y "+Deudor_nombre[2]+" mayor de edad, identificado(a) con "+Deudor_di[2]+" No. "+Deudor_idnumm[2]+" de "+Deudor_ciudad[2]+" ,quienes declaran que se obligan de manera solidaria con EL ARRENDATARIO y frente a EL ARRENDADOR durante el término de duración de este Contrato y hasta que el Inmueble sea devuelto al ARRENDADOR a su entera satisfacción. Manifiestan que se declaran deudores solidarios de EL ARRENDATARIO en forma solidaria e indivisible indicado al inicio de este documento, de todas las cargas y obligaciones contenidas en el presente contrato, tanto durante el término inicialmente pactado como durante sus prórrogas o renovaciones expresas o tácitas y hasta la restitución material del inmueble a EL ARRENDADOR. Responderán por el cumplimiento y pago por concepto de arrendamientos, cuotas de administración, servicios públicos, indemnizaciones, daños en el inmueble, cláusulas penales, gastos de cobranza, costas procesales y cualquier otra derivada del contrato, las cuales podrán ser exigidas por EL ARRENDADOR a cualquiera de los obligados por la vía ejecutiva, sin necesidad de requerimientos privados o judiciales y sin que por esta calidad asuman el caracter de arrendatarios. En caso de abandono del inmueble EL ARRENDATARIO y/o deudores solidarios podrán hacer entrega válidamente del inmueble a EL ARRENDADOR o a quien éste señale en la forma y términos señalados en la cláusula vigésima de este documento. Para este exclusivo efecto EL ARRENDATARIO otorga poder amplio y suficiente a los deudores solidarios en este mismo acto al suscribir el presente contrato.")
            elif Numero_Deudores==3:
                p46.add_run('Para garantizar a ARRENDADOR el cumplimiento de las obligaciones a cargo de EL ARRENDATARIO, tiene como deudores solidarios a '+Deudor_nombre[0]+", mayor de edad, identificado con "+Deudor_di[0]+" No "+Deudor_idnumm[0]+" de "+Deudor_ciudad[0]+" , "+Deudor_nombre[2]+" mayor de edad, identificado(a) con "+Deudor_di[2]+" No. "+Deudor_idnumm[2]+" de "+Deudor_ciudad[2]+" y "+Deudor_nombre[4]+" mayor de edad, identificado(a) con "+Deudor_di[4]+" No. "+Deudor_idnumm[4]+" de "+Deudor_ciudad[4]+" ,quienes declaran que se obligan de manera solidaria con EL ARRENDATARIO y frente a EL ARRENDADOR durante el término de duración de este Contrato y hasta que el Inmueble sea devuelto al ARRENDADOR a su entera satisfacción. Manifiestan que se declaran deudores solidarios de EL ARRENDATARIO en forma solidaria e indivisible indicado al inicio de este documento, de todas las cargas y obligaciones contenidas en el presente contrato, tanto durante el término inicialmente pactado como durante sus prórrogas o renovaciones expresas o tácitas y hasta la restitución material del inmueble a EL ARRENDADOR. Responderán por el cumplimiento y pago por concepto de arrendamientos, cuotas de administración, servicios públicos, indemnizaciones, daños en el inmueble, cláusulas penales, gastos de cobranza, costas procesales y cualquier otra derivada del contrato, las cuales podrán ser exigidas por EL ARRENDADOR a cualquiera de los obligados por la vía ejecutiva, sin necesidad de requerimientos privados o judiciales y sin que por esta calidad asuman el caracter de arrendatarios. En caso de abandono del inmueble EL ARRENDATARIO y/o deudores solidarios podrán hacer entrega válidamente del inmueble a EL ARRENDADOR o a quien éste señale en la forma y términos señalados en la cláusula vigésima de este documento. Para este exclusivo efecto EL ARRENDATARIO otorga poder amplio y suficiente a los deudores solidarios en este mismo acto al suscribir el presente contrato.")
            ##sin admisnistrador
        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            p46 = documentoObjeto.add_paragraph("")
            p46.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            titulo32='CLÁUSULA VIGÉSIMA NOVENA . DEUDORES SOLIDARIOS: '
            p46.add_run(titulo32).bold = True
            if Numero_Deudores==1:
                    p46.add_run('Para garantizar a ARRENDADOR el cumplimiento de las obligaciones a cargo de EL ARRENDATARIO, tiene como deudor solidario a '+Deudor_nombre[0]+", mayor de edad, identificado con "+Deudor_di[0]+" No "+Deudor_idnumm[0]+" de "+Deudor_ciudad[0]+", quien declará que se obliga de manera solidaria con EL ARRENDATARIO y frente a EL ARRENDADOR durante el término de duración de este Contrato y hasta que el Inmueble sea devuelto al ARRENDADOR a su entera satisfacción. Manifiesta que se declara deudor solidario de EL ARRENDATARIO en forma solidaria e indivisible indicado al inicio de este documento, de todas las cargas y obligaciones contenidas en el presente contrato, tanto durante el término inicialmente pactado como durante sus prórrogas o renovaciones expresas o tácitas y hasta la restitución material del inmueble a EL ARRENDADOR. Responderá por el cumplimiento y pago por concepto de arrendamientos, cuotas de administración, servicios públicos, indemnizaciones, daños en el inmueble, cláusulas penales, gastos de cobranza, costas procesales y cualquier otra derivada del contrato, las cuales podrán ser exigidas por EL ARRENDADOR al obligado por la vía ejecutiva, sin necesidad de requerimientos privados o judiciales y sin que por esta calidad asuman el caracter de arrendatarios. En caso de abandono del inmueble EL ARRENDATARIO y/o deudor solidario podrán hacer entrega válidamente del inmueble a EL ARRENDADOR o a quien éste señale en la forma y términos señalados en la cláusula vigésima de este documento. Para este exclusivo efecto EL ARRENDATARIO otorga poder amplio y suficiente al deudor solidario en este mismo acto al suscribir el presente contrato.")
            elif Numero_Deudores==2:
                p46.add_run('Para garantizar a ARRENDADOR el cumplimiento de las obligaciones a cargo de EL ARRENDATARIO, tiene como deudores solidarios a '+Deudor_nombre[0]+", mayor de edad, identificado con "+Deudor_di[0]+" No "+Deudor_idnumm[0]+" de "+Deudor_ciudad[0]+" y "+Deudor_nombre[2]+" mayor de edad, identificado(a) con "+Deudor_di[2]+" No. "+Deudor_idnumm[2]+" de "+Deudor_ciudad[2]+" ,quienes declaran que se obligan de manera solidaria con EL ARRENDATARIO y frente a EL ARRENDADOR durante el término de duración de este Contrato y hasta que el Inmueble sea devuelto al ARRENDADOR a su entera satisfacción. Manifiestan que se declaran deudores solidarios de EL ARRENDATARIO en forma solidaria e indivisible indicado al inicio de este documento, de todas las cargas y obligaciones contenidas en el presente contrato, tanto durante el término inicialmente pactado como durante sus prórrogas o renovaciones expresas o tácitas y hasta la restitución material del inmueble a EL ARRENDADOR. Responderán por el cumplimiento y pago por concepto de arrendamientos, cuotas de administración, servicios públicos, indemnizaciones, daños en el inmueble, cláusulas penales, gastos de cobranza, costas procesales y cualquier otra derivada del contrato, las cuales podrán ser exigidas por EL ARRENDADOR a cualquiera de los obligados por la vía ejecutiva, sin necesidad de requerimientos privados o judiciales y sin que por esta calidad asuman el caracter de arrendatarios. En caso de abandono del inmueble EL ARRENDATARIO y/o deudores solidarios podrán hacer entrega válidamente del inmueble a EL ARRENDADOR o a quien éste señale en la forma y términos señalados en la cláusula vigésima de este documento. Para este exclusivo efecto EL ARRENDATARIO otorga poder amplio y suficiente a los deudores solidarios en este mismo acto al suscribir el presente contrato.")
            elif Numero_Deudores==3:
                p46.add_run('Para garantizar a ARRENDADOR el cumplimiento de las obligaciones a cargo de EL ARRENDATARIO, tiene como deudores solidarios a '+Deudor_nombre[0]+", mayor de edad, identificado con "+Deudor_di[0]+" No "+Deudor_idnumm[0]+" de "+Deudor_ciudad[0]+" , "+Deudor_nombre[2]+" mayor de edad, identificado(a) con "+Deudor_di[2]+" No. "+Deudor_idnumm[2]+" de "+Deudor_ciudad[2]+" y "+Deudor_nombre[4]+" mayor de edad, identificado(a) con "+Deudor_di[4]+" No. "+Deudor_idnumm[4]+" de "+Deudor_ciudad[4]+" ,quienes declaran que se obligan de manera solidaria con EL ARRENDATARIO y frente a EL ARRENDADOR durante el término de duración de este Contrato y hasta que el Inmueble sea devuelto al ARRENDADOR a su entera satisfacción. Manifiestan que se declaran deudores solidarios de EL ARRENDATARIO en forma solidaria e indivisible indicado al inicio de este documento, de todas las cargas y obligaciones contenidas en el presente contrato, tanto durante el término inicialmente pactado como durante sus prórrogas o renovaciones expresas o tácitas y hasta la restitución material del inmueble a EL ARRENDADOR. Responderán por el cumplimiento y pago por concepto de arrendamientos, cuotas de administración, servicios públicos, indemnizaciones, daños en el inmueble, cláusulas penales, gastos de cobranza, costas procesales y cualquier otra derivada del contrato, las cuales podrán ser exigidas por EL ARRENDADOR a cualquiera de los obligados por la vía ejecutiva, sin necesidad de requerimientos privados o judiciales y sin que por esta calidad asuman el caracter de arrendatarios. En caso de abandono del inmueble EL ARRENDATARIO y/o deudores solidarios podrán hacer entrega válidamente del inmueble a EL ARRENDADOR o a quien éste señale en la forma y términos señalados en la cláusula vigésima de este documento. Para este exclusivo efecto EL ARRENDATARIO otorga poder amplio y suficiente a los deudores solidarios en este mismo acto al suscribir el presente contrato.")
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            p46 = documentoObjeto.add_paragraph("")
            p46.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            titulo32='CLÁUSULA TRIGÉSIMA . DEUDORES SOLIDARIOS: '
            p46.add_run(titulo32).bold = True
            
            if Numero_Deudores==1:
                p46.add_run('Para garantizar a ARRENDADOR el cumplimiento de las obligaciones a cargo de EL ARRENDATARIO, tiene como deudor solidario a '+Deudor_nombre[0]+", mayor de edad, identificado con "+Deudor_di[0]+" No "+Deudor_idnumm[0]+" de "+Deudor_ciudad[0]+", quien declará que se obliga de manera solidaria con EL ARRENDATARIO y frente a EL ARRENDADOR durante el término de duración de este Contrato y hasta que el Inmueble sea devuelto al ARRENDADOR a su entera satisfacción. Manifiesta que se declara deudor solidario de EL ARRENDATARIO en forma solidaria e indivisible indicado al inicio de este documento, de todas las cargas y obligaciones contenidas en el presente contrato, tanto durante el término inicialmente pactado como durante sus prórrogas o renovaciones expresas o tácitas y hasta la restitución material del inmueble a EL ARRENDADOR. Responderá por el cumplimiento y pago por concepto de arrendamientos, cuotas de administración, servicios públicos, indemnizaciones, daños en el inmueble, cláusulas penales, gastos de cobranza, costas procesales y cualquier otra derivada del contrato, las cuales podrán ser exigidas por EL ARRENDADOR al obligado por la vía ejecutiva, sin necesidad de requerimientos privados o judiciales y sin que por esta calidad asuman el caracter de arrendatarios. En caso de abandono del inmueble EL ARRENDATARIO y/o deudor solidario podrán hacer entrega válidamente del inmueble a EL ARRENDADOR o a quien éste señale en la forma y términos señalados en la cláusula vigésima de este documento. Para este exclusivo efecto EL ARRENDATARIO otorga poder amplio y suficiente al deudor solidario en este mismo acto al suscribir el presente contrato.")
            elif Numero_Deudores==2:
                p46.add_run('Para garantizar a ARRENDADOR el cumplimiento de las obligaciones a cargo de EL ARRENDATARIO, tiene como deudores solidarios a '+Deudor_nombre[0]+", mayor de edad, identificado con "+Deudor_di[0]+" No "+Deudor_idnumm[0]+" de "+Deudor_ciudad[0]+" y "+Deudor_nombre[2]+" mayor de edad, identificado(a) con "+Deudor_di[2]+" No. "+Deudor_idnumm[2]+" de "+Deudor_ciudad[2]+" ,quienes declaran que se obligan de manera solidaria con EL ARRENDATARIO y frente a EL ARRENDADOR durante el término de duración de este Contrato y hasta que el Inmueble sea devuelto al ARRENDADOR a su entera satisfacción. Manifiestan que se declaran deudores solidarios de EL ARRENDATARIO en forma solidaria e indivisible indicado al inicio de este documento, de todas las cargas y obligaciones contenidas en el presente contrato, tanto durante el término inicialmente pactado como durante sus prórrogas o renovaciones expresas o tácitas y hasta la restitución material del inmueble a EL ARRENDADOR. Responderán por el cumplimiento y pago por concepto de arrendamientos, cuotas de administración, servicios públicos, indemnizaciones, daños en el inmueble, cláusulas penales, gastos de cobranza, costas procesales y cualquier otra derivada del contrato, las cuales podrán ser exigidas por EL ARRENDADOR a cualquiera de los obligados por la vía ejecutiva, sin necesidad de requerimientos privados o judiciales y sin que por esta calidad asuman el caracter de arrendatarios. En caso de abandono del inmueble EL ARRENDATARIO y/o deudores solidarios podrán hacer entrega válidamente del inmueble a EL ARRENDADOR o a quien éste señale en la forma y términos señalados en la cláusula vigésima de este documento. Para este exclusivo efecto EL ARRENDATARIO otorga poder amplio y suficiente a los deudores solidarios en este mismo acto al suscribir el presente contrato. ")
            elif Numero_Deudores==3:
                p46.add_run('Para garantizar a ARRENDADOR el cumplimiento de las obligaciones a cargo de EL ARRENDATARIO, tiene como deudores solidarios a '+Deudor_nombre[0]+", mayor de edad, identificado con "+Deudor_di[0]+" No "+Deudor_idnumm[0]+" de "+Deudor_ciudad[0]+" , "+Deudor_nombre[2]+" mayor de edad, identificado(a) con "+Deudor_di[2]+" No. "+Deudor_idnumm[2]+" de "+Deudor_ciudad[2]+" y "+Deudor_nombre[4]+" mayor de edad, identificado(a) con "+Deudor_di[4]+" No. "+Deudor_idnumm[4]+" de "+Deudor_ciudad[4]+" ,quienes declaran que se obligan de manera solidaria con EL ARRENDATARIO y frente a EL ARRENDADOR durante el término de duración de este Contrato y hasta que el Inmueble sea devuelto al ARRENDADOR a su entera satisfacción. Manifiestan que se declaran deudores solidarios de EL ARRENDATARIO en forma solidaria e indivisible indicado al inicio de este documento, de todas las cargas y obligaciones contenidas en el presente contrato, tanto durante el término inicialmente pactado como durante sus prórrogas o renovaciones expresas o tácitas y hasta la restitución material del inmueble a EL ARRENDADOR. Responderán por el cumplimiento y pago por concepto de arrendamientos, cuotas de administración, servicios públicos, indemnizaciones, daños en el inmueble, cláusulas penales, gastos de cobranza, costas procesales y cualquier otra derivada del contrato, las cuales podrán ser exigidas por EL ARRENDADOR a cualquiera de los obligados por la vía ejecutiva, sin necesidad de requerimientos privados o judiciales y sin que por esta calidad asuman el caracter de arrendatarios. En caso de abandono del inmueble EL ARRENDATARIO y/o deudores solidarios podrán hacer entrega válidamente del inmueble a EL ARRENDADOR o a quien éste señale en la forma y términos señalados en la cláusula vigésima de este documento. Para este exclusivo efecto EL ARRENDATARIO otorga poder amplio y suficiente a los deudores solidarios en este mismo acto al suscribir el presente contrato.")
            ##_____
        else:
            pass

        p47 = documentoObjeto.add_paragraph("")
        p47.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI") and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo33='CLÁUSULA TRIGÉSIMA SEGUNDA. RESTITUCIÓN DEL INMUEBLE: '
            p47.add_run(titulo33).bold = True
            p47.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p47.add_run('Terminado el presente contrato, el ARRENDATARIO o deudor solidario facultado, deberá entregar el precitado inmueble al ARRENDADOR en forma personal o a quien éste autorice para recibirlo, conforme al inventario inicia, en las mismas condiciones en que lo recibió del ARRENDADOR, salvo el deterioro natural causado por el uso legítimo, obligándose a presentar los últimos tres (3)  recibos de cada uno de los servicios públicos debidamente pagados, junto con el respectivo paz y salvo de la administración, con ocho (8) días de antelación a la fecha fijada para la restitución del inmueble en el domicilio del ARRENDADOR. En relación con los servicios públicos pendientes por verificar, el ARRENDATARIO garantiza su pago mediante la provisión proporcional y equivalente al promedio del último consumo según la facturación respectiva. No será válida ni se entenderá como entrega formal y material del inmueble arrendado, la que se realice por medios diferentes a los estipulados en la ley o en el presente contrato. PARÁGRAFO 1. EL ARRENDATARIO se compromete a dejar al ARRENDADOR un valor en dinero en efectivo el cual será utilizado para cubrir las facturas de  servicios públicos pendientes por llegar a la fecha de entrega del predio; para tal fin se establecerá un valor promedio mensual ponderado de los tres últimos meses de cada servicio público cancelados, este valor presentará una re-liquidación 90  días después de ser efectiva la restitución del predio. PARÁGRAFO 2. Al suscribir el presente contrato EL ARRENDATARIO queda plenamente notificado de la visita de inspección que realizara EL ARRENDADOR para hacer verificación del inventario elaborado por las partes y el cual hace parte integral del contrato de arriendo en calidad de anexo 1. PARÁGRAFO 3. EL ARRENDATARIO reconoce y acepta que el predio fue entregado en perfecto estado de pintura, en tales condiciones deberá restituir el predio. PARÁGRAFO 4. No obstante lo anterior, EL ARRENDADOR podrá negarse a recibir el Inmueble, cuando a su juicio existan obligaciones pendientes a cargo de EL ARRENDATARIO que no hayan sido satisfechas en forma debida, caso en el cual se seguirá causando el canon de arrendamiento hasta que EL ARRENDATARIO cumpla con lo que le corresponde. PARÁGRAFO 5: La responsabilidad del ARRENDATARIO subsistirá aún después de restituido el Inmueble, mientras EL ARRENDADOR no haya entregado el paz y salvo correspondiente por escrito a EL ARRENDATARIO.')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo33='CLÁUSULA TRIGÉSIMA TERCERA. RESTITUCIÓN DEL INMUEBLE: '
            p47.add_run(titulo33).bold = True
            p47.add_run('Terminado el presente contrato, el ARRENDATARIO o deudor solidario facultado, deberá entregar el precitado inmueble al ARRENDADOR en forma personal o a quien éste autorice para recibirlo, conforme al inventario inicia, en las mismas condiciones en que lo recibió del ARRENDADOR, salvo el deterioro natural causado por el uso legítimo, obligándose a presentar los últimos tres (3)  recibos de cada uno de los servicios públicos debidamente pagados, junto con el respectivo paz y salvo de la administración, con ocho (8) días de antelación a la fecha fijada para la restitución del inmueble en el domicilio del ARRENDADOR. En relación con los servicios públicos pendientes por verificar, el ARRENDATARIO garantiza su pago mediante la provisión proporcional y equivalente al promedio del último consumo según la facturación respectiva. No será válida ni se entenderá como entrega formal y material del inmueble arrendado, la que se realice por medios diferentes a los estipulados en la ley o en el presente contrato. PARÁGRAFO 1. EL ARRENDATARIO se compromete a dejar al ARRENDADOR un valor en dinero en efectivo el cual será utilizado para cubrir las facturas de  servicios públicos pendientes por llegar a la fecha de entrega del predio; para tal fin se establecerá un valor promedio mensual ponderado de los tres últimos meses de cada servicio público cancelados, este valor presentará una re-liquidación 90  días después de ser efectiva la restitución del predio. PARÁGRAFO 2. Al suscribir el presente contrato EL ARRENDATARIO queda plenamente notificado de la visita de inspección que realizara EL ARRENDADOR para hacer verificación del inventario elaborado por las partes y el cual hace parte integral del contrato de arriendo en calidad de anexo 1. PARÁGRAFO 3. EL ARRENDATARIO reconoce y acepta que el predio fue entregado en perfecto estado de pintura, en tales condiciones deberá restituir el predio. PARÁGRAFO 4. No obstante lo anterior, EL ARRENDADOR podrá negarse a recibir el Inmueble, cuando a su juicio existan obligaciones pendientes a cargo de EL ARRENDATARIO que no hayan sido satisfechas en forma debida, caso en el cual se seguirá causando el canon de arrendamiento hasta que EL ARRENDATARIO cumpla con lo que le corresponde. PARÁGRAFO 5: La responsabilidad del ARRENDATARIO subsistirá aún después de restituido el Inmueble, mientras EL ARRENDADOR no haya entregado el paz y salvo correspondiente por escrito a EL ARRENDATARIO.')
        
        else:
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO" )and (Rta=="No" or Rta=="no"or Rta=="NO" )and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo32='CLÁUSULA TRIGÉSIMA PRIMERA.RESTITUCIÓN DEL INMUEBLE: '
                p47.add_run(titulo32).bold = True
                p47.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p47.add_run('Terminado el presente contrato, el ARRENDATARIO o deudor solidario facultado, deberá entregar el precitado inmueble al ARRENDADOR en forma personal o a quien éste autorice para recibirlo, conforme al inventario inicia, en las mismas condiciones en que lo recibió del ARRENDADOR, salvo el deterioro natural causado por el uso legítimo, obligándose a presentar los últimos tres (3)  recibos de cada uno de los servicios públicos debidamente pagados, junto con el respectivo paz y salvo de la administración, con ocho (8) días de antelación a la fecha fijada para la restitución del inmueble en el domicilio del ARRENDADOR. En relación con los servicios públicos pendientes por verificar, el ARRENDATARIO garantiza su pago mediante la provisión proporcional y equivalente al promedio del último consumo según la facturación respectiva. No será válida ni se entenderá como entrega formal y material del inmueble arrendado, la que se realice por medios diferentes a los estipulados en la ley o en el presente contrato. PARÁGRAFO 1. EL ARRENDATARIO se compromete a dejar al ARRENDADOR un valor en dinero en efectivo el cual será utilizado para cubrir las facturas de  servicios públicos pendientes por llegar a la fecha de entrega del predio; para tal fin se establecerá un valor promedio mensual ponderado de los tres últimos meses de cada servicio público cancelados, este valor presentará una re-liquidación 90  días después de ser efectiva la restitución del predio. PARÁGRAFO 2. Al suscribir el presente contrato EL ARRENDATARIO queda plenamente notificado de la visita de inspección que realizara EL ARRENDADOR para hacer verificación del inventario elaborado por las partes y el cual hace parte integral del contrato de arriendo en calidad de anexo 1. PARÁGRAFO 3. EL ARRENDATARIO reconoce y acepta que el predio fue entregado en perfecto estado de pintura, en tales condiciones deberá restituir el predio. PARÁGRAFO 4. No obstante lo anterior, EL ARRENDADOR podrá negarse a recibir el Inmueble, cuando a su juicio existan obligaciones pendientes a cargo de EL ARRENDATARIO que no hayan sido satisfechas en forma debida, caso en el cual se seguirá causando el canon de arrendamiento hasta que EL ARRENDATARIO cumpla con lo que le corresponde. PARÁGRAFO 5: La responsabilidad del ARRENDATARIO subsistirá aún después de restituido el Inmueble, mientras EL ARRENDADOR no haya entregado el paz y salvo correspondiente por escrito a EL ARRENDATARIO.')
            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="No" or Rta=="no"or Rta=="NO")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo32='CLÁUSULA TRIGÉSIMA SEGUNDA. RESTITUCIÓN DEL INMUEBLE: '
                p47.add_run(titulo32).bold = True
                p47.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p47.add_run('Terminado el presente contrato, el ARRENDATARIO o deudor solidario facultado, deberá entregar el precitado inmueble al ARRENDADOR en forma personal o a quien éste autorice para recibirlo, conforme al inventario inicia, en las mismas condiciones en que lo recibió del ARRENDADOR, salvo el deterioro natural causado por el uso legítimo, obligándose a presentar los últimos tres (3)  recibos de cada uno de los servicios públicos debidamente pagados, junto con el respectivo paz y salvo de la administración, con ocho (8) días de antelación a la fecha fijada para la restitución del inmueble en el domicilio del ARRENDADOR. En relación con los servicios públicos pendientes por verificar, el ARRENDATARIO garantiza su pago mediante la provisión proporcional y equivalente al promedio del último consumo según la facturación respectiva. No será válida ni se entenderá como entrega formal y material del inmueble arrendado, la que se realice por medios diferentes a los estipulados en la ley o en el presente contrato. PARÁGRAFO 1. EL ARRENDATARIO se compromete a dejar al ARRENDADOR un valor en dinero en efectivo el cual será utilizado para cubrir las facturas de  servicios públicos pendientes por llegar a la fecha de entrega del predio; para tal fin se establecerá un valor promedio mensual ponderado de los tres últimos meses de cada servicio público cancelados, este valor presentará una re-liquidación 90  días después de ser efectiva la restitución del predio. PARÁGRAFO 2. Al suscribir el presente contrato EL ARRENDATARIO queda plenamente notificado de la visita de inspección que realizara EL ARRENDADOR para hacer verificación del inventario elaborado por las partes y el cual hace parte integral del contrato de arriendo en calidad de anexo 1. PARÁGRAFO 3. EL ARRENDATARIO reconoce y acepta que el predio fue entregado en perfecto estado de pintura, en tales condiciones deberá restituir el predio. PARÁGRAFO 4. No obstante lo anterior, EL ARRENDADOR podrá negarse a recibir el Inmueble, cuando a su juicio existan obligaciones pendientes a cargo de EL ARRENDATARIO que no hayan sido satisfechas en forma debida, caso en el cual se seguirá causando el canon de arrendamiento hasta que EL ARRENDATARIO cumpla con lo que le corresponde. PARÁGRAFO 5: La responsabilidad del ARRENDATARIO subsistirá aún después de restituido el Inmueble, mientras EL ARRENDADOR no haya entregado el paz y salvo correspondiente por escrito a EL ARRENDATARIO.')
            ##sin administrador
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI") and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo33='CLÁUSULA TRIGÉSIMA . RESTITUCIÓN DEL INMUEBLE: '
            p47.add_run(titulo33).bold = True
            p47.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p47.add_run('Terminado el presente contrato, el ARRENDATARIO o deudor solidario facultado, deberá entregar el precitado inmueble al ARRENDADOR en forma personal o a quien éste autorice para recibirlo, conforme al inventario inicia, en las mismas condiciones en que lo recibió del ARRENDADOR, salvo el deterioro natural causado por el uso legítimo, obligándose a presentar los últimos tres (3)  recibos de cada uno de los servicios públicos debidamente pagados, junto con el respectivo paz y salvo de la administración, con ocho (8) días de antelación a la fecha fijada para la restitución del inmueble en el domicilio del ARRENDADOR. En relación con los servicios públicos pendientes por verificar, el ARRENDATARIO garantiza su pago mediante la provisión proporcional y equivalente al promedio del último consumo según la facturación respectiva. No será válida ni se entenderá como entrega formal y material del inmueble arrendado, la que se realice por medios diferentes a los estipulados en la ley o en el presente contrato. PARÁGRAFO 1. EL ARRENDATARIO se compromete a dejar al ARRENDADOR un valor en dinero en efectivo el cual será utilizado para cubrir las facturas de  servicios públicos pendientes por llegar a la fecha de entrega del predio; para tal fin se establecerá un valor promedio mensual ponderado de los tres últimos meses de cada servicio público cancelados, este valor presentará una re-liquidación 90  días después de ser efectiva la restitución del predio. PARÁGRAFO 2. Al suscribir el presente contrato EL ARRENDATARIO queda plenamente notificado de la visita de inspección que realizara EL ARRENDADOR para hacer verificación del inventario elaborado por las partes y el cual hace parte integral del contrato de arriendo en calidad de anexo 1. PARÁGRAFO 3. EL ARRENDATARIO reconoce y acepta que el predio fue entregado en perfecto estado de pintura, en tales condiciones deberá restituir el predio. PARÁGRAFO 4. No obstante lo anterior, EL ARRENDADOR podrá negarse a recibir el Inmueble, cuando a su juicio existan obligaciones pendientes a cargo de EL ARRENDATARIO que no hayan sido satisfechas en forma debida, caso en el cual se seguirá causando el canon de arrendamiento hasta que EL ARRENDATARIO cumpla con lo que le corresponde. PARÁGRAFO 5: La responsabilidad del ARRENDATARIO subsistirá aún después de restituido el Inmueble, mientras EL ARRENDADOR no haya entregado el paz y salvo correspondiente por escrito a EL ARRENDATARIO.')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo33='CLÁUSULA TRIGÉSIMA PRIMERA. RESTITUCIÓN DEL INMUEBLE: '
            p47.add_run(titulo33).bold = True
            p47.add_run('Terminado el presente contrato, el ARRENDATARIO o deudor solidario facultado, deberá entregar el precitado inmueble al ARRENDADOR en forma personal o a quien éste autorice para recibirlo, conforme al inventario inicia, en las mismas condiciones en que lo recibió del ARRENDADOR, salvo el deterioro natural causado por el uso legítimo, obligándose a presentar los últimos tres (3)  recibos de cada uno de los servicios públicos debidamente pagados, junto con el respectivo paz y salvo de la administración, con ocho (8) días de antelación a la fecha fijada para la restitución del inmueble en el domicilio del ARRENDADOR. En relación con los servicios públicos pendientes por verificar, el ARRENDATARIO garantiza su pago mediante la provisión proporcional y equivalente al promedio del último consumo según la facturación respectiva. No será válida ni se entenderá como entrega formal y material del inmueble arrendado, la que se realice por medios diferentes a los estipulados en la ley o en el presente contrato. PARÁGRAFO 1. EL ARRENDATARIO se compromete a dejar al ARRENDADOR un valor en dinero en efectivo el cual será utilizado para cubrir las facturas de  servicios públicos pendientes por llegar a la fecha de entrega del predio; para tal fin se establecerá un valor promedio mensual ponderado de los tres últimos meses de cada servicio público cancelados, este valor presentará una re-liquidación 90  días después de ser efectiva la restitución del predio. PARÁGRAFO 2. Al suscribir el presente contrato EL ARRENDATARIO queda plenamente notificado de la visita de inspección que realizara EL ARRENDADOR para hacer verificación del inventario elaborado por las partes y el cual hace parte integral del contrato de arriendo en calidad de anexo 1. PARÁGRAFO 3. EL ARRENDATARIO reconoce y acepta que el predio fue entregado en perfecto estado de pintura, en tales condiciones deberá restituir el predio. PARÁGRAFO 4. No obstante lo anterior, EL ARRENDADOR podrá negarse a recibir el Inmueble, cuando a su juicio existan obligaciones pendientes a cargo de EL ARRENDATARIO que no hayan sido satisfechas en forma debida, caso en el cual se seguirá causando el canon de arrendamiento hasta que EL ARRENDATARIO cumpla con lo que le corresponde. PARÁGRAFO 5: La responsabilidad del ARRENDATARIO subsistirá aún después de restituido el Inmueble, mientras EL ARRENDADOR no haya entregado el paz y salvo correspondiente por escrito a EL ARRENDATARIO.')
            
        else:
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO" )and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                titulo32='CLÁUSULA VIGESIMA NOVENA .RESTITUCIÓN DEL INMUEBLE: '
                p47.add_run(titulo32).bold = True
                p47.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p47.add_run('Terminado el presente contrato, el ARRENDATARIO o deudor solidario facultado, deberá entregar el precitado inmueble al ARRENDADOR en forma personal o a quien éste autorice para recibirlo, conforme al inventario inicia, en las mismas condiciones en que lo recibió del ARRENDADOR, salvo el deterioro natural causado por el uso legítimo, obligándose a presentar los últimos tres (3)  recibos de cada uno de los servicios públicos debidamente pagados, junto con el respectivo paz y salvo de la administración, con ocho (8) días de antelación a la fecha fijada para la restitución del inmueble en el domicilio del ARRENDADOR. En relación con los servicios públicos pendientes por verificar, el ARRENDATARIO garantiza su pago mediante la provisión proporcional y equivalente al promedio del último consumo según la facturación respectiva. No será válida ni se entenderá como entrega formal y material del inmueble arrendado, la que se realice por medios diferentes a los estipulados en la ley o en el presente contrato. PARÁGRAFO 1. EL ARRENDATARIO se compromete a dejar al ARRENDADOR un valor en dinero en efectivo el cual será utilizado para cubrir las facturas de  servicios públicos pendientes por llegar a la fecha de entrega del predio; para tal fin se establecerá un valor promedio mensual ponderado de los tres últimos meses de cada servicio público cancelados, este valor presentará una re-liquidación 90  días después de ser efectiva la restitución del predio. PARÁGRAFO 2. Al suscribir el presente contrato EL ARRENDATARIO queda plenamente notificado de la visita de inspección que realizara EL ARRENDADOR para hacer verificación del inventario elaborado por las partes y el cual hace parte integral del contrato de arriendo en calidad de anexo 1. PARÁGRAFO 3. EL ARRENDATARIO reconoce y acepta que el predio fue entregado en perfecto estado de pintura, en tales condiciones deberá restituir el predio. PARÁGRAFO 4. No obstante lo anterior, EL ARRENDADOR podrá negarse a recibir el Inmueble, cuando a su juicio existan obligaciones pendientes a cargo de EL ARRENDATARIO que no hayan sido satisfechas en forma debida, caso en el cual se seguirá causando el canon de arrendamiento hasta que EL ARRENDATARIO cumpla con lo que le corresponde. PARÁGRAFO 5: La responsabilidad del ARRENDATARIO subsistirá aún después de restituido el Inmueble, mientras EL ARRENDADOR no haya entregado el paz y salvo correspondiente por escrito a EL ARRENDATARIO.')
            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si" )and (Rta=="No" or Rta=="no"or Rta=="NO")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                titulo32='CLÁUSULA TRIGÉSIMA PRIMERA. RESTITUCIÓN DEL INMUEBLE: '
                p47.add_run(titulo32).bold = True
                p47.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p47.add_run('Terminado el presente contrato, el ARRENDATARIO o deudor solidario facultado, deberá entregar el precitado inmueble al ARRENDADOR en forma personal o a quien éste autorice para recibirlo, conforme al inventario inicia, en las mismas condiciones en que lo recibió del ARRENDADOR, salvo el deterioro natural causado por el uso legítimo, obligándose a presentar los últimos tres (3)  recibos de cada uno de los servicios públicos debidamente pagados, junto con el respectivo paz y salvo de la administración, con ocho (8) días de antelación a la fecha fijada para la restitución del inmueble en el domicilio del ARRENDADOR. En relación con los servicios públicos pendientes por verificar, el ARRENDATARIO garantiza su pago mediante la provisión proporcional y equivalente al promedio del último consumo según la facturación respectiva. No será válida ni se entenderá como entrega formal y material del inmueble arrendado, la que se realice por medios diferentes a los estipulados en la ley o en el presente contrato. PARÁGRAFO 1. EL ARRENDATARIO se compromete a dejar al ARRENDADOR un valor en dinero en efectivo el cual será utilizado para cubrir las facturas de  servicios públicos pendientes por llegar a la fecha de entrega del predio; para tal fin se establecerá un valor promedio mensual ponderado de los tres últimos meses de cada servicio público cancelados, este valor presentará una re-liquidación 90  días después de ser efectiva la restitución del predio. PARÁGRAFO 2. Al suscribir el presente contrato EL ARRENDATARIO queda plenamente notificado de la visita de inspección que realizara EL ARRENDADOR para hacer verificación del inventario elaborado por las partes y el cual hace parte integral del contrato de arriendo en calidad de anexo 1. PARÁGRAFO 3. EL ARRENDATARIO reconoce y acepta que el predio fue entregado en perfecto estado de pintura, en tales condiciones deberá restituir el predio. PARÁGRAFO 4. No obstante lo anterior, EL ARRENDADOR podrá negarse a recibir el Inmueble, cuando a su juicio existan obligaciones pendientes a cargo de EL ARRENDATARIO que no hayan sido satisfechas en forma debida, caso en el cual se seguirá causando el canon de arrendamiento hasta que EL ARRENDATARIO cumpla con lo que le corresponde. PARÁGRAFO 5: La responsabilidad del ARRENDATARIO subsistirá aún después de restituido el Inmueble, mientras EL ARRENDADOR no haya entregado el paz y salvo correspondiente por escrito a EL ARRENDATARIO.')


        ##_______________________________
        p48 = documentoObjeto.add_paragraph("")
        p48.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo34='CLÁUSULA TRIGÉSIMA TERCERA. COSTOS: '
            p48.add_run(titulo34).bold = True
            p48.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p48.add_run('Los gastos que se causen en este contrato equivalentes al veinte  (20%) más I.V.A. sobre el valor total del primer pago a realizar por el contrato, corresponden a la intermediación y gastos generados por la formalización del contrato, y serán asumidos exclusivamente por EL ARRENDATARIO, de acuerdo a la costumbre mercantil. PARÁGRAFO. El ARRENDADOR no asumirá ningún costo por las diligencias notariales que este contrato pueda originar. ')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo34='CLÁUSULA TRIGÉSIMA CUARTA. COSTOS: '
            p48.add_run(titulo34).bold = True
            p48.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p48.add_run('Los gastos que se causen en este contrato equivalentes al veinte  (20%) más I.V.A. sobre el valor total del primer pago a realizar por el contrato, corresponden a la intermediación y gastos generados por la formalización del contrato, y serán asumidos exclusivamente por EL ARRENDATARIO, de acuerdo a la costumbre mercantil. PARÁGRAFO. El ARRENDADOR no asumirá ningún costo por las diligencias notariales que este contrato pueda originar. ')
        
        else:
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO" )and (Rta=="No" or Rta=="no"or Rta=="NO" )and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo33='CLÁUSULA TRIGÉSIMA SEGUNDA. COSTOS: '
                p48.add_run(titulo33).bold = True
                p48.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p48.add_run('Los gastos que se causen en este contrato equivalentes al veinte  (20%) más I.V.A. sobre el valor total del primer pago a realizar por el contrato, corresponden a la intermediación y gastos generados por la formalización del contrato, y serán asumidos exclusivamente por EL ARRENDATARIO, de acuerdo a la costumbre mercantil. PARÁGRAFO. El ARRENDADOR no asumirá ningún costo por las diligencias notariales que este contrato pueda originar.')
            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="No" or Rta=="no"or Rta=="NO" )and  (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo33='CLÁUSULA TRIGÉSIMA TERCERA. COSTOS: '
                p48.add_run(titulo33).bold = True
                p48.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p48.add_run('Los gastos que se causen en este contrato equivalentes al veinte  (20%) más I.V.A. sobre el valor total del primer pago a realizar por el contrato, corresponden a la intermediación y gastos generados por la formalización del contrato, y serán asumidos exclusivamente por EL ARRENDATARIO, de acuerdo a la costumbre mercantil. PARÁGRAFO. El ARRENDADOR no asumirá ningún costo por las diligencias notariales que este contrato pueda originar.')
            #SIN ADMINISTRADOR
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo34='CLÁUSULA TRIGÉSIMA PRIMERA. COSTOS: '
            p48.add_run(titulo34).bold = True
            p48.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p48.add_run('Los gastos que se causen en este contrato equivalentes al veinte  (20%) más I.V.A. sobre el valor total del primer pago a realizar por el contrato, corresponden a la intermediación y gastos generados por la formalización del contrato, y serán asumidos exclusivamente por EL ARRENDATARIO, de acuerdo a la costumbre mercantil. PARÁGRAFO. El ARRENDADOR no asumirá ningún costo por las diligencias notariales que este contrato pueda originar. ')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo34='CLÁUSULA TRIGÉSIMA SEGUNDA. COSTOS: '
            p48.add_run(titulo34).bold = True
            p48.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p48.add_run('Los gastos que se causen en este contrato equivalentes al veinte  (20%) más I.V.A. sobre el valor total del primer pago a realizar por el contrato, corresponden a la intermediación y gastos generados por la formalización del contrato, y serán asumidos exclusivamente por EL ARRENDATARIO, de acuerdo a la costumbre mercantil. PARÁGRAFO. El ARRENDADOR no asumirá ningún costo por las diligencias notariales que este contrato pueda originar. ')
            
        else:
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO" )and (Rta=="No" or Rta=="no"or Rta=="NO" )and  (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                titulo33='CLÁUSULA TRIGÉSIMA . COSTOS: '
                p48.add_run(titulo33).bold = True
                p48.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p48.add_run('Los gastos que se causen en este contrato equivalentes al veinte  (20%) más I.V.A. sobre el valor total del primer pago a realizar por el contrato, corresponden a la intermediación y gastos generados por la formalización del contrato, y serán asumidos exclusivamente por EL ARRENDATARIO, de acuerdo a la costumbre mercantil. PARÁGRAFO. El ARRENDADOR no asumirá ningún costo por las diligencias notariales que este contrato pueda originar.')
            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                titulo33='CLÁUSULA TRIGÉSIMA SEGUNDA. COSTOS: '
                p48.add_run(titulo33).bold = True
                p48.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p48.add_run('Los gastos que se causen en este contrato equivalentes al veinte  (20%) más I.V.A. sobre el valor total del primer pago a realizar por el contrato, corresponden a la intermediación y gastos generados por la formalización del contrato, y serán asumidos exclusivamente por EL ARRENDATARIO, de acuerdo a la costumbre mercantil. PARÁGRAFO. El ARRENDADOR no asumirá ningún costo por las diligencias notariales que este contrato pueda originar.')


        #____________________
        p49 = documentoObjeto.add_paragraph("")
        p49.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            titulo35='CLÁUSULA TRIGÉSIMA CUARTA. COMPRA DEL INMUEBLE ARRENDADO: '
            p49.add_run(titulo35).bold = True
            p49.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p49.add_run('En caso que el ARRENDATARIO desee adquirir en compra el bien inmueble arrendado, se compromete a efectuar todas las gestiones por intermedio del ARRENDADOR, a quien desde ahora reconocen como intermediario de la venta directa. ')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo35='CLÁUSULA TRIGÉSIMA QUINTA. COMPRA DEL INMUEBLE ARRENDADO: '
            p49.add_run(titulo35).bold = True
            p49.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p49.add_run('En caso que el ARRENDATARIO desee adquirir en compra el bien inmueble arrendado, se compromete a efectuar todas las gestiones por intermedio del ARRENDADOR, a quien desde ahora reconocen como intermediario de la venta directa. ')

        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO" )and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            titulo34='CLÁUSULA TRIGÉSIMA TERCERA. COMPRA DEL INMUEBLE ARRENDADO: '
            p49.add_run(titulo34).bold = True
            p49.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p49.add_run('En caso que el ARRENDATARIO desee adquirir en compra el bien inmueble arrendado, se compromete a efectuar todas las gestiones por intermedio del ARRENDADOR, a quien desde ahora reconocen como intermediario de la venta directa. ')
        
        elif (Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI" )and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo34='CLÁUSULA TRIGÉSIMA CUARTA. COMPRA DEL INMUEBLE ARRENDADO: '
            p49.add_run(titulo34).bold = True
            p49.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p49.add_run('En caso que el ARRENDATARIO desee adquirir en compra el bien inmueble arrendado, se compromete a efectuar todas las gestiones por intermedio del ARRENDADOR, a quien desde ahora reconocen como intermediario de la venta directa.')
        #
        #Sin admiistrador
        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI") and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo35='CLÁUSULA TRIGÉSIMA SEGUNDA. COMPRA DEL INMUEBLE ARRENDADO: '
            p49.add_run(titulo35).bold = True
            p49.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p49.add_run('En caso que el ARRENDATARIO desee adquirir en compra el bien inmueble arrendado, se compromete a efectuar todas las gestiones por intermedio del ARRENDADOR, a quien desde ahora reconocen como intermediario de la venta directa. ')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
            titulo35='CLÁUSULA TRIGÉSIMA TERCERA. COMPRA DEL INMUEBLE ARRENDADO: '
            p49.add_run(titulo35).bold = True
            p49.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p49.add_run('En caso que el ARRENDATARIO desee adquirir en compra el bien inmueble arrendado, se compromete a efectuar todas las gestiones por intermedio del ARRENDADOR, a quien desde ahora reconocen como intermediario de la venta directa. ')
        
        else:
        #
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO" )and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                titulo34='CLÁUSULA TRIGÉSIMA PRIMERA. COMPRA DEL INMUEBLE ARRENDADO: '
                p49.add_run(titulo34).bold = True
                p49.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p49.add_run('En caso que el ARRENDATARIO desee adquirir en compra el bien inmueble arrendado, se compromete a efectuar todas las gestiones por intermedio del ARRENDADOR, a quien desde ahora reconocen como intermediario de la venta directa. ')
                
            elif (Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI" )and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                titulo34='CLÁUSULA TRIGÉSIMA TERCERA. COMPRA DEL INMUEBLE ARRENDADO: '
                p49.add_run(titulo34).bold = True
                p49.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p49.add_run('En caso que el ARRENDATARIO desee adquirir en compra el bien inmueble arrendado, se compromete a efectuar todas las gestiones por intermedio del ARRENDADOR, a quien desde ahora reconocen como intermediario de la venta directa.')

        #_______________________
        p50 = documentoObjeto.add_paragraph("")
        p50.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            titulo36='CLÁUSULA TRIGÉSIMA QUINTA. DEL INCUMPLIMIENTO: '
            p50.add_run(titulo36).bold = True
            p50.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p50.add_run('El incumplimiento del ARRENDATARIO a cualquiera de sus obligaciones legales o contractuales, faculta al ARRENDADOR para ejercer las siguientes acciones, simultáneamente o en el orden que él elija:\t\n 1. Declarar terminado este contrato y reclamar la devolución del inmueble judicial y/o extrajudicialmente.\t\n 2. Exigir y perseguir a través de cualquier medio, judicial o extrajudicialmente al ARRENDATARIO y/o DEUDORES SOLIDARIOS por el monto de los perjuicios resultantes del incumplimiento, así como de la multa por incumplimiento pactada en este Contrato.\n 3. COBRO EXTRAJUDICIAL. Si el incumplimiento de la obligación de cancelar oportunamente los cánones de arrendamiento, servicios públicos, cuotas de administración, o cualquier otra erogación derivada del contrato, diere lugar a alguna diligencia de cobro extrajudicial, EL ARRENDATARIO se obliga a pagar a la entidad encargada de tal gestión,los costos correspondientes.\t\n')
            p50.add_run('4.Realizar el cobro de los intereses moratorios por las sumas adeudadas.')
            p50.add_run('PARÁGRAFO. ').bold = True
            p50.add_run('No obstante lo anterior, las Partes en cualquier tiempo y de común acuerdo podrán dar por terminado el presente Contrato.\t')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
            titulo36='CLÁUSULA TRIGÉSIMA SEXTA. DEL INCUMPLIMIENTO: '
            p50.add_run(titulo36).bold = True
            p50.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p50.add_run('El incumplimiento del ARRENDATARIO a cualquiera de sus obligaciones legales o contractuales, faculta al ARRENDADOR para ejercer las siguientes acciones, simultáneamente o en el orden que él elija:\t\n 1. Declarar terminado este contrato y reclamar la devolución del inmueble judicial y/o extrajudicialmente. \t\n 2. Exigir y perseguir a través de cualquier medio, judicial o extrajudicialmente al ARRENDATARIO y/o DEUDORES SOLIDARIOS por el monto de los perjuicios resultantes del incumplimiento, así como de la multa por incumplimiento pactada en este Contrato.\n 3. COBRO EXTRAJUDICIAL. Si el incumplimiento de la obligación de cancelar oportunamente los cánones de arrendamiento, servicios públicos, cuotas de administración, o cualquier otra erogación derivada del contrato, diere lugar a alguna diligencia de cobro extrajudicial, EL ARRENDATARIO se obliga a pagar a la entidad encargada de tal gestión,los costos correspondientes.\t\n')
            p50.add_run('4.Realizar el cobro de los intereses moratorios por las sumas adeudadas.')
            p50.add_run('PARÁGRAFO. ').bold = True
            p50.add_run('No obstante lo anterior, las Partes en cualquier tiempo y de común acuerdo podrán dar por terminado el presente Contrato.\t')
            
        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            titulo35='CLÁUSULA TRIGÉSIMA CUARTA. DEL INCUMPLIMIENTO: '
            p50.add_run(titulo35).bold = True
            p50.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p50.add_run('El incumplimiento del ARRENDATARIO a cualquiera de sus obligaciones legales o contractuales, faculta al ARRENDADOR para ejercer las siguientes acciones, simultáneamente o en el orden que él elija:\t\n 1. Declarar terminado este contrato y reclamar la devolución del inmueble judicial y/o extrajudicialmente.\t \n 2. Exigir y perseguir a través de cualquier medio, judicial o extrajudicialmente al ARRENDATARIO  por el monto de los perjuicios resultantes del incumplimiento, así como de la multa por incumplimiento pactada en este Contrato.\n 3. COBRO EXTRAJUDICIAL. Si el incumplimiento de la obligación de cancelar oportunamente los cánones de arrendamiento, servicios públicos, cuotas de administración, o cualquier otra erogación derivada del contrato, diere lugar a alguna diligencia de cobro extrajudicial, EL ARRENDATARIO se obliga a pagar a la entidad encargada de tal gestión,los costos correspondientes.\t\n')
            p50.add_run('4.Realizar el cobro de los intereses moratorios por las sumas adeudadas.')
            p50.add_run('PARÁGRAFO. ').bold = True
            p50.add_run('No obstante lo anterior, las Partes en cualquier tiempo y de común acuerdo podrán dar por terminado el presente Contrato.\t')
        elif (Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI")and (Rta=="No" or Rta=="no"or Rta=="NO")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
            titulo35='CLÁUSULA TRIGÉSIMA QUINTA. DEL INCUMPLIMIENTO: '
            p50.add_run(titulo35).bold = True
            p50.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p50.add_run('El incumplimiento del ARRENDATARIO a cualquiera de sus obligaciones legales o contractuales, faculta al ARRENDADOR para ejercer las siguientes acciones, simultáneamente o en el orden que él elija:\t\n 1. Declarar terminado este contrato y reclamar la devolución del inmueble judicial y/o extrajudicialmente.\t \n 2. Exigir y perseguir a través de cualquier medio, judicial o extrajudicialmente al ARRENDATARIO por el monto de los perjuicios resultantes del incumplimiento, así como de la multa por incumplimiento pactada en este Contrato.\n 3. COBRO EXTRAJUDICIAL. Si el incumplimiento de la obligación de cancelar oportunamente los cánones de arrendamiento, servicios públicos, cuotas de administración, o cualquier otra erogación derivada del contrato, diere lugar a alguna diligencia de cobro extrajudicial, EL ARRENDATARIO se obliga a pagar a la entidad encargada de tal gestión,los costos correspondientes.\t\n')
            p50.add_run('4.Realizar el cobro de los intereses moratorios por las sumas adeudadas.')
            p50.add_run('PARÁGRAFO. ').bold = True
            p50.add_run('No obstante lo anterior, las Partes en cualquier tiempo y de común acuerdo podrán dar por terminado el presente Contrato.\t ')
        #Sin administrador
        elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
            titulo36='CLÁUSULA TRIGÉSIMA TERCERA. DEL INCUMPLIMIENTO: '
            p50.add_run(titulo36).bold = True
            p50.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p50.add_run('El incumplimiento del ARRENDATARIO a cualquiera de sus obligaciones legales o contractuales, faculta al ARRENDADOR para ejercer las siguientes acciones, simultáneamente o en el orden que él elija:\t\n 1. Declarar terminado este contrato y reclamar la devolución del inmueble judicial y/o extrajudicialmente.\t \n 2. Exigir y perseguir a través de cualquier medio, judicial o extrajudicialmente al ARRENDATARIO y/o DEUDORES SOLIDARIOS por el monto de los perjuicios resultantes del incumplimiento, así como de la multa por incumplimiento pactada en este Contrato.\n 3. COBRO EXTRAJUDICIAL. Si el incumplimiento de la obligación de cancelar oportunamente los cánones de arrendamiento, servicios públicos, cuotas de administración, o cualquier otra erogación derivada del contrato, diere lugar a alguna diligencia de cobro extrajudicial, EL ARRENDATARIO se obliga a pagar a la entidad encargada de tal gestión,los costos correspondientes.\t\n')
            p50.add_run('4.Realizar el cobro de los intereses moratorios por las sumas adeudadas.')
            p50.add_run('PARÁGRAFO. ').bold = True
            p50.add_run('No obstante lo anterior, las Partes en cualquier tiempo y de común acuerdo podrán dar por terminado el presente Contrato.\t')
            
        elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
            titulo36='CLÁUSULA TRIGÉSIMA CUARTA. DEL INCUMPLIMIENTO: '
            p50.add_run(titulo36).bold = True
            p50.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p50.add_run('El incumplimiento del ARRENDATARIO a cualquiera de sus obligaciones legales o contractuales, faculta al ARRENDADOR para ejercer las siguientes acciones, simultáneamente o en el orden que él elija:\t\n 1. Declarar terminado este contrato y reclamar la devolución del inmueble judicial y/o extrajudicialmente. \t\n 2. Exigir y perseguir a través de cualquier medio, judicial o extrajudicialmente al ARRENDATARIO y/o DEUDORES SOLIDARIOS por el monto de los perjuicios resultantes del incumplimiento, así como de la multa por incumplimiento pactada en este Contrato.\n 3. COBRO EXTRAJUDICIAL. Si el incumplimiento de la obligación de cancelar oportunamente los cánones de arrendamiento, servicios públicos, cuotas de administración, o cualquier otra erogación derivada del contrato, diere lugar a alguna diligencia de cobro extrajudicial, EL ARRENDATARIO se obliga a pagar a la entidad encargada de tal gestión,los costos correspondientes.\t\n')
            p50.add_run('4.Realizar el cobro de los intereses moratorios por las sumas adeudadas.')
            p50.add_run('PARÁGRAFO. ').bold = True
            p50.add_run('No obstante lo anterior, las Partes en cualquier tiempo y de común acuerdo podrán dar por terminado el presente Contrato.\t')
            
        else:
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO")and(Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO")  :
                titulo35='CLÁUSULA TRIGÉSIMA SEGUNDA. DEL INCUMPLIMIENTO: '
                p50.add_run(titulo35).bold = True
                p50.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p50.add_run('El incumplimiento del ARRENDATARIO a cualquiera de sus obligaciones legales o contractuales, faculta al ARRENDADOR para ejercer las siguientes acciones, simultáneamente o en el orden que él elija:\t\n 1. Declarar terminado este contrato y reclamar la devolución del inmueble judicial y/o extrajudicialmente.\t \n 2. Exigir y perseguir a través de cualquier medio, judicial o extrajudicialmente al ARRENDATARIO  por el monto de los perjuicios resultantes del incumplimiento, así como de la multa por incumplimiento pactada en este Contrato.\n 3. COBRO EXTRAJUDICIAL. Si el incumplimiento de la obligación de cancelar oportunamente los cánones de arrendamiento, servicios públicos, cuotas de administración, o cualquier otra erogación derivada del contrato, diere lugar a alguna diligencia de cobro extrajudicial, EL ARRENDATARIO se obliga a pagar a la entidad encargada de tal gestión,los costos correspondientes.\t\n')
                p50.add_run('4.Realizar el cobro de los intereses moratorios por las sumas adeudadas.')
                p50.add_run('PARÁGRAFO. ').bold = True
                p50.add_run('No obstante lo anterior, las Partes en cualquier tiempo y de común acuerdo podrán dar por terminado el presente Contrato.\t')
                
            elif  (Respuesta1=="Si" or Respuesta1=="si"or Respuesta1=="SI")and (Rta=="No" or Rta=="no"or Rta=="NO")and(Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                titulo35='CLÁUSULA TRIGÉSIMA CUARTA. DEL INCUMPLIMIENTO: '
                p50.add_run(titulo35).bold = True
                p50.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p50.add_run('El incumplimiento del ARRENDATARIO a cualquiera de sus obligaciones legales o contractuales, faculta al ARRENDADOR para ejercer las siguientes acciones, simultáneamente o en el orden que él elija:\t\n 1. Declarar terminado este contrato y reclamar la devolución del inmueble judicial y/o extrajudicialmente.\t \n 2. Exigir y perseguir a través de cualquier medio, judicial o extrajudicialmente al ARRENDATARIO por el monto de los perjuicios resultantes del incumplimiento, así como de la multa por incumplimiento pactada en este Contrato.\n 3. COBRO EXTRAJUDICIAL. Si el incumplimiento de la obligación de cancelar oportunamente los cánones de arrendamiento, servicios públicos, cuotas de administración, o cualquier otra erogación derivada del contrato, diere lugar a alguna diligencia de cobro extrajudicial, EL ARRENDATARIO se obliga a pagar a la entidad encargada de tal gestión,los costos correspondientes.\t\n')
                p50.add_run('4.Realizar el cobro de los intereses moratorios por las sumas adeudadas.')
                p50.add_run('PARÁGRAFO. ').bold = True
                p50.add_run('No obstante lo anterior, las Partes en cualquier tiempo y de común acuerdo podrán dar por terminado el presente Contrato.\t ')

        #____________________________________
        if( Clausula== "si" or Clausula=="Si"or Clausula=="SI"):
        
            p54 = documentoObjeto.add_paragraph("")
            p54.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI") and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo37='CLÁUSULA TRIGÉSIMA SEXTA. ESPACIOS EN BLANCO: '
                p54.add_run(titulo37).bold = True
                p54.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p54.add_run('El ARRENDATARIO faculta expresamente al ARRENDADOR para llenar los espacios en blanco del presente contrato. ')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo37='CLÁUSULA TRIGÉSIMA SÉPTIMA. ESPACIOS EN BLANCO: '
                p54.add_run(titulo37).bold = True
                p54.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p54.add_run('El ARRENDATARIO faculta expresamente al ARRENDADOR para llenar los espacios en blanco del presente contrato.')

            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and( Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                    titulo36='CLÁUSULA TRIGÉSIMA QUINTA. ESPACIOS EN BLANCO: '
                    p54.add_run(titulo36).bold = True
                    p54.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p54.add_run('El ARRENDATARIO faculta expresamente al ARRENDADOR para llenar los espacios en blanco del presente contrato. ')

                elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and( Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                    titulo36='CLÁUSULA TRIGÉSIMA SEXTA. ESPACIOS EN BLANCO: '
                    p54.add_run(titulo36).bold = True
                    p54.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p54.add_run('El ARRENDATARIO faculta expresamente al ARRENDADOR para llenar los espacios en blanco del presente contrato.')
                ##SIN ADMINISTRADO
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI") and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                titulo37='CLÁUSULA TRIGÉSIMA CUARTA. ESPACIOS EN BLANCO: '
                p54.add_run(titulo37).bold = True
                p54.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p54.add_run('El ARRENDATARIO faculta expresamente al ARRENDADOR para llenar los espacios en blanco del presente contrato. ')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                titulo37='CLÁUSULA TRIGÉSIMA QUINTA. ESPACIOS EN BLANCO: '
                p54.add_run(titulo37).bold = True
                p54.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p54.add_run('El ARRENDATARIO faculta expresamente al ARRENDADOR para llenar los espacios en blanco del presente contrato.')

            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and( Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                    titulo36='CLÁUSULA TRIGÉSIMA TERCERA. ESPACIOS EN BLANCO: '
                    p54.add_run(titulo36).bold = True
                    p54.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p54.add_run('El ARRENDATARIO faculta expresamente al ARRENDADOR para llenar los espacios en blanco del presente contrato. ')

                elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and( Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                    titulo36='CLÁUSULA TRIGÉSIMA QUINTA. ESPACIOS EN BLANCO: '
                    p54.add_run(titulo36).bold = True
                    p54.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p54.add_run('El ARRENDATARIO faculta expresamente al ARRENDADOR para llenar los espacios en blanco del presente contrato.')
            
        else:
            pass
        #__________________
        p55 = documentoObjeto.add_paragraph("")
        p55.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if( Clausula== "si" or Clausula=="Si"or Clausula=="SI"):

            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")  and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo38='CLÁUSULA TRIGÉSIMA SEPTIMA. NOTIFICACIONES: '
                p55.add_run(titulo38).bold = True
                p55.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p55.add_run('El ARRENDATARIO y sus deudores solidarios, indican que la dirección donde recibirán las notificaciones mientras el presente contrato se encuentre en vigencia, serán las mismas del correo electrónico entregado o en su defecto la dirección del bien inmueble arrendado. ')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo38='CLÁUSULA TRIGÉSIMA OCTAVA. NOTIFICACIONES: '
                p55.add_run(titulo38).bold = True
                p55.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p55.add_run('El ARRENDATARIO y sus deudores solidarios, indican que la dirección donde recibirán las notificaciones mientras el presente contrato se encuentre en vigencia,serán las mismas del correo electrónico entregado o en su defecto la dirección del bien inmueble arrendado.')

            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                    titulo37='CLÁUSULA TRIGÉSIMA SEXTA. NOTIFICACIONES: '
                    p55.add_run(titulo37).bold = True
                    p55.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p55.add_run('El ARRENDATARIO, indica que la dirección donde recibirán las notificaciones mientras el presente contrato se encuentre en vigencia,serán las mismas del correo electrónico entregado o en su defecto la dirección del bien inmueble arrendado.')

                elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI")  :
                    titulo37='CLÁUSULA TRIGÉSIMA SÉPTIMA. NOTIFICACIONES: '
                    p55.add_run(titulo37).bold = True
                    p55.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p55.add_run('El ARRENDATARIO, indica que la dirección donde recibirán las notificaciones mientras el presente contrato se encuentre en vigencia,serán las mismas del correo electrónico entregado o en su defecto la dirección del bien inmueble arrendado.')
                ##SIN ADMINISTRADOR
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")  and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                titulo38='CLÁUSULA TRIGÉSIMA QUINTA. NOTIFICACIONES: '
                p55.add_run(titulo38).bold = True
                p55.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p55.add_run('El ARRENDATARIO y sus deudores solidarios, indican que la dirección donde recibirán las notificaciones mientras el presente contrato se encuentre en vigencia,serán las mismas del correo electrónico entregado o en su defecto la dirección del bien inmueble arrendado.')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                titulo38='CLÁUSULA TRIGÉSIMA SEXTA. NOTIFICACIONES: '
                p55.add_run(titulo38).bold = True
                p55.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p55.add_run('El ARRENDATARIO y sus deudores solidarios, indican que la dirección donde recibirán las notificaciones mientras el presente contrato se encuentre en vigencia,serán las mismas del correo electrónico entregado o en su defecto la dirección del bien inmueble arrendado.')
                
            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO")  :
                    titulo37='CLÁUSULA TRIGÉSIMA CUARTA. NOTIFICACIONES: '
                    p55.add_run(titulo37).bold = True
                    p55.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p55.add_run('El ARRENDATARIO, indica que la dirección donde recibirán las notificaciones mientras el presente contrato se encuentre en vigencia,serán las mismas del correo electrónico entregado o en su defecto la dirección del bien inmueble arrendado.')

                elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                    titulo37='CLÁUSULA TRIGÉSIMA SEXTA. NOTIFICACIONES: '
                    p55.add_run(titulo37).bold = True
                    p55.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p55.add_run('El ARRENDATARIO, indica que la dirección donde recibirán las notificaciones mientras el presente contrato se encuentre en vigencia,serán las mismas del correo electrónico entregado o en su defecto la dirección del bien inmueble arrendado.')
        
        else:

            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")  and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo38='CLÁUSULA TRIGÉSIMA SEXTA. NOTIFICACIONES: '
                p55.add_run(titulo38).bold = True
                p55.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p55.add_run('El ARRENDATARIO y sus deudores solidarios, indican que la dirección donde recibirán las notificaciones mientras el presente contrato se encuentre en vigencia, serán las mismas del correo electrónico entregado o en su defecto la dirección del bien inmueble arrendado. ')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo38='CLÁUSULA TRIGÉSIMA SEPTIMA. NOTIFICACIONES: '
                p55.add_run(titulo38).bold = True
                p55.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p55.add_run('El ARRENDATARIO y sus deudores solidarios, indican que la dirección donde recibirán las notificaciones mientras el presente contrato se encuentre en vigencia,serán las mismas del correo electrónico entregado o en su defecto la dirección del bien inmueble arrendado.')

            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                    titulo37='CLÁUSULA TRIGÉSIMA QUINTA. NOTIFICACIONES: '
                    p55.add_run(titulo37).bold = True
                    p55.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p55.add_run('El ARRENDATARIO, indica que la dirección donde recibirán las notificaciones mientras el presente contrato se encuentre en vigencia,serán las mismas del correo electrónico entregado o en su defecto la dirección del bien inmueble arrendado.')

                elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI")  :
                    titulo37='CLÁUSULA TRIGÉSIMA SEXTA. NOTIFICACIONES: '
                    p55.add_run(titulo37).bold = True
                    p55.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p55.add_run('El ARRENDATARIO, indica que la dirección donde recibirán las notificaciones mientras el presente contrato se encuentre en vigencia,serán las mismas del correo electrónico entregado o en su defecto la dirección del bien inmueble arrendado.')
                ##SIN ADMINISTRADOR
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")  and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                titulo38='CLÁUSULA TRIGÉSIMA CUARTA. NOTIFICACIONES: '
                p55.add_run(titulo38).bold = True
                p55.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p55.add_run('El ARRENDATARIO y sus deudores solidarios, indican que la dirección donde recibirán las notificaciones mientras el presente contrato se encuentre en vigencia,serán las mismas del correo electrónico entregado o en su defecto la dirección del bien inmueble arrendado.')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                titulo38='CLÁUSULA TRIGÉSIMA QUINTA. NOTIFICACIONES: '
                p55.add_run(titulo38).bold = True
                p55.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p55.add_run('El ARRENDATARIO y sus deudores solidarios, indican que la dirección donde recibirán las notificaciones mientras el presente contrato se encuentre en vigencia,serán las mismas del correo electrónico entregado o en su defecto la dirección del bien inmueble arrendado.')
            
            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO")  :
                    titulo37='CLÁUSULA TRIGÉSIMA TERCERA. NOTIFICACIONES: '
                    p55.add_run(titulo37).bold = True
                    p55.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p55.add_run('El ARRENDATARIO, indica que la dirección donde recibirán las notificaciones mientras el presente contrato se encuentre en vigencia,serán las mismas del correo electrónico entregado o en su defecto la dirección del bien inmueble arrendado.')

                elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                    titulo37='CLÁUSULA TRIGÉSIMA QUINTA. NOTIFICACIONES: '
                    p55.add_run(titulo37).bold = True
                    p55.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p55.add_run('El ARRENDATARIO, indica que la dirección donde recibirán las notificaciones mientras el presente contrato se encuentre en vigencia,serán las mismas del correo electrónico entregado o en su defecto la dirección del bien inmueble arrendado.')

        #________________________
        if( Clausula== "si" or Clausula=="Si"or Clausula=="SI"):
            p57 = documentoObjeto.add_paragraph("")
            p57.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                titulo39='CLÁUSULA TRIGÉSIMA OCTAVA. RENUNCIA: '
                p57.add_run(titulo39).bold = True
                p57.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p57.add_run('EL ARRENDATARIO declara que no ha tenido ni tiene posesión del Inmueble, y que renuncia en beneficio de EL ARRENDADOR o de su cesionario, a todo requerimiento para constituirlo en mora en el cumplimiento de las obligaciones a su cargo derivadas de este Contrato.')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo39='CLÁUSULA TRIGÉSIMA NOVENA. RENUNCIA: '
                p57.add_run(titulo39).bold = True
                p57.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p57.add_run('EL ARRENDATARIO declara que no ha tenido ni tiene posesión del Inmueble, y que renuncia en beneficio de EL ARRENDADOR o de su cesionario, a todo requerimiento para constituirlo en mora en el cumplimiento de las obligaciones a su cargo derivadas de este Contrato.')

            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                    titulo38='CLÁUSULA TRIGÉSIMA SÉPTIMA. RENUNCIA: '
                    p57.add_run(titulo38).bold = True
                    p57.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p57.add_run('EL ARRENDATARIO declara que no ha tenido ni tiene posesión del Inmueble, y que renuncia en beneficio de EL ARRENDADOR o de su cesionario, a todo requerimiento para constituirlo en mora en el cumplimiento de las obligaciones a su cargo derivadas de este Contrato.')

                elif(Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI")  :
                    titulo38='CLÁUSULA TRIGÉSIMA OCTAVA. RENUNCIA: '
                    p57.add_run(titulo38).bold = True
                    p57.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p57.add_run('EL ARRENDATARIO declara que no ha tenido ni tiene posesión del Inmueble, y que renuncia en beneficio de EL ARRENDADOR o de su cesionario, a todo requerimiento para constituirlo en mora en el cumplimiento de las obligaciones a su cargo derivadas de este Contrato.')
                ##SIN ADMINISTRADOR
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                titulo39='CLÁUSULA TRIGÉSIMA SEXTA. RENUNCIA: '
                p57.add_run(titulo39).bold = True
                p57.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p57.add_run('EL ARRENDATARIO declara que no ha tenido ni tiene posesión del Inmueble, y que renuncia en beneficio de EL ARRENDADOR o de su cesionario, a todo requerimiento para constituirlo en mora en el cumplimiento de las obligaciones a su cargo derivadas de este Contrato.')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                titulo39='CLÁUSULA TRIGÉSIMA SEPTIMA. RENUNCIA: '
                p57.add_run(titulo39).bold = True
                p57.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p57.add_run('EL ARRENDATARIO declara que no ha tenido ni tiene posesión del Inmueble, y que renuncia en beneficio de EL ARRENDADOR o de su cesionario, a todo requerimiento para constituirlo en mora en el cumplimiento de las obligaciones a su cargo derivadas de este Contrato.')

            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                    titulo38='CLÁUSULA TRIGÉSIMA QUINTA. RENUNCIA: '
                    p57.add_run(titulo38).bold = True
                    p57.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p57.add_run('EL ARRENDATARIO declara que no ha tenido ni tiene posesión del Inmueble, y que renuncia en beneficio de EL ARRENDADOR o de su cesionario, a todo requerimiento para constituirlo en mora en el cumplimiento de las obligaciones a su cargo derivadas de este Contrato.')

                elif(Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                    titulo38='CLÁUSULA TRIGÉSIMA SEPTIMA. RENUNCIA: '
                    p57.add_run(titulo38).bold = True
                    p57.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p57.add_run('EL ARRENDATARIO declara que no ha tenido ni tiene posesión del Inmueble, y que renuncia en beneficio de EL ARRENDADOR o de su cesionario, a todo requerimiento para constituirlo en mora en el cumplimiento de las obligaciones a su cargo derivadas de este Contrato.')
        
        else:
            #Sin clausula blanco
            p57 = documentoObjeto.add_paragraph("")
            p57.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                titulo39='CLÁUSULA TRIGÉSIMA SEPTIMA. RENUNCIA: '
                p57.add_run(titulo39).bold = True
                p57.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p57.add_run('EL ARRENDATARIO declara que no ha tenido ni tiene posesión del Inmueble, y que renuncia en beneficio de EL ARRENDADOR o de su cesionario, a todo requerimiento para constituirlo en mora en el cumplimiento de las obligaciones a su cargo derivadas de este Contrato.')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo39='CLÁUSULA TRIGÉSIMA OCTAVA. RENUNCIA: '
                p57.add_run(titulo39).bold = True
                p57.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p57.add_run('EL ARRENDATARIO declara que no ha tenido ni tiene posesión del Inmueble, y que renuncia en beneficio de EL ARRENDADOR o de su cesionario, a todo requerimiento para constituirlo en mora en el cumplimiento de las obligaciones a su cargo derivadas de este Contrato.')

            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                    titulo38='CLÁUSULA TRIGÉSIMA SEXTA. RENUNCIA: '
                    p57.add_run(titulo38).bold = True
                    p57.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p57.add_run('EL ARRENDATARIO declara que no ha tenido ni tiene posesión del Inmueble, y que renuncia en beneficio de EL ARRENDADOR o de su cesionario, a todo requerimiento para constituirlo en mora en el cumplimiento de las obligaciones a su cargo derivadas de este Contrato.')

                elif(Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI")  :
                    titulo38='CLÁUSULA TRIGÉSIMA SEPTIMA. RENUNCIA: '
                    p57.add_run(titulo38).bold = True
                    p57.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p57.add_run('EL ARRENDATARIO declara que no ha tenido ni tiene posesión del Inmueble, y que renuncia en beneficio de EL ARRENDADOR o de su cesionario, a todo requerimiento para constituirlo en mora en el cumplimiento de las obligaciones a su cargo derivadas de este Contrato.')
                ##SIN ADMINISTRADOR
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                titulo39='CLÁUSULA TRIGÉSIMA QUINTA. RENUNCIA: '
                p57.add_run(titulo39).bold = True
                p57.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p57.add_run('EL ARRENDATARIO declara que no ha tenido ni tiene posesión del Inmueble, y que renuncia en beneficio de EL ARRENDADOR o de su cesionario, a todo requerimiento para constituirlo en mora en el cumplimiento de las obligaciones a su cargo derivadas de este Contrato.')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                titulo39='CLÁUSULA TRIGÉSIMA SEXTA. RENUNCIA: '
                p57.add_run(titulo39).bold = True
                p57.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p57.add_run('EL ARRENDATARIO declara que no ha tenido ni tiene posesión del Inmueble, y que renuncia en beneficio de EL ARRENDADOR o de su cesionario, a todo requerimiento para constituirlo en mora en el cumplimiento de las obligaciones a su cargo derivadas de este Contrato.')

            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                    titulo38='CLÁUSULA TRIGÉSIMA CUARTA. RENUNCIA: '
                    p57.add_run(titulo38).bold = True
                    p57.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p57.add_run('EL ARRENDATARIO declara que no ha tenido ni tiene posesión del Inmueble, y que renuncia en beneficio de EL ARRENDADOR o de su cesionario, a todo requerimiento para constituirlo en mora en el cumplimiento de las obligaciones a su cargo derivadas de este Contrato.')

                elif(Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                    titulo38='CLÁUSULA TRIGÉSIMA SEXTA. RENUNCIA: '
                    p57.add_run(titulo38).bold = True
                    p57.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p57.add_run('EL ARRENDATARIO declara que no ha tenido ni tiene posesión del Inmueble, y que renuncia en beneficio de EL ARRENDADOR o de su cesionario, a todo requerimiento para constituirlo en mora en el cumplimiento de las obligaciones a su cargo derivadas de este Contrato.')

        #______________
        if( Clausula== "si" or Clausula=="Si"or Clausula=="SI"):
            p59 = documentoObjeto.add_paragraph("")
            p59.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                titulo40='CLÁUSULA TRIGÉSIMA NOVENA. VALIDEZ: '
                p59.add_run(titulo40).bold = True
                p59.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p59.add_run('El presente Contrato anula todo convenio anterior relativo al arrendamiento del mismo Inmueble y solamente podrá ser modificado por escrito suscrito por las Partes.  ')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo40='CLÁUSULA CUADRAGÉSIMA. VALIDEZ: '
                p59.add_run(titulo40).bold = True
                p59.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p59.add_run('El presente Contrato anula todo convenio anterior relativo al arrendamiento del mismo Inmueble y solamente podrá ser modificado por escrito suscrito por las Partes. ')

            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                    titulo39='CLÁUSULA TRIGÉSIMA OCTAVA. VALIDEZ: '
                    p59.add_run(titulo39).bold = True
                    p59.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p59.add_run('El presente Contrato anula todo convenio anterior relativo al arrendamiento del mismo Inmueble y solamente podrá ser modificado por escrito suscrito por las Partes. ')

                elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and (Rta=="No" or Rta=="no"or Rta=="NO")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                    titulo39='CLÁUSULA TRIGÉSIMA NOVENA. VALIDEZ: '
                    p59.add_run(titulo39).bold = True
                    p59.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p59.add_run('El presente Contrato anula todo convenio anterior relativo al arrendamiento del mismo Inmueble y solamente podrá ser modificado por escrito suscrito por las Partes. ')
                ##SIN ADMIISTRADOR
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                titulo40='CLÁUSULA TRIGÉSIMA SEPTIMA. VALIDEZ: '
                p59.add_run(titulo40).bold = True
                p59.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p59.add_run('El presente Contrato anula todo convenio anterior relativo al arrendamiento del mismo Inmueble y solamente podrá ser modificado por escrito suscrito por las Partes.  ')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                titulo40='CLÁUSULA TRIGÉSIMA OCTAVA. VALIDEZ: '
                p59.add_run(titulo40).bold = True
                p59.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p59.add_run('El presente Contrato anula todo convenio anterior relativo al arrendamiento del mismo Inmueble y solamente podrá ser modificado por escrito suscrito por las Partes. ')

            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                    titulo39='CLÁUSULA TRIGÉSIMA SEXTA. VALIDEZ: '
                    p59.add_run(titulo39).bold = True
                    p59.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p59.add_run('El presente Contrato anula todo convenio anterior relativo al arrendamiento del mismo Inmueble y solamente podrá ser modificado por escrito suscrito por las Partes. ')

                elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and (Rta=="No" or Rta=="no"or Rta=="NO")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                    titulo39='CLÁUSULA TRIGÉSIMA OCTAVA. VALIDEZ: '
                    p59.add_run(titulo39).bold = True
                    p59.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p59.add_run('El presente Contrato anula todo convenio anterior relativo al arrendamiento del mismo Inmueble y solamente podrá ser modificado por escrito suscrito por las Partes. ')
            
        else:
            #Sin clausula blanco
            p59 = documentoObjeto.add_paragraph("")
            p59.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                titulo40='CLÁUSULA TRIGÉSIMA OCTAVA. VALIDEZ: '
                p59.add_run(titulo40).bold = True
                p59.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p59.add_run('El presente Contrato anula todo convenio anterior relativo al arrendamiento del mismo Inmueble y solamente podrá ser modificado por escrito suscrito por las Partes.  ')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo40='CLÁUSULA TRIGÉSIMA NOVENA. VALIDEZ: '
                p59.add_run(titulo40).bold = True
                p59.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p59.add_run('El presente Contrato anula todo convenio anterior relativo al arrendamiento del mismo Inmueble y solamente podrá ser modificado por escrito suscrito por las Partes. ')

            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                    titulo39='CLÁUSULA TRIGÉSIMA SEPTIMA. VALIDEZ: '
                    p59.add_run(titulo39).bold = True
                    p59.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p59.add_run('El presente Contrato anula todo convenio anterior relativo al arrendamiento del mismo Inmueble y solamente podrá ser modificado por escrito suscrito por las Partes. ')

                elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and (Rta=="No" or Rta=="no"or Rta=="NO")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                    titulo39='CLÁUSULA TRIGÉSIMA OCTAVA. VALIDEZ: '
                    p59.add_run(titulo39).bold = True
                    p59.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p59.add_run('El presente Contrato anula todo convenio anterior relativo al arrendamiento del mismo Inmueble y solamente podrá ser modificado por escrito suscrito por las Partes. ')
                ##SIN ADMIISTRADOR
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                titulo40='CLÁUSULA TRIGÉSIMA SEXTA. VALIDEZ: '
                p59.add_run(titulo40).bold = True
                p59.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p59.add_run('El presente Contrato anula todo convenio anterior relativo al arrendamiento del mismo Inmueble y solamente podrá ser modificado por escrito suscrito por las Partes.  ')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                titulo40='CLÁUSULA TRIGÉSIMA SEPTIMA. VALIDEZ: '
                p59.add_run(titulo40).bold = True
                p59.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p59.add_run('El presente Contrato anula todo convenio anterior relativo al arrendamiento del mismo Inmueble y solamente podrá ser modificado por escrito suscrito por las Partes. ')

            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                    titulo39='CLÁUSULA TRIGÉSIMA QUINTA. VALIDEZ: '
                    p59.add_run(titulo39).bold = True
                    p59.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p59.add_run('El presente Contrato anula todo convenio anterior relativo al arrendamiento del mismo Inmueble y solamente podrá ser modificado por escrito suscrito por las Partes. ')

                elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and (Rta=="No" or Rta=="no"or Rta=="NO")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                    titulo39='CLÁUSULA TRIGÉSIMA CUARTA. VALIDEZ: '
                    p59.add_run(titulo39).bold = True
                    p59.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p59.add_run('El presente Contrato anula todo convenio anterior relativo al arrendamiento del mismo Inmueble y solamente podrá ser modificado por escrito suscrito por las Partes. ')

        #_________________
        p61 = documentoObjeto.add_paragraph("")
        p61.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if( Clausula== "si" or Clausula=="Si"or Clausula=="SI"):

            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                titulo41='CLÁUSULA CUADRAGÉSIMA. PENALIDAD: '
                p61.add_run(titulo41).bold = True
                p61.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p61.add_run('En el evento de incumplimiento de cualquiera de las Partes a las obligaciones a su cargo contenidas en la ley o en este Contrato, la parte incumplida deberá pagar a la otra parte una suma equivalente a tres (3) cánones de arrendamiento y administración vigentes en la fecha del incumplimiento, a título de pena. En el evento que los perjuicios ocasionados por la parte incumplida, excedan el valor de la suma aquí prevista como pena, la Parte incumplida deberá pagar a la otra parte la diferencia entre el valor total de los perjuicios y el valor de la pena prevista en esta Cláusula. En caso de mora o mero retardo en el pago de sus obligaciones dinerarias, contenidas en el presente contrato, El Arrendatario deberá asumir el 100% de los honorarios derivados de la cobranza prejudicial, extrajudicial y /o judicial que se desarrolle con el fin de recuperar la cartera en mora.')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo41='CLÁUSULA CUADRAGÉSIMA PRIMERA. PENALIDAD: '
                p61.add_run(titulo41).bold = True
                p61.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p61.add_run('En el evento de incumplimiento de cualquiera de las Partes a las obligaciones a su cargo contenidas en la ley o en este Contrato, la parte incumplida deberá pagar a la otra parte una suma equivalente a tres (3) cánones de arrendamiento y administración vigentes en la fecha del incumplimiento, a título de pena. En el evento que los perjuicios ocasionados por la parte incumplida, excedan el valor de la suma aquí prevista como pena, la Parte incumplida deberá pagar a la otra parte la diferencia entre el valor total de los perjuicios y el valor de la pena prevista en esta Cláusula. En caso de mora o mero retardo en el pago de sus obligaciones dinerarias, contenidas en el presente contrato, El Arrendatario deberá asumir el 100% de los honorarios derivados de la cobranza prejudicial, extrajudicial y /o judicial que se desarrolle con el fin de recuperar la cartera en mora.')

            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and( Rta=="No" or Rta=="no"or Rta=="NO")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                    titulo40='CLÁUSULA TRIGÉSIMA NOVENA. PENALIDAD: '
                    p61.add_run(titulo40).bold = True
                    p61.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p61.add_run('En el evento de incumplimiento de cualquiera de las Partes a las obligaciones a su cargo contenidas en la ley o en este Contrato, la parte incumplida deberá pagar a la otra parte una suma equivalente a tres (3) cánones de arrendamiento y administración vigentes en la fecha del incumplimiento, a título de pena. En el evento que los perjuicios ocasionados por la parte incumplida, excedan el valor de la suma aquí prevista como pena, la Parte incumplida deberá pagar a la otra parte la diferencia entre el valor total de los perjuicios y el valor de la pena prevista en esta Cláusula. En caso de mora o mero retardo en el pago de sus obligaciones dinerarias, contenidas en el presente contrato, El Arrendatario deberá asumir el 100% de los honorarios derivados de la cobranza prejudicial, extrajudicial y /o judicial que se desarrolle con el fin de recuperar la cartera en mora.')

                elif(Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and( Rta=="No" or Rta=="no"or Rta=="NO")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                    titulo40='CLÁUSULA CUADRAGÉSIMA . PENALIDAD: '
                    p61.add_run(titulo40).bold = True
                    p61.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p61.add_run('En el evento de incumplimiento de cualquiera de las Partes a las obligaciones a su cargo contenidas en la ley o en este Contrato, la parte incumplida deberá pagar a la otra parte una suma equivalente a tres (3) cánones de arrendamiento y administración vigentes en la fecha del incumplimiento, a título de pena. En el evento que los perjuicios ocasionados por la parte incumplida, excedan el valor de la suma aquí prevista como pena, la Parte incumplida deberá pagar a la otra parte la diferencia entre el valor total de los perjuicios y el valor de la pena prevista en esta Cláusula. En caso de mora o mero retardo en el pago de sus obligaciones dinerarias, contenidas en el presente contrato, El Arrendatario deberá asumir el 100% de los honorarios derivados de la cobranza prejudicial, extrajudicial y /o judicial que se desarrolle con el fin de recuperar la cartera en mora.')
                ##SIN ADMINISTRADOR
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                titulo41='CLÁUSULA TRIGÉSIMA OCTAVA. PENALIDAD: '
                p61.add_run(titulo41).bold = True
                p61.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p61.add_run('En el evento de incumplimiento de cualquiera de las Partes a las obligaciones a su cargo contenidas en la ley o en este Contrato, la parte incumplida deberá pagar a la otra parte una suma equivalente a tres (3) cánones de arrendamiento y administración vigentes en la fecha del incumplimiento, a título de pena. En el evento que los perjuicios ocasionados por la parte incumplida, excedan el valor de la suma aquí prevista como pena, la Parte incumplida deberá pagar a la otra parte la diferencia entre el valor total de los perjuicios y el valor de la pena prevista en esta Cláusula. En caso de mora o mero retardo en el pago de sus obligaciones dinerarias, contenidas en el presente contrato, El Arrendatario deberá asumir el 100% de los honorarios derivados de la cobranza prejudicial, extrajudicial y /o judicial que se desarrolle con el fin de recuperar la cartera en mora.')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                titulo41='CLÁUSULA TRIGÉSIMA NOVENA. PENALIDAD: '
                p61.add_run(titulo41).bold = True
                p61.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p61.add_run('En el evento de incumplimiento de cualquiera de las Partes a las obligaciones a su cargo contenidas en la ley o en este Contrato, la parte incumplida deberá pagar a la otra parte una suma equivalente a tres (3) cánones de arrendamiento y administración vigentes en la fecha del incumplimiento, a título de pena. En el evento que los perjuicios ocasionados por la parte incumplida, excedan el valor de la suma aquí prevista como pena, la Parte incumplida deberá pagar a la otra parte la diferencia entre el valor total de los perjuicios y el valor de la pena prevista en esta Cláusula. En caso de mora o mero retardo en el pago de sus obligaciones dinerarias, contenidas en el presente contrato, El Arrendatario deberá asumir el 100% de los honorarios derivados de la cobranza prejudicial, extrajudicial y /o judicial que se desarrolle con el fin de recuperar la cartera en mora.')

            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and( Rta=="No" or Rta=="no"or Rta=="NO")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                    titulo40='CLÁUSULA TRIGÉSIMA SEPTIMA. PENALIDAD: '
                    p61.add_run(titulo40).bold = True
                    p61.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p61.add_run('En el evento de incumplimiento de cualquiera de las Partes a las obligaciones a su cargo contenidas en la ley o en este Contrato, la parte incumplida deberá pagar a la otra parte una suma equivalente a tres (3) cánones de arrendamiento y administración vigentes en la fecha del incumplimiento, a título de pena. En el evento que los perjuicios ocasionados por la parte incumplida, excedan el valor de la suma aquí prevista como pena, la Parte incumplida deberá pagar a la otra parte la diferencia entre el valor total de los perjuicios y el valor de la pena prevista en esta Cláusula. En caso de mora o mero retardo en el pago de sus obligaciones dinerarias, contenidas en el presente contrato, El Arrendatario deberá asumir el 100% de los honorarios derivados de la cobranza prejudicial, extrajudicial y /o judicial que se desarrolle con el fin de recuperar la cartera en mora.')

                elif(Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and( Rta=="No" or Rta=="no"or Rta=="NO")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                    titulo40='CLÁUSULA TRIGÉSIMA NOVENA . PENALIDAD: '
                    p61.add_run(titulo40).bold = True
                    p61.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p61.add_run('En el evento de incumplimiento de cualquiera de las Partes a las obligaciones a su cargo contenidas en la ley o en este Contrato, la parte incumplida deberá pagar a la otra parte una suma equivalente a tres (3) cánones de arrendamiento y administración vigentes en la fecha del incumplimiento, a título de pena. En el evento que los perjuicios ocasionados por la parte incumplida, excedan el valor de la suma aquí prevista como pena, la Parte incumplida deberá pagar a la otra parte la diferencia entre el valor total de los perjuicios y el valor de la pena prevista en esta Cláusula. En caso de mora o mero retardo en el pago de sus obligaciones dinerarias, contenidas en el presente contrato, El Arrendatario deberá asumir el 100% de los honorarios derivados de la cobranza prejudicial, extrajudicial y /o judicial que se desarrolle con el fin de recuperar la cartera en mora.')
            
        else:
            #Sin clausula blanco
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                titulo41='CLÁUSULA TRIGÉSIMA NOVENA. PENALIDAD: '
                p61.add_run(titulo41).bold = True
                p61.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p61.add_run('En el evento de incumplimiento de cualquiera de las Partes a las obligaciones a su cargo contenidas en la ley o en este Contrato, la parte incumplida deberá pagar a la otra parte una suma equivalente a tres (3) cánones de arrendamiento y administración vigentes en la fecha del incumplimiento, a título de pena. En el evento que los perjuicios ocasionados por la parte incumplida, excedan el valor de la suma aquí prevista como pena, la Parte incumplida deberá pagar a la otra parte la diferencia entre el valor total de los perjuicios y el valor de la pena prevista en esta Cláusula. En caso de mora o mero retardo en el pago de sus obligaciones dinerarias, contenidas en el presente contrato, El Arrendatario deberá asumir el 100% de los honorarios derivados de la cobranza prejudicial, extrajudicial y /o judicial que se desarrolle con el fin de recuperar la cartera en mora.')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo41='CLÁUSULA CUADRAGÉSIMA . PENALIDAD: '
                p61.add_run(titulo41).bold = True
                p61.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p61.add_run('En el evento de incumplimiento de cualquiera de las Partes a las obligaciones a su cargo contenidas en la ley o en este Contrato, la parte incumplida deberá pagar a la otra parte una suma equivalente a tres (3) cánones de arrendamiento y administración vigentes en la fecha del incumplimiento, a título de pena. En el evento que los perjuicios ocasionados por la parte incumplida, excedan el valor de la suma aquí prevista como pena, la Parte incumplida deberá pagar a la otra parte la diferencia entre el valor total de los perjuicios y el valor de la pena prevista en esta Cláusula. En caso de mora o mero retardo en el pago de sus obligaciones dinerarias, contenidas en el presente contrato, El Arrendatario deberá asumir el 100% de los honorarios derivados de la cobranza prejudicial, extrajudicial y /o judicial que se desarrolle con el fin de recuperar la cartera en mora.')

            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and( Rta=="No" or Rta=="no"or Rta=="NO")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                    titulo40='CLÁUSULA TRIGÉSIMA OCTAVA. PENALIDAD: '
                    p61.add_run(titulo40).bold = True
                    p61.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p61.add_run('En el evento de incumplimiento de cualquiera de las Partes a las obligaciones a su cargo contenidas en la ley o en este Contrato, la parte incumplida deberá pagar a la otra parte una suma equivalente a tres (3) cánones de arrendamiento y administración vigentes en la fecha del incumplimiento, a título de pena. En el evento que los perjuicios ocasionados por la parte incumplida, excedan el valor de la suma aquí prevista como pena, la Parte incumplida deberá pagar a la otra parte la diferencia entre el valor total de los perjuicios y el valor de la pena prevista en esta Cláusula. En caso de mora o mero retardo en el pago de sus obligaciones dinerarias, contenidas en el presente contrato, El Arrendatario deberá asumir el 100% de los honorarios derivados de la cobranza prejudicial, extrajudicial y /o judicial que se desarrolle con el fin de recuperar la cartera en mora.')

                elif(Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and( Rta=="No" or Rta=="no"or Rta=="NO")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                    titulo40='CLÁUSULA TRIGÉSIMA NOVENA . PENALIDAD: '
                    p61.add_run(titulo40).bold = True
                    p61.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p61.add_run('En el evento de incumplimiento de cualquiera de las Partes a las obligaciones a su cargo contenidas en la ley o en este Contrato, la parte incumplida deberá pagar a la otra parte una suma equivalente a tres (3) cánones de arrendamiento y administración vigentes en la fecha del incumplimiento, a título de pena. En el evento que los perjuicios ocasionados por la parte incumplida, excedan el valor de la suma aquí prevista como pena, la Parte incumplida deberá pagar a la otra parte la diferencia entre el valor total de los perjuicios y el valor de la pena prevista en esta Cláusula. En caso de mora o mero retardo en el pago de sus obligaciones dinerarias, contenidas en el presente contrato, El Arrendatario deberá asumir el 100% de los honorarios derivados de la cobranza prejudicial, extrajudicial y /o judicial que se desarrolle con el fin de recuperar la cartera en mora.')
                ##SIN ADMINISTRADOR
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                titulo41='CLÁUSULA TRIGÉSIMA SEPTIMA. PENALIDAD: '
                p61.add_run(titulo41).bold = True
                p61.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p61.add_run('En el evento de incumplimiento de cualquiera de las Partes a las obligaciones a su cargo contenidas en la ley o en este Contrato, la parte incumplida deberá pagar a la otra parte una suma equivalente a tres (3) cánones de arrendamiento y administración vigentes en la fecha del incumplimiento, a título de pena. En el evento que los perjuicios ocasionados por la parte incumplida, excedan el valor de la suma aquí prevista como pena, la Parte incumplida deberá pagar a la otra parte la diferencia entre el valor total de los perjuicios y el valor de la pena prevista en esta Cláusula. En caso de mora o mero retardo en el pago de sus obligaciones dinerarias, contenidas en el presente contrato, El Arrendatario deberá asumir el 100% de los honorarios derivados de la cobranza prejudicial, extrajudicial y /o judicial que se desarrolle con el fin de recuperar la cartera en mora.')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                titulo41='CLÁUSULA TRIGÉSIMA OCTAVA. PENALIDAD: '
                p61.add_run(titulo41).bold = True
                p61.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p61.add_run('En el evento de incumplimiento de cualquiera de las Partes a las obligaciones a su cargo contenidas en la ley o en este Contrato, la parte incumplida deberá pagar a la otra parte una suma equivalente a tres (3) cánones de arrendamiento y administración vigentes en la fecha del incumplimiento, a título de pena. En el evento que los perjuicios ocasionados por la parte incumplida, excedan el valor de la suma aquí prevista como pena, la Parte incumplida deberá pagar a la otra parte la diferencia entre el valor total de los perjuicios y el valor de la pena prevista en esta Cláusula. En caso de mora o mero retardo en el pago de sus obligaciones dinerarias, contenidas en el presente contrato, El Arrendatario deberá asumir el 100% de los honorarios derivados de la cobranza prejudicial, extrajudicial y /o judicial que se desarrolle con el fin de recuperar la cartera en mora.')

            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and( Rta=="No" or Rta=="no"or Rta=="NO")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                    titulo40='CLÁUSULA TRIGÉSIMA SEXTA. PENALIDAD: '
                    p61.add_run(titulo40).bold = True
                    p61.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p61.add_run('En el evento de incumplimiento de cualquiera de las Partes a las obligaciones a su cargo contenidas en la ley o en este Contrato, la parte incumplida deberá pagar a la otra parte una suma equivalente a tres (3) cánones de arrendamiento y administración vigentes en la fecha del incumplimiento, a título de pena. En el evento que los perjuicios ocasionados por la parte incumplida, excedan el valor de la suma aquí prevista como pena, la Parte incumplida deberá pagar a la otra parte la diferencia entre el valor total de los perjuicios y el valor de la pena prevista en esta Cláusula. En caso de mora o mero retardo en el pago de sus obligaciones dinerarias, contenidas en el presente contrato, El Arrendatario deberá asumir el 100% de los honorarios derivados de la cobranza prejudicial, extrajudicial y /o judicial que se desarrolle con el fin de recuperar la cartera en mora.')

                elif(Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and( Rta=="No" or Rta=="no"or Rta=="NO")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                    titulo40='CLÁUSULA TRIGÉSIMA OCTAVA . PENALIDAD: '
                    p61.add_run(titulo40).bold = True
                    p61.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p61.add_run('En el evento de incumplimiento de cualquiera de las Partes a las obligaciones a su cargo contenidas en la ley o en este Contrato, la parte incumplida deberá pagar a la otra parte una suma equivalente a tres (3) cánones de arrendamiento y administración vigentes en la fecha del incumplimiento, a título de pena. En el evento que los perjuicios ocasionados por la parte incumplida, excedan el valor de la suma aquí prevista como pena, la Parte incumplida deberá pagar a la otra parte la diferencia entre el valor total de los perjuicios y el valor de la pena prevista en esta Cláusula. En caso de mora o mero retardo en el pago de sus obligaciones dinerarias, contenidas en el presente contrato, El Arrendatario deberá asumir el 100% de los honorarios derivados de la cobranza prejudicial, extrajudicial y /o judicial que se desarrolle con el fin de recuperar la cartera en mora.')

        #________________________________
        p62 = documentoObjeto.add_paragraph("")
        p62.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if( Clausula== "si" or Clausula=="Si"or Clausula=="SI"):

            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                titulo42='CLÁUSULA CUADRAGÉSIMA PRIMERA. PAGO MEDIANTE CHEQUE: '
                p62.add_run(titulo42).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Los valores por cánones de arrendamiento o por otros conceptos que pague EL ARRENDATARIO y/o DEUDORES SOLIDARIOS, mediante cheques y sean recibidos por EL ARRENDADOR, sólo se reputará pagado hasta que el cheque se haga efectivo, y cancelado por el Banco girador respectivo. En caso de que el cheque girado por EL ARRENDATARIO y/o DEUDORES SOLIDARIOS, no fuere pagado por el banco, por causa imputable al girador del cheque, se entenderá como no paga la obligación o canon de arrendamiento y EL ARRENDADOR tendrá la facultad para iniciar y llevar hasta su terminación las acciones judiciales pertinentes, tendientes a obtener el pago del valor del cheque impagado y cobrar el 20% sobre el valor del cheque de acuerdo al artículo 731 del Código del Comercio.\t\n')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo42='CLÁUSULA CUADRAGÉSIMA SEGUNDA. PAGO MEDIANTE CHEQUE: '
                p62.add_run(titulo42).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Los valores por cánones de arrendamiento o por otros conceptos que pague EL ARRENDATARIO y/o DEUDORES SOLIDARIOS, mediante cheques y sean recibidos por EL ARRENDADOR, sólo se reputará pagado hasta que el cheque se haga efectivo, y cancelado por el Banco girador respectivo. En caso de que el cheque girado por EL ARRENDATARIO y/o DEUDORES SOLIDARIOS, no fuere pagado por el banco, por causa imputable al girador del cheque, se entenderá como no paga la obligación o canon de arrendamiento y EL ARRENDADOR tendrá la facultad para iniciar y llevar hasta su terminación las acciones judiciales pertinentes, tendientes a obtener el pago del valor del cheque impagado y cobrar el 20% sobre el valor del cheque de acuerdo al artículo 731 del Código del Comercio.\t\n')

            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO" )and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                    titulo41='CLÁUSULA CUADRAGÉSIMA . PAGO MEDIANTE CHEQUE: '
                    p62.add_run(titulo41).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('Los valores por cánones de arrendamiento o por otros conceptos que pague EL ARRENDATARIO , mediante cheques y sean recibidos por EL ARRENDADOR, sólo se reputará pagado hasta que el cheque se haga efectivo, y cancelado por el Banco girador respectivo. En caso de que el cheque girado por EL ARRENDATARIO, no fuere pagado por el banco, por causa imputable al girador del cheque, se entenderá como no paga la obligación o canon de arrendamiento y EL ARRENDADOR tendrá la facultad para iniciar y llevar hasta su terminación las acciones judiciales pertinentes, tendientes a obtener el pago del valor del cheque impagado y cobrar el 20% sobre el valor del cheque de acuerdo al artículo 731 del Código del Comercio.\t\n')

                elif(Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                    titulo41='CLÁUSULA CUADRAGÉSIMA PRIMERA. PAGO MEDIANTE CHEQUE: '
                    p62.add_run(titulo41).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('Los valores por cánones de arrendamiento o por otros conceptos que pague EL ARRENDATARIO , mediante cheques y sean recibidos por EL ARRENDADOR, sólo se reputará pagado hasta que el cheque se haga efectivo, y cancelado por el Banco girador respectivo. En caso de que el cheque girado por EL ARRENDATARIO , no fuere pagado por el banco, por causa imputable al girador del cheque, se entenderá como no paga la obligación o canon de arrendamiento y EL ARRENDADOR tendrá la facultad para iniciar y llevar hasta su terminación las acciones judiciales pertinentes, tendientes a obtener el pago del valor del cheque impagado y cobrar el 20% sobre el valor del cheque de acuerdo al artículo 731 del Código del Comercio.\t\n')
                ##SIN ADMINISTRADOR

            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                titulo42='CLÁUSULA TRIGÉSIMA NOVENA . PAGO MEDIANTE CHEQUE: '
                p62.add_run(titulo42).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Los valores por cánones de arrendamiento o por otros conceptos que pague EL ARRENDATARIO y/o DEUDORES SOLIDARIOS, mediante cheques y sean recibidos por EL ARRENDADOR, sólo se reputará pagado hasta que el cheque se haga efectivo, y cancelado por el Banco girador respectivo. En caso de que el cheque girado por EL ARRENDATARIO y/o DEUDORES SOLIDARIOS, no fuere pagado por el banco, por causa imputable al girador del cheque, se entenderá como no paga la obligación o canon de arrendamiento y EL ARRENDADOR tendrá la facultad para iniciar y llevar hasta su terminación las acciones judiciales pertinentes, tendientes a obtener el pago del valor del cheque impagado y cobrar el 20% sobre el valor del cheque de acuerdo al artículo 731 del Código del Comercio.\t\n')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                titulo42='CLÁUSULA CUADRAGÉSIMA. PAGO MEDIANTE CHEQUE: '
                p62.add_run(titulo42).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Los valores por cánones de arrendamiento o por otros conceptos que pague EL ARRENDATARIO y/o DEUDORES SOLIDARIOS, mediante cheques y sean recibidos por EL ARRENDADOR, sólo se reputará pagado hasta que el cheque se haga efectivo, y cancelado por el Banco girador respectivo. En caso de que el cheque girado por EL ARRENDATARIO y/o DEUDORES SOLIDARIOS, no fuere pagado por el banco, por causa imputable al girador del cheque, se entenderá como no paga la obligación o canon de arrendamiento y EL ARRENDADOR tendrá la facultad para iniciar y llevar hasta su terminación las acciones judiciales pertinentes, tendientes a obtener el pago del valor del cheque impagado y cobrar el 20% sobre el valor del cheque de acuerdo al artículo 731 del Código del Comercio.\t\n')

            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO" )and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                    titulo41='CLÁUSULA TRIGÉSIMA OCTAVA. PAGO MEDIANTE CHEQUE: '
                    p62.add_run(titulo41).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('Los valores por cánones de arrendamiento o por otros conceptos que pague EL ARRENDATARIO , mediante cheques y sean recibidos por EL ARRENDADOR, sólo se reputará pagado hasta que el cheque se haga efectivo, y cancelado por el Banco girador respectivo. En caso de que el cheque girado por EL ARRENDATARIO, no fuere pagado por el banco, por causa imputable al girador del cheque, se entenderá como no paga la obligación o canon de arrendamiento y EL ARRENDADOR tendrá la facultad para iniciar y llevar hasta su terminación las acciones judiciales pertinentes, tendientes a obtener el pago del valor del cheque impagado y cobrar el 20% sobre el valor del cheque de acuerdo al artículo 731 del Código del Comercio.\t\n')

                elif(Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                    titulo41='CLÁUSULA CUADRAGÉSIMA . PAGO MEDIANTE CHEQUE: '
                    p62.add_run(titulo41).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('Los valores por cánones de arrendamiento o por otros conceptos que pague EL ARRENDATARIO , mediante cheques y sean recibidos por EL ARRENDADOR, sólo se reputará pagado hasta que el cheque se haga efectivo, y cancelado por el Banco girador respectivo. En caso de que el cheque girado por EL ARRENDATARIO , no fuere pagado por el banco, por causa imputable al girador del cheque, se entenderá como no paga la obligación o canon de arrendamiento y EL ARRENDADOR tendrá la facultad para iniciar y llevar hasta su terminación las acciones judiciales pertinentes, tendientes a obtener el pago del valor del cheque impagado y cobrar el 20% sobre el valor del cheque de acuerdo al artículo 731 del Código del Comercio.\t\n')
            
        else:
            #Sin clausula blanco

            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                titulo42='CLÁUSULA CUADRAGÉSIMA . PAGO MEDIANTE CHEQUE: '
                p62.add_run(titulo42).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Los valores por cánones de arrendamiento o por otros conceptos que pague EL ARRENDATARIO y/o DEUDORES SOLIDARIOS, mediante cheques y sean recibidos por EL ARRENDADOR, sólo se reputará pagado hasta que el cheque se haga efectivo, y cancelado por el Banco girador respectivo. En caso de que el cheque girado por EL ARRENDATARIO y/o DEUDORES SOLIDARIOS, no fuere pagado por el banco, por causa imputable al girador del cheque, se entenderá como no paga la obligación o canon de arrendamiento y EL ARRENDADOR tendrá la facultad para iniciar y llevar hasta su terminación las acciones judiciales pertinentes, tendientes a obtener el pago del valor del cheque impagado y cobrar el 20% sobre el valor del cheque de acuerdo al artículo 731 del Código del Comercio.\t\n')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo42='CLÁUSULA CUADRAGÉSIMA PRIMERA. PAGO MEDIANTE CHEQUE: '
                p62.add_run(titulo42).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Los valores por cánones de arrendamiento o por otros conceptos que pague EL ARRENDATARIO y/o DEUDORES SOLIDARIOS, mediante cheques y sean recibidos por EL ARRENDADOR, sólo se reputará pagado hasta que el cheque se haga efectivo, y cancelado por el Banco girador respectivo. En caso de que el cheque girado por EL ARRENDATARIO y/o DEUDORES SOLIDARIOS, no fuere pagado por el banco, por causa imputable al girador del cheque, se entenderá como no paga la obligación o canon de arrendamiento y EL ARRENDADOR tendrá la facultad para iniciar y llevar hasta su terminación las acciones judiciales pertinentes, tendientes a obtener el pago del valor del cheque impagado y cobrar el 20% sobre el valor del cheque de acuerdo al artículo 731 del Código del Comercio.\t\n')

            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO" )and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                    titulo41='CLÁUSULA TRIGÉSIMA NOVENA . PAGO MEDIANTE CHEQUE: '
                    p62.add_run(titulo41).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('Los valores por cánones de arrendamiento o por otros conceptos que pague EL ARRENDATARIO , mediante cheques y sean recibidos por EL ARRENDADOR, sólo se reputará pagado hasta que el cheque se haga efectivo, y cancelado por el Banco girador respectivo. En caso de que el cheque girado por EL ARRENDATARIO, no fuere pagado por el banco, por causa imputable al girador del cheque, se entenderá como no paga la obligación o canon de arrendamiento y EL ARRENDADOR tendrá la facultad para iniciar y llevar hasta su terminación las acciones judiciales pertinentes, tendientes a obtener el pago del valor del cheque impagado y cobrar el 20% sobre el valor del cheque de acuerdo al artículo 731 del Código del Comercio.\t\n')

                elif(Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                    titulo41='CLÁUSULA CUADRAGÉSIMA . PAGO MEDIANTE CHEQUE: '
                    p62.add_run(titulo41).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('Los valores por cánones de arrendamiento o por otros conceptos que pague EL ARRENDATARIO , mediante cheques y sean recibidos por EL ARRENDADOR, sólo se reputará pagado hasta que el cheque se haga efectivo, y cancelado por el Banco girador respectivo. En caso de que el cheque girado por EL ARRENDATARIO , no fuere pagado por el banco, por causa imputable al girador del cheque, se entenderá como no paga la obligación o canon de arrendamiento y EL ARRENDADOR tendrá la facultad para iniciar y llevar hasta su terminación las acciones judiciales pertinentes, tendientes a obtener el pago del valor del cheque impagado y cobrar el 20% sobre el valor del cheque de acuerdo al artículo 731 del Código del Comercio.\t\n')
                ##SIN ADMINISTRADOR

            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                titulo42='CLÁUSULA TRIGÉSIMA OCTAVA . PAGO MEDIANTE CHEQUE: '
                p62.add_run(titulo42).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Los valores por cánones de arrendamiento o por otros conceptos que pague EL ARRENDATARIO y/o DEUDORES SOLIDARIOS, mediante cheques y sean recibidos por EL ARRENDADOR, sólo se reputará pagado hasta que el cheque se haga efectivo, y cancelado por el Banco girador respectivo. En caso de que el cheque girado por EL ARRENDATARIO y/o DEUDORES SOLIDARIOS, no fuere pagado por el banco, por causa imputable al girador del cheque, se entenderá como no paga la obligación o canon de arrendamiento y EL ARRENDADOR tendrá la facultad para iniciar y llevar hasta su terminación las acciones judiciales pertinentes, tendientes a obtener el pago del valor del cheque impagado y cobrar el 20% sobre el valor del cheque de acuerdo al artículo 731 del Código del Comercio.\t\n')

            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                titulo42='CLÁUSULA TRIGÉSIMA NOVENA. PAGO MEDIANTE CHEQUE: '
                p62.add_run(titulo42).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Los valores por cánones de arrendamiento o por otros conceptos que pague EL ARRENDATARIO y/o DEUDORES SOLIDARIOS, mediante cheques y sean recibidos por EL ARRENDADOR, sólo se reputará pagado hasta que el cheque se haga efectivo, y cancelado por el Banco girador respectivo. En caso de que el cheque girado por EL ARRENDATARIO y/o DEUDORES SOLIDARIOS, no fuere pagado por el banco, por causa imputable al girador del cheque, se entenderá como no paga la obligación o canon de arrendamiento y EL ARRENDADOR tendrá la facultad para iniciar y llevar hasta su terminación las acciones judiciales pertinentes, tendientes a obtener el pago del valor del cheque impagado y cobrar el 20% sobre el valor del cheque de acuerdo al artículo 731 del Código del Comercio.\t\n')

            else:
                if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO" )and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                    titulo41='CLÁUSULA TRIGÉSIMA SEPTIMA. PAGO MEDIANTE CHEQUE: '
                    p62.add_run(titulo41).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('Los valores por cánones de arrendamiento o por otros conceptos que pague EL ARRENDATARIO , mediante cheques y sean recibidos por EL ARRENDADOR, sólo se reputará pagado hasta que el cheque se haga efectivo, y cancelado por el Banco girador respectivo. En caso de que el cheque girado por EL ARRENDATARIO, no fuere pagado por el banco, por causa imputable al girador del cheque, se entenderá como no paga la obligación o canon de arrendamiento y EL ARRENDADOR tendrá la facultad para iniciar y llevar hasta su terminación las acciones judiciales pertinentes, tendientes a obtener el pago del valor del cheque impagado y cobrar el 20% sobre el valor del cheque de acuerdo al artículo 731 del Código del Comercio.\t\n')

                elif(Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and (Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                    titulo41='CLÁUSULA TRIGÉSIMA NOVENA . PAGO MEDIANTE CHEQUE: '
                    p62.add_run(titulo41).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('Los valores por cánones de arrendamiento o por otros conceptos que pague EL ARRENDATARIO , mediante cheques y sean recibidos por EL ARRENDADOR, sólo se reputará pagado hasta que el cheque se haga efectivo, y cancelado por el Banco girador respectivo. En caso de que el cheque girado por EL ARRENDATARIO , no fuere pagado por el banco, por causa imputable al girador del cheque, se entenderá como no paga la obligación o canon de arrendamiento y EL ARRENDADOR tendrá la facultad para iniciar y llevar hasta su terminación las acciones judiciales pertinentes, tendientes a obtener el pago del valor del cheque impagado y cobrar el 20% sobre el valor del cheque de acuerdo al artículo 731 del Código del Comercio.\t\n')

        #____________________________
        #surcursal lago
        if(Surcursal==2)and( Clausula== "si" or Clausula=="Si"or Clausula=="SI"):
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and( Clausula== "si" or Clausula=="Si"or Clausula=="SI")and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                titulo43='CLÁUSULA CUADRAGÉSIMA SEGUNDA. GASTOS DE COBRANZA: '
                p62.add_run(titulo43).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('EL ARRENDATARIO y/o DEUDORES SOLIDARIOS aceptan desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago,el 4% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 6% sobre el valor del canon; sin perjuicio de las acciones judiciales correspondientes.\t\n ')
            #__
                titulo44='CLÁUSULA CUADRAGÉSIMA TERCERA. FIRMAS: '
                p62.add_run(titulo44).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
            
            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and( Clausula== "si" or Clausula=="Si"or Clausula=="SI") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo43='CLÁUSULA CUADRAGÉSIMA TERCERA. GASTOS DE COBRANZA: '
                p62.add_run(titulo43).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('EL ARRENDATARIO y/o DEUDORES SOLIDARIOS aceptan desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago,el 4% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 6% sobre el valor del canon; sin perjuicio de las acciones judiciales correspondientes.\t\n')
            #__
                titulo44='CLÁUSULA CUADRAGÉSIMA CUARTA. FIRMAS: '
                p62.add_run(titulo44).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
            
            elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO" )and( Clausula== "si" or Clausula=="Si"or Clausula=="SI")and (Rta=="No" or Rta=="no"or Rta=="NO" ) and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                titulo42='CLÁUSULA CUADRAGÉSIMA PRIMERA. GASTOS DE COBRANZA: '
                p62.add_run(titulo42).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('EL ARRENDATARIO  acepta desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago,el 4% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 6% sobre el valor del canon; sin perjuicio de las acciones judiciales correspondientes.\t\n')     
            #__
                titulo44='CLÁUSULA CUADRAGÉSIMA SEGUNDA. FIRMAS: '
                p62.add_run(titulo44).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
            
            elif(Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and( Clausula== "si" or Clausula=="Si"or Clausula=="SI")and (Rta=="No" or Rta=="no"or Rta=="NO" ) and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo43='CLÁUSULA CUADRAGÉSIMA SEGUNDA. GASTOS DE COBRANZA: '
                p62.add_run(titulo43).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('EL ARRENDATARIO  acepta desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago,el 4% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 6% sobre el valor del canon; sin perjuicio de las acciones judiciales correspondientes.\t\n') 
            #__
                titulo44='CLÁUSULA CUADRAGÉSIMA TERCERA. FIRMAS: '
                p62.add_run(titulo44).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
            
        else:
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and( Clausula== "si" or Clausula=="Si"or Clausula=="SI")and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                titulo43='CLÁUSULA CUADRAGÉSIMA SEGUNDA. GASTOS DE COBRANZA: '
                p62.add_run(titulo43).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('EL ARRENDATARIO y/o DEUDORES SOLIDARIOS aceptan desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago, el 3% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 4% sobre el valor del canon; c) Si no se hace el pago en las fechas anteriormente estipuladas en los literales a y b, el recargo será del 6%; sin perjuicio de las acciones judiciales correspondientes.')
        #__
                titulo44='CLÁUSULA CUADRAGÉSIMA TERCERA. FIRMAS: '
                p62.add_run(titulo44).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
            
            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and( Clausula== "si" or Clausula=="Si"or Clausula=="SI")and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo43='CLÁUSULA CUADRAGÉSIMA TERCERA. GASTOS DE COBRANZA: '
                p62.add_run(titulo43).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('EL ARRENDATARIO y/o DEUDORES SOLIDARIOS aceptan desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago, el 3% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 4% sobre el valor del canon; c) Si no se hace el pago en las fechas anteriormente estipuladas en los literales a y b, el recargo será del 6%; sin perjuicio de las acciones judiciales correspondientes.')
        #_________
                titulo44='CLÁUSULA CUADRAGÉSIMA CUARTA. FIRMAS: '
                p62.add_run(titulo44).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
                
            elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and( Clausula== "si" or Clausula=="Si"or Clausula=="SI")and ( Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo43='CLÁUSULA CUADRAGÉSIMA PRIMERA. GASTOS DE COBRANZA: '
                p62.add_run(titulo43).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('EL ARRENDATARIO acepta desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago, el 3% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 4% sobre el valor del canon; c) Si no se hace el pago en las fechas anteriormente estipuladas en los literales a y b, el recargo será del 6%; sin perjuicio de las acciones judiciales correspondientes.\t\n')
        #_________
                titulo44='CLÁUSULA CUADRAGÉSIMA SEGUNDA. FIRMAS: '
                p62.add_run(titulo44).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
                
            elif(Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and( Clausula== "si" or Clausula=="Si"or Clausula=="SI")and ( Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                titulo43='CLÁUSULA CUADRAGÉSIMA SEGUNDA. GASTOS DE COBRANZA: '
                p62.add_run(titulo43).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('EL ARRENDATARIO acepta desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago, el 3% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 4% sobre el valor del canon; c) Si no se hace el pago en las fechas anteriormente estipuladas en los literales a y b, el recargo será del 6%; sin perjuicio de las acciones judiciales correspondientes.\t\n')
        #_________
                titulo44='CLÁUSULA CUADRAGÉSIMA TERCERA. FIRMAS: '
                p62.add_run(titulo44).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
                
            ##SIN ADMINISTRADOR
        if(Surcursal==2):
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and( Clausula== "si" or Clausula=="Si"or Clausula=="SI")and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                    titulo43='CLÁUSULA CUADRAGÉSIMA . GASTOS DE COBRANZA: '
                    p62.add_run(titulo43).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('EL ARRENDATARIO y/o DEUDORES SOLIDARIOS aceptan desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago,el 4% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 6% sobre el valor del canon; sin perjuicio de las acciones judiciales correspondientes.\t\n')
            #_________
                    titulo44='CLÁUSULA CUADRAGÉSIMA PRIMERA. FIRMAS: '
                    p62.add_run(titulo44).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
                    
            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and( Clausula== "si" or Clausula=="Si"or Clausula=="SI")and (Rta=="Si" or Rta=="si"or Rta=="SI")and(Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                    titulo43='CLÁUSULA CUADRAGÉSIMA PRIMERA. GASTOS DE COBRANZA: '
                    p62.add_run(titulo43).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('EL ARRENDATARIO y/o DEUDORES SOLIDARIOS aceptan desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago,el 4% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 6% sobre el valor del canon; sin perjuicio de las acciones judiciales correspondientes.\t\n')
            #_________
                    titulo44='CLÁUSULA CUADRAGÉSIMA SEGUNDA. FIRMAS: '
                    p62.add_run(titulo44).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
                    
            elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO" )and( Clausula== "si" or Clausula=="Si"or Clausula=="SI")and (Rta=="No" or Rta=="no"or Rta=="NO" ) and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                            titulo42='CLÁUSULA TRIGÉSIMA NOVENA . GASTOS DE COBRANZA: '
                            p62.add_run(titulo42).bold = True
                            p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p62.add_run('EL ARRENDATARIO  acepta desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago,el 4% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 6% sobre el valor del canon; sin perjuicio de las acciones judiciales correspondientes.\t\n')
            #_________
                            titulo44='CLÁUSULA CUADRAGÉSIMA . FIRMAS: '
                            p62.add_run(titulo44).bold = True
                            p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
                                
            elif(Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and( Clausula== "si" or Clausula=="Si"or Clausula=="SI")and (Rta=="No" or Rta=="no"or Rta=="NO" ) and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                        titulo42='CLÁUSULA CUADRAGÉSIMA PRIMERA. GASTOS DE COBRANZA: '
                        p62.add_run(titulo43).bold = True
                        p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p62.add_run('EL ARRENDATARIO  acepta desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago,el 4% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 6% sobre el valor del canon; sin perjuicio de las acciones judiciales correspondientes.\t\n')
            #_________
                        titulo44='CLÁUSULA CUADRAGÉSIMA SEGUNDA. FIRMAS: '
                        p62.add_run(titulo44).bold = True
                        p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
                                
        else:
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and( Clausula== "si" or Clausula=="Si"or Clausula=="SI")and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                    titulo43='CLÁUSULA CUADRAGÉSIMA . GASTOS DE COBRANZA: '
                    p62.add_run(titulo43).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('EL ARRENDATARIO y/o DEUDORES SOLIDARIOS aceptan desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago, el 3% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 4% sobre el valor del canon; c) Si no se hace el pago en las fechas anteriormente estipuladas en los literales a y b, el recargo será del 6%; sin perjuicio de las acciones judiciales correspondientes.\t\n\nPara constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
        #_________
                    titulo44='CLÁUSULA CUADRAGÉSIMA PRIMERA. FIRMAS: '
                    p62.add_run(titulo44).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
        
            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and( Clausula== "si" or Clausula=="Si"or Clausula=="SI")and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                    titulo43='CLÁUSULA CUADRAGÉSIMA PRIMERA. GASTOS DE COBRANZA: '
                    p62.add_run(titulo43).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('EL ARRENDATARIO y/o DEUDORES SOLIDARIOS aceptan desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago, el 3% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 4% sobre el valor del canon; c) Si no se hace el pago en las fechas anteriormente estipuladas en los literales a y b, el recargo será del 6%; sin perjuicio de las acciones judiciales correspondientes.')
        #_________
                    titulo44='CLÁUSULA CUADRAGÉSIMA SEGUNDA. FIRMAS: '
                    p62.add_run(titulo44).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
        
            elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and( Clausula== "si" or Clausula=="Si"or Clausula=="SI") and ( Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                        titulo43='CLÁUSULA CUADRAGÉSIMA . GASTOS DE COBRANZA: '
                        p62.add_run(titulo43).bold = True
                        p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p62.add_run('EL ARRENDATARIO acepta desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago, el 3% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 4% sobre el valor del canon; c) Si no se hace el pago en las fechas anteriormente estipuladas en los literales a y b, el recargo será del 6%; sin perjuicio de las acciones judiciales correspondientes.\t\n')
        #_________
                        titulo44='CLÁUSULA CUADRAGÉSIMA PRIMERA. FIRMAS: '
                        p62.add_run(titulo44).bold = True
                        p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
            
            elif(Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and( Clausula== "si" or Clausula=="Si"or Clausula=="SI")and ( Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                        titulo43='CLÁUSULA CUADRAGÉSIMA PRIMERA. GASTOS DE COBRANZA: '
                        p62.add_run(titulo43).bold = True
                        p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p62.add_run('EL ARRENDATARIO acepta desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago, el 3% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 4% sobre el valor del canon; c) Si no se hace el pago en las fechas anteriormente estipuladas en los literales a y b, el recargo será del 6%; sin perjuicio de las acciones judiciales correspondientes.\t\n')
        #_________
                        titulo44='CLÁUSULA CUADRAGÉSIMA SEGUNDA. FIRMAS: '
                        p62.add_run(titulo44).bold = True
                        p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
            
        #sin clausula
        if(Surcursal==2)and( Clausula== "no" or Clausula=="No"or Clausula=="NO"):
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and(Clausula== "no" or Clausula=="No"or Clausula=="NO")and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                titulo43='CLÁUSULA CUADRAGÉSIMA PRIMERA. GASTOS DE COBRANZA: '
                p62.add_run(titulo43).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('EL ARRENDATARIO y/o DEUDORES SOLIDARIOS aceptan desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago,el 4% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 6% sobre el valor del canon; sin perjuicio de las acciones judiciales correspondientes.\t\n ')
            #_________
                titulo44='CLÁUSULA CUADRAGÉSIMA SEGUNDA. FIRMAS: '
                p62.add_run(titulo44).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
                
            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and(Clausula== "no" or Clausula=="No"or Clausula=="NO") and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo43='CLÁUSULA CUADRAGÉSIMA SEGUNDA. GASTOS DE COBRANZA: '
                p62.add_run(titulo43).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('EL ARRENDATARIO y/o DEUDORES SOLIDARIOS aceptan desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago,el 4% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 6% sobre el valor del canon; sin perjuicio de las acciones judiciales correspondientes.\t\n')
            #_________
                titulo44='CLÁUSULA CUADRAGÉSIMA TERCERA. FIRMAS: '
                p62.add_run(titulo44).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
                
            elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO" )and( Clausula== "no" or Clausula=="No"or Clausula=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO" ) and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                titulo42='CLÁUSULA CUADRAGÉSIMA. GASTOS DE COBRANZA: '
                p62.add_run(titulo42).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('EL ARRENDATARIO  acepta desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago,el 4% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 6% sobre el valor del canon; sin perjuicio de las acciones judiciales correspondientes.\t\n')     
            #_________
                titulo44='CLÁUSULA CUADRAGÉSIMA PRIMERA. FIRMAS: '
                p62.add_run(titulo44).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
                
            elif(Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and( Clausula== "no" or Clausula=="No"or Clausula=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO" ) and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo43='CLÁUSULA CUADRAGÉSIMA PRIMERA. GASTOS DE COBRANZA: '
                p62.add_run(titulo43).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('EL ARRENDATARIO  acepta desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago,el 4% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 6% sobre el valor del canon; sin perjuicio de las acciones judiciales correspondientes.\t\n') 
            #_________
                titulo44='CLÁUSULA CUADRAGÉSIMA SEGUNDA. FIRMAS: '
                p62.add_run(titulo44).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
                
        else:
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and( Clausula== "no" or Clausula=="No"or Clausula=="NO")and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                titulo43='CLÁUSULA CUADRAGÉSIMA PRIMERA. GASTOS DE COBRANZA: '
                p62.add_run(titulo43).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('EL ARRENDATARIO y/o DEUDORES SOLIDARIOS aceptan desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago, el 3% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 4% sobre el valor del canon; c) Si no se hace el pago en las fechas anteriormente estipuladas en los literales a y b, el recargo será del 6%; sin perjuicio de las acciones judiciales correspondientes.')
        #_________
                titulo44='CLÁUSULA CUADRAGÉSIMA SEGUNDA. FIRMAS: '
                p62.add_run(titulo44).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
        
            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and( Clausula== "no" or Clausula=="No"or Clausula=="NO")and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo43='CLÁUSULA CUADRAGÉSIMA SEGUNDA. GASTOS DE COBRANZA: '
                p62.add_run(titulo43).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('EL ARRENDATARIO y/o DEUDORES SOLIDARIOS aceptan desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago, el 3% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 4% sobre el valor del canon; c) Si no se hace el pago en las fechas anteriormente estipuladas en los literales a y b, el recargo será del 6%; sin perjuicio de las acciones judiciales correspondientes.')
        #_________
                titulo44='CLÁUSULA CUADRAGÉSIMA TERCERA. FIRMAS: '
                p62.add_run(titulo44).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
            
            elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and( Clausula== "no" or Clausula=="No"or Clausula=="NO")and ( Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI"):
                titulo43='CLÁUSULA CUADRAGÉSIMA. GASTOS DE COBRANZA: '
                p62.add_run(titulo43).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('EL ARRENDATARIO acepta desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago, el 3% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 4% sobre el valor del canon; c) Si no se hace el pago en las fechas anteriormente estipuladas en los literales a y b, el recargo será del 6%; sin perjuicio de las acciones judiciales correspondientes.\t\n')
        #_________
                titulo44='CLÁUSULA CUADRAGÉSIMA PRIMERA. FIRMAS: '
                p62.add_run(titulo44).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
        
            elif(Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and( Clausula== "no" or Clausula=="No"or Clausula=="NO")and ( Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "Si" or Respuesta3=="si"or Respuesta3=="SI") :
                titulo43='CLÁUSULA CUADRAGÉSIMA PRIMERA. GASTOS DE COBRANZA: '
                p62.add_run(titulo43).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('EL ARRENDATARIO acepta desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago, el 3% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 4% sobre el valor del canon; c) Si no se hace el pago en las fechas anteriormente estipuladas en los literales a y b, el recargo será del 6%; sin perjuicio de las acciones judiciales correspondientes.\t\n')
        #_________
                titulo44='CLÁUSULA CUADRAGÉSIMA SEGUNDA. FIRMAS: '
                p62.add_run(titulo44).bold = True
                p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
        
            ##SIN ADMINISTRADOR
        if(Surcursal==2):
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and( Clausula== "no" or Clausula=="No"or Clausula=="NO")and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                    titulo43='CLÁUSULA TRIGÉSIMA NOVENA. GASTOS DE COBRANZA: '
                    p62.add_run(titulo43).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('EL ARRENDATARIO y/o DEUDORES SOLIDARIOS aceptan desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago,el 4% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 6% sobre el valor del canon; sin perjuicio de las acciones judiciales correspondientes.\t\n')
            #_________
                    titulo44='CLÁUSULA CUADRAGÉSIMA . FIRMAS: '
                    p62.add_run(titulo44).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
            
            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and( Clausula== "no" or Clausula=="No"or Clausula=="NO")and (Rta=="Si" or Rta=="si"or Rta=="SI")and(Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                    titulo43='CLÁUSULA CUADRAGÉSIMA. GASTOS DE COBRANZA: '
                    p62.add_run(titulo43).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('EL ARRENDATARIO y/o DEUDORES SOLIDARIOS aceptan desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago,el 4% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 6% sobre el valor del canon; sin perjuicio de las acciones judiciales correspondientes.\t\n')
            #_________
                    titulo44='CLÁUSULA CUADRAGÉSIMA PRIMERA. FIRMAS: '
                    p62.add_run(titulo44).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
            
            elif(Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO" )and(Clausula== "no" or Clausula=="No"or Clausula=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO" ) and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                            titulo42='CLÁUSULA TRIGÉSIMA OCTAVA. GASTOS DE COBRANZA: '
                            p62.add_run(titulo42).bold = True
                            p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p62.add_run('EL ARRENDATARIO  acepta desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago,el 4% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 6% sobre el valor del canon; sin perjuicio de las acciones judiciales correspondientes.\t\n')
            #_________
                            titulo44='CLÁUSULA TRIGÉSIMA NOVENA. FIRMAS: '
                            p62.add_run(titulo44).bold = True
                            p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
                        
            elif(Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si")and(Clausula== "no" or Clausula=="No"or Clausula=="NO")and (Rta=="No" or Rta=="no"or Rta=="NO" ) and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                        titulo42='CLÁUSULA CUADRAGÉSIMA. GASTOS DE COBRANZA: '
                        p62.add_run(titulo42).bold = True
                        p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p62.add_run('EL ARRENDATARIO  acepta desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago,el 4% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 6% sobre el valor del canon; sin perjuicio de las acciones judiciales correspondientes.\t\n')
            #_________
                        titulo44='CLÁUSULA CUADRAGÉSIMA PRIMERA. FIRMAS: '
                        p62.add_run(titulo44).bold = True
                        p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
                    
        else:
            if (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO") and( Clausula== "no" or Clausula=="No"or Clausula=="NO")and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                    titulo43='CLÁUSULA TRIGÉSIMA NOVENA. GASTOS DE COBRANZA: '
                    p62.add_run(titulo43).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('EL ARRENDATARIO y/o DEUDORES SOLIDARIOS aceptan desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago, el 3% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 4% sobre el valor del canon; c) Si no se hace el pago en las fechas anteriormente estipuladas en los literales a y b, el recargo será del 6%; sin perjuicio de las acciones judiciales correspondientes.\t\n\nPara constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
        #_________
                    titulo44='CLÁUSULA CUADRAGÉSIMA . FIRMAS: '
                    p62.add_run(titulo44).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
                
            elif (Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and( Clausula== "no" or Clausula=="No"or Clausula=="NO")and (Rta=="Si" or Rta=="si"or Rta=="SI")and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                    titulo43='CLÁUSULA CUADRAGÉSIMA. GASTOS DE COBRANZA: '
                    p62.add_run(titulo43).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('EL ARRENDATARIO y/o DEUDORES SOLIDARIOS aceptan desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago, el 3% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 4% sobre el valor del canon; c) Si no se hace el pago en las fechas anteriormente estipuladas en los literales a y b, el recargo será del 6%; sin perjuicio de las acciones judiciales correspondientes.')
        #_________
                    titulo44='CLÁUSULA CUADRAGÉSIMA PRIMERA. FIRMAS: '
                    p62.add_run(titulo44).bold = True
                    p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
                    
            elif (Respuesta1=="No" or Respuesta1=="no"or Respuesta1=="NO")and( Clausula== "no" or Clausula=="No"or Clausula=="NO") and ( Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO"):
                        titulo43='CLÁUSULA TRIGÉSIMA NOVENA. GASTOS DE COBRANZA: '
                        p62.add_run(titulo43).bold = True
                        p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p62.add_run('EL ARRENDATARIO acepta desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago, el 3% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 4% sobre el valor del canon; c) Si no se hace el pago en las fechas anteriormente estipuladas en los literales a y b, el recargo será del 6%; sin perjuicio de las acciones judiciales correspondientes.\t\n')
        #_________
                        titulo44='CLÁUSULA CUADRAGÉSIMA . FIRMAS: '
                        p62.add_run(titulo44).bold = True
                        p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
                
            elif(Respuesta1=="SI" or Respuesta1=="si"or Respuesta1=="Si") and(Clausula== "no" or Clausula=="No"or Clausula=="NO")and ( Rta=="No" or Rta=="no"or Rta=="NO" )and (Respuesta3== "No" or Respuesta3=="no"or Respuesta3=="NO") :
                        titulo43='CLÁUSULA CUADRAGÉSIMA. GASTOS DE COBRANZA: '
                        p62.add_run(titulo43).bold = True
                        p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p62.add_run('EL ARRENDATARIO acepta desde ahora que serán de su cargo los gastos de cobranza. generados por el incumplimiento en el pago del arrendamiento. Tales retardos en el pago del canon de arrendamiento dentro del término previsto en el presente contrato, ocasionará recargos en dinero de la siguiente manera: a) dentro de los cinco (5) días siguientes al vencimiento de la fecha oportuna de pago, el 3% sobre el valor del canon; b) dentro de los cinco (5) días siguientes a la segunda fecha estipulada en el literal “a”, el recargo será del 4% sobre el valor del canon; c) Si no se hace el pago en las fechas anteriormente estipuladas en los literales a y b, el recargo será del 6%; sin perjuicio de las acciones judiciales correspondientes.\t\n\t\n')
        #_________
                        titulo44='CLÁUSULA CUADRAGÉSIMA PRIMERA. FIRMAS: '
                        p62.add_run(titulo44).bold = True
                        p62.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p62.add_run('Las partes aclaran que el presente documento será firmado de forma electronica conforme el articulo 7 de la ley 527 de 1999 y soportado por el Decreto 2364, este contrato podrá ser firmado de manera digital por la plataforma Autentic y/o la que el Arrendador tenga para tal fin o podra gozar de una firma manuscrita sin que eso menoscabe su autenticidad ni validez; En el evento que sea firmado por medio electronico a este se anexará el certificado de aceptación de firma que la plataforma para tal fin genera.\t\n Para constancia el presente Contrato, una vez leído y aprobado por las partes, se firma en la ciudad de Bogotá D.C., al '+ Fecha_Firma_Contrato+' en '+Numero_de_Contrato+' ejemplares del mismo valor. \t \n\n ')
                    
        #___________________________________
        p67 = documentoObjeto.add_paragraph('')
        if Numero_Deudores==1:
            p67.add_run('ARRENDADOR: \n\n\n PROMOTORA INMOBILIARIA R&G S.A.S.\n N.I.T.: 800.239.928-9\n R. LEGAL: '+Nombre_Promotora+"\n C.C. No "+cc_rep_promotora+"\nDIRECCIÓN DE CONTACTO: "+direccion_promotora+"\n TELÉFONO DE CONTACTO: "+telefono_promotora+"\n CORREO ELECTRÓNICO: "+correo_promotora+"\n\n ARRENDATARIO: \n\n\n"+nombre_arrendatario+"\n"+Numero_id_arrendatario+" "+ciudadexpediciondiarrendatario+"\nDIRECCIÓN DE CONTACTO: "+direccion_arrendatario+"\nTELÉFONO DE CONTACTO: "+telefono_arrendatario+"\nCORREO ELECTRÓNICO: "+correo_arrendatario+"\n\n DEUDOR SOLIDARIO:\n\n\n"+Deudor_nombre[0]+"\n "+tipo_di_deudor_1_no_+":"+Deudor_idnumm[0]+"\n DIRECCIÓN DE CONTACTO:."+Deudor_direccion[0]+"\n TELÉFONO DE CONTACTO: "+Deudor_telefono[0]+"\n CORREO ELECTRÓNICO: "+Deudor_correo[0]+"\n").bold =True

        elif Numero_Deudores==2:
            p67.add_run('ARRENDADOR: \n\n\n PROMOTORA INMOBILIARIA R&G S.A.S.\n N.I.T.: 800.239.928-9\n R. LEGAL: '+Nombre_Promotora+"\n C.C. No "+cc_rep_promotora+"\nDIRECCIÓN DE CONTACTO: "+direccion_promotora+"\n TELÉFONO DE CONTACTO: "+telefono_promotora+"\n CORREO ELECTRÓNICO: "+correo_promotora+"\n\n ARRENDATARIO: \n\n\n"+nombre_arrendatario+"\n"+Numero_id_arrendatario+" "+ciudadexpediciondiarrendatario+"\nDIRECCIÓN DE CONTACTO: "+direccion_arrendatario+"\nTELÉFONO DE CONTACTO: "+telefono_arrendatario+"\nCORREO ELECTRÓNICO: "+correo_arrendatario+"\n\n DEUDOR SOLIDARIO:\n\n\n"+Deudor_nombre[0]+"\n "+tipo_di_deudor_1_no_+":"+Deudor_idnumm[0]+"\n DIRECCIÓN DE CONTACTO:."+Deudor_direccion[0]+"\n TELÉFONO DE CONTACTO: "+Deudor_telefono[0]+"\n CORREO ELECTRÓNICO: "+Deudor_correo[0]+"\n"+"\n\n DEUDOR SOLIDARIO:\n\n\n"+Deudor_nombre[2]+"\n "+tipo_di_deudor_2_no_+":"+Deudor_idnumm[2]+"\n DIRECCIÓN DE CONTACTO:."+Deudor_direccion[2]+"\n TELÉFONO DE CONTACTO: "+Deudor_telefono[2]+"\n CORREO ELECTRÓNICO: "+Deudor_correo[2]+"\n").bold =True
            
        elif Numero_Deudores==3:
            p67.add_run('ARRENDADOR: \n\n\n PROMOTORA INMOBILIARIA R&G S.A.S.\n N.I.T.: 800.239.928-9\n R. LEGAL: '+Nombre_Promotora+"\n C.C. No "+cc_rep_promotora+"\nDIRECCIÓN DE CONTACTO: "+direccion_promotora+"\n TELÉFONO DE CONTACTO: "+telefono_promotora+"\n CORREO ELECTRÓNICO: "+correo_promotora+"\n\n ARRENDATARIO: \n\n\n"+nombre_arrendatario+"\n"+Numero_id_arrendatario+" "+ciudadexpediciondiarrendatario+"\nDIRECCIÓN DE CONTACTO: "+direccion_arrendatario+"\nTELÉFONO DE CONTACTO: "+telefono_arrendatario+"\nCORREO ELECTRÓNICO: "+correo_arrendatario+"\n\n DEUDOR SOLIDARIO:\n\n\n"+Deudor_nombre[0]+"\n "+tipo_di_deudor_1_no_+":"+Deudor_idnumm[0]+"\n DIRECCIÓN DE CONTACTO:."+Deudor_direccion[0]+"\n TELÉFONO DE CONTACTO: "+Deudor_telefono[0]+"\n CORREO ELECTRÓNICO: "+Deudor_correo[0]+"\n"+"\n\n DEUDOR SOLIDARIO:\n\n\n"+Deudor_nombre[2]+"\n "+tipo_di_deudor_2_no_+":"+Deudor_idnumm[2]+"\n DIRECCIÓN DE CONTACTO:."+Deudor_direccion[2]+"\n TELÉFONO DE CONTACTO: "+Deudor_telefono[2]+"\n CORREO ELECTRÓNICO: "+Deudor_correo[2]+"\n"+"\n\n DEUDOR SOLIDARIO:\n\n\n"+Deudor_nombre[4]+"\n "+tipo_di_deudor_3_no_+":"+Deudor_idnumm[4]+"\n DIRECCIÓN DE CONTACTO:."+Deudor_direccion[4]+"\n TELÉFONO DE CONTACTO: "+Deudor_telefono[4]+"\n CORREO ELECTRÓNICO: "+Deudor_correo[4]+"\n").bold =True
            
        else:
            p67.add_run('ARRENDADOR: \n\n\n PROMOTORA INMOBILIARIA R&G S.A.S.\n N.I.T.: 800.239.928-9\n R. LEGAL: '+Nombre_Promotora+"\n C.C. No "+cc_rep_promotora+"\nDIRECCIÓN DE CONTACTO: "+direccion_promotora+"\n TELÉFONO DE CONTACTO: "+telefono_promotora+"\n CORREO ELECTRÓNICO: "+correo_promotora+"\n\n ARRENDATARIO: \n\n\n"+nombre_arrendatario+"\n"+Numero_id_arrendatario+" "+ciudadexpediciondiarrendatario+"\nDIRECCIÓN DE CONTACTO: "+direccion_arrendatario+"\nTELÉFONO DE CONTACTO: "+telefono_arrendatario+"\nCORREO ELECTRÓNICO: "+correo_arrendatario+"\n").bold =True

        try:
            # Intenta usar .get() si el objeto tiene este método
            nombre_archivo = Nombre_Archivo.get().strip()
        except AttributeError:
            # Si no tiene .get(), conviértelo a cadena
            nombre_archivo = str(Nombre_Archivo).strip()

        # Asegurarse de que el nombre no esté vacío
        if not nombre_archivo:
            print("Error: El nombre del archivo no puede estar vacío.")
        else:
            # Validar que termine con .docx; si no, añadir la extensión
            if not nombre_archivo.endswith('.docx'):
                nombre_archivo += '.docx'
            documentoObjeto.save(nombre_archivo)

        ventana.destroy()
#______________________________________________________

    except Exception as e:
        print(f'An exception occurred: {e}')
        traceback.print_exc()  # Imprime la traza completa del error
        messagebox.showinfo("Alert", "Verifique todos los espacios", icon="warning", parent=None)
        input("Terminar de crear Ariba Manualmente")
        print("**VOLVER A CORRER CODIGO**")


#__________________________________________________________________________________________________________________
def navegar_siguiente_pestania():
    # Obtener el índice de la pestaña actual
    indice_actual = notebook.index("current")
    # Navegar a la siguiente pestaña
    nuevo_indice = (indice_actual + 1) % notebook.index("end")  # Considera el ciclo de las pestañas
    notebook.select(nuevo_indice)
def navegar_si_existe_o_mostrar_arriba(tab_pages):
        # Crear el StringVar para controlar el CTkEntry
    Nombre_Archivo_var = customtkinter.StringVar()

    # Asociar la función on_text_change a los cambios en el StringVar
    Nombre_Archivo_var.trace_add("write", on_text_change)
    rta_value = Rta_.get()
    global Matricula_Garaje_,Chip_Garaje_,Numero_Garaje_,Deposito,Numero_Deposito_
    global nombre_deudor_1_,tipo_di_deudor_1_no__,Numero_id_deudor1_,di_deudor_1_ciudad_expedicion_di_deudor_1_,direccion_deudor_1_,telefono_deudor_1_,correo_deudor_1_,Numero_Deudores
    global nombre_deudor_2_,tipo_di_deudor_2_no__,Numero_id_deudor2_,di_deudor_2_ciudad_expedicion_di_deudor_2_,direccion_deudor_2_,telefono_deudor_2_,correo_deudor_2_
    global nombre_deudor_3_,tipo_di_deudor_3_no__,Numero_id_deudor3_,di_deudor_3_ciudad_expedicion_di_deudor_2_,direccion_deudor_3_,telefono_deudor_3_,correo_deudor_3_,Nombre_Archivo
    print("nombre", Nombre_Archivo_var.get())
    Surcursal_letra=Sucursal_.get()
    Surcursal2_letra=Representante.get()
    Respuesta1=Garaje.get()
    depo=Deposito.get()
    Respuesta3=Admin.get()
    Direccion_Inmueble=Direccion.get()
    Apartamento_y_torre=Apartamento_Torre.get()
    Nombre_edificio=NOMBRE_EDIFICIO.get()
    Ciudad=Ciudad_.get()
    valor_canon=Valor_Canon_.get()
    Canon_Letras=Valor_Canon_letra.get()
    valor_administracion=Valor_Admin_des.get()
    Administracion_Letras=Valor_Admin_des_letra.get()
    valor_total=valor_total_.get()
    valor_total_letra=valor_total_letras.get()
    numero_matricula=numero_matricula_.get()
    Chip=Chip_.get()
    valor_pleno_administracion=valor_pleno_administracion_.get()
    administracion_plena=administracion_plena_.get()
    nombre_arrendatario=nombre_arrendatario_.get()
    tipodiarrendatariono_=tipodiarrendatariono__.get()
    Numero_id_arrendatario=Numero_id_arrendatario_.get()
    ciudadexpediciondiarrendatario=ciudadexpediciondiarrendatario_.get()
    direccion_arrendatario =direccion_arrendatario_.get()
    telefono_arrendatario =telefono_arrendatario_.get()
    correo_arrendatario =correo_arrendatario_.get() 
    Cuenta_Acueducto=Cuenta_Acueducto_.get()
    Cuenta_Alcatarillado=Cuenta_Alcatarillado_.get()
    Cuenta_Basuras=Cuenta_Basuras_.get()
    Cuenta_Energia=Cuenta_Energia_.get()
    Cuenta_Gas=Cuenta_Gas_.get()
    Rta=Rta_.get()
    print("Cambio",rta_value)
    if  rta_value=="Si":
        notebook.select(tab_pages["Deudores"])
        Nombre_Archivo_var = customtkinter.StringVar()

    # Asociar la función on_text_change a los cambios en el StringVar
        Nombre_Archivo_var.trace_add("write", on_text_change)
    else:
        rta_value = Rta_.get()
        print("Cambio2",rta_value)
        Nombre_Archivo_var = customtkinter.StringVar()
    if 'Numero_Deposito_' in globals() and hasattr(Numero_Deposito_, 'get'):
        Numero_deposito = Numero_Deposito_.get()
    else:
        Numero_deposito = 0

    # Asociar la función on_text_change a los cambios en el StringVar
        Nombre_Archivo_var.trace_add("write", on_text_change)
        mostras_Ariba()
# Crear ventana principal
ventana = ctk.CTk()
ventana.title("Contratos")
ventana.geometry("600x400")
ventana.minsize(587, 475)
ventana.maxsize(587, 475)
ventana.resizable(0, 0)

# Crear Notebook para las pestañas
notebook = ttk.Notebook(ventana)
notebook.pack(expand=True, fill="both")

# Crear la primera pestaña
primera_pestania = ttk.Frame(notebook)
notebook.add(primera_pestania, text="Pestaña 1")
#
logo_ico=os.path.join(ruta_assets,'Logo-_1_.ico')
ventana.iconbitmap(logo_ico)
ventana.attributes()
#
#fondo2
# Load the image
image_fondo=os.path.join(ruta_assets,"Fondo1.jpg")
image=Image.open(image_fondo)
# Resize the image in the given (width, height)
original_width,original_height=image.size
image=image.resize((900,1000))
fondo2=ImageTk.PhotoImage(image)
l2=ctk.CTkLabel(
                            master=primera_pestania,
                            image=fondo2
)
l2.place(
        relx=0.5,
        rely=0.5,
        relwidth=1,
        relheight=1,
        anchor=tkinter.CENTER,

)

#creating custom frame
# Ajusta el canal alfa del color de fondo para hacerlo transparente
frame = customtkinter.CTkFrame(
    master=l2,
    width=520,
    height=350,
    corner_radius=15,
    fg_color="#1B1A1A",
    bg_color="#1B1A1A")
frame.place(
            relx=0.5,
            rely=0.5,
            anchor=tkinter.CENTER)

l3=customtkinter.CTkLabel(
                        master=frame,
                        text="""Creacion de contratos 2024""",
                        font=('SLBSans-Regular.woff',15),
                        justify='center',
)
l3.place(x=145, y=15)
#Crear espacio de texto usuario 4
#1)
l4=customtkinter.CTkLabel(
                        master=frame,
                        text="""Representantes Legales""",
                        font=('SLBSans-Regular.woff',13))
l4.place(x=22, y=42)
Representante=customtkinter.CTkComboBox(master=frame,
                                     values=['FREDDY HERNANDO GUERRERO RIVERA',
                                             "NELSON ALEJANDRO GUERRERO RIVERA",
                                             "CINDY JOHANNA GUERRERO RIVERA"],
                                     )
Representante.place(x=22, y=70)
##
#2)

#________
#3)
######################___
#1)
l7=customtkinter.CTkLabel(
                        master=frame,
                        text="""Sucursal""",
                        font=('SLBSans-Regular.woff',13))
l7.place(x=22, y=100)
Sucursal_=customtkinter.CTkComboBox(master=frame,
                                     values=['Sucursal salitre',
                                             "Sucursal Lago",
                                             "Sucursal Unicentro ","Sucursal multicentro"],
                                     )
Sucursal_.place(x=22, y=130)

l6=customtkinter.CTkLabel(
                        master=frame,
                        text="""Cuenta con garaje ?""",
                        font=('SLBSans-Regular.woff',13))
l6.place(x=22, y=160)
def Garajes(event):
    global Matricula_Garaje_,Chip_Garaje_,Numero_Garaje_,Deposito,Garaje
    if Garaje.get()=="Si":
        l11=customtkinter.CTkLabel(
                        master=frame,
                        text="""Matricula del garaje""",
                        font=('SLBSans-Regular.woff',13))
        l11.place(x=190, y=42)

        Matricula_Garaje_=customtkinter.CTkEntry(
                                master=frame,
                                width=140,
                                placeholder_text='Matricula Garaje')
        Matricula_Garaje_.place(x=190, y=70)

        l8=customtkinter.CTkLabel(
                                master=frame,
                                text="""Chip del garaje""",
                                font=('SLBSans-Regular.woff',13))
        l8.place(x=190, y=100)
        Chip_Garaje_=customtkinter.CTkEntry(
                                master=frame,
                                width=140,
                                placeholder_text='Chip del garaje')
        Chip_Garaje_.place(x=190, y=130)
        #________
        #3)
        l9=customtkinter.CTkLabel(
                                master=frame,
                                text="""Numero del garaje""",
                                font=('SLBSans-Regular.woff',13))
        l9.place(x=190, y=160)
        Numero_Garaje_=customtkinter.CTkEntry(
                                master=frame,
                                width=140,
                                placeholder_text='Numero del garaje')
        Numero_Garaje_.place(x=190, y=190)
        
        l10=customtkinter.CTkLabel(
                                master=frame,
                                text="""Cuenta con Dedposito?""",
                                font=('SLBSans-Regular.woff',13))
        l10.place(x=190, y=220)
        Deposito=customtkinter.CTkComboBox(master=frame,
                                     values=['Si/No','Si', 'No'],
                                     command=Numero_deposito,
                                     )
        Deposito.place(x=190, y=250)
    elif Garaje.get() == "No":
          Deposito=Garaje
          Numero_Deposito_ = ctk.StringVar(value=0) 
          print("1")
          frame2 = customtkinter.CTkFrame(
                master=frame,
                width=150,
                height=175,
                corner_radius=15,
                fg_color="transparent",
                bg_color="transparent")
          frame2.place(
                        relx=0.5,
                        rely=0.38,
                        anchor=tkinter.CENTER)
    else:
        print("F")
def Numero_deposito(event):
    global Deposito,Numero_Deposito_
    Numero_Deposito_ = ctk.StringVar(value=0)
    if Deposito.get()=="Si":
        l11=customtkinter.CTkLabel(
                                master=frame,
                                text="""Numero del Deposito""",
                                font=('SLBSans-Regular.woff',13))
        l11.place(x=360, y=220)
        Numero_Deposito_=customtkinter.CTkEntry(
                                master=frame,
                                width=140,
                                placeholder_text='Numero del deposito')
        Numero_Deposito_.place(x=360, y=250)

    elif Deposito.get() == "No":
        print("2")
        Numero_Deposito_ = ctk.StringVar(value=0)
        frame4 = customtkinter.CTkFrame(
                master=frame,
                width=170,
                height=70,
                corner_radius=15,
                fg_color="transparent",
                bg_color="transparent")
        frame4.place(
                        relx=0.8,
                        rely=0.72,
                        anchor=tkinter.CENTER)
        
Garaje=customtkinter.CTkComboBox(master=frame,
                                     values=['Si/No','Si', 'No'],
                                     command=Garajes,
                                     )
Garaje.place(x=22, y=190)
#____________________

#____________________

l12 = customtkinter.CTkLabel(
    master=frame,
    text="""Cuenta con Adminisrador?""",
    font=('SLBSans-Regular.woff', 13))
l12.place(x=360, y=42)
Admin = customtkinter.CTkComboBox(master=frame,
                                              values=['Si/No','Si', 'No'],
                                             #command=valor,
                                              )
Admin.place(x=360, y=70)
#____________________
l10=customtkinter.CTkLabel(
                        master=frame,
                        text="""Direccion del inmueble""",
                        font=('SLBSans-Regular.woff',13))
l10.place(x=360, y=100)
Direccion=customtkinter.CTkEntry(
                                master=frame,
                                width=140,
                                placeholder_text='Direccion inmueble')
Direccion.place(x=360, y=130)

l5=customtkinter.CTkLabel(
                        master=frame,
                        text="""Apartamento Y Torre""",
                        font=('SLBSans-Regular.woff',13))
l5.place(x=360, y=160)
Apartamento_Torre=customtkinter.CTkEntry(
                                master=frame,
                                width=140,
                                placeholder_text='APARTA/Torre')
Apartamento_Torre.place(x=360, y=190)
##
l13=customtkinter.CTkLabel(
                        master=frame,
                        text="""Nombre Edificio""",
                        font=('SLBSans-Regular.woff',13))
l13.place(x=22, y=220)
NOMBRE_EDIFICIO=customtkinter.CTkEntry(
                                master=frame,
                                width=140,
                                placeholder_text='NOMBRE EDIFICIO')
NOMBRE_EDIFICIO.place(x=22, y=250)
#Nueva pagina

##___________________
#Create custom button
button1 = customtkinter.CTkButton(
                                master=frame,
                                width=240,
                                text="Siguiente",
                                command=navegar_siguiente_pestania,
                                corner_radius=6,
                                fg_color="blue")
button1.place(x=140, y=300)
#_____________2 pestaña::::::::::::::

# Crear la primera pestaña
Segunda_pestania = ttk.Frame(notebook)
notebook.add(Segunda_pestania, text="Pestaña 2")
l20=ctk.CTkLabel(
                            master=Segunda_pestania,
                            image=fondo2
)
l20.place(
        relx=0.5,
        rely=0.5,
        relwidth=1,
        relheight=1,
        anchor=tkinter.CENTER,

)

#creating custom frame
# Ajusta el canal alfa del color de fondo para hacerlo transparente
frame3 = customtkinter.CTkFrame(
    master=l20,
    width=520,
    height=350,
    corner_radius=15,
    fg_color="#1B1A1A",
    bg_color="#1B1A1A")
frame3.place(
            relx=0.5,
            rely=0.5,
            anchor=tkinter.CENTER)

l22=customtkinter.CTkLabel(
                        master=frame3,
                        text="""Creacion de contratos 2024""",
                        font=('SLBSans-Regular.woff',15),
                        justify='center',
)
l22.place(x=145, y=15)
#Crear espacio de texto usuario 4
#1)
l23=customtkinter.CTkLabel(
                        master=frame3,
                        text="""Digite la Ciudad:""",
                        font=('SLBSans-Regular.woff',13))
l23.place(x=22, y=42)
Ciudad_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='Ciudad')
Ciudad_.place(x=22, y=70)

l24=customtkinter.CTkLabel(
                        master=frame3,
                        text="""valor canon""",
                        font=('SLBSans-Regular.woff',13))
l24.place(x=22, y=100)
Valor_Canon_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='valor canon')
Valor_Canon_.place(x=22, y=130)

l25=customtkinter.CTkLabel(
                        master=frame3,
                        text="""Cannon letra""",
                        font=('SLBSans-Regular.woff',13))
l25.place(x=22, y=160)
Valor_Canon_letra=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='Cannon letra')
Valor_Canon_letra.place(x=22, y=190)
l26=customtkinter.CTkLabel(
                        master=frame3,
                        text="""valor administracion""",
                        font=('SLBSans-Regular.woff',13))
l26.place(x=190, y=42)

Valor_Admin_des=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='administracion descuento')
Valor_Admin_des.place(x=190, y=70)

l8=customtkinter.CTkLabel(
                                master=frame3,
                                text="""valor administracion""",
                                font=('SLBSans-Regular.woff',13))
l8.place(x=190, y=100)
Valor_Admin_des_letra=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='administracion letras')
Valor_Admin_des_letra.place(x=190, y=130)
        #________
        #3)
l9=customtkinter.CTkLabel(
                                master=frame3,
                                text="""valor total""",
                                font=('SLBSans-Regular.woff',13))
l9.place(x=190, y=160)
valor_total_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='valor total')
valor_total_.place(x=190, y=190)

#____________________

l12 = customtkinter.CTkLabel(
    master=frame3,
    text="""valor total letras""",
    font=('SLBSans-Regular.woff', 13))
l12.place(x=360, y=42)
valor_total_letras=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='valor total letras')
valor_total_letras.place(x=360, y=70)
#____________________
l10=customtkinter.CTkLabel(
                        master=frame3,
                        text="""# Matircula inmueble""",
                        font=('SLBSans-Regular.woff',13))
l10.place(x=360, y=100)
numero_matricula_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='# Matircula inmueble')
numero_matricula_.place(x=360, y=130)

l5=customtkinter.CTkLabel(
                        master=frame3,
                        text="""# Del chip""",
                        font=('SLBSans-Regular.woff',13))
l5.place(x=360, y=160)
Chip_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='# Del chip')
Chip_.place(x=360, y=190)
##
l13=customtkinter.CTkLabel(
                        master=frame3,
                        text="""valor pleno administrativo""",
                        font=('SLBSans-Regular.woff',13))
l13.place(x=190, y=220)
valor_pleno_administracion_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='valor pleno')
valor_pleno_administracion_.place(x=190, y=250)
#Nueva pagina

##_____
button1 = customtkinter.CTkButton(
                                master=frame3,
                                width=240,
                                text="Siguiente",
                                command=navegar_siguiente_pestania,
                                corner_radius=6,
                                fg_color="blue")
button1.place(x=140, y=300)
#_____________3 pestaña::::::::::::::

# Crear la primera pestaña
Tercera_pestania = ttk.Frame(notebook)
notebook.add(Tercera_pestania, text="Pestaña 3")
l20=ctk.CTkLabel(
                            master=Tercera_pestania,
                            image=fondo2
)
l20.place(
        relx=0.5,
        rely=0.5,
        relwidth=1,
        relheight=1,
        anchor=tkinter.CENTER,

)

#creating custom frame
# Ajusta el canal alfa del color de fondo para hacerlo transparente
frame3 = customtkinter.CTkFrame(
    master=l20,
    width=520,
    height=350,
    corner_radius=15,
    fg_color="#1B1A1A",
    bg_color="#1B1A1A")
frame3.place(
            relx=0.5,
            rely=0.5,
            anchor=tkinter.CENTER)

l22=customtkinter.CTkLabel(
                        master=frame3,
                        text="""Creacion de contratos 2024""",
                        font=('SLBSans-Regular.woff',15),
                        justify='center',
)
l22.place(x=145, y=15)
#Crear espacio de texto usuario 4
#1)
l23=customtkinter.CTkLabel(
                        master=frame3,
                        text="""Administracion plena letra""",
                        font=('SLBSans-Regular.woff',13))
l23.place(x=22, y=42)
administracion_plena_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='Administracion plena')
administracion_plena_.place(x=22, y=70)

l24=customtkinter.CTkLabel(
                        master=frame3,
                        text="""Nombre arrendatario""",
                        font=('SLBSans-Regular.woff',13))
l24.place(x=22, y=100)
nombre_arrendatario_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='Nombre arrendatario')
nombre_arrendatario_.place(x=22, y=130)

l25=customtkinter.CTkLabel(
                        master=frame3,
                        text="""Tipo di arrendatario""",
                        font=('SLBSans-Regular.woff',13))
l25.place(x=22, y=160)
tipodiarrendatariono__=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='Tipo di arrendatario')
tipodiarrendatariono__.place(x=22, y=190)
l26=customtkinter.CTkLabel(
                        master=frame3,
                        text="""# Di arrendatario""",
                        font=('SLBSans-Regular.woff',13))
l26.place(x=190, y=42)

Numero_id_arrendatario_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='# Di arrendatario')
Numero_id_arrendatario_.place(x=190, y=70)

l8=customtkinter.CTkLabel(
                                master=frame3,
                                text="""Ciudad expedicion""",
                                font=('SLBSans-Regular.woff',13))
l8.place(x=190, y=100)
ciudadexpediciondiarrendatario_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='Expedicion arrendatario')
ciudadexpediciondiarrendatario_.place(x=190, y=130)
        #________
        #3)
l9=customtkinter.CTkLabel(
                                master=frame3,
                                text="""Direccion arrendatario""",
                                font=('SLBSans-Regular.woff',13))
l9.place(x=190, y=160)
direccion_arrendatario_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='Direccion arrendatario')
direccion_arrendatario_.place(x=190, y=190)

#____________________

l12 = customtkinter.CTkLabel(
    master=frame3,
    text="""Telefono arrendatario""",
    font=('SLBSans-Regular.woff', 13))
l12.place(x=360, y=42)
telefono_arrendatario_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='Telefono arrendatario')
telefono_arrendatario_.place(x=360, y=70)
#____________________
l10=customtkinter.CTkLabel(
                        master=frame3,
                        text="""Correo arrendatario""",
                        font=('SLBSans-Regular.woff',13))
l10.place(x=360, y=100)
correo_arrendatario_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='Correo arrendatario')
correo_arrendatario_.place(x=360, y=130)

l5=customtkinter.CTkLabel(
                        master=frame3,
                        text="""# Cuenta Acueducto""",
                        font=('SLBSans-Regular.woff',13))
l5.place(x=360, y=160)
Cuenta_Acueducto_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='# Cuenta Acueducto')
Cuenta_Acueducto_.place(x=360, y=190)
##
l13=customtkinter.CTkLabel(
                        master=frame3,
                        text="""#Cuenta Alcantarillado""",
                        font=('SLBSans-Regular.woff',13))
l13.place(x=190, y=220)
Cuenta_Alcatarillado_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='#Cuenta Alcantarillado')
Cuenta_Alcatarillado_.place(x=190, y=250)
#Nueva pagina

##_____
button1 = customtkinter.CTkButton(
                                master=frame3,
                                width=240,
                                text="Siguiente",
                                command=navegar_siguiente_pestania,
                                corner_radius=6,
                                fg_color="blue")
button1.place(x=140, y=300)
#_____________4 pestaña::::::::::::::

# Crear la primera pestaña
Cuarta_pestania = ttk.Frame(notebook)
notebook.add(Cuarta_pestania, text="Pestaña 4")
l20=ctk.CTkLabel(
                            master=Cuarta_pestania,
                            image=fondo2
)
l20.place(
        relx=0.5,
        rely=0.5,
        relwidth=1,
        relheight=1,
        anchor=tkinter.CENTER,

)

#creating custom frame
# Ajusta el canal alfa del color de fondo para hacerlo transparente
frame3 = customtkinter.CTkFrame(
    master=l20,
    width=520,
    height=350,
    corner_radius=15,
    fg_color="#1B1A1A",
    bg_color="#1B1A1A")
frame3.place(
            relx=0.5,
            rely=0.5,
            anchor=tkinter.CENTER)

l22=customtkinter.CTkLabel(
                        master=frame3,
                        text="""Creacion de contratos 2024""",
                        font=('SLBSans-Regular.woff',15),
                        justify='center',
)
l22.place(x=145, y=15)
#Crear espacio de texto usuario 4
#1)
l23=customtkinter.CTkLabel(
                        master=frame3,
                        text="""#Cuenta de Basuras""",
                        font=('SLBSans-Regular.woff',13))
l23.place(x=22, y=42)
Cuenta_Basuras_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='#Cuenta de Basuras')
Cuenta_Basuras_.place(x=22, y=70)

l24=customtkinter.CTkLabel(
                        master=frame3,
                        text="""#Cuenta energia electrica""",
                        font=('SLBSans-Regular.woff',13))
l24.place(x=22, y=100)
Cuenta_Energia_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='#Cuenta energia electrica')
Cuenta_Energia_.place(x=22, y=130)

l25=customtkinter.CTkLabel(
                        master=frame3,
                        text="""#Cuenta del Gas""",
                        font=('SLBSans-Regular.woff',13))
l25.place(x=22, y=160)
Cuenta_Gas_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='#Cuenta del Gas')
Cuenta_Gas_.place(x=22, y=190)
l26=customtkinter.CTkLabel(
                        master=frame3,
                        text="""Cuenta con deudores?""",
                        font=('SLBSans-Regular.woff',13))
l26.place(x=190, y=42) 
tab_pages = {}
def valor(event):
    global nombre_deudor_1_,tipo_di_deudor_1_no__,Numero_id_deudor1_,di_deudor_1_ciudad_expedicion_di_deudor_1_,direccion_deudor_1_,telefono_deudor_1_,correo_deudor_1_,Numero_Deudores_
    global nombre_deudor_2_,tipo_di_deudor_2_no__,Numero_id_deudor2_,di_deudor_2_ciudad_expedicion_di_deudor_2_,direccion_deudor_2_,telefono_deudor_2_,correo_deudor_2_
    global nombre_deudor_3_,tipo_di_deudor_3_no__,Numero_id_deudor3_,di_deudor_3_ciudad_expedicion_di_deudor_2_,direccion_deudor_3_,telefono_deudor_3_,correo_deudor_3_,Nombre_Archivo
    
    if Rta_.get()=="Si":  
        # Crear la primera pestaña
        print("entro")
                                        # Crear el StringVar para controlar el CTkEntry
        Nombre_Archivo_var = customtkinter.StringVar()

                # Asociar la función on_text_change a los cambios en el StringVar
        Nombre_Archivo_var.trace_add("write", on_text_change)
        sexta_pestania = ttk.Frame(notebook)
        notebook.add(sexta_pestania, text="Deudores")
        l20=ctk.CTkLabel(
                                    master=sexta_pestania,
                                    image=fondo2
        )
        l20.place(
                relx=0.5,
                rely=0.5,
                relwidth=1,
                relheight=1,
                anchor=tkinter.CENTER,

        )

        #creating custom frame
        # Ajusta el canal alfa del color de fondo para hacerlo transparente
        frame3 = customtkinter.CTkFrame(
            master=l20,
            width=520,
            height=350,
            corner_radius=15,
            fg_color="#1B1A1A",
            bg_color="#1B1A1A")
        frame3.place(
                    relx=0.5,
                    rely=0.5,
                    anchor=tkinter.CENTER)

        l22=customtkinter.CTkLabel(
                                master=frame3,
                                text="""Creacion de contratos 2024""",
                                font=('SLBSans-Regular.woff',15),
                                justify='center',
        )
        l22.place(x=145, y=15)
                # Agregar la nueva pestaña al diccionario de pestañas
        tab_pages["Deudores"] = sexta_pestania
        l8=customtkinter.CTkLabel(
                                        master=frame3,
                                        text="""Numero de deudores""",
                                        font=('SLBSans-Regular.woff',13))
        l8.place(x=22, y=42)
        def numeros_deu(event):
            global nombre_deudor_1_,tipo_di_deudor_1_no__,Numero_id_deudor1_,di_deudor_1_ciudad_expedicion_di_deudor_1_,direccion_deudor_1_,telefono_deudor_1_,correo_deudor_1_,Numero_Deudores_
            global nombre_deudor_2_,tipo_di_deudor_2_no__,Numero_id_deudor2_,di_deudor_2_ciudad_expedicion_di_deudor_2_,direccion_deudor_2_,telefono_deudor_2_,correo_deudor_2_
            global nombre_deudor_3_,tipo_di_deudor_3_no__,Numero_id_deudor3_,di_deudor_3_ciudad_expedicion_di_deudor_2_,direccion_deudor_3_,telefono_deudor_3_,correo_deudor_3_,Nombre_Archivo
            
            if Numero_Deudores_.get()=='1':
                                                # Crear el StringVar para controlar el CTkEntry
                Nombre_Archivo_var = customtkinter.StringVar()

                # Asociar la función on_text_change a los cambios en el StringVar
                Nombre_Archivo_var.trace_add("write", on_text_change)
                print("1")
                if "Deudor 2" in tab_pages:
                    # Obtener el índice de la pestaña 'Deudores'
                    index = notebook.index(tab_pages["Deudor 2"])
                    # Olvidar (eliminar) la pestaña 'Deudores' del notebook
                    notebook.forget(index)
                    # Eliminar la referencia de la pestaña 'Deudores' del diccionario
                    del tab_pages["Deudor 2"]
                    # Obtener el índice de la pestaña 'Deudores'
                    index2 = notebook.index(tab_pages["Deudor 3"])
                    # Olvidar (eliminar) la pestaña 'Deudores' del notebook
                    notebook.forget(index2)
                    # Eliminar la referencia de la pestaña 'Deudores' del diccionario
                    del tab_pages["Deudor 3"]
                l24=customtkinter.CTkLabel(
                        master=frame3,
                        text="""Nombre del deudor""",
                        font=('SLBSans-Regular.woff',13))
                l24.place(x=22, y=100)
                nombre_deudor_1_=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='Nombre deudor')
                nombre_deudor_1_.place(x=22, y=130)

                l25=customtkinter.CTkLabel(
                                        master=frame3,
                                        text="""Tipo di deudor""",
                                        font=('SLBSans-Regular.woff',13))
                l25.place(x=22, y=160)
                tipo_di_deudor_1_no__=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='Tipo di deudor')
                tipo_di_deudor_1_no__.place(x=22, y=190)
                l26=customtkinter.CTkLabel(
                                        master=frame3,
                                        text="""#DI deudor""",
                                        font=('SLBSans-Regular.woff',13))
                l26.place(x=190, y=42)

                Numero_id_deudor1_=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='#DI deudor')
                Numero_id_deudor1_.place(x=190, y=70)

                l8=customtkinter.CTkLabel(
                                                master=frame3,
                                                text="""Ciudad de expedicion di""",
                                                font=('SLBSans-Regular.woff',13))
                l8.place(x=190, y=100)
                di_deudor_1_ciudad_expedicion_di_deudor_1_=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='Expedicion di deudor')
                di_deudor_1_ciudad_expedicion_di_deudor_1_.place(x=190, y=130)
                        #________
                        #3)
                l9=customtkinter.CTkLabel(
                                                master=frame3,
                                                text="""Direccion deudor""",
                                                font=('SLBSans-Regular.woff',13))
                l9.place(x=190, y=160)
                direccion_deudor_1_=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='Direccion deudor')
                direccion_deudor_1_.place(x=190, y=190)

                #____________________

                l12 = customtkinter.CTkLabel(
                    master=frame3,
                    text="""Telefono deudor""",
                    font=('SLBSans-Regular.woff', 13))
                l12.place(x=360, y=42)
                telefono_deudor_1_=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='Telefono deudor')
                telefono_deudor_1_.place(x=360, y=70)
                #____________________
                l10=customtkinter.CTkLabel(
                                        master=frame3,
                                        text="""Correo deudor""",
                                        font=('SLBSans-Regular.woff',13))
                l10.place(x=360, y=100)
                correo_deudor_1_=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='Correo deudor')
                correo_deudor_1_.place(x=360, y=130)
                l13=customtkinter.CTkLabel(
                        master=frame3,
                        text="""Nombre Archivo""",
                        font=('SLBSans-Regular.woff',13))
                l13.place(x=190, y=220)
                Nombre_Archivo=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='Nombre Archivo')
                Nombre_Archivo.place(x=190, y=250)
                button1 = customtkinter.CTkButton(
                    master=frame3,
                    width=240,
                    text="Crear Word",
                    command=lambda: mostras_Ariba(),
                    corner_radius=6,
                fg_color="blue")
                button1.place(x=140, y=300)
            elif Numero_Deudores_.get()=='2':
                print("2")
                                                # Crear el StringVar para controlar el CTkEntry
                Nombre_Archivo_var = customtkinter.StringVar()

                # Asociar la función on_text_change a los cambios en el StringVar
                Nombre_Archivo_var.trace_add("write", on_text_change)
                if "Deudor 3" in tab_pages:
                    # Obtener el índice de la pestaña 'Deudores'
                    index = notebook.index(tab_pages["Deudor 3"])
                    # Olvidar (eliminar) la pestaña 'Deudores' del notebook
                    notebook.forget(index)
                    # Eliminar la referencia de la pestaña 'Deudores' del diccionario
                    del tab_pages["Deudor 3"]
                l24=customtkinter.CTkLabel(
                        master=frame3,
                        text="""Nombre del deudor""",
                        font=('SLBSans-Regular.woff',13))
                l24.place(x=22, y=100)
                nombre_deudor_1_=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='Nombre deudor')
                nombre_deudor_1_.place(x=22, y=130)

                l25=customtkinter.CTkLabel(
                                        master=frame3,
                                        text="""Tipo di deudor""",
                                        font=('SLBSans-Regular.woff',13))
                l25.place(x=22, y=160)
                tipo_di_deudor_1_no__=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='Tipo di deudor')
                tipo_di_deudor_1_no__.place(x=22, y=190)
                l26=customtkinter.CTkLabel(
                                        master=frame3,
                                        text="""#DI deudor""",
                                        font=('SLBSans-Regular.woff',13))
                l26.place(x=190, y=42)

                Numero_id_deudor1_=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='#DI deudor')
                Numero_id_deudor1_.place(x=190, y=70)

                l8=customtkinter.CTkLabel(
                                                master=frame3,
                                                text="""Ciudad de expedicion di""",
                                                font=('SLBSans-Regular.woff',13))
                l8.place(x=190, y=100)
                di_deudor_1_ciudad_expedicion_di_deudor_1_=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='Expedicion di deudor')
                di_deudor_1_ciudad_expedicion_di_deudor_1_.place(x=190, y=130)
                        #________
                        #3)
                l9=customtkinter.CTkLabel(
                                                master=frame3,
                                                text="""Direccion deudor""",
                                                font=('SLBSans-Regular.woff',13))
                l9.place(x=190, y=160)
                direccion_deudor_1_=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='Direccion deudor')
                direccion_deudor_1_.place(x=190, y=190)

                #____________________

                l12 = customtkinter.CTkLabel(
                    master=frame3,
                    text="""Telefono deudor""",
                    font=('SLBSans-Regular.woff', 13))
                l12.place(x=360, y=42)
                telefono_deudor_1_=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='Telefono deudor')
                telefono_deudor_1_.place(x=360, y=70)
                #____________________
                l10=customtkinter.CTkLabel(
                                        master=frame3,
                                        text="""Correo deudor""",
                                        font=('SLBSans-Regular.woff',13))
                l10.place(x=360, y=100)
                correo_deudor_1_=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='Correo deudor')
                correo_deudor_1_.place(x=360, y=130)
                button1 = customtkinter.CTkButton(
                master=frame3,
                width=240,
                text="siguiente",
                command=lambda: navegar_siguiente_pestania(),
                corner_radius=6,
                fg_color="blue")
                button1.place(x=140, y=300)
# #""deduor 2""______________________________________________________________________________________
                septima_pestania = ttk.Frame(notebook)
                notebook.add(septima_pestania, text="Deudor 2")
                tab_pages["Deudor 2"] = septima_pestania
                l20=ctk.CTkLabel(
                                            master=septima_pestania,
                                            image=fondo2
                )
                l20.place(
                        relx=0.5,
                        rely=0.5,
                        relwidth=1,
                        relheight=1,
                        anchor=tkinter.CENTER,

                )
                frame6 = customtkinter.CTkFrame(
                    master=l20,
                    width=520,
                    height=350,
                    corner_radius=15,
                    fg_color="#1B1A1A",
                    bg_color="#1B1A1A")
                frame6.place(
                            relx=0.5,
                            rely=0.5,
                            anchor=tkinter.CENTER)

                l22=customtkinter.CTkLabel(
                                        master=frame6,
                                        text="""Creacion de contratos 2024""",
                                        font=('SLBSans-Regular.woff',15),
                                        justify='center',
                )
                l22.place(x=145, y=15)
                l24=customtkinter.CTkLabel(
                        master=frame6,
                        text="""Nombre del deudor 2""",
                        font=('SLBSans-Regular.woff',13))
                l24.place(x=22, y=42)
                nombre_deudor_2_=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='Nombre deudor 2')
                nombre_deudor_2_.place(x=22, y=70)

                l25=customtkinter.CTkLabel(
                                        master=frame6,
                                        text="""Tipo di deudor 2""",
                                        font=('SLBSans-Regular.woff',13))
                l25.place(x=22, y=100)
                tipo_di_deudor_2_no__=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='Tipo di deudor 2')
                tipo_di_deudor_2_no__.place(x=22, y=130)
                l26=customtkinter.CTkLabel(
                                        master=frame6,
                                        text="""#DI deudor 2""",
                                        font=('SLBSans-Regular.woff',13))
                l26.place(x=190, y=42)

                Numero_id_deudor2_=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='#DI deudor 2')
                Numero_id_deudor2_.place(x=190, y=70)

                l8=customtkinter.CTkLabel(
                                                master=frame6,
                                                text="""Ciudad de expedicion di""",
                                                font=('SLBSans-Regular.woff',13))
                l8.place(x=190, y=100)
                di_deudor_2_ciudad_expedicion_di_deudor_2_=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='Expedicion di deudor')
                di_deudor_2_ciudad_expedicion_di_deudor_2_.place(x=190, y=130)
                        #________
                        #3)
                l9=customtkinter.CTkLabel(
                                                master=frame6,
                                                text="""Direccion deudor 2""",
                                                font=('SLBSans-Regular.woff',13))
                l9.place(x=190, y=160)
                direccion_deudor_2_=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='Direccion deudor 2')
                direccion_deudor_2_.place(x=190, y=190)

                #____________________

                l12 = customtkinter.CTkLabel(
                    master=frame6,
                    text="""Telefono deudor 2""",
                    font=('SLBSans-Regular.woff', 13))
                l12.place(x=360, y=42)
                telefono_deudor_2_=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='Telefono deudor 2')
                telefono_deudor_2_.place(x=360, y=70)
                #____________________
                l10=customtkinter.CTkLabel(
                                        master=frame6,
                                        text="""Correo deudor 2""",
                                        font=('SLBSans-Regular.woff',13))
                l10.place(x=360, y=100)
                correo_deudor_2_=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='Correo deudor 2')
                correo_deudor_2_.place(x=360, y=130)
                l13=customtkinter.CTkLabel(
                        master=frame6,
                        text="""Nombre Archivo""",
                        font=('SLBSans-Regular.woff',13))
                l13.place(x=190, y=220)
                Nombre_Archivo=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='Nombre Archivo')
                Nombre_Archivo.place(x=190, y=250)
                button1 = customtkinter.CTkButton(
                master=frame6,
                width=240,
                text="Crear Word",
                command=lambda: mostras_Ariba(),
                corner_radius=6,
                fg_color="blue")
                button1.place(x=140, y=300)
            elif Numero_Deudores_.get()=='3':
                print("3")
                                # Crear el StringVar para controlar el CTkEntry
                Nombre_Archivo_var = customtkinter.StringVar()

                # Asociar la función on_text_change a los cambios en el StringVar
                Nombre_Archivo_var.trace_add("write", on_text_change)
                if "Deudor 2" in tab_pages:
                    # Obtener el índice de la pestaña 'Deudores'
                    index = notebook.index(tab_pages["Deudor 2"])
                    # Olvidar (eliminar) la pestaña 'Deudores' del notebook
                    notebook.forget(index)
                    # Eliminar la referencia de la pestaña 'Deudores' del diccionario
                    del tab_pages["Deudor 2"]
                l24=customtkinter.CTkLabel(
                        master=frame3,
                        text="""Nombre del deudor""",
                        font=('SLBSans-Regular.woff',13))
                l24.place(x=22, y=100)
                nombre_deudor_1_=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='Nombre deudor')
                nombre_deudor_1_.place(x=22, y=130)

                l25=customtkinter.CTkLabel(
                                        master=frame3,
                                        text="""Tipo di deudor""",
                                        font=('SLBSans-Regular.woff',13))
                l25.place(x=22, y=160)
                tipo_di_deudor_1_no__=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='Tipo di deudor')
                tipo_di_deudor_1_no__.place(x=22, y=190)
                l26=customtkinter.CTkLabel(
                                        master=frame3,
                                        text="""#DI deudor""",
                                        font=('SLBSans-Regular.woff',13))
                l26.place(x=190, y=42)

                Numero_id_deudor1_=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='#DI deudor')
                Numero_id_deudor1_.place(x=190, y=70)

                l8=customtkinter.CTkLabel(
                                                master=frame3,
                                                text="""Ciudad de expedicion di""",
                                                font=('SLBSans-Regular.woff',13))
                l8.place(x=190, y=100)
                di_deudor_1_ciudad_expedicion_di_deudor_1_=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='Expedicion di deudor')
                di_deudor_1_ciudad_expedicion_di_deudor_1_.place(x=190, y=130)
                        #________
                        #3)
                l9=customtkinter.CTkLabel(
                                                master=frame3,
                                                text="""Direccion deudor""",
                                                font=('SLBSans-Regular.woff',13))
                l9.place(x=190, y=160)
                direccion_deudor_1_=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='Direccion deudor')
                direccion_deudor_1_.place(x=190, y=190)

                #____________________

                l12 = customtkinter.CTkLabel(
                    master=frame3,
                    text="""Telefono deudor""",
                    font=('SLBSans-Regular.woff', 13))
                l12.place(x=360, y=42)
                telefono_deudor_1_=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='Telefono deudor')
                telefono_deudor_1_.place(x=360, y=70)
                #____________________
                l10=customtkinter.CTkLabel(
                                        master=frame3,
                                        text="""Correo deudor""",
                                        font=('SLBSans-Regular.woff',13))
                l10.place(x=360, y=100)
                correo_deudor_1_=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='Correo deudor')
                correo_deudor_1_.place(x=360, y=130)
                button1 = customtkinter.CTkButton(
                master=frame3,
                width=240,
                text="siguiente",
                command=lambda: navegar_siguiente_pestania(),
                corner_radius=6,
                fg_color="blue")
                button1.place(x=140, y=300)
# #""deduor 2""______________________________________________________________________________________
                septima_pestania = ttk.Frame(notebook)
                notebook.add(septima_pestania, text="Deudor 2")
                tab_pages["Deudor 2"] = septima_pestania
                l20=ctk.CTkLabel(
                                            master=septima_pestania,
                                            image=fondo2
                )
                l20.place(
                        relx=0.5,
                        rely=0.5,
                        relwidth=1,
                        relheight=1,
                        anchor=tkinter.CENTER,

                )
                frame6 = customtkinter.CTkFrame(
                    master=l20,
                    width=520,
                    height=350,
                    corner_radius=15,
                    fg_color="#1B1A1A",
                    bg_color="#1B1A1A")
                frame6.place(
                            relx=0.5,
                            rely=0.5,
                            anchor=tkinter.CENTER)

                l22=customtkinter.CTkLabel(
                                        master=frame6,
                                        text="""Creacion de contratos 2024""",
                                        font=('SLBSans-Regular.woff',15),
                                        justify='center',
                )
                l22.place(x=145, y=15)
                l24=customtkinter.CTkLabel(
                        master=frame6,
                        text="""Nombre del deudor 2""",
                        font=('SLBSans-Regular.woff',13))
                l24.place(x=22, y=42)
                nombre_deudor_2_=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='Nombre deudor 2')
                nombre_deudor_2_.place(x=22, y=70)

                l25=customtkinter.CTkLabel(
                                        master=frame6,
                                        text="""Tipo di deudor 2""",
                                        font=('SLBSans-Regular.woff',13))
                l25.place(x=22, y=100)
                tipo_di_deudor_2_no__=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='Tipo di deudor 2')
                tipo_di_deudor_2_no__.place(x=22, y=130)
                l26=customtkinter.CTkLabel(
                                        master=frame6,
                                        text="""#DI deudor 2""",
                                        font=('SLBSans-Regular.woff',13))
                l26.place(x=190, y=42)

                Numero_id_deudor2_=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='#DI deudor 2')
                Numero_id_deudor2_.place(x=190, y=70)

                l8=customtkinter.CTkLabel(
                                                master=frame6,
                                                text="""Ciudad de expedicion di""",
                                                font=('SLBSans-Regular.woff',13))
                l8.place(x=190, y=100)
                di_deudor_2_ciudad_expedicion_di_deudor_2_=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='Expedicion di deudor')
                di_deudor_2_ciudad_expedicion_di_deudor_2_.place(x=190, y=130)
                        #________
                        #3)
                l9=customtkinter.CTkLabel(
                                                master=frame6,
                                                text="""Direccion deudor 2""",
                                                font=('SLBSans-Regular.woff',13))
                l9.place(x=190, y=160)
                direccion_deudor_2_=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='Direccion deudor 2')
                direccion_deudor_2_.place(x=190, y=190)

                #____________________

                l12 = customtkinter.CTkLabel(
                    master=frame6,
                    text="""Telefono deudor 2""",
                    font=('SLBSans-Regular.woff', 13))
                l12.place(x=360, y=42)
                telefono_deudor_2_=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='Telefono deudor 2')
                telefono_deudor_2_.place(x=360, y=70)
                #____________________
                l10=customtkinter.CTkLabel(
                                        master=frame6,
                                        text="""Correo deudor 2""",
                                        font=('SLBSans-Regular.woff',13))
                l10.place(x=360, y=100)
                correo_deudor_2_=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='Correo deudor 2')
                correo_deudor_2_.place(x=360, y=130)
                button1 = customtkinter.CTkButton(
                master=frame6,
                width=240,
                text="Siguiente",
                command=lambda: navegar_siguiente_pestania(),
                corner_radius=6,
                fg_color="blue")
                button1.place(x=140, y=300)
                # #""deduor 3""______________________________________________________________________________________
                Octava_pestania = ttk.Frame(notebook)
                notebook.add(Octava_pestania, text="Deudor 3")
                tab_pages["Deudor 3"] = Octava_pestania
                l20=ctk.CTkLabel(
                                            master=Octava_pestania,
                                            image=fondo2
                )
                l20.place(
                        relx=0.5,
                        rely=0.5,
                        relwidth=1,
                        relheight=1,
                        anchor=tkinter.CENTER,

                )
                frame6 = customtkinter.CTkFrame(
                    master=l20,
                    width=520,
                    height=350,
                    corner_radius=15,
                    fg_color="#1B1A1A",
                    bg_color="#1B1A1A")
                frame6.place(
                            relx=0.5,
                            rely=0.5,
                            anchor=tkinter.CENTER)

                l22=customtkinter.CTkLabel(
                                        master=frame6,
                                        text="""Creacion de contratos 2024""",
                                        font=('SLBSans-Regular.woff',15),
                                        justify='center',
                )
                l22.place(x=145, y=15)
                l24=customtkinter.CTkLabel(
                        master=frame6,
                        text="""Nombre del deudor 3""",
                        font=('SLBSans-Regular.woff',13))
                l24.place(x=22, y=42)
                nombre_deudor_3_=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='Nombre deudor 3')
                nombre_deudor_3_.place(x=22, y=70)

                l25=customtkinter.CTkLabel(
                                        master=frame6,
                                        text="""Tipo di deudor 3""",
                                        font=('SLBSans-Regular.woff',13))
                l25.place(x=22, y=100)
                tipo_di_deudor_3_no__=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='Tipo di deudor 3')
                tipo_di_deudor_3_no__.place(x=22, y=130)
                l26=customtkinter.CTkLabel(
                                        master=frame6,
                                        text="""#DI deudor 3""",
                                        font=('SLBSans-Regular.woff',13))
                l26.place(x=190, y=42)

                Numero_id_deudor3_=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='#DI deudor 3')
                Numero_id_deudor3_.place(x=190, y=70)

                l8=customtkinter.CTkLabel(
                                                master=frame6,
                                                text="""Ciudad de expedicion di""",
                                                font=('SLBSans-Regular.woff',13))
                l8.place(x=190, y=100)
                di_deudor_3_ciudad_expedicion_di_deudor_2_=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='Expedicion di deudor')
                di_deudor_3_ciudad_expedicion_di_deudor_2_.place(x=190, y=130)
                        #________
                        #3)
                l9=customtkinter.CTkLabel(
                                                master=frame6,
                                                text="""Direccion deudor 3""",
                                                font=('SLBSans-Regular.woff',13))
                l9.place(x=190, y=160)
                direccion_deudor_3_=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='Direccion deudor 3')
                direccion_deudor_3_.place(x=190, y=190)

                #____________________

                l12 = customtkinter.CTkLabel(
                    master=frame6,
                    text="""Telefono deudor 3""",
                    font=('SLBSans-Regular.woff', 13))
                l12.place(x=360, y=42)
                telefono_deudor_3_=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='Telefono deudor 3')
                telefono_deudor_3_.place(x=360, y=70)
                #____________________
                l10=customtkinter.CTkLabel(
                                        master=frame6,
                                        text="""Correo deudor 3""",
                                        font=('SLBSans-Regular.woff',13))
                l10.place(x=360, y=100)
                correo_deudor_3_=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='Correo deudor 3')
                correo_deudor_3_.place(x=360, y=130)
                l13=customtkinter.CTkLabel(
                        master=frame6,
                        text="""Nombre Archivo""",
                        font=('SLBSans-Regular.woff',13))
                l13.place(x=190, y=220)
                Nombre_Archivo=customtkinter.CTkEntry(
                                                master=frame6,
                                                width=140,
                                                placeholder_text='Nombre Archivo')
                Nombre_Archivo.place(x=190, y=250)
                button1 = customtkinter.CTkButton(
                master=frame6,
                width=240,
                text="Crear Word",
                command=lambda: mostras_Ariba(),
                corner_radius=6,
                fg_color="blue")
                button1.place(x=140, y=300)
            else:
                 print("1")                             # Crear el StringVar para controlar el CTkEntry
                 Nombre_Archivo_var = customtkinter.StringVar()

                # Asociar la función on_text_change a los cambios en el StringVar
                 Nombre_Archivo_var.trace_add("write", on_text_change)
        Numero_Deudores_= customtkinter.CTkComboBox(master=frame3,
                                                    values=['#Deudores','1', '2','3'],
                                                    command=numeros_deu,
                                                    )
        Numero_Deudores_.place(x=22, y=70)
    else:
        print("Borrando pestaña 'Deudores'")
        # Verificar si la pestaña 'Deudores' está en el diccionario de pestañas
                                  # Crear el StringVar para controlar el CTkEntry
        Nombre_Archivo_var = customtkinter.StringVar()
                # Asociar la función on_text_change a los cambios en el StringVar
        Nombre_Archivo_var.trace_add("write", on_text_change)
        if "Deudores" in tab_pages:
            # Obtener el índice de la pestaña 'Deudores'
            index = notebook.index(tab_pages["Deudores"])
            # Olvidar (eliminar) la pestaña 'Deudores' del notebook
            notebook.forget(index)
            # Eliminar la referencia de la pestaña 'Deudores' del diccionario
            del tab_pages["Deudores"]


Rta_= customtkinter.CTkComboBox(master=frame3,
                                              values=['Si/No','Si', 'No'],
                                              command=valor,
                                              )
Rta_.place(x=190, y=70)

l8=customtkinter.CTkLabel(
                                        master=frame3,
                                        text="""#Copias del contrato""",
                                        font=('SLBSans-Regular.woff',13))
l8.place(x=190, y=100)
Numero_de_Contrato_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='#Copias del contrato')
Numero_de_Contrato_.place(x=190, y=130)
l9=customtkinter.CTkLabel(
                                master=frame3,
                                text="""Fecha inicio contrato""",
                                font=('SLBSans-Regular.woff',13))
l9.place(x=190, y=160)
Fecha_inicio_Contrato_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='01/23/2000')
Fecha_inicio_Contrato_.place(x=190, y=190)

#____________________

l12 = customtkinter.CTkLabel(
    master=frame3,
    text="""Fecha firma contrato""",
    font=('SLBSans-Regular.woff', 13))
l12.place(x=360, y=42)
Fecha_Firma_Contrato_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='01/23/2000')
Fecha_Firma_Contrato_.place(x=360, y=70)
#____________________
l10=customtkinter.CTkLabel(
                        master=frame3,
                        text="""vigencia del contrato""",
                        font=('SLBSans-Regular.woff',13))
l10.place(x=360, y=100)
Vigencia_Contrato_=customtkinter.CTkEntry(
                                master=frame3,
                                width=140,
                                placeholder_text='vigencia del contrato')
Vigencia_Contrato_.place(x=360, y=130)

l5=customtkinter.CTkLabel(
                        master=frame3,
                        text="""Incluir clausula\n'ESPACIOS EN BLANCO'""",
                        font=('SLBSans-Regular.woff',13))
l5.place(x=360, y=160)
Clausula_= customtkinter.CTkComboBox(master=frame3,
                                              values=['Si/No','Si', 'No'],
                                              )
Clausula_.place(x=360, y=200)
##
def on_text_change(*args):
    global Nombre_Archivo,Deposito
    print("El texto ha cambiado a:", Nombre_Archivo_var.get())
    print("El texto ha cambiado a:", Deposito.get())
    Nombre_Archivo=Nombre_Archivo_var.get()
# Crear el StringVar para controlar el CTkEntry
Nombre_Archivo_var = customtkinter.StringVar()

# Asociar la función on_text_change a los cambios en el StringVar
Nombre_Archivo_var.trace_add("write", on_text_change)

l13=customtkinter.CTkLabel(
                        master=frame3,
                        text="""Nombre Archivo""",
                        font=('SLBSans-Regular.woff',13))
l13.place(x=190, y=220)
Nombre_Archivo=customtkinter.CTkEntry(
                                                master=frame3,
                                                width=140,
                                                placeholder_text='Nombre Archivo',
    textvariable=Nombre_Archivo_var)
Nombre_Archivo.place(x=190, y=250)
#Nueva pagina

##_____
button1 = customtkinter.CTkButton(
            master=frame3,
            width=240,
            text="siguiente",
            command=lambda: navegar_si_existe_o_mostrar_arriba(tab_pages),
            corner_radius=6,
            fg_color="blue")
button1.place(x=140, y=300)

ventana.mainloop()

