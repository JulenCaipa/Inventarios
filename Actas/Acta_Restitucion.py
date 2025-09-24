import customtkinter as ctk
import tkinter
from tkinter import messagebox
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")

class RestitucionApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Acta de Restitución")
        self.geometry("700x900")
        self.minsize(700, 900)
        self.resizable(True, True)

        # Crear el sistema de pestañas
        self.tabview = ctk.CTkTabview(self, width=680, height=860)
        self.tabview.pack(fill="both", expand=True, padx=10, pady=10)

        # Crear las pestañas
        self.tabview.add("Datos Generales")
        self.tabview.add("Detalles y Entrega")

        # --- PESTAÑA 1: DATOS GENERALES ---
        tab1 = self.tabview.tab("Datos Generales")
        
        self.label_title1 = ctk.CTkLabel(tab1, text="ACTA DE RESTITUCIÓN", font=("Arial Narrow", 22, "bold"))
        self.label_title1.pack(pady=(20, 15))
        
        self.nombre_archivo = ctk.CTkEntry(tab1, placeholder_text="Nombre del archivo (ej: ACTA_RESTITUCION_FINAL.docx)")
        self.nombre_archivo.pack(pady=5, padx=15, fill="x")
        
        self.sucursal_var = ctk.StringVar()
        self.label_sucursal = ctk.CTkLabel(tab1, text="Representante:")
        self.label_sucursal.pack(anchor="w", padx=15, pady=(10,0))
        self.sucursal_dropdown = ctk.CTkComboBox(tab1, variable=self.sucursal_var, values=[
            "FREDDY HERNANDO GUERRERO RIVERA",
            "NELSON ALEJANDRO GUERRERO RIVERA",
            "CINDY JOHANNA GUERRERO RIVERA"
        ])
        self.sucursal_dropdown.pack(pady=5, padx=15, fill="x")
        
        self.ciudad = ctk.CTkEntry(tab1, placeholder_text="Ciudad")
        self.ciudad.pack(pady=5, padx=15, fill="x")
        
        self.fecha = ctk.CTkEntry(tab1, placeholder_text="Fecha (ej: diecinueve (19) de agosto del año 2025)")
        self.fecha.pack(pady=5, padx=15, fill="x")
        
        self.direccion = ctk.CTkEntry(tab1, placeholder_text="Dirección del inmueble")
        self.direccion.pack(pady=5, padx=15, fill="x")
        
        self.arrendatario = ctk.CTkEntry(tab1, placeholder_text="Nombre arrendatario")
        self.arrendatario.pack(pady=5, padx=15, fill="x")
        
        self.tipo_doc_arr = ctk.CTkEntry(tab1, placeholder_text="Tipo documento arrendatario")
        self.tipo_doc_arr.pack(pady=5, padx=15, fill="x")
        
        self.cc_arr = ctk.CTkEntry(tab1, placeholder_text="Número documento arrendatario")
        self.cc_arr.pack(pady=5, padx=15, fill="x")
        
        self.ciudad_doc_arr = ctk.CTkEntry(tab1, placeholder_text="Ciudad de expedición documento arrendatario")
        self.ciudad_doc_arr.pack(pady=5, padx=15, fill="x")
        
        self.autorizado_nombre = ctk.CTkEntry(tab1, placeholder_text="Nombre autorizado (si aplica)")
        self.autorizado_nombre.pack(pady=5, padx=15, fill="x")
        
        self.tipo_doc_aut = ctk.CTkEntry(tab1, placeholder_text="Tipo documento autorizado")
        self.tipo_doc_aut.pack(pady=5, padx=15, fill="x")
        
        self.autorizado_cc = ctk.CTkEntry(tab1, placeholder_text="Número documento autorizado")
        self.autorizado_cc.pack(pady=5, padx=15, fill="x")

        # --- PESTAÑA 2: DETALLES Y ENTREGA ---
        tab2 = self.tabview.tab("Detalles y Entrega")
        
        self.label_title2 = ctk.CTkLabel(tab2, text="DETALLES DE ENTREGA", font=("Arial Narrow", 20, "bold"))
        self.label_title2.pack(pady=(20, 15))
        
        self.juegos_llaves = ctk.CTkEntry(tab2, placeholder_text="Número de juegos de llaves convencionales")
        self.juegos_llaves.pack(pady=5, padx=15, fill="x")
        
        self.llaves_seguridad = ctk.CTkEntry(tab2, placeholder_text="Número de llaves de seguridad")
        self.llaves_seguridad.pack(pady=5, padx=15, fill="x")
        
        self.otros_entregas = ctk.CTkEntry(tab2, placeholder_text="Otros entregados (pines, tarjetas, etc.)")
        self.otros_entregas.pack(pady=5, padx=15, fill="x")
        
        self.traslado_internet_var = ctk.StringVar()
        self.label_traslado_internet = ctk.CTkLabel(tab2, text="Traslado internet y similares (SI/NO):")
        self.label_traslado_internet.pack(anchor="w", padx=15, pady=(10,0))
        self.traslado_internet = ctk.CTkComboBox(tab2, variable=self.traslado_internet_var, values=["SI", "NO"])
        self.traslado_internet.pack(pady=5, padx=15, fill="x")
        
        self.observaciones_internet = ctk.CTkEntry(tab2, placeholder_text="Observaciones internet")
        self.observaciones_internet.pack(pady=5, padx=15, fill="x")
        
        self.traslado_creditos_var = ctk.StringVar()
        self.label_traslado_creditos = ctk.CTkLabel(tab2, text="Traslado créditos servicios públicos (SI/NO):")
        self.label_traslado_creditos.pack(anchor="w", padx=15, pady=(10,0))
        self.traslado_creditos = ctk.CTkComboBox(tab2, variable=self.traslado_creditos_var, values=["SI", "NO"])
        self.traslado_creditos.pack(pady=5, padx=15, fill="x")
        
        self.observaciones_creditos = ctk.CTkEntry(tab2, placeholder_text="Observaciones créditos")
        self.observaciones_creditos.pack(pady=5, padx=15, fill="x")
        
        self.resanado_var = ctk.StringVar()
        self.label_resanado = ctk.CTkLabel(tab2, text="¿El inmueble se recibe resanado? (SI/NO):")
        self.label_resanado.pack(anchor="w", padx=15, pady=(10,0))
        self.resanado = ctk.CTkComboBox(tab2, variable=self.resanado_var, values=["SI", "NO"])
        self.resanado.pack(pady=5, padx=15, fill="x")
        
        self.pintado_var = ctk.StringVar()
        self.label_pintado = ctk.CTkLabel(tab2, text="¿El inmueble se recibe recién pintado? (SI/NO):")
        self.label_pintado.pack(anchor="w", padx=15, pady=(10,0))
        self.pintado = ctk.CTkComboBox(tab2, variable=self.pintado_var, values=["SI", "NO"])
        self.pintado.pack(pady=5, padx=15, fill="x")
        
        self.aseado_var = ctk.StringVar()
        self.label_aseado = ctk.CTkLabel(tab2, text="¿El inmueble se recibe aseado? (SI/NO):")
        self.label_aseado.pack(anchor="w", padx=15, pady=(10,0))
        self.aseado = ctk.CTkComboBox(tab2, variable=self.aseado_var, values=["SI", "NO"])
        self.aseado.pack(pady=5, padx=15, fill="x")
        
        self.observaciones_finales = ctk.CTkEntry(tab2, placeholder_text="Novedades / Observaciones adicionales")
        self.observaciones_finales.pack(pady=5, padx=15, fill="x")
        
        self.nombre_asesor = ctk.CTkEntry(tab2, placeholder_text="Nombre asesor inmobiliario")
        self.nombre_asesor.pack(pady=5, padx=15, fill="x")
        
        # Botón de generar al final de la segunda pestaña
        self.btn_generar = ctk.CTkButton(tab2, text="Generar Acta", width=150, height=35, 
                                        font=("Arial", 14, "bold"), command=self.generar_acta)
        self.btn_generar.pack(pady=30)

    def generar_acta(self):
        archivo = self.nombre_archivo.get().strip()
        if archivo and not archivo.lower().endswith('.docx'):
            archivo += '.docx'
        Surcursal2_letra = self.sucursal_var.get().strip()
        ciudad = self.ciudad.get().strip()
        fecha = self.fecha.get().strip()
        direccion_inmueble = self.direccion.get().strip()
        arrendatario = self.arrendatario.get().strip()
        tipo_doc_arrendatario = self.tipo_doc_arr.get().strip()
        cc_arrendatario = self.cc_arr.get().strip()
        ciudad_doc_arrendatario = self.ciudad_doc_arr.get().strip()
        autorizado_nombre = self.autorizado_nombre.get().strip()
        tipo_doc_autorizado = self.tipo_doc_aut.get().strip()
        autorizado_cc = self.autorizado_cc.get().strip()
        juegos_llaves = self.juegos_llaves.get().strip()
        llaves_seguridad = self.llaves_seguridad.get().strip()
        otros_entregas = self.otros_entregas.get().strip()
        traslado_internet = self.traslado_internet_var.get().strip()
        observaciones_internet = self.observaciones_internet.get().strip()
        traslado_creditos = self.traslado_creditos_var.get().strip()
        observaciones_creditos = self.observaciones_creditos.get().strip()
        resanado = self.resanado_var.get().strip()
        pintado = self.pintado_var.get().strip()
        aseado = self.aseado_var.get().strip()
        observaciones_finales = self.observaciones_finales.get().strip()
        asesor = self.nombre_asesor.get().strip()

        if not archivo:
            messagebox.showerror("Error", "Debes ingresar el nombre del archivo.")
            return
        if not Surcursal2_letra or not ciudad or not fecha or not direccion_inmueble or not arrendatario or not tipo_doc_arrendatario or not cc_arrendatario or not ciudad_doc_arrendatario or not asesor:
            messagebox.showerror("Error", "Todos los campos obligatorios deben estar completos.")
            return

        # Lógica de asignación de representante
        if Surcursal2_letra == 'FREDDY HERNANDO GUERRERO RIVERA':
            representante = "FREDY HERNANDO GUERRERO RIVERA"
            cc_representante = "79.509.071"
        elif Surcursal2_letra == "NELSON ALEJANDRO GUERRERO RIVERA":
            representante = "NELSON ALEJANDRO GUERRERO RIVERA"
            cc_representante = "79.946.997"
        elif Surcursal2_letra == "CINDY JOHANNA GUERRERO RIVERA":
            representante = "CINDY JOHANNA GUERRERO RIVERA"
            cc_representante = "1.030.552.945"
        else:
            messagebox.showerror("Error", "Selecciona un representante válido.")
            return

        tipo_doc_representante = "CC. "

        try:
            documento = docx.Document()
            style = documento.styles['Normal']
            style.font.name = 'Arial Narrow'
            style.font.size = Pt(11)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            sections = documento.sections
            for section in sections:
                section.top_margin = Cm(3)
                section.bottom_margin = Cm(2.5)
                section.left_margin = Cm(3)
                section.right_margin = Cm(3)

            # Título
            p_titulo = documento.add_paragraph()
            p_titulo.add_run("ACTA DE RESTITUCIÓN").bold = True
            p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Párrafo inicial
            p0 = documento.add_paragraph()
            p0.add_run(f"En {ciudad}, el día {fecha}, ").bold = False
            p0.add_run("PROMOTORA INMOBILIARIA R&G S.A.S. ").bold = True
            p0.add_run("identificados con ").bold = False
            p0.add_run("NIT: 800.239.928-9").bold = True
            p0.add_run(", sociedad representada legalmente por ")
            p0.add_run(representante).bold = True
            p0.add_run(f", identificado con {tipo_doc_representante} No. ")
            p0.add_run(cc_representante).bold=True
            p0.add_run(f" de {ciudad}, recibe el inmueble ubicado en ")
            p0.add_run(direccion_inmueble).bold = True
            p0.add_run(f", en la ciudad de {ciudad}, a ")
            p0.add_run(arrendatario).bold = True
            p0.add_run(f", identificado(a) con {tipo_doc_arrendatario} No. {cc_arrendatario} {ciudad_doc_arrendatario}.")

            # Texto previo a viñetas con ARRENDATARIO en negrilla
            p1 = documento.add_paragraph()
            p1.add_run("En caso de ausencia, El ")
            p1.add_run("ARRENDATARIO").bold = True
            p1.add_run(" de manera explícita autoriza a la diligencia de restitución de inmueble a la siguiente persona:")

            # Autorización y entregas en viñetas
            def agregar_vineta(textos):
                p = documento.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.first_line_indent = Cm(-0.5)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = p.add_run("•\t")
                for texto, bold in textos:
                    run = p.add_run(texto)
                    run.bold = bold
                return p

            agregar_vineta([
                (f"El(la) señor(a) {autorizado_nombre}, identificado(a) con {tipo_doc_autorizado} No. {autorizado_cc}, "
                "a quien se le recibe real y material del inmueble, de igual forma también se reciben:", False)
            ])
            agregar_vineta([(f"Juegos de llaves convencionales {juegos_llaves}", False)])
            agregar_vineta([(f"Llaves de seguridad {llaves_seguridad}", False)])
            agregar_vineta([(f"Otros (pines, tarjetas, controles) {otros_entregas}", False)])

            # Nota completa
            nota = documento.add_paragraph()
            nota.add_run("NOTA: ").bold = True
            nota.add_run(
                "Sr Arrendatario (a), recuerde que previo a esta restitución se hizo una pre visita, "
                "y si en su caso quedaron recomendaciones o ajustes por realizar, en caso de no encontrar el inmueble a conformidad, "
                "el asesor podrá retirarse del inmueble, por tanto, el curso de los días seguiría su cobro hasta que finalmente "
                "se restituya el inmueble debidamente organizado."
            )

            # Paz y salvo / traslados con SI/NO
            p = documento.add_paragraph()
            p.add_run("SE RECIBE PAZ Y SALVO Y/O TRASLADO DE LINEA TELEFONICA INTERNET Y SIMILARES:").bold = True
            p = documento.add_paragraph()
            p.add_run(self.marcar_si_no(traslado_internet))
            p = documento.add_paragraph()
            p.add_run("OBSERVACIONES:").bold = True
            documento.add_paragraph(observaciones_internet)
            p = documento.add_paragraph()
            p.add_run("SE RECIBE PAZ Y SALVO Y/O TRASLADO DE LOS CREDITOS TOMADOS Y CARGADOS A ALGUN SERVICIO PUBLICO:").bold = True
            p = documento.add_paragraph()
            p.add_run(self.marcar_si_no(traslado_creditos))
            p = documento.add_paragraph()
            p.add_run("OBSERVACIONES:").bold = True
            documento.add_paragraph(observaciones_creditos)

            # Condiciones inmueble con X
            documento.add_paragraph(f"¿El inmueble se recibe resanado?   {self.marcar_si_no(resanado)}")
            documento.add_paragraph(f"¿El inmueble se recibe recién pintado?   {self.marcar_si_no(pintado)}")
            documento.add_paragraph(f"¿El inmueble se recibe aseado?   {self.marcar_si_no(aseado)}")

            # Observaciones adicionales
            documento.add_paragraph(
                "Además de las observaciones especificadas anteriormente, se dejan por escrito las siguientes novedades evidenciadas que tienen que ver con la restitución del inmueble:"
            )
            documento.add_paragraph(observaciones_finales)

            def quitar_margenes_celda(cell, top=0, left=0, bottom=0, right=0):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcMar = OxmlElement('w:tcMar')
                for m, v in zip(("top", "left", "bottom", "right"), (top, left, bottom, right)):
                    node = OxmlElement(f"w:{m}")
                    node.set(qn('w:w'), str(v))
                    node.set(qn('w:type'), 'dxa')
                    tcMar.append(node)
                tcPr.append(tcMar)

            def quitar_espaciado_parrafo(paragraph):
                p_fmt = paragraph.paragraph_format
                p_fmt.space_before = Pt(0)
                p_fmt.space_after = Pt(0)
                p_fmt.line_spacing = 1.0

            # --- Firmas en tabla ---
            tabla = documento.add_table(rows=3, cols=2)
            tabla.autofit = True
            for row in tabla.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = OxmlElement('w:tcBorders')
                    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
                        border = OxmlElement(f"w:{border_name}")
                        border.set(qn("w:val"), "nil")
                        tcBorders.append(border)
                    tcPr.append(tcBorders)
                    quitar_margenes_celda(cell, top=50, left=50, bottom=50, right=50)

            p1 = tabla.cell(0,0).add_paragraph("QUIEN ENTREGA")
            p1.runs[0].bold = True
            p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            quitar_espaciado_parrafo(p1)
            p2 = tabla.cell(0,1).add_paragraph("QUIEN RECIBE")
            p2.runs[0].bold = True
            p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            quitar_espaciado_parrafo(p2)
            p3 = tabla.cell(1,0).add_paragraph(arrendatario)
            p3.runs[0].bold = True
            p3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            quitar_espaciado_parrafo(p3)
            p4 = tabla.cell(1,1).add_paragraph(asesor)
            p4.runs[0].bold = True
            p4.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            quitar_espaciado_parrafo(p4)
            p5 = tabla.cell(2,0).add_paragraph(f"{tipo_doc_arrendatario} {cc_arrendatario}")
            p5.runs[0].bold = True
            p5.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            quitar_espaciado_parrafo(p5)
            cell = tabla.cell(2,1)
            p6 = cell.add_paragraph("ASESOR COMERCIAL - SUCURSAL UNICENTRO")
            p6.runs[0].bold = True
            p6.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            quitar_espaciado_parrafo(p6)
            p7 = cell.add_paragraph("PROMOTORA INMOBILIARIA R&G S.A.S.")
            p7.runs[0].bold = True
            p7.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            quitar_espaciado_parrafo(p7)
            p8 = cell.add_paragraph("NIT: 800.239.928-9")
            p8.runs[0].bold = True
            p8.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            quitar_espaciado_parrafo(p8)
            for row in tabla.rows:
                for cell in row.cells:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            import os
            print(f"[DEBUG] Intentando guardar el archivo en: {os.path.abspath(archivo)}")
            documento.save(archivo)
            print(f"[DEBUG] Archivo guardado exitosamente en: {os.path.abspath(archivo)}")
            messagebox.showinfo("Éxito", f"Acta generada correctamente como {archivo}")
        except Exception as e:
            import traceback
            print("Error al guardar el archivo:", e)
            traceback.print_exc()
            messagebox.showerror("Error", f"No se pudo guardar el archivo: {e}")

    def marcar_si_no(self, valor):
        if valor.upper() == "SI":
            return "SI  X     NO  ____"
        elif valor.upper() == "NO":
            return "SI  ____  NO  X"
        else:
            return "SI  ____  NO  ____"

if __name__ == "__main__":
    app = RestitucionApp()
    app.mainloop()
