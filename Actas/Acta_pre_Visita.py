
# Interfaz gráfica para Acta de Pre-Visita
import customtkinter as ctk
import tkinter
from tkinter import messagebox
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")

class PreVisitaApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Acta de Pre-Visita")
        self.geometry("600x700")
        self.resizable(False, False)

        # Frame principal
        self.frame = ctk.CTkFrame(self)
        self.frame.pack(padx=20, pady=20, fill="both", expand=True)

        # Título
        self.label_title = ctk.CTkLabel(self.frame, text="ACTA DE PRE-VISITA", font=("Arial Narrow", 22, "bold"))
        self.label_title.pack(pady=(10, 20))

        # Campos de entrada
        self.nombre_archivo = ctk.CTkEntry(self.frame, placeholder_text="Nombre del archivo (ej: ACTA_PRE_VISITA_FINAL.docx)")
        self.nombre_archivo.pack(pady=5, fill="x")

        self.sucursal_var = ctk.StringVar()
        self.label_sucursal = ctk.CTkLabel(self.frame, text="Representante:")
        self.label_sucursal.pack(anchor="w")
        self.sucursal_dropdown = ctk.CTkComboBox(self.frame, variable=self.sucursal_var, values=[
            "FREDDY HERNANDO GUERRERO RIVERA",
            "NELSON ALEJANDRO GUERRERO RIVERA",
            "CINDY JOHANNA GUERRERO RIVERA"
        ])
        self.sucursal_dropdown.pack(pady=5, fill="x")

        self.fecha = ctk.CTkEntry(self.frame, placeholder_text="Fecha de pre-visita")
        self.fecha.pack(pady=5, fill="x")
        self.direccion = ctk.CTkEntry(self.frame, placeholder_text="Dirección")
        self.direccion.pack(pady=5, fill="x")
        self.ciudad = ctk.CTkEntry(self.frame, placeholder_text="Ciudad")
        self.ciudad.pack(pady=5, fill="x")
        self.nombre_persona = ctk.CTkEntry(self.frame, placeholder_text="Nombre de la persona")
        self.nombre_persona.pack(pady=5, fill="x")
        self.tipo_doc = ctk.CTkEntry(self.frame, placeholder_text="Tipo de documento")
        self.tipo_doc.pack(pady=5, fill="x")
        self.num_doc = ctk.CTkEntry(self.frame, placeholder_text="Número de documento")
        self.num_doc.pack(pady=5, fill="x")
        self.correo1 = ctk.CTkEntry(self.frame, placeholder_text="Correo principal")
        self.correo1.pack(pady=5, fill="x")
        self.correo2 = ctk.CTkEntry(self.frame, placeholder_text="Correo administrativo")
        self.correo2.pack(pady=5, fill="x")
        self.observaciones = ctk.CTkEntry(self.frame, placeholder_text="Observaciones")
        self.observaciones.pack(pady=5, fill="x")
        self.nombre_asesor = ctk.CTkEntry(self.frame, placeholder_text="Nombre asesor")
        self.nombre_asesor.pack(pady=5, fill="x")

        self.btn_generar = ctk.CTkButton(self.frame, text="Generar Acta", command=self.generar_acta)
        self.btn_generar.pack(pady=20)

    def generar_acta(self):
        # Obtener valores
        archivo = self.nombre_archivo.get().strip()
        if archivo and not archivo.lower().endswith('.docx'):
            archivo += '.docx'
        Surcursal2_letra = self.sucursal_var.get().strip()
        Fecha_previsita = self.fecha.get().strip()
        Direccion_previsita = self.direccion.get().strip()
        Ciudad_previsita = self.ciudad.get().strip()
        Nombre_previsita = self.nombre_persona.get().strip()
        tipo_cedula_previsita = self.tipo_doc.get().strip()
        cedula_previsita = self.num_doc.get().strip()
        correo1 = self.correo1.get().strip()
        correo2 = self.correo2.get().strip()
        observaciones = self.observaciones.get().strip()
        nombre_asesor = self.nombre_asesor.get().strip()

        if not archivo:
            messagebox.showerror("Error", "Debes ingresar el nombre del archivo.")
            return
        if not Surcursal2_letra or not Fecha_previsita or not Direccion_previsita or not Ciudad_previsita or not Nombre_previsita or not tipo_cedula_previsita or not cedula_previsita or not correo1 or not correo2 or not nombre_asesor:
            messagebox.showerror("Error", "Todos los campos son obligatorios.")
            return

        # Lógica de asignación de representante
        if Surcursal2_letra == 'FREDDY HERNANDO GUERRERO RIVERA':
            Nombre_Promotora = "FREDY HERNANDO GUERRERO RIVERA"
            cc_rep_promotora = "79.509.071"
        elif Surcursal2_letra == "NELSON ALEJANDRO GUERRERO RIVERA":
            Nombre_Promotora = "NELSON ALEJANDRO GUERRERO RIVERA"
            cc_rep_promotora = "79.946.997"
        elif Surcursal2_letra == "CINDY JOHANNA GUERRERO RIVERA":
            Nombre_Promotora = "CINDY JOHANNA GUERRERO RIVERA"
            cc_rep_promotora = "1.030.552.945"
        else:
            messagebox.showerror("Error", "Selecciona un representante válido.")
            return

        # Crear documento
        documentoObjeto = docx.Document()
        style = documentoObjeto.styles['Normal']
        style.font.name = 'Arial Narrow'
        style.font.size = Pt(11)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        sections = documentoObjeto.sections
        for section in sections:
            section.top_margin = Cm(3)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(3)

        # Título
        parrafo = documentoObjeto.add_paragraph('')
        parrafo.add_run('ACTA DE PRE-VISITA').bold = True
        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Párrafo inicial
        p0 = documentoObjeto.add_paragraph()
        p0.add_run('En ' + Ciudad_previsita + ' el día ' + Fecha_previsita + ', ')
        p0.add_run('PROMOTORA INMOBILIARIA R&G S.A.S. ').bold = True
        p0.add_run('identificada con NIT: 800.239.928-9, sociedad representada legalmente por ')
        p0.add_run(Nombre_Promotora).bold = True
        p0.add_run(', identificado con cédula de ciudadanía ')
        p0.add_run(cc_rep_promotora).bold = True
        p0.add_run(' de Bogotá D.C. se realiza la pre-visita al inmueble ubicado en la ')
        p0.add_run(Direccion_previsita).bold = True
        p0.add_run(', en la ciudad de ' + Ciudad_previsita + ', a la persona ')
        p0.add_run(Nombre_previsita).bold = True
        p0.add_run(', identificada con número de ' + tipo_cedula_previsita + ' ')
        p0.add_run(cedula_previsita).bold = True
        p0.add_run(' ' + Ciudad_previsita)

        p1 = documentoObjeto.add_paragraph(
            "Mediante el siguiente documento se hace una verificación actual de las condiciones reales en las que se encuentra el inmueble, "
            "con el fin de verificar que todo esté acorde al inventario inicial y, en caso de presentar inconsistencias en el mismo, "
            "el arrendatario se compromete a hacer los arreglos y ajustes pertinentes del inmueble para poder programar la restitución definitiva."
        )
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        def agregar_vineta(textos):
            p = documentoObjeto.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run("•\t")
            run.bold = False
            for texto, bold in textos:
                run = p.add_run(texto)
                run.bold = bold
            return p

        agregar_vineta([
            ("Si para la restitución, el/la arrendatario(a) ", False),
            ("NO", True),
            (" puede estar presente, debe enviar al correo ", False),
            (correo1, True),
            (" la debida autorización, informando los datos básicos de la persona que hará dicha diligencia, de ", False),
            ("NO", True),
            (" enviar esta autorización no se podrá recibir el inmueble.", False),
        ])
        agregar_vineta([
            ("Recuerde que una vez teniendo la fecha tentativa de restitución, debe enviar al correo ", False),
            (correo2, True),
            (", con antelación de ", False),
            ("OCHO (08)", True),
            (" días calendario, los recibos de cada servicio público que cuente el inmueble, estos recibos deben enviarse legibles, escaneados, por ambas caras de los TRES (03) últimos meses, con su respectivo soporte de pago, para que pueda tener un estimado de los valores que debe dejar por provisión de servicios.", False),
        ])
        agregar_vineta([
            ("Es ", False),
            ("OBLIGACIÓN", True),
            (" del arrendatario estar al día con todos los pagos; (días de arrendamiento, canon total de arrendamiento, cláusula penal, ", False),
            ("(según sea el caso)", True),
            (") y la debida provisión de servicios), UN (01) día antes de la fecha oficial de restitución de inmueble, si no se evidencia el pago no podrá ser recibido el inmueble.", False),
        ])
        agregar_vineta([
            ("El arrendatario ", False),
            ("SE OBLIGA Y SE COMPROMETE", True),
            (" a gestionar y presentar ante la inmobiliaria el certificado y/o radicado del ", False),
            ("TRASLADO Y/O SUSPENSIÓN DEFINITIVA", True),
            (", de los servicios derivados del internet o línea telefónica que haya solicitado en el inmueble, por lo que se hace ", False),
            ("ÉNFASIS", True),
            (" en que se debe gestionar lo pertinente ", False),
            ("con anticipación", True),
            (" ante las empresas concernientes, si a la fecha de la restitución no se han tramitado dichas solicitudes, hasta no tener soporte pertinente de traslado y/o suspensión definitiva, se seguirán causando los días de arrendamiento.", False),
        ])
        agregar_vineta([
            ("EL ARRENDATARIO, acepta y reconoce que son claras las especificaciones y recomendaciones consignadas en este documento, "
             "las cuales se compromete a cumplir a cabalidad para evitar retrasos en la restitución de inmueble que le conlleven gastos adicionales.", False),
        ])
        agregar_vineta([
            ("Sr ARRENDATARIO (A), tenga en cuenta que, dentro de su contrato de arrendamiento hay una cláusula que dice RESTITUCIÓN DE INMUEBLE", False),
        ])
        agregar_vineta([
            ("Tenga en cuenta que, posterior a la restitución de inmueble, se contarán cinco (05) días calendario con el fin que la PROPIETARIA (O) pueda verificar "
             "las condiciones en que fue restituido, por eso, hacemos énfasis en que el inmueble debe estar en óptimas condiciones de aseo, óptimas condiciones de pintura, "
             "esto incluye (paredes generales, paredes dentro de los closets, baños cuando aplique), en general todo debidamente pintado de blanco.", False),
        ])

        documentoObjeto.add_paragraph(
            "Se enfatiza que la restitución del inmueble en todos los casos la deben hacer en las mismas buenas condiciones que lo recibió inicialmente."
        )

        p2 = documentoObjeto.add_paragraph()
        p2.add_run('OBSERVACIONES ADICIONALES PARA TENER EN CUENTA ENCONTRADAS EN LA PRE-VISITA:\n').bold = True
        p2.add_run(observaciones)

        def quitar_margenes_celda(cell, top=0, left=0, bottom=0, right=0):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for m, v in zip(("top", "left", "bottom", "right"), (top, left, bottom, right)):
                node = OxmlElement(f"w:{m}")
                node.set(qn('w:w'), str(v))
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)

        def quitar_espaciado_parrafo(paragraph):
            p_fmt = paragraph.paragraph_format
            p_fmt.space_before = Pt(0)
            p_fmt.space_after = Pt(0)
            p_fmt.line_spacing = 1.0

        tabla = documentoObjeto.add_table(rows=3, cols=2)
        tabla.autofit = True
        for row in tabla.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    border = OxmlElement(f"w:{border_name}")
                    border.set(qn("w:val"), "nil")
                    tcBorders.append(border)
                tcPr.append(tcBorders)
                quitar_margenes_celda(cell, top=50, left=50, bottom=50, right=50)

        p1 = tabla.cell(0,0).add_paragraph("QUIEN RECIBE LA VISITA")
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        quitar_espaciado_parrafo(p1)
        p2 = tabla.cell(0,1).add_paragraph("ASESOR INMOBILIARIO")
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        quitar_espaciado_parrafo(p2)
        p3 = tabla.cell(1,0).add_paragraph(Nombre_previsita)
        p3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        quitar_espaciado_parrafo(p3)
        p4 = tabla.cell(1,1).add_paragraph(nombre_asesor)
        p4.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        quitar_espaciado_parrafo(p4)
        p5 = tabla.cell(2,0).add_paragraph(tipo_cedula_previsita+" "+cedula_previsita)
        p5.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        quitar_espaciado_parrafo(p5)

        # --- Fila 3, columna 2: datos de empresa/asesor ---
        cell = tabla.cell(2,1)
        p6 = cell.add_paragraph("ASESOR COMERCIAL - SUCURSAL UNICENTRO")
        p6.runs[0].bold = True
        p6.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        quitar_espaciado_parrafo(p6)

        p7 = cell.add_paragraph("PROMOTORA INMOBILIARIA R&G S.A.S.")
        p7.runs[0].bold = True
        p7.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        quitar_espaciado_parrafo(p7)

        p8 = cell.add_paragraph("NIT: 800.239.928-9")
        p8.runs[0].bold = True
        p8.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        quitar_espaciado_parrafo(p8)

        for row in tabla.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        import os
        print(f"[DEBUG] Intentando guardar el archivo en: {os.path.abspath(archivo)}")
        try:
            documentoObjeto.save(archivo)
            print(f"[DEBUG] Archivo guardado exitosamente en: {os.path.abspath(archivo)}")
            messagebox.showinfo("Éxito", f"Acta generada correctamente como {archivo}")
        except Exception as e:
            import traceback
            print("Error al guardar el archivo:", e)
            traceback.print_exc()
            messagebox.showerror("Error", f"No se pudo guardar el archivo: {e}")


if __name__ == "__main__":
    app = PreVisitaApp()
    app.mainloop()

