import customtkinter as ctk
import tkinter
from tkinter import messagebox
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")

class EntregaApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Acta de Entrega")
        self.geometry("700x900")
        self.minsize(700, 900)
        self.resizable(True, True)

        # Crear el sistema de pestañas
        self.tabview = ctk.CTkTabview(self, width=680, height=860)
        self.tabview.pack(fill="both", expand=True, padx=10, pady=10)

        # Crear las pestañas
        self.tabview.add("Datos Generales")
        self.tabview.add("Detalles y Entrega")

        # --- PESTAÑA 1: DATOS GENERALES ---
        tab1 = self.tabview.tab("Datos Generales")
        self.label_title1 = ctk.CTkLabel(tab1, text="ACTA DE ENTREGA", font=("Arial Narrow", 22, "bold"))
        self.label_title1.pack(pady=(20, 15))

        self.nombre_archivo = ctk.CTkEntry(tab1, placeholder_text="Nombre del archivo (ej: ACTA_ENTREGA_FINAL.docx)")
        self.nombre_archivo.pack(pady=5, padx=15, fill="x")

        self.sucursal_var = ctk.StringVar()
        self.label_sucursal = ctk.CTkLabel(tab1, text="Representante:")
        self.label_sucursal.pack(anchor="w", padx=15, pady=(10,0))
        self.sucursal_dropdown = ctk.CTkComboBox(tab1, variable=self.sucursal_var, values=[
            "FREDDY HERNANDO GUERRERO RIVERA",
            "NELSON ALEJANDRO GUERRERO RIVERA",
            "CINDY JOHANNA GUERRERO RIVERA"
        ])
        self.sucursal_dropdown.pack(pady=5, padx=15, fill="x")

        self.ciudad = ctk.CTkEntry(tab1, placeholder_text="Ciudad")
        self.ciudad.pack(pady=5, padx=15, fill="x")

        self.fecha = ctk.CTkEntry(tab1, placeholder_text="Fecha (ej: cinco (05) de agosto del año 2025)")
        self.fecha.pack(pady=5, padx=15, fill="x")

        self.direccion = ctk.CTkEntry(tab1, placeholder_text="Dirección del inmueble")
        self.direccion.pack(pady=5, padx=15, fill="x")

        self.arrendatario = ctk.CTkEntry(tab1, placeholder_text="Nombre arrendatario")
        self.arrendatario.pack(pady=5, padx=15, fill="x")

        self.tipo_doc_arr = ctk.CTkEntry(tab1, placeholder_text="Tipo documento arrendatario")
        self.tipo_doc_arr.pack(pady=5, padx=15, fill="x")

        self.cc_arr = ctk.CTkEntry(tab1, placeholder_text="Número documento arrendatario")
        self.cc_arr.pack(pady=5, padx=15, fill="x")

        self.ciudad_doc_arr = ctk.CTkEntry(tab1, placeholder_text="Ciudad de expedición documento arrendatario")
        self.ciudad_doc_arr.pack(pady=5, padx=15, fill="x")

        self.autorizado_nombre = ctk.CTkEntry(tab1, placeholder_text="Nombre autorizado (si aplica)")
        self.autorizado_nombre.pack(pady=5, padx=15, fill="x")

        self.tipo_doc_aut = ctk.CTkEntry(tab1, placeholder_text="Tipo documento autorizado")
        self.tipo_doc_aut.pack(pady=5, padx=15, fill="x")

        self.autorizado_cc = ctk.CTkEntry(tab1, placeholder_text="Número documento autorizado")
        self.autorizado_cc.pack(pady=5, padx=15, fill="x")

        # --- PESTAÑA 2: DETALLES Y ENTREGA ---
        tab2 = self.tabview.tab("Detalles y Entrega")
        self.label_title2 = ctk.CTkLabel(tab2, text="DETALLES DE ENTREGA", font=("Arial Narrow", 20, "bold"))
        self.label_title2.pack(pady=(20, 15))

        self.juegos_llaves = ctk.CTkEntry(tab2, placeholder_text="Número de juegos de llaves convencionales")
        self.juegos_llaves.pack(pady=5, padx=15, fill="x")

        self.llaves_seguridad = ctk.CTkEntry(tab2, placeholder_text="Número de llaves de seguridad")
        self.llaves_seguridad.pack(pady=5, padx=15, fill="x")

        self.llaves_otras = ctk.CTkEntry(tab2, placeholder_text="Otros entregados (pines, tarjetas, etc.)")
        self.llaves_otras.pack(pady=5, padx=15, fill="x")

        self.resanado_var = ctk.StringVar()
        self.label_resanado = ctk.CTkLabel(tab2, text="¿El inmueble se entrega resanado? (SI/NO):")
        self.label_resanado.pack(anchor="w", padx=15, pady=(10,0))
        self.resanado = ctk.CTkComboBox(tab2, variable=self.resanado_var, values=["SI", "NO"])
        self.resanado.pack(pady=5, padx=15, fill="x")

        self.pintado_var = ctk.StringVar()
        self.label_pintado = ctk.CTkLabel(tab2, text="¿El inmueble se entrega recién pintado? (SI/NO):")
        self.label_pintado.pack(anchor="w", padx=15, pady=(10,0))
        self.pintado = ctk.CTkComboBox(tab2, variable=self.pintado_var, values=["SI", "NO"])
        self.pintado.pack(pady=5, padx=15, fill="x")

        self.aseado_var = ctk.StringVar()
        self.label_aseado = ctk.CTkLabel(tab2, text="¿El inmueble se entrega aseado? (SI/NO):")
        self.label_aseado.pack(anchor="w", padx=15, pady=(10,0))
        self.aseado = ctk.CTkComboBox(tab2, variable=self.aseado_var, values=["SI", "NO"])
        self.aseado.pack(pady=5, padx=15, fill="x")

        self.observaciones_finales = ctk.CTkEntry(tab2, placeholder_text="Observaciones adicionales")
        self.observaciones_finales.pack(pady=5, padx=15, fill="x")

        self.nombre_asesor = ctk.CTkEntry(tab2, placeholder_text="Nombre asesor inmobiliario")
        self.nombre_asesor.pack(pady=5, padx=15, fill="x")

        self.btn_generar = ctk.CTkButton(tab2, text="Generar Acta", width=150, height=35, font=("Arial", 14, "bold"), command=self.generar_acta)
        self.btn_generar.pack(pady=30)

    def generar_acta(self):
        archivo = self.nombre_archivo.get().strip()
        if archivo and not archivo.lower().endswith('.docx'):
            archivo += '.docx'
        Surcursal2_letra = self.sucursal_var.get().strip()
        ciudad = self.ciudad.get().strip()
        fecha = self.fecha.get().strip()
        direccion_inmueble = self.direccion.get().strip()
        arrendatario = self.arrendatario.get().strip()
        tipo_doc_arrendatario = self.tipo_doc_arr.get().strip()
        cc_arrendatario = self.cc_arr.get().strip()
        ciudad_doc_arrendatario = self.ciudad_doc_arr.get().strip()
        autorizado_nombre = self.autorizado_nombre.get().strip()
        tipo_doc_autorizado = self.tipo_doc_aut.get().strip()
        autorizado_cc = self.autorizado_cc.get().strip()
        juegos_llaves = self.juegos_llaves.get().strip()
        llaves_seguridad = self.llaves_seguridad.get().strip()
        llaves_otras = self.llaves_otras.get().strip()
        resanado = self.resanado_var.get().strip()
        pintado = self.pintado_var.get().strip()
        aseado = self.aseado_var.get().strip()
        observaciones_finales = self.observaciones_finales.get().strip()
        asesor = self.nombre_asesor.get().strip()

        if not archivo:
            messagebox.showerror("Error", "Debes ingresar el nombre del archivo.")
            return
        if not Surcursal2_letra or not ciudad or not fecha or not direccion_inmueble or not arrendatario or not tipo_doc_arrendatario or not cc_arrendatario or not ciudad_doc_arrendatario or not asesor:
            messagebox.showerror("Error", "Todos los campos obligatorios deben estar completos.")
            return

        # Lógica de asignación de representante
        if Surcursal2_letra == 'FREDDY HERNANDO GUERRERO RIVERA':
            representante = "FREDY HERNANDO GUERRERO RIVERA"
            cc_representante = "79.509.071"
        elif Surcursal2_letra == "NELSON ALEJANDRO GUERRERO RIVERA":
            representante = "NELSON ALEJANDRO GUERRERO RIVERA"
            cc_representante = "79.946.997"
        elif Surcursal2_letra == "CINDY JOHANNA GUERRERO RIVERA":
            representante = "CINDY JOHANNA GUERRERO RIVERA"
            cc_representante = "1.030.552.945"
        else:
            messagebox.showerror("Error", "Selecciona un representante válido.")
            return

        tipo_doc_representante = "CC."

        try:
            documento = docx.Document()
            style = documento.styles['Normal']
            style.font.name = 'Arial Narrow'
            style.font.size = Pt(11)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            sections = documento.sections
            for section in sections:
                section.top_margin = Cm(3)
                section.bottom_margin = Cm(2.5)
                section.left_margin = Cm(3)
                section.right_margin = Cm(3)

            # Título
            p_titulo = documento.add_paragraph()
            p_titulo.add_run("ACTA DE ENTREGA").bold = True
            p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Párrafo inicial
            p0 = documento.add_paragraph()
            p0.add_run(f"En {ciudad}, el día {fecha}, ").bold = False
            p0.add_run("PROMOTORA INMOBILIARIA R&G S.A.S. ").bold = True
            p0.add_run("identificada con NIT: ").bold = False
            p0.add_run("800.239.928-9").bold = True
            p0.add_run(", sociedad representada legalmente por ")
            p0.add_run(representante).bold = True
            p0.add_run(f", identificado con {tipo_doc_representante} No. {cc_representante} de {ciudad}, realiza la entrega del inmueble ubicado en ")
            p0.add_run(direccion_inmueble).bold = True
            p0.add_run(f", en la ciudad de {ciudad}, a ")
            p0.add_run(arrendatario).bold = True
            p0.add_run(f", identificado(a) con {tipo_doc_arrendatario} No. {cc_arrendatario} {ciudad_doc_arrendatario}.")

            # Texto previo a viñetas
            p1 = documento.add_paragraph()
            p1.add_run("En caso de ausencia, el ")
            p1.add_run("ARRENDATARIO/A").bold = True
            p1.add_run(" de manera explícita autoriza a la diligencia de entrega de inmueble a la siguiente persona:")

            # Autorización y entregas en viñetas
            def agregar_vineta(textos):
                p = documento.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.first_line_indent = Cm(-0.5)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = p.add_run("•\t")
                for texto, bold in textos:
                    run = p.add_run(texto)
                    run.bold = bold
                return p

            agregar_vineta([(f"El(la) señor(a) {autorizado_nombre}, identificado(a) con {tipo_doc_autorizado} No. {autorizado_cc}, a quien se le entrega real y material del inmueble, de igual forma también se entregan:", False)])
            agregar_vineta([(f"Juegos de llaves convencionales {juegos_llaves}", False)])
            agregar_vineta([(f"Llaves de seguridad {llaves_seguridad}", False)])
            agregar_vineta([(f"Otros (pines, tarjetas, controles) {llaves_otras}", False)])

            # Texto adicional
            documento.add_paragraph(
                "Adicionalmente, se anexan a este documento, fotos y videos de registro como evidencia en forma digital."
            )

            # Texto previo a viñetas
            p2 = documento.add_paragraph()
            p2.add_run("SEÑOR ").bold = True
            p2.add_run("ARRENDATARIO (A),").bold = True
            p2.add_run(" leer por favor detenidamente el proceso que a continuación se menciona, con el fin de que pueda dirigir sus solicitudes a las áreas pertinentes.").bold = True

            # Viñetas del proceso
            agregar_vineta([
                ("El inmueble se entrega sin línea telefónica, plan de internet y similares; sin embargo, el arrendatario está autorizado a instalarlas en el inmueble siempre y cuando las retire, o realice su traslado cuando restituya el inmueble. Por tal motivo al momento de restituir el inmueble debe presentar los PAZ y SALVOS correspondientes o el traslado con soporte.", False)
            ])
            agregar_vineta([
                ("Los recibos que lleguen de periodos anteriores a la fecha de la entrega del inmueble, deberán ser cancelados por el arrendatario y este, a su vez, los enviará debidamente escaneados y legibles por ambas caras con el soporte de pago, al correo ", False),
                ("unicentro.administrativo@rginmobiliaria.com.co", True),
                (", para que se puedan verificar los periodos correspondientes; los valores que corresponda, solo serán descontados, si EL ARRENDATARIO envía la información antes del 28 del mes que se encuentre en curso, puesto que de no recibirla, pasado un (01) mes posterior a la entrega NO se harán ajustes o devoluciones, siendo responsabilidad del ARRENDATARIO enviar según lo mencionado en este documento para el proceso, solo se tendrán en cuenta los recibos que hagan llegar al correo, vía WhatsApp, no serán tenidos en cuenta.", False)
            ])
            agregar_vineta([
                ("Toda correspondencia que llegue al inmueble y que corresponda a éste tales como: impuestos, notificaciones, información de asambleas, entre otras que correspondan al propietario, deberán ser enviadas y/o comunicadas por el Arrendatario a las oficinas del Arrendador, mediante correo certificado.", False)
            ])
            agregar_vineta([
                ("El inmueble se entrega en buen estado, funcional, el inmueble ", False),
                ("NO", True),
                (", por tanto, los posibles arreglos que surjan no están en total obligatoriedad por parte del propietario de realizarlos.", False)
            ])
            agregar_vineta([
                ("El inmueble debe ser restituido en su momento en las mismas buenas condiciones que le fue entregado, por lo que se recomienda tener en cuenta estas observaciones y en su momento no generar contratiempos.", False)
            ])
            agregar_vineta([
                ("Se entrega, VOLANTE DE INSTRUCTIVOS pertinentes, para el pago mensual de la obligación contraída.", False)
            ])
            agregar_vineta([
                ("Se entrega Volante informativo de BIENVENIDA de cada una de las áreas, con el fin de que usted como arrendatario sepa a qué área dirigirse en caso de alguna novedad, duda e inquietud.", False)
            ])
            agregar_vineta([
                ("Está PROHIBIDO, transferir cualquier tipo de crédito de consumo o similar a cualquier servicio público; dentro del contrato de arrendamiento se considera incumplimiento, y en caso de evidenciar omisión a la información se ejecutará el contrato de la manera jurídica que corresponda.", False)
            ])

            # Condiciones inmueble con SI/NO
            documento.add_paragraph(f"¿El inmueble se entrega resanado?   {self.marcar_si_no(resanado)}")
            documento.add_paragraph(f"¿El inmueble se entrega recién pintado?   {self.marcar_si_no(pintado)}")
            documento.add_paragraph(f"¿El inmueble se entrega aseado?   {self.marcar_si_no(aseado)}")

            # Observaciones adicionales
            documento.add_paragraph("Observaciones Adicionales Encontradas:")
            documento.add_paragraph(observaciones_finales)

            def quitar_margenes_celda(cell, top=0, left=0, bottom=0, right=0):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcMar = OxmlElement('w:tcMar')
                for m, v in zip(("top", "left", "bottom", "right"), (top, left, bottom, right)):
                    node = OxmlElement(f"w:{m}")
                    node.set(qn('w:w'), str(v))
                    node.set(qn('w:type'), 'dxa')
                    tcMar.append(node)
                tcPr.append(tcMar)

            def quitar_espaciado_parrafo(paragraph):
                p_fmt = paragraph.paragraph_format
                p_fmt.space_before = Pt(0)
                p_fmt.space_after = Pt(0)
                p_fmt.line_spacing = 1.0

            tabla = documento.add_table(rows=3, cols=2)
            tabla.autofit = True
            for row in tabla.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = OxmlElement('w:tcBorders')
                    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
                        border = OxmlElement(f"w:{border_name}")
                        border.set(qn("w:val"), "nil")
                        tcBorders.append(border)
                    tcPr.append(tcBorders)
                    quitar_margenes_celda(cell, top=50, left=50, bottom=50, right=50)

            p1 = tabla.cell(0,0).add_paragraph("QUIEN RECIBE")
            p1.runs[0].bold = True
            quitar_espaciado_parrafo(p1)
            p2 = tabla.cell(0,1).add_paragraph("QUIEN ENTREGA")
            p2.runs[0].bold = True
            quitar_espaciado_parrafo(p2)
            p3 = tabla.cell(1,0).add_paragraph(arrendatario)
            p3.runs[0].bold = True
            quitar_espaciado_parrafo(p3)
            p4 = tabla.cell(1,1).add_paragraph(asesor)
            p4.runs[0].bold = True
            quitar_espaciado_parrafo(p4)
            p5 = tabla.cell(2,0).add_paragraph(f"{tipo_doc_arrendatario} {cc_arrendatario}")
            p5.runs[0].bold = True
            quitar_espaciado_parrafo(p5)
            cell = tabla.cell(2,1)
            p6 = cell.add_paragraph("ASESOR COMERCIAL - SUCURSAL UNICENTRO")
            p6.runs[0].bold = True
            quitar_espaciado_parrafo(p6)
            p7 = cell.add_paragraph("PROMOTORA INMOBILIARIA R&G S.A.S.")
            p7.runs[0].bold = True
            quitar_espaciado_parrafo(p7)
            p8 = cell.add_paragraph("NIT: 800.239.928-9")
            p8.runs[0].bold = True
            quitar_espaciado_parrafo(p8)
            for row in tabla.rows:
                for cell in row.cells:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            import os
            print(f"[DEBUG] Intentando guardar el archivo en: {os.path.abspath(archivo)}")
            documento.save(archivo)
            print(f"[DEBUG] Archivo guardado exitosamente en: {os.path.abspath(archivo)}")
            messagebox.showinfo("Éxito", f"Acta generada correctamente como {archivo}")
        except Exception as e:
            import traceback
            print("Error al guardar el archivo:", e)
            traceback.print_exc()
            messagebox.showerror("Error", f"No se pudo guardar el archivo: {e}")

    def marcar_si_no(self, valor):
        if valor.upper() == "SI":
            return "SI  X     NO  ____"
        elif valor.upper() == "NO":
            return "SI  ____  NO  X"
        else:
            return "SI  ____  NO  ____"

if __name__ == "__main__":
    app = EntregaApp()
    app.mainloop()
